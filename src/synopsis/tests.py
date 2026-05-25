from datetime import date, datetime, timedelta
import importlib
import io
import json
import re
from urllib.parse import urlparse
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

from django.conf import settings
import jwt
from django.contrib.auth.models import Group, User, AnonymousUser
from django.core import mail
from django.core.cache import cache
from django.core.mail import EmailMessage
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings, RequestFactory, SimpleTestCase
from django.contrib.messages import get_messages
from django.contrib.messages.storage.fallback import FallbackStorage
from django.contrib.sessions.middleware import SessionMiddleware
from django.core.exceptions import PermissionDenied
from django.core.management.base import CommandError
from django.urls import reverse

import shutil
import tempfile
import textwrap
from django.utils import timezone
from django.core.management import call_command

from .models import (
    AdvisoryBoardInvitation,
    AdvisoryBoardMember,
    AdvisoryBoardCustomField,
    AdvisoryBoardCustomFieldValueHistory,
    Funder,
    IUCNCategory,
    Project,
    ProjectChangeLog,
    ProtocolFeedback,
    Protocol,
    ProtocolRevision,
    ActionList,
    ActionListRevision,
    ActionListFeedback,
    CollaborativeSession,
    UserRole,
    LibraryReference,
    LibraryReferenceFolderHistory,
    LibraryImportBatch,
    ReferenceSourceBatch,
    ReferenceSourceBatchNoteHistory,
    Reference,
    ReferenceSummary,
    ReferenceSummaryComment,
    ReferenceActionSummary,
    SynopsisChapter,
    SynopsisSubheading,
    SynopsisIntervention,
    SynopsisInterventionKeyMessage,
    SynopsisAssignment,
    SynopsisFeedback,
)
from .forms import (
    ActionListReminderScheduleForm,
    ActionListSendForm,
    AdvisoryBulkInviteForm,
    AdvisoryBoardMemberForm,
    AdvisoryInviteForm,
    AdvisoryMemberCustomDataForm,
    FunderForm,
    FunderContactFormSet,
    ProtocolSendForm,
    ProtocolReminderScheduleForm,
    ProjectDeleteForm,
    ProjectSettingsForm,
    ReminderScheduleForm,
    IUCN_ACTION_CHOICES,
    IUCN_HABITAT_CHOICES,
    IUCN_THREAT_CHOICES,
    ReferenceSummaryDraftForm,
    ReferenceSummaryUpdateForm,
)
from .email_backends import AttachmentSummaryConsoleEmailBackend
from .utils import (
    BRAND,
    GLOBAL_GROUPS,
    default_action_list_review_message,
    default_advisory_invitation_message,
    default_protocol_review_message,
    default_synopsis_review_message,
    email_subject,
    ensure_global_groups,
    reference_summary_effective_citation,
    reference_hash,
    reply_to_list,
    split_inline_italic_markup,
)
from .views import (
    _advisory_board_context,
    _apply_revision_to_action_list,
    _apply_revision_to_protocol,
    _build_advisory_invitation_email,
    _build_onlyoffice_config,
    _download_onlyoffice_file,
    _parse_onlyoffice_callback,
    _create_protocol_feedback,
    _format_deadline,
    _format_value,
    _funder_contact_label,
    _log_project_change,
    _normalise_import_record,
    _user_can_confirm_phase,
    _user_is_manager,
    _user_can_edit_project,
    protocol_delete_revision,
    action_list_delete_revision,
    _parse_plaintext_references,
    _parse_endnote_xml,
    project_synopsis_structure,
    _intervention_reference_numbering,
    _format_reference_number_ranges,
    _generate_synopsis_docx,
    _link_library_references_to_project,
    _reference_export_citation,
    _reference_summary_paragraph,
    _structured_summary_paragraph,
    _project_reference_summary_ids,
)

# TODO: #25 Split this test module into smaller files once the current workflow areas stop moving around so it stays easier to navigate.


class EmailSubjectTests(TestCase):
    def setUp(self):
        self.project = SimpleNamespace(title="Coastal Restoration")

    def test_invite_subject_includes_due_date(self):
        due = timezone.make_aware(datetime(2025, 5, 1, 9, 30))
        expected_due = timezone.localtime(due).strftime("%d %b %Y %H:%M")
        subject = email_subject("invite", self.project, due)
        self.assertEqual(
            subject,
            f"[{BRAND}] Invitation to advise on {self.project.title} (reply by {expected_due})",
        )


@override_settings(EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend")
class AuthenticationFlowTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(
            username="author@example.com",
            email="author@example.com",
            password="StrongPass123!",
            first_name="Author",
        )
        self.manager = User.objects.create_user(
            username="manager@example.com",
            email="manager@example.com",
            password="StrongPass123!",
            is_staff=True,
        )
        self.project = Project.objects.create(title="Creation Assignment Project")

    def test_login_remember_me_controls_session_expiry(self):
        response = self.client.post(
            reverse("synopsis:login"),
            {"username": "author@example.com", "password": "StrongPass123!"},
        )
        self.assertRedirects(response, reverse("synopsis:dashboard"))
        self.assertTrue(self.client.session.get_expire_at_browser_close())

        self.client.post(reverse("synopsis:logout"))

        response = self.client.post(
            reverse("synopsis:login"),
            {
                "username": "author@example.com",
                "password": "StrongPass123!",
                "remember_me": "on",
            },
        )
        self.assertRedirects(response, reverse("synopsis:dashboard"))
        self.assertFalse(self.client.session.get_expire_at_browser_close())

    def test_login_allows_standard_django_username_for_superuser_style_accounts(self):
        root_user = User.objects.create_user(
            username="admin",
            email="admin@example.com",
            password="RootPass123!",
            is_staff=True,
            is_superuser=True,
        )

        response = self.client.post(
            reverse("synopsis:login"),
            {"username": "admin", "password": "RootPass123!"},
        )

        self.assertRedirects(response, reverse("synopsis:dashboard"))
        self.assertEqual(int(self.client.session["_auth_user_id"]), root_user.id)

    def test_logout_requires_post(self):
        self.client.login(username="author@example.com", password="StrongPass123!")
        response = self.client.get(reverse("synopsis:logout"))
        self.assertEqual(response.status_code, 405)

    def test_password_reset_request_sends_email(self):
        response = self.client.post(
            reverse("synopsis:password_reset"),
            {"email": "author@example.com"},
        )

        self.assertRedirects(response, reverse("synopsis:password_reset_done"))
        self.assertEqual(len(mail.outbox), 1)
        self.assertIn("/accounts/reset/", mail.outbox[0].body)
        self.assertEqual(mail.outbox[0].to, ["author@example.com"])

    def test_manager_create_user_sends_account_setup_email_and_allows_password_setup(self):
        self.client.login(username="manager@example.com", password="StrongPass123!")

        response = self.client.post(
            reverse("synopsis:user_create"),
            {
                "first_name": "New",
                "last_name": "Author",
                "email": "new.author@example.com",
                "global_role": "author",
            },
        )

        self.assertRedirects(response, reverse("synopsis:manager_dashboard"))
        created_user = User.objects.get(username="new.author@example.com")
        self.assertFalse(created_user.has_usable_password())
        self.assertEqual(created_user.email, "new.author@example.com")
        self.assertTrue(created_user.groups.filter(name="author").exists())
        self.assertEqual(len(mail.outbox), 1)
        self.assertEqual(mail.outbox[0].to, ["new.author@example.com"])
        self.assertIn("Set up your CE Synopsis Portal account", mail.outbox[0].subject)

        match = re.search(r"http://testserver(/accounts/reset/\S+)", mail.outbox[0].body)
        self.assertIsNotNone(match)
        reset_path = match.group(1)

        response = self.client.get(reset_path, follow=True)
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Set password")
        confirm_path = response.request["PATH_INFO"]

        response = self.client.post(
            confirm_path,
            {
                "new_password1": "EvenStrongerPass123!",
                "new_password2": "EvenStrongerPass123!",
            },
        )
        self.assertRedirects(response, reverse("synopsis:password_reset_complete"))

        created_user.refresh_from_db()
        self.assertTrue(created_user.has_usable_password())
        self.client.post(reverse("synopsis:logout"))
        self.assertTrue(
            self.client.login(
                username="new.author@example.com",
                password="EvenStrongerPass123!",
            )
        )

    def test_manager_can_create_external_author_with_assigned_synopsis(self):
        self.client.login(username="manager@example.com", password="StrongPass123!")

        response = self.client.post(
            reverse("synopsis:user_create"),
            {
                "first_name": "External",
                "last_name": "Author",
                "email": "external.author@example.com",
                "global_role": "external_collaborator",
                "assigned_projects": [str(self.project.id)],
            },
        )

        self.assertRedirects(response, reverse("synopsis:manager_dashboard"))
        created_user = User.objects.get(username="external.author@example.com")
        self.assertTrue(
            created_user.groups.filter(name="external_collaborator").exists()
        )
        self.assertTrue(
            UserRole.objects.filter(
                user=created_user, project=self.project, role="author"
            ).exists()
        )


@override_settings(EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend")
class ManagerUserManagementTests(TestCase):
    def setUp(self):
        ensure_global_groups()
        self.manager = User.objects.create_user(
            username="manager@example.com",
            email="manager@example.com",
            password="StrongPass123!",
            is_staff=True,
        )
        self.manager.groups.add(Group.objects.get(name="manager"))
        self.target = User.objects.create_user(
            username="target@example.com",
            email="target@example.com",
            password="StrongPass123!",
            first_name="Target",
            last_name="User",
        )
        self.target.groups.add(Group.objects.get(name="author"))
        self.pending_user = User.objects.create_user(
            username="pending@example.com",
            email="pending@example.com",
            first_name="Pending",
        )
        self.pending_user.set_unusable_password()
        self.pending_user.save(update_fields=["password"])
        self.pending_user.groups.add(Group.objects.get(name="external_collaborator"))
        self.project = Project.objects.create(title="Seagrass Pilot")
        self.superuser = User.objects.create_superuser(
            username="root",
            email="root@example.com",
            password="StrongPass123!",
        )
        self.client.login(username="manager@example.com", password="StrongPass123!")

    def test_manager_dashboard_removes_staff_column_and_shows_manage_actions(self):
        response = self.client.get(reverse("synopsis:manager_dashboard"))

        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, "Staff?")
        self.assertContains(response, "Global role")
        self.assertContains(response, "Access")
        self.assertContains(response, "Manage user")
        self.assertContains(response, "Protected")

    def test_manager_can_update_global_role_and_account_status(self):
        response = self.client.post(
            reverse("synopsis:manager_user_edit", args=[self.target.id]),
            {
                "action": "update_user",
                "first_name": "Updated",
                "last_name": "User",
                "email": "updated.target@example.com",
                "global_role": "external_collaborator",
                "is_active": "",
            },
        )

        self.assertRedirects(
            response, reverse("synopsis:manager_user_edit", args=[self.target.id])
        )
        self.target.refresh_from_db()
        self.assertEqual(self.target.username, "updated.target@example.com")
        self.assertEqual(self.target.email, "updated.target@example.com")
        self.assertFalse(self.target.is_active)
        self.assertFalse(self.target.is_staff)
        self.assertTrue(self.target.groups.filter(name="external_collaborator").exists())
        self.assertFalse(self.target.groups.filter(name="author").exists())

    def test_manager_can_send_password_reset_email_for_existing_account(self):
        response = self.client.post(
            reverse("synopsis:manager_user_edit", args=[self.target.id]),
            {"action": "send_access_email"},
        )

        self.assertRedirects(
            response, reverse("synopsis:manager_user_edit", args=[self.target.id])
        )
        self.assertEqual(len(mail.outbox), 1)
        self.assertIn("password reset", mail.outbox[0].subject.lower())
        self.assertEqual(mail.outbox[0].to, ["target@example.com"])

    def test_manager_can_resend_setup_email_for_pending_account(self):
        response = self.client.post(
            reverse("synopsis:manager_user_edit", args=[self.pending_user.id]),
            {"action": "send_access_email"},
        )

        self.assertRedirects(
            response,
            reverse("synopsis:manager_user_edit", args=[self.pending_user.id]),
        )
        self.assertEqual(len(mail.outbox), 1)
        self.assertIn("set up your ce synopsis portal account", mail.outbox[0].subject.lower())
        self.assertEqual(mail.outbox[0].to, ["pending@example.com"])

    def test_manager_can_assign_synopses_to_external_author(self):
        response = self.client.post(
            reverse("synopsis:manager_user_edit", args=[self.pending_user.id]),
            {
                "action": "update_user",
                "first_name": "Pending",
                "last_name": "",
                "email": "pending@example.com",
                "global_role": "external_collaborator",
                "is_active": "on",
                "assigned_projects": [str(self.project.id)],
            },
        )

        self.assertRedirects(
            response,
            reverse("synopsis:manager_user_edit", args=[self.pending_user.id]),
        )
        self.assertTrue(
            UserRole.objects.filter(
                user=self.pending_user, project=self.project, role="author"
            ).exists()
        )

    def test_manager_can_delete_user_with_email_confirmation(self):
        response = self.client.post(
            reverse("synopsis:manager_user_edit", args=[self.target.id]),
            {
                "action": "delete_user",
                "confirm_email": "target@example.com",
            },
        )

        self.assertRedirects(response, reverse("synopsis:manager_dashboard"))
        self.assertFalse(User.objects.filter(pk=self.target.id).exists())

    def test_manager_cannot_delete_own_account(self):
        response = self.client.post(
            reverse("synopsis:manager_user_edit", args=[self.manager.id]),
            {
                "action": "delete_user",
                "confirm_email": "manager@example.com",
            },
            follow=True,
        )

        self.assertRedirects(
            response, reverse("synopsis:manager_user_edit", args=[self.manager.id])
        )
        self.assertTrue(User.objects.filter(pk=self.manager.id).exists())
        self.assertContains(response, "You cannot delete your own account.")

    def test_superuser_accounts_are_protected_from_manager_edit_screen(self):
        response = self.client.get(
            reverse("synopsis:manager_user_edit", args=[self.superuser.id]),
            follow=True,
        )

        self.assertRedirects(response, reverse("synopsis:manager_dashboard"))
        self.assertContains(
            response,
            "System admin accounts are managed outside this screen.",
        )


class EmailSubjectFormattingTests(TestCase):
    def setUp(self):
        self.project = SimpleNamespace(title="Coastal Restoration")

    def test_invite_reminder_with_date(self):
        due = date(2025, 5, 10)
        subject = email_subject("invite_reminder", self.project, due)
        self.assertEqual(
            subject,
            f"[Reminder] {self.project.title} — please reply by {due.strftime('%d %b %Y')}",
        )

    def test_fallback_subject_for_unknown_kind(self):
        subject = email_subject("unknown", self.project)
        self.assertEqual(subject, f"[{BRAND}] {self.project.title}")

    def test_invite_without_due_date(self):
        subject = email_subject("invite", self.project)
        self.assertEqual(
            subject,
            f"[{BRAND}] Invitation to advise on {self.project.title}",
        )

    def test_protocol_reminder_subject(self):
        due = timezone.make_aware(datetime(2025, 6, 5, 18, 0))
        formatted = timezone.localtime(due).strftime("%d %b %Y %H:%M")
        subject = email_subject("protocol_reminder", self.project, due)
        self.assertEqual(
            subject,
            f"[Reminder] Protocol feedback due for {self.project.title} ({formatted})",
        )

    def test_protocol_review_subject(self):
        subject = email_subject("protocol_review", self.project)
        self.assertEqual(
            subject,
            f"[Action requested] Protocol for review — {self.project.title}",
        )

    def test_advisory_invitation_email_escapes_urls_in_html_hrefs(self):
        text, html_body = _build_advisory_invitation_email(
            project=self.project,
            recipient_name="Will",
            due_date=date(2025, 5, 10),
            yes_url="https://example.com/yes?x=1&y='two'",
            no_url='https://example.com/no?x=1&y="three"',
            attachment_lines=[
                ("Action list", "https://files.example.com/doc?version=1&lang='en'")
            ],
        )

        self.assertIn("https://example.com/yes?x=1&y='two'", text)
        self.assertIn('https://example.com/no?x=1&y="three"', text)
        self.assertIn(
            "href='https://example.com/yes?x=1&amp;y=&#x27;two&#x27;'",
            html_body,
        )
        self.assertIn(
            "href='https://example.com/no?x=1&amp;y=&quot;three&quot;'",
            html_body,
        )
        self.assertIn(
            "href='https://files.example.com/doc?version=1&amp;lang=&#x27;en&#x27;'",
            html_body,
        )
        self.assertIn("Privacy notice:", text)
        self.assertIn("Lawful basis:", text)
        self.assertIn("ICO:", text)
        self.assertIn("<strong>Privacy notice</strong>", html_body)


class AdvisoryBoardMemberFormTests(TestCase):
    def test_add_member_form_marks_core_fields_and_placeholders(self):
        form = AdvisoryBoardMemberForm()

        self.assertFalse(form.fields["title"].required)
        self.assertTrue(form.fields["first_name"].required)
        self.assertTrue(form.fields["last_name"].required)
        self.assertTrue(form.fields["email"].required)
        self.assertEqual(
            form.fields["first_name"].widget.attrs["placeholder"],
            "First name (required)",
        )
        self.assertEqual(
            form.fields["country"].widget.attrs["placeholder"],
            "Country (optional)",
        )
        self.assertEqual(
            form.fields["notes"].widget.attrs["placeholder"],
            "Notes for this member (optional)",
        )


class AdvisoryBoardMemberAddUiTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Advisory UX")
        self.user = User.objects.create_user(username="advisor-owner", password="pw")
        UserRole.objects.create(user=self.user, project=self.project, role="author")
        self.client.force_login(self.user)
        self.url = reverse("synopsis:advisory_board_list", args=[self.project.id])

    def test_board_page_explains_required_fields_and_review_step(self):
        response = self.client.get(self.url)

        self.assertContains(
            response,
            "Only first name, last name and email are required.",
        )
        self.assertContains(response, "Required to add member")
        self.assertContains(response, "Optional details")
        self.assertContains(response, "Review member details")

    def test_invalid_add_member_reopens_modal_with_error_guidance(self):
        response = self.client.post(
            self.url,
            {
                "action": "add_member",
                "title": "",
                "first_name": "",
                "middle_name": "",
                "last_name": "",
                "organisation": "",
                "email": "",
                "country": "",
                "continent": "",
                "notes": "",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="addMemberModal"')
        self.assertContains(response, 'window.bootstrap.Modal.getOrCreateInstance')
        self.assertContains(response, "This field is required.")

    def test_edit_details_from_confirm_page_reopens_add_member_modal(self):
        response = self.client.post(
            self.url,
            {
                "action": "add_member_back",
                "title": "Dr",
                "first_name": "Amira",
                "middle_name": "",
                "last_name": "Shah",
                "organisation": "Conservation Lab",
                "email": "amira@example.com",
                "country": "UK",
                "continent": "Europe",
                "notes": "Review protocol section.",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="addMemberModal"')
        self.assertContains(response, 'window.bootstrap.Modal.getOrCreateInstance')
        self.assertContains(response, 'value="Amira"')
        self.assertContains(response, 'value="amira@example.com"')
        self.assertTrue(response.context["open_add_member_modal"])


class ReplyToListTests(TestCase):
    @override_settings(DEFAULT_FROM_EMAIL="fallback@example.com")
    def test_prefers_inviter_email(self):
        self.assertEqual(reply_to_list("inviter@example.com"), ["inviter@example.com"])

    @override_settings(DEFAULT_FROM_EMAIL="fallback@example.com")
    def test_uses_fallback_when_inviter_missing(self):
        self.assertEqual(reply_to_list(None), ["fallback@example.com"])


class EnsureGlobalGroupsTests(TestCase):
    def test_creates_all_expected_groups_once(self):
        ensure_global_groups()
        ensure_global_groups()
        names = set(Group.objects.values_list("name", flat=True))
        self.assertTrue(set(GLOBAL_GROUPS).issubset(names))
        self.assertEqual(
            Group.objects.filter(name__in=GLOBAL_GROUPS).count(),
            len(GLOBAL_GROUPS),
        )


class ProjectPhaseTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Marine Study")

    def test_defaults_to_draft_protocol_without_manual_phase(self):
        self.assertEqual(self.project.phase, "draft_protocol")
        self.assertEqual(self.project.get_phase_display(), "Draft protocol")

    def test_manual_phase_does_not_regress(self):
        self.project.phase_manual = "draft_protocol"
        self.project.save(update_fields=["phase_manual"])
        self.assertEqual(self.project.phase, "draft_protocol")

    def test_manual_phase_can_advance(self):
        self.project.phase_manual = "summary_writing"
        self.project.save(update_fields=["phase_manual"])
        self.assertEqual(self.project.phase, "summary_writing")
        self.assertEqual(
            self.project.get_phase_display(),
            "Summary writing",
        )

    def test_computed_phase_advances_after_protocol_upload(self):
        Protocol.objects.create(
            project=self.project,
            document=SimpleUploadedFile("protocol.docx", b"protocol"),
        )

        self.assertEqual(self.project.phase, "invite_advisory_board")

    def test_computed_phase_does_not_regress_existing_project_without_manual_phase(self):
        Protocol.objects.create(
            project=self.project,
            document=SimpleUploadedFile("protocol.docx", b"protocol"),
        )
        AdvisoryBoardInvitation.objects.create(
            project=self.project,
            email="advisor@example.com",
            accepted=True,
        )
        AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Advisor",
            email="advisor@example.com",
            response="Y",
            feedback_on_protocol_received=timezone.localdate(),
        )

        self.assertEqual(self.project.phase_manual, None)
        self.assertEqual(self.project.phase, "draft_chapters")

    def test_manual_phase_cannot_regress_below_computed_progress(self):
        Protocol.objects.create(
            project=self.project,
            document=SimpleUploadedFile("protocol.docx", b"protocol"),
        )
        AdvisoryBoardInvitation.objects.create(
            project=self.project,
            email="advisor@example.com",
            accepted=True,
        )
        AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Advisor",
            email="advisor@example.com",
            response="Y",
            feedback_on_protocol_received=timezone.localdate(),
        )
        self.project.phase_manual = "draft_protocol"
        self.project.save(update_fields=["phase_manual"])

        self.assertEqual(self.project.phase, "draft_chapters")

    def test_computed_phase_respects_disabled_protocol_and_advisory_steps(self):
        self.project.protocol_relevant = False
        self.project.advisory_board_relevant = False
        self.project.save(update_fields=["protocol_relevant", "advisory_board_relevant"])

        self.assertEqual(self.project.phase, "references_screening")


class ProjectAuthorUsersTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Author Demo")
        self.author = User.objects.create_user(username="zoe")
        self.coauthor = User.objects.create_user(username="adam")
        self.manager = User.objects.create_user(username="manager")

        UserRole.objects.create(user=self.author, project=self.project, role="author")
        UserRole.objects.create(user=self.coauthor, project=self.project, role="author")
        UserRole.objects.create(user=self.manager, project=self.project, role="manager")

    def test_returns_authors_sorted_by_username(self):
        usernames = list(self.project.author_users.values_list("username", flat=True))
        self.assertEqual(usernames, ["adam", "zoe"])


class SynopsisStructureTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Synopsis Demo")
        self.user = User.objects.create_user(username="manager", password="pw", is_staff=True)
        UserRole.objects.create(user=self.user, project=self.project, role="manager")
        self.client.force_login(self.user)
        # Attach a reference/summary for assignment flows
        self.batch = ReferenceSourceBatch.objects.create(
            project=self.project,
            label="Batch 1",
            source_type="manual_upload",
            uploaded_by=self.user,
        )
        self.reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Test ref",
            hash_key="hash-test-ref",
            screening_status="included",
        )
        self.summary = ReferenceSummary.objects.create(
            project=self.project, reference=self.reference
        )

    def test_apply_preset_creates_chapters_once(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        response = self.client.post(
            url,
            {"action": "apply-preset", "preset_key": "standard_ce_toc"},
        )
        self.assertEqual(response.status_code, 302)
        self.assertTrue(SynopsisChapter.objects.filter(project=self.project).exists())
        front_chapter = SynopsisChapter.objects.get(
            project=self.project, title="Advisory Board"
        )
        self.assertEqual(front_chapter.chapter_type, SynopsisChapter.TYPE_TEXT)
        evidence_chapter = SynopsisChapter.objects.get(
            project=self.project,
            title="2. Threat: Residential and commercial development",
        )
        self.assertEqual(
            evidence_chapter.chapter_type, SynopsisChapter.TYPE_EVIDENCE
        )
        count_after_first = SynopsisChapter.objects.filter(project=self.project).count()
        # Second apply should be blocked because outline is not empty
        response = self.client.post(
            url,
            {"action": "apply-preset", "preset_key": "standard_ce_toc"},
        )
        self.assertEqual(response.status_code, 302)
        count_after_second = SynopsisChapter.objects.filter(project=self.project).count()
        self.assertEqual(count_after_first, count_after_second)

    def test_reset_structure_clears_outline(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        SynopsisChapter.objects.create(project=self.project, title="Tmp", position=1)
        response = self.client.post(url, {"action": "reset-structure"})
        self.assertEqual(response.status_code, 302)
        self.assertFalse(SynopsisChapter.objects.filter(project=self.project).exists())

    def test_create_intervention_without_subheading_creates_default(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(project=self.project, title="Ch", position=1)
        response = self.client.post(
            url,
            {
                "action": "create-intervention",
                "chapter_id": chapter.id,
                "title": "Intervention A",
            },
        )
        self.assertEqual(response.status_code, 302)
        subheadings = SynopsisSubheading.objects.filter(chapter=chapter)
        self.assertEqual(subheadings.count(), 1)
        intervention_titles = list(
            SynopsisIntervention.objects.filter(subheading__chapter=chapter).values_list("title", flat=True)
        )
        self.assertIn("Intervention A", intervention_titles)

    def test_move_intervention_to_another_subheading(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        general = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="General",
            position=1,
        )
        arable = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Arable",
            position=2,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=general,
            title="Mow more frequently",
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "move-intervention-to-subheading",
                "intervention_id": intervention.id,
                "target_subheading_id": arable.id,
            },
        )

        self.assertEqual(response.status_code, 302)
        intervention.refresh_from_db()
        self.assertEqual(intervention.subheading, arable)
        self.assertEqual(intervention.position, 1)
        self.assertEqual(general.interventions.count(), 0)
        self.assertEqual(arable.interventions.count(), 1)

    def test_structure_page_explains_intervention_group_linking(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Arable",
            position=1,
        )
        SynopsisIntervention.objects.create(
            subheading=subheading,
            title="Mow more frequently",
            position=1,
        )
        self.summary.synopsis_draft = (
            "A replicated study found that mowing more frequently increased arable plant richness."
        )
        self.summary.use_custom_synopsis_draft = True
        self.summary.save(update_fields=["synopsis_draft", "use_custom_synopsis_draft"])
        SynopsisAssignment.objects.create(
            intervention=SynopsisIntervention.objects.get(title="Mow more frequently"),
            reference_summary=self.summary,
            position=1,
        )

        response = self.client.get(url)

        self.assertContains(response, "Add intervention to Arable")
        self.assertContains(response, "Intervention group")
        self.assertContains(response, "Move to group")
        self.assertContains(response, "Edit metadata, background and key messages")
        self.assertContains(response, "Metadata")
        self.assertContains(response, "Background")
        self.assertContains(response, "Key messages")
        self.assertContains(response, "Assigned summaries")
        self.assertContains(response, "Additional intervention text")
        self.assertContains(response, "Most content should be added as background, key messages or assigned summaries.")
        self.assertContains(response, "Assigned study summaries to review")
        self.assertContains(response, "review the assigned summaries here first")
        self.assertContains(response, "mowing more frequently increased arable plant richness")

    def test_structure_page_renders_restore_state_hooks(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Arable",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="Mow more frequently",
            position=1,
        )

        response = self.client.get(url)

        self.assertContains(response, 'id="synopsis-structure-page"', html=False)
        self.assertContains(
            response,
            f'id="subheading-{subheading.id}"',
            html=False,
        )
        self.assertContains(
            response,
            f'id="intervention-{intervention.id}"',
            html=False,
        )
        self.assertContains(
            response,
            f'id="intervention-editor-{intervention.id}"',
            html=False,
        )
        self.assertContains(response, "cePreservePageState({", html=False)
        self.assertContains(
            response,
            f'"synopsis-structure-state-evidence-{self.project.id}"',
            html=False,
        )
        self.assertContains(response, 'closest("details[id]")', html=False)

    def test_text_chapter_blocks_subheading_and_intervention(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        text_chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="Advisory Board",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=1,
        )
        response = self.client.post(
            url,
            {
                "action": "create-subheading",
                "chapter_id": text_chapter.id,
                "title": "Should fail",
            },
        )
        self.assertEqual(response.status_code, 302)
        self.assertFalse(SynopsisSubheading.objects.filter(chapter=text_chapter).exists())
        response = self.client.post(
            url,
            {
                "action": "create-intervention",
                "chapter_id": text_chapter.id,
                "title": "Should also fail",
            },
        )
        self.assertEqual(response.status_code, 302)
        self.assertFalse(
            SynopsisIntervention.objects.filter(subheading__chapter=text_chapter).exists()
        )

    def test_update_intervention_synthesis_fields(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "update-intervention-synthesis",
                "intervention_id": intervention.id,
                "ce_action_url": "https://www.conservationevidence.com/actions/4018",
                "evidence_status": SynopsisIntervention.EVIDENCE_STATUS_NO_STUDIES,
                "synthesis_text": "No direct studies were identified in the searched evidence base.",
            },
        )
        self.assertEqual(response.status_code, 302)
        intervention.refresh_from_db()
        self.assertEqual(
            intervention.ce_action_url,
            "https://www.conservationevidence.com/actions/4018",
        )
        self.assertEqual(
            intervention.evidence_status,
            SynopsisIntervention.EVIDENCE_STATUS_NO_STUDIES,
        )
        self.assertIn("No direct studies", intervention.synthesis_text)

    def test_add_and_update_key_message(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        second_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Second test ref",
            hash_key="hash-second-test-ref",
            screening_status="included",
        )
        second_summary = ReferenceSummary.objects.create(
            project=self.project, reference=second_reference
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=self.summary,
            position=1,
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=second_summary,
            position=2,
        )

        response = self.client.post(
            url,
            {
                "action": "add-key-message",
                "intervention_id": intervention.id,
                "response_group": SynopsisInterventionKeyMessage.GROUP_POPULATION,
                "outcome_label": "Abundance/Cover",
                "study_count": 3,
                "statement": "Three studies found increased coral cover after intervention.",
                "supporting_summaries": [str(self.summary.id)],
            },
        )
        self.assertEqual(response.status_code, 302)
        key_message = intervention.key_messages.first()
        self.assertIsNotNone(key_message)
        self.assertEqual(key_message.study_count, 3)
        self.assertEqual(
            list(
                key_message.supporting_summaries.order_by("id").values_list("id", flat=True)
            ),
            [self.summary.id],
        )

        response = self.client.post(
            url,
            {
                "action": "update-key-message",
                "intervention_id": intervention.id,
                "key_message_id": key_message.id,
                "response_group": SynopsisInterventionKeyMessage.GROUP_COMMUNITY,
                "outcome_label": "Richness/diversity",
                "study_count": 5,
                "statement": "Five studies found no clear community-level change.",
                "supporting_summaries": [str(second_summary.id), str(self.summary.id)],
            },
        )
        self.assertEqual(response.status_code, 302)
        key_message.refresh_from_db()
        self.assertEqual(
            key_message.response_group,
            SynopsisInterventionKeyMessage.GROUP_COMMUNITY,
        )
        self.assertEqual(key_message.outcome_label, "Richness/diversity")
        self.assertEqual(key_message.study_count, 5)
        self.assertEqual(
            list(
                key_message.supporting_summaries.order_by("id").values_list("id", flat=True)
            ),
            sorted([self.summary.id, second_summary.id]),
        )

    def test_intervention_reference_numbering_uses_oldest_first_order(self):
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        older_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Older study",
            authors="Alpha A.",
            publication_year=2001,
            hash_key="hash-old-study",
            screening_status="included",
        )
        newer_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Newer study",
            authors="Beta B.",
            publication_year=2018,
            hash_key="hash-new-study",
            screening_status="included",
        )
        older_summary = ReferenceSummary.objects.create(
            project=self.project, reference=older_reference
        )
        newer_summary = ReferenceSummary.objects.create(
            project=self.project, reference=newer_reference
        )
        SynopsisAssignment.objects.create(
            intervention=intervention, reference_summary=newer_summary, position=1
        )
        SynopsisAssignment.objects.create(
            intervention=intervention, reference_summary=older_summary, position=2
        )

        assignments = list(intervention.assignments.all())
        ordered_assignments, summary_numbers, ordered_references = (
            _intervention_reference_numbering(assignments)
        )

        self.assertEqual(
            [assignment.reference_summary_id for assignment in ordered_assignments],
            [older_summary.id, newer_summary.id],
        )
        self.assertEqual(summary_numbers[older_summary.id], 1)
        self.assertEqual(summary_numbers[newer_summary.id], 2)
        self.assertEqual(
            [reference.id for _, reference in ordered_references],
            [older_reference.id, newer_reference.id],
        )
        self.assertEqual(
            [numbers for numbers, _ in ordered_references],
            [[1], [2]],
        )

    def test_intervention_reference_numbering_groups_duplicate_references(self):
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Group duplicate paper studies",
            position=2,
        )
        shared_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Shared study paper",
            authors="Gamma G.",
            publication_year=2009,
            hash_key="hash-shared-study-paper",
            screening_status="included",
        )
        first_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=shared_reference,
            action_description="Study A",
        )
        second_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=shared_reference,
            action_description="Study B",
        )
        SynopsisAssignment.objects.create(
            intervention=intervention, reference_summary=first_summary, position=1
        )
        SynopsisAssignment.objects.create(
            intervention=intervention, reference_summary=second_summary, position=2
        )

        assignments = list(intervention.assignments.all())
        ordered_assignments, summary_numbers, ordered_references = (
            _intervention_reference_numbering(assignments)
        )

        self.assertEqual(
            [assignment.reference_summary_id for assignment in ordered_assignments],
            [first_summary.id, second_summary.id],
        )
        self.assertEqual(summary_numbers[first_summary.id], 1)
        self.assertEqual(summary_numbers[second_summary.id], 2)
        self.assertEqual(len(ordered_references), 1)
        self.assertEqual(ordered_references[0][1].id, shared_reference.id)
        self.assertEqual(ordered_references[0][0], [1, 2])
        self.assertEqual(_format_reference_number_ranges([1, 2]), "1-2")

    def test_key_message_rejects_unassigned_supporting_summary(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=self.summary,
            position=1,
        )
        outside_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Outside intervention summary",
            hash_key="hash-outside-summary",
            screening_status="included",
        )
        outside_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=outside_reference,
        )

        response = self.client.post(
            url,
            {
                "action": "add-key-message",
                "intervention_id": intervention.id,
                "response_group": SynopsisInterventionKeyMessage.GROUP_RESPONSE,
                "statement": "Message with invalid supporting study link.",
                "supporting_summaries": [str(outside_summary.id)],
            },
        )
        self.assertEqual(response.status_code, 302)
        self.assertFalse(intervention.key_messages.exists())

    def test_evidence_page_groups_summary_tabs_under_reference(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        self.reference.authors = "Rebecca Smith"
        self.reference.publication_year = 2024
        self.reference.save(update_fields=["authors", "publication_year", "updated_at"])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        alt_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=self.reference,
            action_description="Study B",
        )
        self.summary.action_description = "Study A"
        self.summary.save(update_fields=["action_description"])
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=self.summary,
            position=1,
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=alt_summary,
            position=2,
        )

        response = self.client.get(url)

        self.assertEqual(response.status_code, 200)
        grouped = response.context["reference_summary_groups"]
        self.assertEqual(len(grouped), 1)
        self.assertEqual(grouped[0]["reference_heading"], "Rebecca Smith · 2024")
        self.assertEqual(grouped[0]["reference_context"], "Test ref")
        self.assertEqual(
            [item["summary_display"] for item in grouped[0]["summaries"]],
            ["D1000.a — Study A", "D1000.b — Study B"],
        )
        self.assertContains(response, "Assign summary tabs, not whole papers.")
        self.assertContains(response, "Grouped by reference author.")
        self.assertContains(
            response,
            "Choose a summary tab to preview its reference and tab label.",
        )
        self.assertContains(response, "Rebecca Smith · 2024")
        self.assertContains(response, "D1000.a — Study A")
        self.assertContains(response, "D1000.b — Study B")
        self.assertContains(response, "Same source paper")
        self.assertContains(response, "shared reference line (1-2)")
        self.assertContains(response, "Compilation preview")
        self.assertContains(response, "study paragraphs")
        self.assertContains(response, "source paper")
        self.assertContains(response, "2 summary tabs from the same paper")

    def test_evidence_workspace_uses_action_only_iucn_categories(self):
        response = self.client.get(
            reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        )

        context_categories = list(response.context["iucn_categories"])
        self.assertTrue(context_categories)
        self.assertTrue(
            all(category.kind == IUCNCategory.KIND_ACTION for category in context_categories)
        )
        self.assertIn(
            "Land/water protection-Area protection",
            [category.name for category in context_categories],
        )
        self.assertNotIn(
            "Residential & commercial development",
            [category.name for category in context_categories],
        )

        form_categories = list(
            response.context["intervention_form"].fields["iucn_category"].queryset
        )
        self.assertTrue(form_categories)
        self.assertTrue(
            all(category.kind == IUCNCategory.KIND_ACTION for category in form_categories)
        )

    def test_update_intervention_metadata_rejects_threat_category(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        threat_category = IUCNCategory.objects.filter(
            kind=IUCNCategory.KIND_THREAT,
            is_active=True,
        ).first()

        self.assertIsNotNone(threat_category)

        response = self.client.post(
            url,
            {
                "action": "update-intervention-metadata",
                "intervention_id": intervention.id,
                "iucn_category": str(threat_category.id),
            },
        )

        self.assertEqual(response.status_code, 302)
        intervention.refresh_from_db()
        self.assertIsNone(intervention.iucn_category)

    def test_delete_assignment_removes_supporting_links_from_key_messages(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        assignment = SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=self.summary,
            position=1,
        )
        key_message = SynopsisInterventionKeyMessage.objects.create(
            intervention=intervention,
            response_group=SynopsisInterventionKeyMessage.GROUP_RESPONSE,
            statement="Linked message",
            position=1,
        )
        key_message.supporting_summaries.add(self.summary)

        response = self.client.post(
            url,
            {
                "action": "delete-assignment",
                "assignment_id": assignment.id,
            },
        )
        self.assertEqual(response.status_code, 302)
        key_message.refresh_from_db()
        self.assertFalse(key_message.supporting_summaries.exists())

    def test_export_citation_and_reference_identifier_override(self):
        reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Corallivorous snail removal",
            authors="Miller M.",
            publication_year=2001,
            journal="Coral Reefs",
            volume="19",
            pages="293-295",
            doi="10.1007/PL00006963",
            hash_key="hash-citation-study",
            screening_status="included",
        )
        summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=reference,
            study_design="replicated, controlled study",
            year_range="1999",
            summary_of_results="snail removal reduced live tissue loss.",
            reference_identifier="X",
        )

        citation = _reference_export_citation(reference)
        self.assertIn("Miller M. (2001)", citation)
        self.assertIn("Corallivorous snail removal.", citation)
        self.assertIn("Coral Reefs, 19, 293-295.", citation)
        self.assertIn("https://doi.org/10.1007/PL00006963", citation)

        paragraph = _structured_summary_paragraph(
            summary, reference_identifier_override="3"
        )
        self.assertIn("(3)", paragraph)
        self.assertNotIn("(X)", paragraph)

    def test_generate_docx_collapses_duplicate_reference_lines_but_keeps_study_paragraphs(self):
        from docx import Document

        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        shared_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Shared study paper",
            authors="Gamma G.",
            publication_year=2009,
            hash_key="hash-docx-shared-paper",
            screening_status="included",
        )
        first_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=shared_reference,
            study_design="replicated, controlled study",
            year_range="2009",
            summary_of_results="first finding improved coral cover.",
        )
        second_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=shared_reference,
            study_design="replicated, controlled study",
            year_range="2010",
            summary_of_results="second finding improved coral recruitment.",
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=first_summary,
            position=1,
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=second_summary,
            position=2,
        )

        payload = _generate_synopsis_docx(self.project)
        document = Document(io.BytesIO(payload))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

        self.assertTrue(
            any("(1)" in paragraph and "first finding improved coral cover" in paragraph for paragraph in paragraphs)
        )
        self.assertTrue(
            any("(2)" in paragraph and "second finding improved coral recruitment" in paragraph for paragraph in paragraphs)
        )
        collapsed_reference_lines = [
            paragraph for paragraph in paragraphs if paragraph.startswith("(1-2) ")
        ]
        self.assertEqual(len(collapsed_reference_lines), 1)
        self.assertIn("Shared study paper.", collapsed_reference_lines[0])
        self.assertFalse(
            any(
                paragraph.startswith("(1) Gamma G. (2009) Shared study paper.")
                or paragraph.startswith("(2) Gamma G. (2009) Shared study paper.")
                for paragraph in paragraphs
            )
        )

    def test_generate_docx_uses_summary_citation_override_with_italics(self):
        from docx import Document

        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Shared study paper",
            authors="Gamma G.",
            publication_year=2009,
            journal="Journal of Marine Trials",
            volume="12",
            pages="34-40",
            hash_key="hash-docx-citation-override",
            screening_status="included",
        )
        summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=reference,
            study_design="replicated study",
            year_range="2009",
            summary_of_results="kelp cover improved.",
            citation="Gamma G. (2009) <i>Glipa</i> restoration note.",
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=summary,
            position=1,
        )

        payload = _generate_synopsis_docx(self.project)
        document = Document(io.BytesIO(payload))
        reference_paragraph = next(
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text.startswith("(1) Gamma G. (2009) Glipa restoration note.")
        )

        self.assertEqual(
            reference_paragraph.text,
            "(1) Gamma G. (2009) Glipa restoration note.",
        )
        self.assertTrue(any(run.text == "Glipa" and run.italic for run in reference_paragraph.runs))

    def test_workspace_routes_load(self):
        narrative_url = reverse(
            "synopsis:project_synopsis_narrative", args=[self.project.id]
        )
        evidence_url = reverse(
            "synopsis:project_synopsis_evidence", args=[self.project.id]
        )
        structure_url = reverse(
            "synopsis:project_synopsis_structure", args=[self.project.id]
        )
        self.assertEqual(self.client.get(narrative_url).status_code, 200)
        self.assertEqual(self.client.get(evidence_url).status_code, 200)
        self.assertEqual(self.client.get(structure_url).status_code, 200)

    def test_narrative_workspace_post_redirects_to_narrative_route(self):
        url = reverse("synopsis:project_synopsis_narrative", args=[self.project.id])
        response = self.client.post(
            url,
            {
                "action": "create-chapter",
                "title": "Advisory Board",
                "chapter_type": SynopsisChapter.TYPE_TEXT,
            },
        )
        self.assertEqual(response.status_code, 302)
        self.assertEqual(
            urlparse(response["Location"]).path,
            urlparse(url).path,
        )

    def test_narrative_workspace_renders_restore_state_hooks(self):
        url = reverse("synopsis:project_synopsis_narrative", args=[self.project.id])
        SynopsisChapter.objects.create(
            project=self.project,
            title="Advisory Board",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=1,
        )

        response = self.client.get(url)

        self.assertContains(response, 'id="synopsis-narrative-page"', html=False)
        self.assertContains(response, "cePreservePageState({", html=False)
        self.assertContains(
            response,
            f'"synopsis-narrative-state-{self.project.id}"',
            html=False,
        )


class AdvisoryDeadlineValidationTests(TestCase):
    def test_invite_due_date_rejects_today_and_past_dates(self):
        today = timezone.localdate()
        yesterday = today - timedelta(days=1)

        today_form = AdvisoryInviteForm(
            data={"email": "ibrahim@example.com", "due_date": today.isoformat()}
        )
        self.assertFalse(today_form.is_valid())
        self.assertIn("due_date", today_form.errors)

        yesterday_form = AdvisoryBulkInviteForm(
            data={"due_date": yesterday.isoformat()}
        )
        self.assertFalse(yesterday_form.is_valid())
        self.assertIn("due_date", yesterday_form.errors)

    def test_schedule_forms_expose_minimum_dates_in_widgets(self):
        tomorrow = timezone.localdate() + timedelta(days=1)
        expected_date_min = tomorrow.isoformat()
        expected_datetime_min = f"{expected_date_min}T00:00"

        self.assertEqual(
            AdvisoryInviteForm().fields["due_date"].widget.attrs.get("min"),
            expected_date_min,
        )
        self.assertEqual(
            ReminderScheduleForm().fields["reminder_date"].widget.attrs.get("min"),
            expected_date_min,
        )
        self.assertEqual(
            ProtocolReminderScheduleForm().fields["deadline"].widget.attrs.get("min"),
            expected_datetime_min,
        )
        self.assertEqual(
            ActionListReminderScheduleForm().fields["deadline"].widget.attrs.get("min"),
            expected_datetime_min,
        )

    @override_settings(
        ADVISORY_INVITE_RESPONSE_WINDOW_DAYS=14,
        ADVISORY_DOCUMENT_FEEDBACK_WINDOW_DAYS=21,
    )
    def test_advisory_help_text_uses_runtime_settings(self):
        self.assertIn(
            "Defaults to 14 days from today.",
            AdvisoryInviteForm().fields["due_date"].help_text,
        )
        self.assertIn(
            "Defaults to 14 days from today",
            AdvisoryBulkInviteForm().fields["due_date"].help_text,
        )
        self.assertIn(
            "Defaults to 14 days from today.",
            ReminderScheduleForm().fields["reminder_date"].help_text,
        )
        self.assertIn(
            "Defaults to 21 days from today",
            ProtocolSendForm().fields["due_date"].help_text,
        )
        self.assertIn(
            "Defaults to 21 days from today",
            ActionListSendForm().fields["due_date"].help_text,
        )
        self.assertIn(
            "Defaults to 21 days from today.",
            ProtocolReminderScheduleForm().fields["deadline"].help_text,
        )
        self.assertIn(
            "Defaults to 21 days from today.",
            ActionListReminderScheduleForm().fields["deadline"].help_text,
        )

    def test_protocol_deadline_rejects_same_day_datetime(self):
        same_day_value = f"{timezone.localdate().isoformat()}T12:00"
        form = ProtocolReminderScheduleForm(data={"deadline": same_day_value})
        self.assertFalse(form.is_valid())
        self.assertIn("deadline", form.errors)


class AttachmentSummaryConsoleEmailBackendTests(SimpleTestCase):
    def test_prints_body_and_attachment_summary_without_payload(self):
        stream = io.StringIO()
        backend = AttachmentSummaryConsoleEmailBackend(stream=stream)
        message = EmailMessage(
            "Demo link",
            "Open this link: http://example.com/demo-token",
            "from@example.com",
            ["to@example.com"],
        )
        message.attach(
            "synopsis.docx",
            b"raw document bytes that should not be printed",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        sent = backend.send_messages([message])

        output = stream.getvalue()
        self.assertEqual(sent, 1)
        self.assertIn("Open this link: http://example.com/demo-token", output)
        self.assertIn("synopsis.docx", output)
        self.assertIn("content not printed", output)
        self.assertNotIn("raw document bytes", output)


@override_settings(DEFAULT_FROM_EMAIL="reminders@example.com")
class SendDueRemindersTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Reminder Project")
        self.invite_member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Ina",
            email="invite@example.com",
            invite_sent=True,
            response_date=date(2025, 1, 10),
        )
        aware_now = timezone.now()
        self.protocol_deadline = aware_now + timedelta(days=5)
        self.protocol_member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Proto",
            email="proto@example.com",
            invite_sent=True,
            response="Y",
            sent_protocol_at=aware_now,
            feedback_on_protocol_deadline=self.protocol_deadline,
        )
        self.action_list_deadline = aware_now + timedelta(days=6)
        self.action_member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Al",
            email="action@example.com",
            invite_sent=True,
            response="Y",
            sent_action_list_at=aware_now,
            feedback_on_action_list_deadline=self.action_list_deadline,
        )
        self.synopsis_deadline = aware_now + timedelta(days=7)
        self.synopsis_member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Syn",
            email="synopsis@example.com",
            invite_sent=True,
            response="Y",
            sent_synopsis_at=aware_now,
            feedback_on_synopsis_deadline=self.synopsis_deadline,
        )

    @patch("synopsis.management.commands.send_due_reminders.EmailMultiAlternatives")
    @patch("synopsis.management.commands.send_due_reminders.minus_business_days")
    @patch("synopsis.management.commands.send_due_reminders.timezone")
    def test_sends_due_reminders_for_active_streams(
        self, mock_timezone, mock_minus, mock_email
    ):
        today = date(2025, 1, 8)
        real_now = timezone.now()
        mock_timezone.localdate.return_value = today
        mock_timezone.now.return_value = real_now
        mock_timezone.localtime.side_effect = lambda value: value

        def minus_side_effect(deadline, offset):
            self.assertEqual(offset, 2)
            return today

        mock_minus.side_effect = minus_side_effect

        email_calls = []

        def build_email(*args, **kwargs):
            instance = MagicMock()
            email_calls.append((args, kwargs, instance))
            return instance

        mock_email.side_effect = build_email

        call_command("send_due_reminders")

        self.assertEqual(len(email_calls), 4)
        for _, _, instance in email_calls:
            instance.send.assert_called_once()

        self.invite_member.refresh_from_db()
        self.protocol_member.refresh_from_db()
        self.action_member.refresh_from_db()
        self.synopsis_member.refresh_from_db()

        self.assertTrue(self.invite_member.reminder_sent)
        self.assertTrue(self.protocol_member.protocol_reminder_sent)
        self.assertTrue(self.action_member.action_list_reminder_sent)
        self.assertTrue(self.synopsis_member.synopsis_reminder_sent)
        self.assertIsNotNone(self.invite_member.reminder_sent_at)
        self.assertIsNotNone(self.protocol_member.protocol_reminder_sent_at)
        self.assertIsNotNone(self.action_member.action_list_reminder_sent_at)
        self.assertIsNotNone(self.synopsis_member.synopsis_reminder_sent_at)

        subjects = [args[0] for args, _, _ in email_calls]
        self.assertIn(
            email_subject("invite_reminder", self.project, self.invite_member.response_date),
            subjects,
        )
        self.assertIn(
            email_subject("protocol_reminder", self.project, self.protocol_deadline),
            subjects,
        )
        self.assertIn(
            email_subject(
                "action_list_reminder", self.project, self.action_list_deadline
            ),
            subjects,
        )
        self.assertIn(
            email_subject("synopsis_reminder", self.project, self.synopsis_deadline),
            subjects,
        )

    @override_settings(ADVISORY_REMINDER_LEAD_BUSINESS_DAYS=-1)
    def test_rejects_negative_reminder_lead_business_days(self):
        with self.assertRaisesMessage(
            CommandError,
            "ADVISORY_REMINDER_LEAD_BUSINESS_DAYS must be a non-negative integer",
        ):
            call_command("send_due_reminders")

    @override_settings(ADVISORY_REMINDER_LEAD_BUSINESS_DAYS="tomorrow")
    def test_rejects_non_integer_reminder_lead_business_days(self):
        with self.assertRaisesMessage(
            CommandError,
            "ADVISORY_REMINDER_LEAD_BUSINESS_DAYS must be an integer",
        ):
            call_command("send_due_reminders")


class MemberReminderUpdateTests(TestCase):
    def setUp(self):
        self.media_dir = tempfile.mkdtemp()
        self.addCleanup(lambda: shutil.rmtree(self.media_dir, ignore_errors=True))
        override = override_settings(MEDIA_ROOT=self.media_dir)
        override.enable()
        self.addCleanup(override.disable)

        self.project = Project.objects.create(title="Inline Reminders")
        self.user = User.objects.create_user(username="owner", password="pw")
        UserRole.objects.create(user=self.user, project=self.project, role="author")
        self.client.force_login(self.user)
        self.board_url = reverse("synopsis:advisory_board_list", args=[self.project.id])

    def test_update_response_deadline(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Randy",
            email="randy@example.com",
            reminder_sent=True,
            reminder_sent_at=timezone.now(),
        )
        target_date = timezone.localdate() + timedelta(days=7)
        response = self.client.post(
            reverse(
                "synopsis:advisory_member_set_deadline",
                args=[self.project.id, member.id, "invite"],
            ),
            {"reminder_date": target_date.strftime("%Y-%m-%d")},
        )
        self.assertRedirects(response, self.board_url)
        member.refresh_from_db()
        self.assertEqual(member.response_date, target_date)
        self.assertFalse(member.reminder_sent)
        self.assertIsNone(member.reminder_sent_at)
        self.assertTrue(
            ProjectChangeLog.objects.filter(
                project=self.project, action="Updated invite reminder"
            ).exists()
        )

    def test_board_page_sets_minimum_response_deadline_on_inline_input(self):
        AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Randy",
            email="randy@example.com",
        )
        tomorrow = timezone.localdate() + timedelta(days=1)
        response = self.client.get(self.board_url)
        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            f'min="{tomorrow.strftime("%Y-%m-%d")}"',
        )

    def test_board_page_shows_accepted_participation_note(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Ness",
            email="ness@example.com",
            response="Y",
            participation_confirmed=True,
            participation_statement="Happy to help, but I may be slower next week.",
        )

        response = self.client.get(self.board_url)

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "View note")
        self.assertContains(response, f"accepted-message-{member.id}")
        self.assertContains(
            response,
            "Happy to help, but I may be slower next week.",
        )

    def test_board_page_greys_out_response_deadline_for_accepted_member(self):
        accepted_date = timezone.localdate() + timedelta(days=4)
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Amina",
            last_name="Accepted",
            email="amina@example.com",
            response="Y",
            participation_confirmed=True,
            response_date=accepted_date,
        )

        response = self.client.get(self.board_url)

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, f'value="{accepted_date.strftime("%Y-%m-%d")}"')
        self.assertContains(response, 'aria-label="Response deadline for Amina"')
        self.assertContains(response, "Locked")
        self.assertContains(response, "Accepted")
        self.assertNotContains(
            response,
            reverse(
                "synopsis:advisory_member_set_deadline",
                args=[self.project.id, member.id, "invite"],
            ),
        )

    def test_board_page_shows_protocol_feedback_modal(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Pat",
            last_name="Protocol",
            email="pat@example.com",
            response="Y",
            participation_confirmed=True,
            sent_protocol_at=timezone.now(),
        )
        feedback = ProtocolFeedback.objects.create(
            project=self.project,
            member=member,
            content="Please tighten the scope in the introduction.",
            submitted_at=timezone.now(),
            uploaded_document=SimpleUploadedFile(
                "pat-protocol-comments.docx",
                b"protocol comments",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        )

        response = self.client.get(self.board_url)

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "View feedback")
        self.assertContains(response, "View document")
        self.assertContains(response, f'aria-label="View protocol feedback for {member.first_name} {member.last_name}"')
        self.assertContains(response, f"protocol-feedback-{member.id}")
        self.assertContains(response, f"protocol-feedback-document-{member.id}")
        self.assertContains(response, "Please tighten the scope in the introduction.")
        self.assertContains(response, feedback.latest_document_label())

    def test_board_page_shows_action_list_feedback_modal(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Will",
            last_name="Actions",
            email="will@example.com",
            response="Y",
            participation_confirmed=True,
            sent_action_list_at=timezone.now(),
        )
        feedback = ActionListFeedback.objects.create(
            project=self.project,
            member=member,
            content="I suggest splitting habitat creation and restoration.",
            submitted_at=timezone.now(),
            uploaded_document=SimpleUploadedFile(
                "will-action-comments.docx",
                b"action comments",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        )

        response = self.client.get(self.board_url)

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "View feedback")
        self.assertContains(response, "View document")
        self.assertContains(response, f'aria-label="View action list feedback for {member.first_name} {member.last_name}"')
        self.assertContains(response, f"action-list-feedback-{member.id}")
        self.assertContains(response, f"action-list-feedback-document-{member.id}")
        self.assertContains(
            response,
            "I suggest splitting habitat creation and restoration.",
        )
        self.assertContains(response, feedback.latest_document_label())

    def test_board_page_shows_synopsis_feedback_modal(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Sam",
            last_name="Synopsis",
            email="sam@example.com",
            response="Y",
            participation_confirmed=True,
            sent_synopsis_at=timezone.now(),
        )
        feedback = SynopsisFeedback.objects.create(
            project=self.project,
            member=member,
            content="Please expand the key messages.",
            submitted_at=timezone.now(),
            uploaded_document=SimpleUploadedFile(
                "sam-synopsis-comments.pdf",
                b"synopsis comments",
                content_type="application/pdf",
            ),
        )

        response = self.client.get(self.board_url)

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "View feedback")
        self.assertContains(response, "View document")
        self.assertContains(response, f"synopsis-feedback-{member.id}")
        self.assertContains(response, f"synopsis-feedback-document-{member.id}")
        self.assertContains(response, "Please expand the key messages.")
        self.assertContains(response, feedback.latest_document_label())

    def test_board_page_shows_action_list_feedback_deadline(self):
        deadline = timezone.now().replace(second=0, microsecond=0) + timedelta(days=5)
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="ibrahim",
            last_name="Deadline",
            email="ibrahim@hotmail.com",
            response="Y",
            participation_confirmed=True,
            sent_action_list_at=timezone.now(),
            feedback_on_action_list_deadline=deadline,
        )

        response = self.client.get(self.board_url)

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            timezone.localtime(member.feedback_on_action_list_deadline).strftime(
                "%Y-%m-%d %H:%M"
            ),
        )

    def test_board_page_shows_updated_document_tracking_headers_without_guidance(self):
        AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Header",
            last_name="Check",
            email="header@example.com",
            response="Y",
            participation_confirmed=True,
        )

        response = self.client.get(self.board_url)

        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, "GUIDANCE FEEDBACK")
        self.assertNotContains(response, "GUIDANCE SENT")
        self.assertContains(response, "FEEDBACK DOCUMENT", count=3)
        self.assertContains(response, "AUTHOR REPLIED", count=3)
        self.assertContains(response, "REMINDER SENT", count=4)
        self.assertContains(response, 'th class="ab-status-action text-start"', count=7)
        self.assertContains(response, 'td class="ab-status-action"', count=7)
        self.assertContains(response, 'th class="ab-status-protocol text-start"', count=7)
        self.assertContains(response, 'td class="ab-status-protocol"', count=7)
        self.assertContains(response, 'th class="ab-status-synopsis text-start"', count=7)
        self.assertContains(response, 'td class="ab-status-synopsis"', count=7)

    def test_can_toggle_action_list_author_replied_from_board(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Reyhan",
            last_name="Reply",
            email="reyhan@example.com",
            response="Y",
            participation_confirmed=True,
            sent_action_list_at=timezone.now(),
            feedback_on_action_list_received=timezone.localdate(),
        )

        response = self.client.post(
            reverse(
                "synopsis:advisory_member_set_action_list_flag",
                args=[self.project.id, member.id, "author-replied"],
            ),
            {"value": "1"},
        )

        self.assertRedirects(response, self.board_url)
        member.refresh_from_db()
        self.assertTrue(member.wm_replied)
        self.assertTrue(
            ProjectChangeLog.objects.filter(
                project=self.project,
                action="Updated action list tracking",
                details__contains="Marked author replied",
            ).exists()
        )

    def test_cannot_toggle_action_list_author_replied_before_feedback_received(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Early",
            last_name="Reply",
            email="early@example.com",
            response="Y",
            participation_confirmed=True,
            sent_action_list_at=timezone.now(),
        )

        response = self.client.post(
            reverse(
                "synopsis:advisory_member_set_action_list_flag",
                args=[self.project.id, member.id, "author-replied"],
            ),
            {"value": "1"},
            follow=True,
        )

        self.assertRedirects(response, self.board_url)
        member.refresh_from_db()
        self.assertFalse(member.wm_replied)
        self.assertContains(
            response,
            "Action list tracking can only be updated after feedback has been received.",
        )

    def test_action_list_tracking_no_longer_exposes_guidance_feedback_flag(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Giad",
            last_name="Guidance",
            email="giad@example.com",
            response="Y",
            participation_confirmed=True,
            sent_action_list_at=timezone.now(),
            feedback_on_action_list_received=timezone.localdate(),
        )

        response = self.client.post(
            reverse(
                "synopsis:advisory_member_set_action_list_flag",
                args=[self.project.id, member.id, "guidance-feedback"],
            ),
            {"value": "1"},
        )

        self.assertEqual(response.status_code, 400)

    def test_board_page_shows_readable_action_list_tracking_controls(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Raed",
            last_name="Readable",
            email="raed@example.com",
            response="Y",
            participation_confirmed=True,
            sent_action_list_at=timezone.now(),
            feedback_on_action_list_received=timezone.localdate(),
        )

        response = self.client.get(self.board_url)

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Not yet")
        self.assertContains(response, "Mark replied")
        self.assertContains(response, "Not added")
        self.assertContains(response, "Mark added")
        self.assertNotContains(response, "Not marked")
        self.assertNotContains(response, "Mark guidance")

    def test_can_toggle_protocol_tracking_flags_from_board(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Pia",
            last_name="Protocol",
            email="pia@example.com",
            response="Y",
            participation_confirmed=True,
            sent_protocol_at=timezone.now(),
            feedback_on_protocol_received=timezone.localdate(),
        )

        author_response = self.client.post(
            reverse(
                "synopsis:advisory_member_set_protocol_flag",
                args=[self.project.id, member.id, "author-replied"],
            ),
            {"value": "1"},
        )
        added_response = self.client.post(
            reverse(
                "synopsis:advisory_member_set_protocol_flag",
                args=[self.project.id, member.id, "added-to-doc"],
            ),
            {"value": "1"},
        )
        self.assertRedirects(author_response, self.board_url)
        self.assertRedirects(added_response, self.board_url)
        member.refresh_from_db()
        self.assertTrue(member.protocol_author_replied)
        self.assertTrue(member.added_to_protocol_doc)

    def test_protocol_tracking_no_longer_exposes_guidance_feedback_flag(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Pia",
            last_name="Protocol",
            email="pia@example.com",
            response="Y",
            participation_confirmed=True,
            sent_protocol_at=timezone.now(),
            feedback_on_protocol_received=timezone.localdate(),
        )

        response = self.client.post(
            reverse(
                "synopsis:advisory_member_set_protocol_flag",
                args=[self.project.id, member.id, "guidance-feedback"],
            ),
            {"value": "1"},
        )

        self.assertEqual(response.status_code, 400)

    def test_can_toggle_synopsis_tracking_flags_from_board(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Sia",
            last_name="Synopsis",
            email="sia@example.com",
            response="Y",
            participation_confirmed=True,
            sent_synopsis_at=timezone.now(),
            feedback_on_synopsis_received=timezone.localdate(),
        )

        author_response = self.client.post(
            reverse(
                "synopsis:advisory_member_set_synopsis_flag",
                args=[self.project.id, member.id, "author-replied"],
            ),
            {"value": "1"},
        )
        added_response = self.client.post(
            reverse(
                "synopsis:advisory_member_set_synopsis_flag",
                args=[self.project.id, member.id, "added-to-doc"],
            ),
            {"value": "1"},
        )

        self.assertRedirects(author_response, self.board_url)
        self.assertRedirects(added_response, self.board_url)
        member.refresh_from_db()
        self.assertTrue(member.synopsis_author_replied)
        self.assertTrue(member.added_to_synopsis_doc)

    def test_can_mark_synopsis_feedback_received_from_board(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Sia",
            last_name="Feedback",
            email="sia-feedback@example.com",
            response="Y",
            participation_confirmed=True,
            sent_synopsis_at=timezone.now(),
        )

        response = self.client.post(
            reverse(
                "synopsis:advisory_member_set_synopsis_flag",
                args=[self.project.id, member.id, "feedback-received"],
            ),
            {"value": "1"},
        )

        self.assertRedirects(response, self.board_url)
        member.refresh_from_db()
        self.assertEqual(member.feedback_on_synopsis_received, timezone.localdate())

    def test_synopsis_feedback_link_accepts_comments_and_document(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Sia",
            last_name="Reviewer",
            email="sia-reviewer@example.com",
            response="Y",
            participation_confirmed=True,
            sent_synopsis_at=timezone.now(),
            feedback_on_synopsis_deadline=timezone.now() + timedelta(days=4),
        )
        feedback = SynopsisFeedback.objects.create(
            project=self.project,
            member=member,
            email=member.email,
            feedback_deadline_at=member.feedback_on_synopsis_deadline,
        )
        get_response = self.client.get(
            reverse("synopsis:synopsis_feedback", args=[str(feedback.token)])
        )
        self.assertEqual(get_response.status_code, 200)
        self.assertContains(get_response, "How your information is used")
        self.assertContains(get_response, "authorised project authors and managers")
        uploaded_doc = SimpleUploadedFile(
            "synopsis-comments.pdf",
            b"annotated synopsis",
            content_type="application/pdf",
        )

        response = self.client.post(
            reverse("synopsis:synopsis_feedback", args=[str(feedback.token)]),
            {
                "content": "Please clarify the evidence summary.",
                "uploaded_document": uploaded_doc,
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Thank you")
        feedback.refresh_from_db()
        member.refresh_from_db()
        self.assertEqual(feedback.content, "Please clarify the evidence summary.")
        self.assertTrue(feedback.uploaded_document.name)
        self.assertEqual(member.feedback_on_synopsis_received, timezone.localdate())

    def test_board_page_shows_readable_protocol_and_synopsis_tracking_controls(self):
        protocol_member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Paul",
            email="paul@example.com",
            response="Y",
            participation_confirmed=True,
            sent_protocol_at=timezone.now(),
            feedback_on_protocol_received=timezone.localdate(),
        )
        synopsis_member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Sara",
            email="sara@example.com",
            response="Y",
            participation_confirmed=True,
            sent_synopsis_at=timezone.now(),
            feedback_on_synopsis_received=timezone.localdate(),
        )

        response = self.client.get(self.board_url)

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            reverse(
                "synopsis:advisory_member_set_protocol_flag",
                args=[self.project.id, protocol_member.id, "author-replied"],
            ),
        )
        self.assertContains(
            response,
            reverse(
                "synopsis:advisory_member_set_synopsis_flag",
                args=[self.project.id, synopsis_member.id, "author-replied"],
            ),
        )
        self.assertNotContains(
            response,
            "synopsis:advisory_member_set_guidance_flag",
        )
        self.assertContains(response, "Mark replied", count=2)
        self.assertContains(response, "Not added", count=2)

    def test_board_page_shows_action_list_author_replied_as_awaiting_before_feedback(self):
        AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Wait",
            last_name="Feedback",
            email="wait@example.com",
            response="Y",
            participation_confirmed=True,
            sent_action_list_at=timezone.now(),
        )

        response = self.client.get(self.board_url)

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Awaiting feedback")

    def test_board_page_shows_synopsis_send_and_tracking_controls(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Sian",
            email="sian@example.com",
            response="Y",
            participation_confirmed=True,
            sent_synopsis_at=timezone.now(),
        )

        response = self.client.get(self.board_url)

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            reverse("synopsis:advisory_send_synopsis_compose_all", args=[self.project.id]),
        )
        self.assertContains(
            response,
            reverse(
                "synopsis:advisory_member_set_deadline",
                args=[self.project.id, member.id, "synopsis"],
            ),
        )
        self.assertContains(
            response,
            reverse(
                "synopsis:advisory_member_set_synopsis_flag",
                args=[self.project.id, member.id, "feedback-received"],
            ),
        )
        self.assertContains(response, "Mark received")

    def test_clear_response_deadline(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Clara",
            email="clara@example.com",
            response_date=date(2024, 12, 1),
            reminder_sent=True,
            reminder_sent_at=timezone.now(),
        )
        response = self.client.post(
            reverse(
                "synopsis:advisory_member_set_deadline",
                args=[self.project.id, member.id, "invite"],
            ),
            {"clear_deadline": "1"},
        )
        self.assertRedirects(response, self.board_url)
        member.refresh_from_db()
        self.assertIsNone(member.response_date)
        self.assertFalse(member.reminder_sent)
        self.assertIsNone(member.reminder_sent_at)

    def test_same_day_response_deadline_is_rejected(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Clara",
            email="clara@example.com",
        )
        today = timezone.localdate()
        response = self.client.post(
            reverse(
                "synopsis:advisory_member_set_deadline",
                args=[self.project.id, member.id, "invite"],
            ),
            {"reminder_date": today.strftime("%Y-%m-%d")},
            follow=True,
        )
        self.assertRedirects(response, self.board_url)
        member.refresh_from_db()
        self.assertIsNone(member.response_date)
        self.assertContains(response, "at least one day in the future")

    def test_update_protocol_deadline(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Paula",
            email="paula@example.com",
            response="Y",
            sent_protocol_at=timezone.now(),
            protocol_reminder_sent=True,
            protocol_reminder_sent_at=timezone.now(),
        )
        ProtocolFeedback.objects.create(project=self.project, member=member)
        deadline = timezone.now().replace(second=0, microsecond=0) + timedelta(days=4)
        local_deadline = timezone.localtime(deadline)
        response = self.client.post(
            reverse(
                "synopsis:advisory_member_set_deadline",
                args=[self.project.id, member.id, "protocol"],
            ),
            {"deadline": local_deadline.strftime("%Y-%m-%dT%H:%M")},
        )
        self.assertRedirects(response, self.board_url)
        member.refresh_from_db()
        self.assertIsNotNone(member.feedback_on_protocol_deadline)
        self.assertAlmostEqual(
            member.feedback_on_protocol_deadline.timestamp(), deadline.timestamp(), delta=1
        )
        self.assertFalse(member.protocol_reminder_sent)
        self.assertIsNone(member.protocol_reminder_sent_at)
        feedback = ProtocolFeedback.objects.get(project=self.project, member=member)
        self.assertIsNotNone(feedback.feedback_deadline_at)
        self.assertAlmostEqual(
            feedback.feedback_deadline_at.timestamp(), deadline.timestamp(), delta=1
        )
        self.assertTrue(
            ProjectChangeLog.objects.filter(
                project=self.project, action="Updated protocol reminder"
            ).exists()
        )

    def test_same_day_protocol_deadline_is_rejected(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Paula",
            email="paula@example.com",
            response="Y",
            sent_protocol_at=timezone.now(),
        )
        ProtocolFeedback.objects.create(project=self.project, member=member)
        today_local = timezone.localtime(timezone.now()).replace(second=0, microsecond=0)
        response = self.client.post(
            reverse(
                "synopsis:advisory_member_set_deadline",
                args=[self.project.id, member.id, "protocol"],
            ),
            {"deadline": today_local.strftime("%Y-%m-%dT%H:%M")},
            follow=True,
        )
        self.assertRedirects(response, self.board_url)
        member.refresh_from_db()
        self.assertIsNone(member.feedback_on_protocol_deadline)
        self.assertContains(response, "at least one day in the future")

    def test_protocol_page_deadline_warns_before_protocol_is_sent(self):
        Protocol.objects.create(
            project=self.project,
            document=SimpleUploadedFile("protocol.docx", b"protocol"),
        )
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Paula",
            email="paula@example.com",
            response="Y",
            participation_confirmed=True,
        )
        local_deadline = timezone.localtime(timezone.now() + timedelta(days=5)).replace(
            second=0,
            microsecond=0,
        )

        response = self.client.post(
            reverse("synopsis:advisory_schedule_protocol_reminders", args=[self.project.id]),
            {"deadline": local_deadline.strftime("%Y-%m-%dT%H:%M")},
            follow=True,
        )

        self.assertRedirects(
            response, reverse("synopsis:protocol_detail", args=[self.project.id])
        )
        member.refresh_from_db()
        self.assertIsNone(member.feedback_on_protocol_deadline)
        self.assertContains(
            response,
            "No protocol deadline was updated because no accepted advisory board member has been sent the protocol yet.",
        )

    def test_update_action_list_deadline(self):
        action_list = ActionList.objects.create(
            project=self.project,
            document=SimpleUploadedFile("action-list.txt", b"test"),
        )
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Vanessa",
            email="vanessa@example.com",
            response="Y",
            sent_action_list_at=timezone.now(),
            action_list_reminder_sent=True,
            action_list_reminder_sent_at=timezone.now(),
        )
        ActionListFeedback.objects.create(
            project=self.project, member=member, action_list=action_list
        )
        deadline = timezone.now().replace(second=0, microsecond=0) + timedelta(days=6)
        local_deadline = timezone.localtime(deadline)
        response = self.client.post(
            reverse(
                "synopsis:advisory_member_set_deadline",
                args=[self.project.id, member.id, "action-list"],
            ),
            {"deadline": local_deadline.strftime("%Y-%m-%dT%H:%M")},
        )
        self.assertRedirects(response, self.board_url)
        member.refresh_from_db()
        self.assertIsNotNone(member.feedback_on_action_list_deadline)
        self.assertAlmostEqual(
            member.feedback_on_action_list_deadline.timestamp(),
            deadline.timestamp(),
            delta=1,
        )
        self.assertFalse(member.action_list_reminder_sent)
        self.assertIsNone(member.action_list_reminder_sent_at)
        feedback = ActionListFeedback.objects.get(project=self.project, member=member)
        self.assertIsNotNone(feedback.feedback_deadline_at)
        self.assertAlmostEqual(
            feedback.feedback_deadline_at.timestamp(),
            deadline.timestamp(),
            delta=1,
        )
        self.assertTrue(
            ProjectChangeLog.objects.filter(
                project=self.project, action="Updated action list reminder"
            ).exists()
        )

    def test_update_synopsis_deadline(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Sam",
            email="sam@example.com",
            response="Y",
            sent_synopsis_at=timezone.now(),
            synopsis_reminder_sent=True,
            synopsis_reminder_sent_at=timezone.now(),
        )
        deadline = timezone.now().replace(second=0, microsecond=0) + timedelta(days=5)
        local_deadline = timezone.localtime(deadline)

        response = self.client.post(
            reverse(
                "synopsis:advisory_member_set_deadline",
                args=[self.project.id, member.id, "synopsis"],
            ),
            {"deadline": local_deadline.strftime("%Y-%m-%dT%H:%M")},
        )

        self.assertRedirects(response, self.board_url)
        member.refresh_from_db()
        self.assertIsNotNone(member.feedback_on_synopsis_deadline)
        self.assertAlmostEqual(
            member.feedback_on_synopsis_deadline.timestamp(),
            deadline.timestamp(),
            delta=1,
        )
        self.assertFalse(member.synopsis_reminder_sent)
        self.assertIsNone(member.synopsis_reminder_sent_at)
        self.assertTrue(
            ProjectChangeLog.objects.filter(
                project=self.project, action="Updated synopsis reminder"
            ).exists()
        )

    def test_action_list_page_deadline_warns_before_action_list_is_sent(self):
        ActionList.objects.create(
            project=self.project,
            document=SimpleUploadedFile("action-list.docx", b"action list"),
        )
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Vanessa",
            email="vanessa@example.com",
            response="Y",
            participation_confirmed=True,
        )
        local_deadline = timezone.localtime(timezone.now() + timedelta(days=5)).replace(
            second=0,
            microsecond=0,
        )

        response = self.client.post(
            reverse(
                "synopsis:advisory_schedule_action_list_reminders",
                args=[self.project.id],
            ),
            {"deadline": local_deadline.strftime("%Y-%m-%dT%H:%M")},
            follow=True,
        )

        self.assertRedirects(
            response, reverse("synopsis:action_list_detail", args=[self.project.id])
        )
        member.refresh_from_db()
        self.assertIsNone(member.feedback_on_action_list_deadline)
        self.assertContains(
            response,
            "No action list deadline was updated because no accepted advisory board member has been sent the action list yet.",
        )

    def test_declined_member_cannot_schedule_invite(self):
        original_date = date(2024, 11, 15)
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Dana",
            email="dana@example.com",
            response="N",
            response_date=original_date,
            reminder_sent=True,
            reminder_sent_at=timezone.now(),
        )
        response = self.client.post(
            reverse(
                "synopsis:advisory_member_set_deadline",
                args=[self.project.id, member.id, "invite"],
            ),
            {"reminder_date": "2025-01-01"},
        )
        self.assertRedirects(response, self.board_url)
        member.refresh_from_db()
        self.assertEqual(member.response_date, original_date)
        self.assertTrue(member.reminder_sent)

    def test_decline_captures_optional_reason(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Iggy",
            email="iggy@example.com",
        )
        inv = AdvisoryBoardInvitation.objects.create(
            project=self.project,
            member=member,
            email=member.email,
        )
        url = reverse("synopsis:advisory_invite_reply", args=[str(inv.token), "no"])
        response = self.client.get(url)
        self.assertEqual(response.status_code, 200)

        decline_reason = "Unavailable this cycle"
        response = self.client.post(url, data={"reason": decline_reason})
        self.assertEqual(response.status_code, 200)

        member.refresh_from_db()
        inv.refresh_from_db()
        self.assertEqual(member.response, "N")
        self.assertEqual(member.participation_statement, decline_reason)
        self.assertFalse(inv.accepted)
        self.assertIsNotNone(inv.responded_at)
        self.assertFalse(inv.accepted)

    def test_edit_member_details(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            title="Dr",
            first_name="Rebecca",
            last_name="Smith",
            email="rebecca@example.com",
            organisation="Org",
        )

        payload = {
            "title": "Prof",
            "first_name": "Rebecca",
            "middle_name": "A",
            "last_name": "Thornton",
            "organisation": "Updated Org",
            "email": "rebecca@example.com",
            "country": "United Kingdom",
            "continent": "Europe",
            "notes": "Updated notes",
        }

        url = reverse(
            "synopsis:advisory_member_edit",
            args=[self.project.id, member.id],
        )
        response = self.client.post(url, data=payload)

        self.assertRedirects(response, self.board_url)
        member.refresh_from_db()
        self.assertEqual(member.title, "Prof")
        self.assertEqual(member.last_name, "Thornton")
        self.assertEqual(member.organisation, "Updated Org")
        self.assertEqual(member.country, "United Kingdom")
        self.assertTrue(
            ProjectChangeLog.objects.filter(
                project=self.project,
                action="Updated advisory member",
            ).exists()
        )

    def test_edit_member_page_includes_delete_action(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            title="Dr",
            first_name="Rebecca",
            last_name="Smith",
            email="rebecca@example.com",
        )

        url = reverse(
            "synopsis:advisory_member_edit",
            args=[self.project.id, member.id],
        )
        response = self.client.get(url)

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Delete advisory board member")
        self.assertContains(response, 'name="action" value="delete_member"')

    def test_delete_member_from_edit_page(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            title="Dr",
            first_name="Rebecca",
            last_name="Smith",
            email="rebecca@example.com",
        )
        custom_field = AdvisoryBoardCustomField.objects.create(
            project=self.project,
            name="Expertise",
            data_type=AdvisoryBoardCustomField.TYPE_TEXT,
        )
        custom_field.set_value_for_member(member, "Wetlands", changed_by=self.user)
        invitation = AdvisoryBoardInvitation.objects.create(
            project=self.project,
            member=member,
            email=member.email,
            invited_by=self.user,
        )
        protocol_feedback = ProtocolFeedback.objects.create(
            project=self.project,
            member=member,
            email=member.email,
            content="Protocol notes",
        )
        action_list_feedback = ActionListFeedback.objects.create(
            project=self.project,
            member=member,
            email=member.email,
            content="Action list notes",
        )

        url = reverse(
            "synopsis:advisory_member_edit",
            args=[self.project.id, member.id],
        )
        response = self.client.post(url, data={"action": "delete_member"})

        self.assertRedirects(response, self.board_url)
        self.assertFalse(
            AdvisoryBoardMember.objects.filter(pk=member.id).exists()
        )
        self.assertFalse(
            AdvisoryBoardInvitation.objects.filter(pk=invitation.id).exists()
        )
        self.assertFalse(
            custom_field.values.filter(member_id=member.id).exists()
        )
        protocol_feedback.refresh_from_db()
        action_list_feedback.refresh_from_db()
        self.assertIsNone(protocol_feedback.member)
        self.assertIsNone(action_list_feedback.member)
        self.assertTrue(
            ProjectChangeLog.objects.filter(
                project=self.project,
                action="Deleted advisory member",
                details__contains="Rebecca Smith",
            ).exists()
        )

    def test_decline_reason_length_validation(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Lee",
            email="lee@example.com",
        )
        inv = AdvisoryBoardInvitation.objects.create(
            project=self.project,
            member=member,
            email=member.email,
        )
        url = reverse("synopsis:advisory_invite_reply", args=[str(inv.token), "no"])
        long_reason = "x" * 201
        response = self.client.post(url, data={"reason": long_reason})

        self.assertEqual(response.status_code, 200)
        form = response.context.get("form")
        self.assertIsNotNone(form)
        self.assertTrue(form.errors)
        self.assertIn("reason", form.errors)

        member.refresh_from_db()
        inv.refresh_from_db()
        self.assertEqual(member.participation_statement, "")
        self.assertIsNone(inv.accepted)


class AdvisoryInviteFlowTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Invite Flow")
        self.user = User.objects.create_user(
            username="author", password="pw", email="author@example.com"
        )
        UserRole.objects.create(user=self.user, project=self.project, role="author")
        self.client.force_login(self.user)
        self.board_url = reverse("synopsis:advisory_board_list", args=[self.project.id])

    def assert_public_nav_actions_hidden(self, response):
        self.assertNotContains(response, ">Home</a>")
        self.assertNotContains(response, ">Login</a>")
        self.assertNotContains(response, ">Logout</button>")
        self.assertNotContains(response, ">Create New Synopsis</a>")

    def test_public_invitation_pages_hide_author_navigation_buttons(self):
        self.client.logout()
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Ibrahim",
            email="ibrahim@example.com",
        )
        invitation = AdvisoryBoardInvitation.objects.create(
            project=self.project,
            member=member,
            email=member.email,
        )

        response = self.client.get(
            reverse(
                "synopsis:advisory_invite_reply",
                args=[str(invitation.token), "yes"],
            )
        )
        self.assertEqual(response.status_code, 200)
        self.assert_public_nav_actions_hidden(response)

        response = self.client.post(
            reverse(
                "synopsis:advisory_invite_reply",
                args=[str(invitation.token), "yes"],
            ),
            {"confirm_participation": "on", "statement": "Happy to help."},
        )
        self.assertEqual(response.status_code, 200)
        self.assert_public_nav_actions_hidden(response)

    def test_public_feedback_pages_hide_author_navigation_buttons(self):
        self.client.logout()
        protocol_feedback = ProtocolFeedback.objects.create(
            project=self.project,
            email="reviewer@example.com",
        )
        action_list_feedback = ActionListFeedback.objects.create(
            project=self.project,
            email="reviewer@example.com",
        )

        response = self.client.get(
            reverse(
                "synopsis:protocol_feedback",
                args=[str(protocol_feedback.token)],
            )
        )
        self.assertEqual(response.status_code, 200)
        self.assert_public_nav_actions_hidden(response)
        self.assertContains(response, "How your information is used")
        self.assertContains(response, "authorised project authors and managers")

        response = self.client.get(
            reverse(
                "synopsis:action_list_feedback",
                args=[str(action_list_feedback.token)],
            )
        )
        self.assertEqual(response.status_code, 200)
        self.assert_public_nav_actions_hidden(response)
        self.assertContains(response, "How your information is used")
        self.assertContains(response, "authorised project authors and managers")

    def test_legacy_accept_link_redirects_to_participation_confirmation(self):
        self.client.logout()
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Ibrahim",
            email="ibrahim@example.com",
        )
        invitation = AdvisoryBoardInvitation.objects.create(
            project=self.project,
            member=member,
            email=member.email,
        )

        response = self.client.get(
            reverse("synopsis:advisory_invite_accept", args=[str(invitation.token)])
        )

        self.assertRedirects(
            response,
            reverse(
                "synopsis:advisory_invite_reply",
                args=[str(invitation.token), "yes"],
            )
            + "?source=legacy_accept",
        )
        invitation.refresh_from_db()
        member.refresh_from_db()
        self.assertIsNone(invitation.accepted)
        self.assertFalse(member.participation_confirmed)
        self.assertTrue(
            ProjectChangeLog.objects.filter(
                project=self.project,
                action="Opened advisory invite link",
                details__contains="Source: legacy accept link",
            ).exists()
        )

    def test_participation_confirmation_page_explains_two_step_flow(self):
        self.client.logout()
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Ibrahim",
            email="ibrahim@example.com",
        )
        invitation = AdvisoryBoardInvitation.objects.create(
            project=self.project,
            member=member,
            email=member.email,
        )

        response = self.client.get(
            reverse(
                "synopsis:advisory_invite_reply",
                args=[str(invitation.token), "yes"],
            )
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Step 1 of 2")
        self.assertContains(
            response,
            "Your response is not recorded until you submit the form below.",
        )
        self.assertContains(
            response,
            "you will see a separate thank-you page confirming that your response has been recorded",
        )
        self.assertContains(response, "How your information is used")
        self.assertContains(response, "Lawful basis used by Conservation Evidence")

    @patch("synopsis.views.EmailMultiAlternatives")
    def test_single_invite_sets_due_date_and_resets_flags(self, mock_email):
        mock_email.return_value = MagicMock()
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Ibrahim",
            email="ibrahim@example.com",
            reminder_sent=True,
            reminder_sent_at=timezone.now(),
        )
        due = timezone.localdate() + timedelta(days=14)
        url = reverse(
            "synopsis:advisory_invite_create_for_member",
            args=[self.project.id, member.id],
        )
        response = self.client.post(
            url,
            {
                "email": member.email,
                "due_date": due.strftime("%Y-%m-%d"),
                "message": "Welcome aboard",
            },
        )
        self.assertRedirects(response, self.board_url)
        member.refresh_from_db()
        self.assertTrue(member.invite_sent)
        self.assertEqual(member.response_date, due)
        self.assertFalse(member.reminder_sent)
        self.assertIsNone(member.reminder_sent_at)
        self.assertEqual(
            AdvisoryBoardInvitation.objects.filter(member=member).count(), 1
        )
        self.assertEqual(mock_email.call_count, 1)

    @patch("synopsis.views.EmailMultiAlternatives")
    def test_single_invite_defaults_due_date_to_configured_window_when_blank(self, mock_email):
        mock_email.return_value = MagicMock()
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Ibrahim",
            email="ibrahim@example.com",
        )
        url = reverse(
            "synopsis:advisory_invite_create_for_member",
            args=[self.project.id, member.id],
        )

        response = self.client.post(
            url,
            {
                "email": member.email,
                "due_date": "",
                "message": "Welcome aboard",
            },
        )

        expected_due = timezone.localdate() + timedelta(
            days=settings.ADVISORY_INVITE_RESPONSE_WINDOW_DAYS
        )
        self.assertRedirects(response, self.board_url)
        member.refresh_from_db()
        invitation = AdvisoryBoardInvitation.objects.get(member=member)
        self.assertEqual(member.response_date, expected_due)
        self.assertEqual(invitation.due_date, expected_due)
        self.assertEqual(mock_email.call_count, 1)

    @patch("synopsis.views.EmailMultiAlternatives")
    def test_single_invite_uses_default_standard_message_when_project_has_no_custom_message(
        self, mock_email
    ):
        email_instance = MagicMock()
        mock_email.return_value = email_instance
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Ibrahim",
            email="ibrahim@example.com",
        )
        due = timezone.localdate() + timedelta(days=10)

        response = self.client.post(
            reverse(
                "synopsis:advisory_invite_create_for_member",
                args=[self.project.id, member.id],
            ),
            {
                "email": member.email,
                "due_date": due.strftime("%Y-%m-%d"),
                "message": "",
            },
        )

        self.assertRedirects(response, self.board_url)
        args, _kwargs = mock_email.call_args
        self.assertIn(default_advisory_invitation_message(), args[1])
        html_body = email_instance.attach_alternative.call_args[0][0]
        self.assertIn(default_advisory_invitation_message(), html_body)
        self.project.refresh_from_db()
        self.assertEqual(self.project.advisory_invitation_message, "")

    def test_single_invite_form_shows_preview_with_default_standard_message(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Ibrahim",
            email="ibrahim@example.com",
        )

        response = self.client.get(
            reverse(
                "synopsis:advisory_invite_create_for_member",
                args=[self.project.id, member.id],
            )
        )

        self.assertContains(response, "Invitation preview")
        self.assertContains(response, default_advisory_invitation_message())

    @patch("synopsis.views.EmailMultiAlternatives")
    def test_single_invite_missing_standard_message_field_keeps_saved_project_message(
        self, mock_email
    ):
        email_instance = MagicMock()
        mock_email.return_value = email_instance
        self.project.advisory_invitation_message = "Saved standard message"
        self.project.save(update_fields=["advisory_invitation_message"])
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Ibrahim",
            email="ibrahim@example.com",
        )
        due = timezone.localdate() + timedelta(days=14)

        response = self.client.post(
            reverse(
                "synopsis:advisory_invite_create_for_member",
                args=[self.project.id, member.id],
            ),
            {
                "email": member.email,
                "due_date": due.strftime("%Y-%m-%d"),
                "message": "Welcome aboard",
            },
        )

        self.assertRedirects(response, self.board_url)
        self.project.refresh_from_db()
        self.assertEqual(
            self.project.advisory_invitation_message, "Saved standard message"
        )
        args, _kwargs = mock_email.call_args
        self.assertIn("Saved standard message", args[1])
        html_body = email_instance.attach_alternative.call_args[0][0]
        self.assertIn("Saved standard message", html_body)

    @patch("synopsis.views.EmailMultiAlternatives")
    def test_bulk_invite_skips_members_with_existing_invites(self, mock_email):
        mock_email.return_value = MagicMock()
        already_invited = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Iris",
            email="iris@example.com",
            invite_sent=True,
            invite_sent_at=timezone.now(),
            response_date=date(2025, 10, 1),
        )
        new_member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Liam",
            email="liam@example.com",
        )
        due = timezone.localdate() + timedelta(days=21)
        response = self.client.post(
            reverse("synopsis:advisory_send_invites_bulk", args=[self.project.id]),
            {
                "due_date": due.strftime("%Y-%m-%d"),
                "message": "Bulk kickoff",
            },
        )
        self.assertRedirects(response, self.board_url)
        new_member.refresh_from_db()
        already_invited.refresh_from_db()
        self.assertTrue(new_member.invite_sent)
        self.assertEqual(new_member.response_date, due)
        self.assertEqual(already_invited.response_date, date(2025, 10, 1))
        self.assertEqual(
            AdvisoryBoardInvitation.objects.filter(project=self.project).count(), 1
        )
        self.assertEqual(mock_email.call_count, 1)
        args, kwargs = mock_email.call_args
        self.assertEqual(kwargs["to"], [new_member.email])

    @patch("synopsis.views.EmailMultiAlternatives")
    def test_bulk_invite_saves_custom_standard_message_and_keeps_optional_message(
        self, mock_email
    ):
        email_instance = MagicMock()
        mock_email.return_value = email_instance
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Liam",
            email="liam@example.com",
        )
        due = timezone.localdate() + timedelta(days=21)

        response = self.client.post(
            reverse("synopsis:advisory_send_invites_bulk", args=[self.project.id]),
            {
                "due_date": due.strftime("%Y-%m-%d"),
                "standard_message": "This is the saved team invite message.",
                "message": "Bulk kickoff note",
            },
        )

        self.assertRedirects(response, self.board_url)
        self.project.refresh_from_db()
        self.assertEqual(
            self.project.advisory_invitation_message,
            "This is the saved team invite message.",
        )
        args, kwargs = mock_email.call_args
        self.assertEqual(kwargs["to"], [member.email])
        self.assertIn("This is the saved team invite message.", args[1])
        self.assertIn("Bulk kickoff note", args[1])
        html_body = email_instance.attach_alternative.call_args[0][0]
        self.assertIn("This is the saved team invite message.", html_body)
        self.assertIn("Bulk kickoff note", html_body)

    def test_bulk_invite_form_shows_preview_with_default_standard_message(self):
        response = self.client.get(
            reverse("synopsis:advisory_send_invites_bulk", args=[self.project.id])
        )

        self.assertContains(response, "Invitation preview")
        self.assertContains(response, default_advisory_invitation_message())

    @patch("synopsis.views.EmailMultiAlternatives")
    def test_bulk_invite_missing_standard_message_field_keeps_saved_project_message(
        self, mock_email
    ):
        email_instance = MagicMock()
        mock_email.return_value = email_instance
        self.project.advisory_invitation_message = "Saved standard message"
        self.project.save(update_fields=["advisory_invitation_message"])
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Liam",
            email="liam@example.com",
        )
        due = timezone.localdate() + timedelta(days=21)

        response = self.client.post(
            reverse("synopsis:advisory_send_invites_bulk", args=[self.project.id]),
            {
                "due_date": due.strftime("%Y-%m-%d"),
                "message": "Bulk kickoff",
            },
        )

        self.assertRedirects(response, self.board_url)
        self.project.refresh_from_db()
        self.assertEqual(
            self.project.advisory_invitation_message, "Saved standard message"
        )
        args, kwargs = mock_email.call_args
        self.assertEqual(kwargs["to"], [member.email])
        self.assertIn("Saved standard message", args[1])
        html_body = email_instance.attach_alternative.call_args[0][0]
        self.assertIn("Saved standard message", html_body)


class AdvisoryDocumentSendDefaultDeadlineTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Document Deadlines")
        self.user = User.objects.create_user(
            username="author", password="pw", email="author@example.com"
        )
        UserRole.objects.create(user=self.user, project=self.project, role="author")
        self.client.force_login(self.user)
        self.protocol = Protocol.objects.create(
            project=self.project,
            document=SimpleUploadedFile("protocol.txt", b"protocol"),
        )
        self.action_list = ActionList.objects.create(
            project=self.project,
            document=SimpleUploadedFile("action-list.txt", b"action list"),
        )
        self.member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Rebecca",
            last_name="Smith",
            email="rebecca@example.com",
            response="Y",
            participation_confirmed=True,
        )
        self.expected_due = timezone.localdate() + timedelta(
            days=settings.ADVISORY_DOCUMENT_FEEDBACK_WINDOW_DAYS
        )

    def test_document_send_pages_show_email_previews(self):
        pages = [
            (
                reverse(
                    "synopsis:advisory_send_protocol_compose_member",
                    args=[self.project.id, self.member.id],
                ),
                "Protocol email preview",
                "data-document-kind=\"protocol\"",
            ),
            (
                reverse(
                    "synopsis:advisory_send_action_list_compose_member",
                    args=[self.project.id, self.member.id],
                ),
                "Action list email preview",
                "data-document-kind=\"action list\"",
            ),
            (
                reverse(
                    "synopsis:advisory_send_synopsis_compose_member",
                    args=[self.project.id, self.member.id],
                ),
                "Synopsis email preview",
                "data-document-kind=\"synopsis\"",
            ),
        ]

        for url, heading, kind_attr in pages:
            with self.subTest(url=url):
                response = self.client.get(url)
                self.assertEqual(response.status_code, 200)
                self.assertContains(response, heading)
                self.assertContains(response, kind_attr)
                self.assertContains(response, "data-preview-subject")
                self.assertContains(response, "data-preview-body")
                self.assertContains(response, "Standard message")
                self.assertContains(response, "Rebecca Smith")

    @patch("synopsis.views.EmailMultiAlternatives")
    def test_protocol_member_send_defaults_deadline_to_configured_window(self, mock_email):
        mock_email.return_value = MagicMock()
        response = self.client.post(
            reverse(
                "synopsis:advisory_send_protocol_compose_member",
                args=[self.project.id, self.member.id],
            ),
            {
                "due_date": "",
                "message": "",
                "include_protocol_document": "on",
            },
        )

        self.assertRedirects(
            response, reverse("synopsis:advisory_board_list", args=[self.project.id])
        )
        self.member.refresh_from_db()
        self.assertEqual(self.member.feedback_on_protocol_deadline.date(), self.expected_due)
        self.assertEqual(self.member.feedback_on_protocol_deadline.hour, 23)
        self.assertEqual(self.member.feedback_on_protocol_deadline.minute, 59)
        feedback = ProtocolFeedback.objects.get(project=self.project, member=self.member)
        self.assertEqual(feedback.feedback_deadline_at, self.member.feedback_on_protocol_deadline)
        email_body = mock_email.call_args[0][1]
        self.assertIn("Dear Rebecca Smith", email_body)
        self.assertIn(default_protocol_review_message(), email_body)

    @patch("synopsis.views.EmailMultiAlternatives")
    def test_protocol_bulk_send_uses_generic_greeting(self, mock_email):
        mock_email.return_value = MagicMock()

        response = self.client.post(
            reverse(
                "synopsis:advisory_send_protocol_compose_all",
                args=[self.project.id],
            ),
            {
                "due_date": "",
                "message": "",
                "include_protocol_document": "on",
            },
        )

        self.assertRedirects(
            response, reverse("synopsis:advisory_board_list", args=[self.project.id])
        )
        email_body = mock_email.call_args[0][1]
        self.assertIn("Dear advisory board member", email_body)
        self.assertNotIn("Dear Rebecca Smith", email_body)

    @patch("synopsis.views.EmailMultiAlternatives")
    def test_action_list_member_send_defaults_deadline_to_configured_window(self, mock_email):
        mock_email.return_value = MagicMock()
        response = self.client.post(
            reverse(
                "synopsis:advisory_send_action_list_compose_member",
                args=[self.project.id, self.member.id],
            ),
            {
                "due_date": "",
                "message": "",
                "include_action_list_document": "on",
            },
        )

        self.assertRedirects(
            response, reverse("synopsis:advisory_board_list", args=[self.project.id])
        )
        self.member.refresh_from_db()
        self.assertEqual(
            self.member.feedback_on_action_list_deadline.date(), self.expected_due
        )
        self.assertEqual(self.member.feedback_on_action_list_deadline.hour, 23)
        self.assertEqual(self.member.feedback_on_action_list_deadline.minute, 59)
        feedback = ActionListFeedback.objects.get(project=self.project, member=self.member)
        self.assertEqual(
            feedback.feedback_deadline_at,
            self.member.feedback_on_action_list_deadline,
        )
        email_body = mock_email.call_args[0][1]
        self.assertIn("Dear Rebecca Smith", email_body)
        self.assertIn(default_action_list_review_message(), email_body)

    @patch("synopsis.views.EmailMultiAlternatives")
    def test_action_list_bulk_send_uses_generic_greeting(self, mock_email):
        mock_email.return_value = MagicMock()

        response = self.client.post(
            reverse(
                "synopsis:advisory_send_action_list_compose_all",
                args=[self.project.id],
            ),
            {
                "due_date": "",
                "message": "",
                "include_action_list_document": "on",
            },
        )

        self.assertRedirects(
            response, reverse("synopsis:advisory_board_list", args=[self.project.id])
        )
        email_body = mock_email.call_args[0][1]
        self.assertIn("Dear advisory board member", email_body)
        self.assertNotIn("Dear Rebecca Smith", email_body)

    @patch("synopsis.views._generate_synopsis_docx", return_value=b"docx")
    @patch("synopsis.views.EmailMultiAlternatives")
    def test_synopsis_member_send_defaults_deadline_to_configured_window(
        self, mock_email, mock_generate
    ):
        email_instance = MagicMock()
        mock_email.return_value = email_instance

        response = self.client.post(
            reverse(
                "synopsis:advisory_send_synopsis_compose_member",
                args=[self.project.id, self.member.id],
            ),
            {
                "due_date": "",
                "message": "",
            },
        )

        self.assertRedirects(
            response, reverse("synopsis:advisory_board_list", args=[self.project.id])
        )
        self.member.refresh_from_db()
        self.assertIsNotNone(self.member.sent_synopsis_at)
        self.assertEqual(
            self.member.feedback_on_synopsis_deadline.date(), self.expected_due
        )
        self.assertEqual(self.member.feedback_on_synopsis_deadline.hour, 23)
        self.assertEqual(self.member.feedback_on_synopsis_deadline.minute, 59)
        feedback = SynopsisFeedback.objects.get(project=self.project, member=self.member)
        self.assertEqual(
            feedback.feedback_deadline_at,
            self.member.feedback_on_synopsis_deadline,
        )
        email_body = mock_email.call_args[0][1]
        self.assertIn("Dear Rebecca Smith", email_body)
        self.assertIn(default_synopsis_review_message(), email_body)
        self.assertIn(str(feedback.token), email_body)
        self.assertIn("Provide feedback:", email_body)
        mock_generate.assert_called_once_with(self.project)
        email_instance.attach.assert_called_once()

    @patch("synopsis.views._generate_synopsis_docx", return_value=b"docx")
    @patch("synopsis.views.EmailMultiAlternatives")
    def test_synopsis_bulk_send_uses_generic_greeting(
        self, mock_email, mock_generate
    ):
        email_instance = MagicMock()
        mock_email.return_value = email_instance

        response = self.client.post(
            reverse(
                "synopsis:advisory_send_synopsis_compose_all",
                args=[self.project.id],
            ),
            {
                "due_date": "",
                "message": "",
            },
        )

        self.assertRedirects(
            response, reverse("synopsis:advisory_board_list", args=[self.project.id])
        )
        email_body = mock_email.call_args[0][1]
        self.assertIn("Dear advisory board member", email_body)
        self.assertNotIn("Dear Rebecca Smith", email_body)
        mock_generate.assert_called_once_with(self.project)
        email_instance.attach.assert_called_once()

    @patch("synopsis.views._generate_synopsis_docx", return_value=b"generated")
    @patch("synopsis.views.EmailMultiAlternatives")
    def test_synopsis_member_send_can_use_uploaded_attachment(
        self, mock_email, mock_generate
    ):
        email_instance = MagicMock()
        mock_email.return_value = email_instance
        uploaded_doc = SimpleUploadedFile(
            "review-draft.pdf",
            b"uploaded synopsis",
            content_type="application/pdf",
        )

        response = self.client.post(
            reverse(
                "synopsis:advisory_send_synopsis_compose_member",
                args=[self.project.id, self.member.id],
            ),
            {
                "due_date": "",
                "standard_message": "Please review this final draft.",
                "message": "Please review this version.",
                "synopsis_document": uploaded_doc,
            },
        )

        self.assertRedirects(
            response, reverse("synopsis:advisory_board_list", args=[self.project.id])
        )
        mock_generate.assert_not_called()
        email_body = mock_email.call_args[0][1]
        self.assertIn("Please review this final draft.", email_body)
        self.assertIn("Please review this version.", email_body)
        email_instance.attach.assert_called_once_with(
            "review-draft.pdf",
            b"uploaded synopsis",
            "application/pdf",
        )

class FunderUtilityTests(TestCase):
    def test_build_display_name_prefers_organisation(self):
        name = Funder.build_display_name("Org Inc", "Dr", "Ann", "Thornton")
        self.assertEqual(name, "Org Inc")

    def test_build_display_name_from_names(self):
        name = Funder.build_display_name(None, "Dr", "Ann", "Thornton")
        self.assertEqual(name, "Dr Ann Thornton")

    def test_build_display_name_default(self):
        self.assertEqual(Funder.build_display_name(None, None, None, None), "(Funder)")


class FunderFormTests(TestCase):
    def test_valid_with_only_organisation(self):
        form = FunderForm(data={"organisation": "Ocean Trust"})
        self.assertTrue(form.is_valid())
        self.assertTrue(form.has_identity_fields())
        self.assertTrue(form.has_meaningful_input())

    def test_empty_form_has_no_meaningful_input(self):
        form = FunderForm(data={})
        self.assertTrue(form.is_valid())
        self.assertFalse(form.has_meaningful_input())

    def test_notes_count_as_meaningful_input(self):
        form = FunderForm(data={"organisation_details": "Focuses on wetlands"})
        self.assertTrue(form.is_valid())
        self.assertTrue(form.has_meaningful_input())

    def test_start_date_cannot_be_after_end_date(self):
        form = FunderForm(
            data={
                "organisation": "Ocean Trust",
                "fund_start_date": "2025-02-01",
                "fund_end_date": "2025-01-01",
            }
        )
        self.assertFalse(form.is_valid())
        self.assertIn(
            "Start date cannot be after the end date.",
            form.errors.get("fund_start_date", []),
        )
        self.assertIn(
            "Start date cannot be after the end date.",
            form.errors.get("fund_end_date", []),
        )

    def test_start_end_date_valid_when_ordered(self):
        form = FunderForm(
            data={
                "organisation": "Ocean Trust",
                "fund_start_date": "2025-01-01",
                "fund_end_date": "2025-02-01",
            }
        )
        self.assertTrue(form.is_valid())


class FunderContactFormSetTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Project", start_date=date(2025, 1, 1))
        self.funder = Funder.objects.create(project=self.project, name="Seed")

    def _formset_payload(self, overrides=None):
        base = {
            "contacts-TOTAL_FORMS": "1",
            "contacts-INITIAL_FORMS": "0",
            "contacts-MIN_NUM_FORMS": "0",
            "contacts-MAX_NUM_FORMS": "1000",
            "contacts-0-title": "",
            "contacts-0-first_name": "Will",
            "contacts-0-last_name": "Morgan",
            "contacts-0-email": "",
            "contacts-0-is_primary": "",
            "contacts-0-DELETE": "",
        }
        if overrides:
            base.update(overrides)
        return base

    def test_primary_auto_selected_when_missing(self):
        payload = self._formset_payload()
        formset = FunderContactFormSet(
            data=payload, instance=self.funder, prefix="contacts"
        )
        self.assertTrue(formset.is_valid())
        self.assertTrue(formset.forms[0].cleaned_data.get("is_primary"))

    def test_primary_contact_email_optional(self):
        payload = self._formset_payload({"contacts-0-is_primary": "on", "contacts-0-email": ""})
        formset = FunderContactFormSet(
            data=payload, instance=self.funder, prefix="contacts"
        )
        self.assertTrue(formset.is_valid())
        self.assertTrue(formset.forms[0].cleaned_data.get("is_primary"))

    def test_valid_primary_contact(self):
        payload = self._formset_payload(
            {"contacts-0-is_primary": "on", "contacts-0-email": "will@example.com"}
        )
        formset = FunderContactFormSet(
            data=payload, instance=self.funder, prefix="contacts"
        )
        self.assertTrue(formset.is_valid())


class AdvisoryBoardCustomColumnsDynamicTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Dynamic Columns")
        self.editor = User.objects.create_user(username="editor")
        self.accepted = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Ann",
            email="ann@example.com",
            response="Y",
        )
        self.pending = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Vanessa",
            email="vanessa@example.com",
            response="",
        )
        self.general_field = AdvisoryBoardCustomField.objects.create(
            project=self.project,
            name="General note",
            data_type=AdvisoryBoardCustomField.TYPE_TEXT,
        )
        self.pending_only_field = AdvisoryBoardCustomField.objects.create(
            project=self.project,
            name="Follow-up date",
            data_type=AdvisoryBoardCustomField.TYPE_DATE,
            sections=[AdvisoryBoardCustomField.SECTION_PENDING],
        )
        self.general_field.set_value_for_member(self.accepted, "Confirmed")
        self.general_field.set_value_for_member(self.pending, "Need reply")
        self.pending_only_field.set_value_for_member(self.pending, date(2025, 5, 1))

    def test_section_fields_match_custom_field_configuration(self):
        context = _advisory_board_context(self.project)
        sections = {section["key"]: section for section in context["member_sections"]}

        accepted_field_ids = [
            field.id
            for field in sections[AdvisoryBoardCustomField.SECTION_ACCEPTED]["fields"]
        ]
        pending_field_ids = [
            field.id
            for field in sections[AdvisoryBoardCustomField.SECTION_PENDING]["fields"]
        ]

        self.assertEqual(accepted_field_ids, [self.general_field.id])
        self.assertEqual(
            pending_field_ids, [self.general_field.id, self.pending_only_field.id]
        )

    def test_member_rows_include_formatted_custom_values(self):
        context = _advisory_board_context(self.project)
        sections = {section["key"]: section for section in context["member_sections"]}

        accepted_member = sections[AdvisoryBoardCustomField.SECTION_ACCEPTED][
            "members"
        ][0]
        pending_member = sections[AdvisoryBoardCustomField.SECTION_PENDING]["members"][
            0
        ]

        self.assertEqual(
            accepted_member.custom_field_values[self.general_field.id], "Confirmed"
        )
        self.assertEqual(
            pending_member.custom_field_values[self.general_field.id], "Need reply"
        )
        self.assertEqual(
            pending_member.custom_field_values[self.pending_only_field.id], "2025-05-01"
        )
        self.assertNotIn(
            self.pending_only_field.id, accepted_member.custom_field_values
        )

    def test_custom_fields_list_exposes_all_configured_fields(self):
        context = _advisory_board_context(self.project)
        custom_field_ids = [field.id for field in context["custom_fields"]]
        self.assertEqual(
            custom_field_ids, [self.general_field.id, self.pending_only_field.id]
        )

    def test_custom_fields_are_grouped_by_display_area_in_board_context(self):
        invite_field = AdvisoryBoardCustomField.objects.create(
            project=self.project,
            name="Invite progress",
            data_type=AdvisoryBoardCustomField.TYPE_TEXT,
            display_group=AdvisoryBoardCustomField.DISPLAY_GROUP_INVITATION,
        )
        synopsis_field = AdvisoryBoardCustomField.objects.create(
            project=self.project,
            name="Synopsis note",
            data_type=AdvisoryBoardCustomField.TYPE_TEXT,
            sections=[AdvisoryBoardCustomField.SECTION_PENDING],
            display_group=AdvisoryBoardCustomField.DISPLAY_GROUP_SYNOPSIS,
        )

        context = _advisory_board_context(self.project)
        sections = {section["key"]: section for section in context["member_sections"]}
        pending_groups = sections[AdvisoryBoardCustomField.SECTION_PENDING][
            "fields_by_group"
        ]
        accepted_groups = sections[AdvisoryBoardCustomField.SECTION_ACCEPTED][
            "fields_by_group"
        ]

        self.assertIn(
            invite_field.id,
            [field.id for field in pending_groups[AdvisoryBoardCustomField.DISPLAY_GROUP_INVITATION]],
        )
        self.assertIn(
            invite_field.id,
            [field.id for field in accepted_groups[AdvisoryBoardCustomField.DISPLAY_GROUP_INVITATION]],
        )
        self.assertEqual(
            [field.id for field in pending_groups[AdvisoryBoardCustomField.DISPLAY_GROUP_SYNOPSIS]],
            [synopsis_field.id],
        )
        self.assertEqual(
            [field.id for field in pending_groups[AdvisoryBoardCustomField.DISPLAY_GROUP_CUSTOM]],
            [self.general_field.id, self.pending_only_field.id],
        )
        self.assertEqual(
            [field.id for field in accepted_groups[AdvisoryBoardCustomField.DISPLAY_GROUP_CUSTOM]],
            [self.general_field.id],
        )

    def test_move_custom_field_action_updates_display_group(self):
        self.client.force_login(self.editor)
        url = reverse("synopsis:advisory_board_list", args=[self.project.id])
        response = self.client.post(
            url,
            {
                "action": "custom_field_move",
                "field_id": self.general_field.id,
                "display_group": AdvisoryBoardCustomField.DISPLAY_GROUP_PROTOCOL,
            },
        )
        self.assertEqual(response.status_code, 302)
        self.general_field.refresh_from_db()
        self.assertEqual(
            self.general_field.display_group,
            AdvisoryBoardCustomField.DISPLAY_GROUP_PROTOCOL,
        )

    def test_history_records_updates(self):
        base_count = AdvisoryBoardCustomFieldValueHistory.objects.filter(
            field=self.general_field, member=self.accepted
        ).count()

        self.general_field.set_value_for_member(
            self.accepted, "Updated note", changed_by=self.editor
        )
        self.general_field.set_value_for_member(
            self.accepted, "", changed_by=self.editor
        )

        history = AdvisoryBoardCustomFieldValueHistory.objects.filter(
            field=self.general_field, member=self.accepted
        ).order_by("-created_at")

        self.assertEqual(history.count(), base_count + 2)
        latest = history.first()
        self.assertTrue(latest.is_cleared)
        previous = history[1]
        self.assertEqual(previous.value, "Updated note")
        self.assertEqual(previous.changed_by, self.editor)

    def test_history_shows_current_value_first(self):
        self.general_field.set_value_for_member(
            self.accepted, "First", changed_by=self.editor
        )
        self.general_field.set_value_for_member(
            self.accepted, "Second", changed_by=self.editor
        )

        history = list(
            AdvisoryBoardCustomFieldValueHistory.objects.filter(
                field=self.general_field, member=self.accepted
            )
        )

        self.assertGreaterEqual(len(history), 2)
        self.assertEqual(history[0].value, "Second")


class AdvisoryMemberCustomDataFormTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Form Columns")
        self.shared_field = AdvisoryBoardCustomField.objects.create(
            project=self.project,
            name="Notes",
            data_type=AdvisoryBoardCustomField.TYPE_TEXT,
        )
        self.pending_field = AdvisoryBoardCustomField.objects.create(
            project=self.project,
            name="Reminder",
            data_type=AdvisoryBoardCustomField.TYPE_BOOLEAN,
            sections=[AdvisoryBoardCustomField.SECTION_PENDING],
        )

    def test_form_includes_only_fields_for_member_section(self):
        initial_values = {self.shared_field.id: "hello"}
        accepted_form = AdvisoryMemberCustomDataForm(
            [self.shared_field, self.pending_field],
            AdvisoryBoardCustomField.SECTION_ACCEPTED,
            initial_values,
        )
        accepted_field_ids = [field.id for field, _ in accepted_form.iter_fields()]
        self.assertEqual(accepted_field_ids, [self.shared_field.id])

        pending_form = AdvisoryMemberCustomDataForm(
            [self.shared_field, self.pending_field],
            AdvisoryBoardCustomField.SECTION_PENDING,
            initial_values,
        )
        pending_field_ids = [field.id for field, _ in pending_form.iter_fields()]
        self.assertEqual(
            pending_field_ids, [self.shared_field.id, self.pending_field.id]
        )

    def test_initial_values_are_parsed_for_form_fields(self):
        initial_values = {
            self.shared_field.id: "value",
            self.pending_field.id: "true",
        }
        form = AdvisoryMemberCustomDataForm(
            [self.shared_field, self.pending_field],
            AdvisoryBoardCustomField.SECTION_PENDING,
            initial_values,
        )
        key_shared = form._field_key(self.shared_field)
        key_pending = form._field_key(self.pending_field)
        self.assertEqual(form.initial[key_shared], "value")
        self.assertTrue(form.initial[key_pending])


class AdvisoryMemberCustomDataViewTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="View Columns")
        self.manager = User.objects.create_user(username="manager", password="x")
        self.manager.is_staff = True
        self.manager.save(update_fields=["is_staff"])
        self.member = AdvisoryBoardMember.objects.create(
            project=self.project,
            email="member@example.com",
            first_name="Mia",
            response="Y",
        )
        self.field_one = AdvisoryBoardCustomField.objects.create(
            project=self.project,
            name="Engagement",
            data_type=AdvisoryBoardCustomField.TYPE_TEXT,
        )
        self.field_two = AdvisoryBoardCustomField.objects.create(
            project=self.project,
            name="Notes",
            data_type=AdvisoryBoardCustomField.TYPE_TEXT,
        )
        self.field_one.set_value_for_member(
            self.member, "Initial", changed_by=self.manager
        )
        self.field_two.set_value_for_member(
            self.member, "Aux", changed_by=self.manager
        )
        self.url = reverse(
            "synopsis:advisory_member_custom_data",
            args=[self.project.id, self.member.id],
        )
        self.client.force_login(self.manager)

    def test_focus_field_filters_form_and_history(self):
        response = self.client.get(self.url, {"field": self.field_one.id})
        self.assertEqual(response.status_code, 200)
        fields = response.context["fields"]
        self.assertEqual(len(fields), 1)
        self.assertEqual(fields[0].id, self.field_one.id)
        form_fields = list(response.context["form"].fields.keys())
        self.assertEqual(form_fields, [f"field_{self.field_one.id}"])
        history_map = response.context["history_map"]
        self.assertEqual(list(history_map.keys()), [self.field_one.id])

    def test_without_focus_shows_all_fields(self):
        response = self.client.get(self.url)
        self.assertEqual(response.status_code, 200)
        field_ids = [field.id for field in response.context["fields"]]
        self.assertCountEqual(field_ids, [self.field_one.id, self.field_two.id])


class OnlyOfficeDownloadTests(TestCase):
    def setUp(self):
        from . import views

        self.views = views
        self.original_settings = views.ONLYOFFICE_SETTINGS
        views.ONLYOFFICE_SETTINGS = {
            "base_url": "https://onlyoffice.example.com/office",
            "callback_timeout": 7,
        }
        self.addCleanup(self._restore_settings)

    def _restore_settings(self):
        self.views.ONLYOFFICE_SETTINGS = self.original_settings

    @patch("synopsis.views.requests.get")
    def test_download_allows_trusted_host(self, mock_get):
        mock_response = MagicMock()
        mock_response.content = b"doc"
        mock_response.raise_for_status.return_value = None
        mock_get.return_value = mock_response

        url = "https://onlyoffice.example.com/office/storage/doc.docx"
        content = self.views._download_onlyoffice_file(url)

        self.assertEqual(content, b"doc")
        mock_get.assert_called_once_with(url, timeout=7)

    @patch("synopsis.views.requests.get")
    def test_download_rejects_untrusted_host(self, mock_get):
        url = "https://files.example.com/storage/doc.docx"
        with self.assertRaisesMessage(ValueError, "Untrusted OnlyOffice download URL"):
            self.views._download_onlyoffice_file(url)
        mock_get.assert_not_called()

    @patch("synopsis.views.requests.get")
    def test_download_rejects_untrusted_path(self, mock_get):
        url = "https://onlyoffice.example.com/other/doc.docx"
        with self.assertRaisesMessage(ValueError, "Untrusted OnlyOffice download URL"):
            self.views._download_onlyoffice_file(url)
        mock_get.assert_not_called()

    @patch("synopsis.views.requests.get")
    def test_download_uses_internal_onlyoffice_url_when_base_is_browser_host(self, mock_get):
        self.views.ONLYOFFICE_SETTINGS = {
            "base_url": "http://localhost:8080",
            "internal_url": "http://onlyoffice",
            "callback_timeout": 7,
            "trusted_download_urls": ["http://onlyoffice"],
        }
        mock_response = MagicMock()
        mock_response.content = b"doc"
        mock_response.raise_for_status.return_value = None
        mock_get.return_value = mock_response

        content = _download_onlyoffice_file(
            "http://localhost:8080/cache/files/data/demo/output.docx?x=1"
        )

        self.assertEqual(content, b"doc")
        mock_get.assert_called_once_with(
            "http://onlyoffice/cache/files/data/demo/output.docx?x=1",
            timeout=7,
        )


class OnlyOfficeCallbackParsingTests(TestCase):
    def setUp(self):
        from . import views

        self.views = views
        self.original_settings = views.ONLYOFFICE_SETTINGS
        views.ONLYOFFICE_SETTINGS = {
            "jwt_secret": "change-me",
        }
        self.addCleanup(self._restore_settings)
        self.factory = RequestFactory()

    def _restore_settings(self):
        self.views.ONLYOFFICE_SETTINGS = self.original_settings

    def test_parse_callback_unwraps_nested_payload_from_jwt(self):
        payload = {
            "payload": {
                "status": 2,
                "url": "http://localhost:8080/cache/files/data/demo/output.docx",
            }
        }
        token = jwt.encode(payload, "change-me", algorithm="HS256")
        request = self.factory.post(
            "/",
            data=json.dumps({}),
            content_type="application/json",
            HTTP_AUTHORIZATION=f"Bearer {token}",
        )

        parsed = _parse_onlyoffice_callback(request)

        self.assertEqual(parsed["status"], 2)
        self.assertEqual(
            parsed["url"],
            "http://localhost:8080/cache/files/data/demo/output.docx",
        )


class CollaborativeClosureTests(TestCase):
    def setUp(self):
        from . import views

        self.views = views
        self.original_settings = views.ONLYOFFICE_SETTINGS
        views.ONLYOFFICE_SETTINGS = {
            "base_url": "https://onlyoffice.example.com/office",
            "callback_timeout": 5,
        }
        self.addCleanup(self._restore_settings)

        self.factory = RequestFactory()
        self.project = Project.objects.create(title="Collaborative Close")
        self.manager = User.objects.create_user(username="manager", password="pw")
        self.manager.is_staff = True
        self.manager.save(update_fields=["is_staff"])
        UserRole.objects.create(user=self.manager, project=self.project, role="manager")
        self.client.force_login(self.manager)
        self.protocol = Protocol.objects.create(
            project=self.project,
            document=SimpleUploadedFile(
                "protocol.docx",
                b"test-protocol",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        )

    def _restore_settings(self):
        self.views.ONLYOFFICE_SETTINGS = self.original_settings

    def _build_request(self):
        request = self.factory.get("/", HTTP_HOST="testserver")
        request.user = self.manager
        return request

    @patch("synopsis.views._download_onlyoffice_file", return_value=b"updated-doc")
    def test_collaborative_save_keeps_clean_filename_from_current_revision(self, mock_download):
        revision = ProtocolRevision.objects.create(
            protocol=self.protocol,
            file=SimpleUploadedFile(
                "protocol.docx",
                b"original-doc",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            stage="draft",
            original_name="protocol.docx",
            uploaded_by=self.manager,
            file_size=len(b"original-doc"),
        )
        self.protocol.current_revision = revision
        self.protocol.document.name = (
            f"protocols/{self.project.id}/"
            "11111111-1111-1111-1111-111111111111_"
            "22222222-2222-2222-2222-222222222222_protocol.docx"
        )
        self.protocol.save(update_fields=["current_revision", "document"])

        session = CollaborativeSession.objects.create(
            project=self.project,
            document_type=CollaborativeSession.DOCUMENT_PROTOCOL,
            started_by=self.manager,
            last_activity_at=timezone.now(),
            initial_protocol_revision=revision,
        )

        success = self.views._handle_collaborative_save(
            self.project,
            CollaborativeSession.DOCUMENT_PROTOCOL,
            self.protocol,
            session,
            {"url": "https://onlyoffice.example.com/office/storage/protocol.docx"},
            2,
        )

        self.assertTrue(success)
        self.protocol.refresh_from_db()
        self.assertEqual(self.protocol.current_revision.original_name, "protocol.docx")
        self.assertNotIn(
            "22222222-2222-2222-2222-222222222222_protocol.docx",
            self.protocol.current_revision.original_name,
        )
        self.assertTrue(self.protocol.document.name.endswith("_protocol.docx"))
        self.assertNotIn(
            "11111111-1111-1111-1111-111111111111_22222222-2222-2222-2222-222222222222_protocol.docx",
            self.protocol.document.name,
        )
        mock_download.assert_called_once()

    def test_closing_protocol_disables_collaborative_session(self):
        request = self._build_request()
        url = self.views._ensure_collaborative_invite_link(
            request,
            self.project,
            CollaborativeSession.DOCUMENT_PROTOCOL,
            None,
        )
        self.assertTrue(url)
        session = CollaborativeSession.objects.get(
            project=self.project, document_type=CollaborativeSession.DOCUMENT_PROTOCOL
        )
        self.assertTrue(session.is_active)

        close_url = reverse(
            "synopsis:advisory_protocol_feedback_close", args=[self.project.id]
        )
        response = self.client.post(close_url, {"message": "Window closed"})
        self.assertRedirects(
            response,
            reverse("synopsis:protocol_detail", args=[self.project.id]),
        )

        self.project = Project.objects.get(id=self.project.id)
        self.protocol = self.project.protocol

        session.refresh_from_db()
        self.assertFalse(session.is_active)
        self.assertIsNotNone(session.ended_at)

        request = self._build_request()
        disabled_url = self.views._ensure_collaborative_invite_link(
            request,
            self.project,
            CollaborativeSession.DOCUMENT_PROTOCOL,
            None,
        )
        self.assertEqual(disabled_url, "")
        self.assertEqual(
            CollaborativeSession.objects.filter(
                project=self.project,
                document_type=CollaborativeSession.DOCUMENT_PROTOCOL,
                is_active=True,
            ).count(),
            0,
        )

        response = self.client.get(url)
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Window closed")
        self.assertTemplateUsed(response, "synopsis/collaborative_editor.html")

        start_url = reverse(
            "synopsis:collaborative_start",
            args=[self.project.id, "protocol"],
        )
        response = self.client.post(start_url)
        self.assertRedirects(
            response,
            reverse("synopsis:protocol_detail", args=[self.project.id]),
        )
        self.assertEqual(
            CollaborativeSession.objects.filter(
                project=self.project,
                document_type=CollaborativeSession.DOCUMENT_PROTOCOL,
                is_active=True,
            ).count(),
            0,
        )

        reopen_response = self.client.post(close_url, {"action": "reopen"})
        self.assertRedirects(
            reopen_response,
            reverse("synopsis:protocol_detail", args=[self.project.id]),
        )
        self.project.refresh_from_db()
        self.assertIsNone(self.project.protocol.feedback_closed_at)

        response = self.client.get(url)
        self.assertEqual(response.status_code, 302)
        self.assertNotEqual(response["Location"], url)
        parsed = urlparse(response["Location"])
        new_token = parsed.path.rstrip("/").split("/")[-1]
        new_session = CollaborativeSession.objects.get(token=new_token)
        self.assertTrue(new_session.is_active)


class ProtocolUploadFlowTests(TestCase):
    def setUp(self):
        from . import views

        self.views = views
        self.original_settings = views.ONLYOFFICE_SETTINGS
        views.ONLYOFFICE_SETTINGS = {
            "base_url": "https://onlyoffice.example.com/office",
            "callback_timeout": 7,
        }
        self.addCleanup(self._restore_settings)

        self.media_dir = tempfile.mkdtemp()
        self.addCleanup(lambda: shutil.rmtree(self.media_dir, ignore_errors=True))
        override = override_settings(MEDIA_ROOT=self.media_dir)
        override.enable()
        self.addCleanup(override.disable)

        self.project = Project.objects.create(title="Ibrahim Protocol Pilot")
        self.ibrahim = User.objects.create_user(username="ibrahim", password="pw")
        UserRole.objects.create(user=self.ibrahim, project=self.project, role="author")
        self.client.force_login(self.ibrahim)

    def _restore_settings(self):
        self.views.ONLYOFFICE_SETTINGS = self.original_settings

    def _docx_upload(self, name, content):
        return SimpleUploadedFile(
            name,
            content,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def test_initial_protocol_upload_creates_revision_and_redirects(self):
        response = self.client.post(
            reverse("synopsis:protocol_detail", args=[self.project.id]),
            {
                "stage": "draft",
                "change_reason": "",
                "version_label": "v1.0",
                "document": self._docx_upload("ibrahim-protocol.docx", b"protocol"),
            },
        )

        self.assertRedirects(
            response,
            reverse("synopsis:protocol_detail", args=[self.project.id]),
        )
        protocol = Protocol.objects.get(project=self.project)
        self.assertTrue(protocol.document.name.endswith(".docx"))
        self.assertIsNotNone(protocol.current_revision)
        self.assertEqual(protocol.current_revision.version_label, "v1.0")

    def test_protocol_can_reupload_same_filename_after_delete_file(self):
        detail_url = reverse("synopsis:protocol_detail", args=[self.project.id])

        self.client.post(
            detail_url,
            {
                "stage": "draft",
                "change_reason": "",
                "version_label": "v1",
                "document": self._docx_upload("draft-protocol.docx", b"first"),
            },
        )
        protocol = Protocol.objects.get(project=self.project)
        original_document_path = protocol.document.name
        original_revision_id = protocol.current_revision_id

        response = self.client.post(
            reverse("synopsis:protocol_delete_file", args=[self.project.id])
        )
        self.assertRedirects(response, detail_url)
        protocol.refresh_from_db()
        self.assertFalse(protocol.document)
        self.assertIsNone(protocol.current_revision)

        response = self.client.post(
            detail_url,
            {
                "stage": "draft",
                "change_reason": "Replacing deleted draft",
                "version_label": "v2",
                "document": self._docx_upload("draft-protocol.docx", b"second"),
            },
        )
        self.assertRedirects(response, detail_url)

        protocol.refresh_from_db()
        self.assertTrue(protocol.document)
        self.assertNotEqual(protocol.document.name, original_document_path)
        self.assertIsNotNone(protocol.current_revision)
        self.assertNotEqual(protocol.current_revision_id, original_revision_id)
        self.assertEqual(protocol.current_revision.original_name, "draft-protocol.docx")
        self.assertEqual(protocol.current_revision.version_label, "v2")
        self.assertEqual(ProtocolRevision.objects.filter(protocol=protocol).count(), 2)

    def test_protocol_missing_file_after_delete_shows_clear_error(self):
        detail_url = reverse("synopsis:protocol_detail", args=[self.project.id])
        self.client.post(
            detail_url,
            {
                "stage": "draft",
                "change_reason": "",
                "version_label": "v1",
                "document": self._docx_upload("draft-protocol.docx", b"first"),
            },
        )
        self.client.post(reverse("synopsis:protocol_delete_file", args=[self.project.id]))

        response = self.client.post(
            detail_url,
            {
                "stage": "draft",
                "change_reason": "",
                "version_label": "",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            "Choose a protocol file to upload. You can reuse the same filename as a file you deleted.",
        )

    def test_action_list_can_reupload_same_filename_after_delete_file(self):
        detail_url = reverse("synopsis:action_list_detail", args=[self.project.id])

        self.client.post(
            detail_url,
            {
                "stage": "draft",
                "change_reason": "",
                "version_label": "v1",
                "document": self._docx_upload("draft-action-list.docx", b"first"),
            },
        )
        action_list = ActionList.objects.get(project=self.project)
        original_document_path = action_list.document.name
        original_revision_id = action_list.current_revision_id

        response = self.client.post(
            reverse("synopsis:action_list_delete_file", args=[self.project.id])
        )
        self.assertRedirects(response, detail_url)
        action_list.refresh_from_db()
        self.assertFalse(action_list.document)
        self.assertIsNone(action_list.current_revision)

        response = self.client.post(
            detail_url,
            {
                "stage": "draft",
                "change_reason": "Replacing deleted draft",
                "version_label": "v2",
                "document": self._docx_upload("draft-action-list.docx", b"second"),
            },
        )
        self.assertRedirects(response, detail_url)

        action_list.refresh_from_db()
        self.assertTrue(action_list.document)
        self.assertNotEqual(action_list.document.name, original_document_path)
        self.assertIsNotNone(action_list.current_revision)
        self.assertNotEqual(action_list.current_revision_id, original_revision_id)
        self.assertEqual(
            action_list.current_revision.original_name, "draft-action-list.docx"
        )
        self.assertEqual(action_list.current_revision.version_label, "v2")
        self.assertEqual(
            ActionListRevision.objects.filter(action_list=action_list).count(), 2
        )

    def test_action_list_missing_file_after_delete_shows_clear_error(self):
        detail_url = reverse("synopsis:action_list_detail", args=[self.project.id])
        self.client.post(
            detail_url,
            {
                "stage": "draft",
                "change_reason": "",
                "version_label": "v1",
                "document": self._docx_upload("draft-action-list.docx", b"first"),
            },
        )
        self.client.post(
            reverse("synopsis:action_list_delete_file", args=[self.project.id])
        )

        response = self.client.post(
            detail_url,
            {
                "stage": "draft",
                "change_reason": "",
                "version_label": "",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            "Choose an action list file to upload. You can reuse the same filename as a file you deleted.",
        )

    def test_current_document_links_use_download_current_document_labels(self):
        protocol_detail_url = reverse("synopsis:protocol_detail", args=[self.project.id])
        action_list_detail_url = reverse(
            "synopsis:action_list_detail", args=[self.project.id]
        )
        project_hub_url = reverse("synopsis:project_hub", args=[self.project.id])

        self.client.post(
            protocol_detail_url,
            {
                "stage": "draft",
                "change_reason": "",
                "version_label": "v1",
                "document": self._docx_upload("draft-protocol.docx", b"protocol"),
            },
        )
        self.client.post(
            action_list_detail_url,
            {
                "stage": "draft",
                "change_reason": "",
                "version_label": "v1",
                "document": self._docx_upload("draft-action-list.docx", b"action"),
            },
        )

        project_hub_response = self.client.get(project_hub_url)
        protocol_detail_response = self.client.get(protocol_detail_url)
        action_list_detail_response = self.client.get(action_list_detail_url)

        protocol_view_url = reverse(
            "synopsis:document_view", args=[self.project.id, "protocol"]
        )
        action_list_view_url = reverse(
            "synopsis:document_view", args=[self.project.id, "action-list"]
        )

        self.assertContains(project_hub_response, protocol_view_url)
        self.assertContains(project_hub_response, action_list_view_url)
        self.assertContains(protocol_detail_response, protocol_view_url)
        self.assertContains(action_list_detail_response, action_list_view_url)
        self.assertContains(project_hub_response, "Download current document")
        self.assertContains(protocol_detail_response, "Download current document")
        self.assertContains(action_list_detail_response, "Download current document")
        self.assertNotContains(protocol_detail_response, "Open in new tab")
        self.assertNotContains(action_list_detail_response, "Open in new tab")

    def test_document_view_route_returns_latest_document_inline(self):
        protocol_detail_url = reverse("synopsis:protocol_detail", args=[self.project.id])
        self.client.post(
            protocol_detail_url,
            {
                "stage": "draft",
                "change_reason": "",
                "version_label": "v1",
                "document": self._docx_upload("draft-protocol.docx", b"protocol"),
            },
        )

        response = self.client.get(
            reverse("synopsis:document_view", args=[self.project.id, "protocol"])
        )
        protocol = Protocol.objects.get(project=self.project)
        expected_filename = protocol.document.name.rsplit("/", 1)[-1]

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response["Content-Disposition"],
            f'inline; filename="{expected_filename}"',
        )
        self.assertEqual(
            response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertEqual(b"".join(response.streaming_content), b"protocol")

    def test_protocol_and_action_list_danger_zones_explain_permanent_deletion(self):
        protocol_url = reverse("synopsis:protocol_detail", args=[self.project.id])
        action_list_url = reverse("synopsis:action_list_detail", args=[self.project.id])

        self.client.post(
            protocol_url,
            {
                "stage": "draft",
                "change_reason": "",
                "version_label": "v1",
                "document": self._docx_upload("draft-protocol.docx", b"first"),
            },
        )
        self.client.post(
            action_list_url,
            {
                "stage": "draft",
                "change_reason": "",
                "version_label": "v1",
                "document": self._docx_upload("draft-action-list.docx", b"first"),
            },
        )

        protocol_response = self.client.get(protocol_url)
        action_list_response = self.client.get(action_list_url)

        self.assertContains(protocol_response, "Danger zone: permanent deletion")
        self.assertContains(
            protocol_response,
            "These actions are final and destructive.",
        )
        self.assertContains(protocol_response, "Permanently delete file")
        self.assertContains(protocol_response, "Permanently delete protocol")
        self.assertContains(
            protocol_response,
            "This action is final and cannot be undone from the portal.",
        )

        self.assertContains(action_list_response, "Danger zone: permanent deletion")
        self.assertContains(
            action_list_response,
            "These actions are final and destructive.",
        )
        self.assertContains(action_list_response, "Permanently delete file")
        self.assertContains(action_list_response, "Permanently delete action list")
        self.assertContains(
            action_list_response,
            "This action is final and cannot be undone from the portal.",
        )

    def test_protocol_revision_history_uses_clear_current_and_earlier_sections(self):
        detail_url = reverse("synopsis:protocol_detail", args=[self.project.id])
        self.client.post(
            detail_url,
            {
                "stage": "draft",
                "change_reason": "",
                "version_label": "v1",
                "document": self._docx_upload("draft-protocol.docx", b"first"),
            },
        )
        self.client.post(
            detail_url,
            {
                "stage": "draft",
                "change_reason": "Updated methods section",
                "version_label": "v2",
                "document": self._docx_upload("draft-protocol.docx", b"second"),
            },
        )

        response = self.client.get(detail_url)

        self.assertContains(response, "Current live version")
        self.assertContains(response, "Earlier saved versions")
        self.assertContains(response, "Working draft")
        self.assertContains(response, "Revision note:")
        self.assertContains(response, "Restore as current")
        self.assertContains(response, "Delete revision")

    def test_action_list_revision_history_uses_clear_current_and_earlier_sections(self):
        detail_url = reverse("synopsis:action_list_detail", args=[self.project.id])
        self.client.post(
            detail_url,
            {
                "stage": "draft",
                "change_reason": "",
                "version_label": "v1",
                "document": self._docx_upload("draft-action-list.docx", b"first"),
            },
        )
        self.client.post(
            detail_url,
            {
                "stage": "draft",
                "change_reason": "Added missing interventions",
                "version_label": "v2",
                "document": self._docx_upload("draft-action-list.docx", b"second"),
            },
        )

        response = self.client.get(detail_url)

        self.assertContains(response, "Current live version")
        self.assertContains(response, "Earlier saved versions")
        self.assertContains(response, "Working draft")
        self.assertContains(response, "Revision note:")
        self.assertContains(response, "Restore as current")
        self.assertContains(response, "Delete revision")

    def test_protocol_delete_closes_stale_collaborative_session_before_reupload(self):
        detail_url = reverse("synopsis:protocol_detail", args=[self.project.id])
        self.client.post(
            detail_url,
            {
                "stage": "draft",
                "change_reason": "",
                "version_label": "v1",
                "document": self._docx_upload("draft-protocol.docx", b"first"),
            },
        )
        session = CollaborativeSession.objects.create(
            project=self.project,
            document_type=CollaborativeSession.DOCUMENT_PROTOCOL,
            started_by=self.ibrahim,
            last_activity_at=timezone.now(),
        )

        response = self.client.post(
            reverse("synopsis:protocol_delete", args=[self.project.id])
        )
        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.project.id])
        )
        session.refresh_from_db()
        self.assertFalse(session.is_active)
        self.assertEqual(session.end_reason, "Protocol deleted")

        response = self.client.post(
            detail_url,
            {
                "stage": "draft",
                "change_reason": "",
                "version_label": "v2",
                "document": self._docx_upload("draft-protocol.docx", b"second"),
            },
        )
        self.assertRedirects(response, detail_url)

        response = self.client.get(detail_url)
        self.assertContains(response, "Start collaborative edit")
        self.assertNotContains(response, "Open editor")

    def test_action_list_delete_closes_stale_collaborative_session_before_reupload(self):
        detail_url = reverse("synopsis:action_list_detail", args=[self.project.id])
        self.client.post(
            detail_url,
            {
                "stage": "draft",
                "change_reason": "",
                "version_label": "v1",
                "document": self._docx_upload("draft-action-list.docx", b"first"),
            },
        )
        session = CollaborativeSession.objects.create(
            project=self.project,
            document_type=CollaborativeSession.DOCUMENT_ACTION_LIST,
            started_by=self.ibrahim,
            last_activity_at=timezone.now(),
        )

        response = self.client.post(
            reverse("synopsis:action_list_delete", args=[self.project.id])
        )
        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.project.id])
        )
        session.refresh_from_db()
        self.assertFalse(session.is_active)
        self.assertEqual(session.end_reason, "Action list deleted")

        response = self.client.post(
            detail_url,
            {
                "stage": "draft",
                "change_reason": "",
                "version_label": "v2",
                "document": self._docx_upload("draft-action-list.docx", b"second"),
            },
        )
        self.assertRedirects(response, detail_url)

        response = self.client.get(detail_url)
        self.assertContains(response, "Start collaborative edit")
        self.assertNotContains(response, "Open editor")


class OnlyOfficeConfigTests(TestCase):
    def setUp(self):
        from . import views

        self.views = views
        self.original_settings = views.ONLYOFFICE_SETTINGS
        views.ONLYOFFICE_SETTINGS = {
            "base_url": "http://localhost:8080",
            "internal_url": "http://onlyoffice",
            "app_base_url": "http://web:8000",
            "jwt_secret": "change-me",
            "callback_timeout": 10,
            "trusted_download_urls": [
                "http://localhost:8080",
                "http://onlyoffice",
            ],
        }
        self.addCleanup(self._restore_settings)

        self.media_dir = tempfile.mkdtemp()
        self.addCleanup(lambda: shutil.rmtree(self.media_dir, ignore_errors=True))
        override = override_settings(MEDIA_ROOT=self.media_dir)
        override.enable()
        self.addCleanup(override.disable)

        self.factory = RequestFactory()
        self.project = Project.objects.create(title="Will Collaboration Test")
        self.ibrahim = User.objects.create_user(username="ibrahim-editor", password="pw")
        UserRole.objects.create(user=self.ibrahim, project=self.project, role="author")
        self.protocol = Protocol.objects.create(
            project=self.project,
            document=SimpleUploadedFile(
                "will-protocol.docx",
                b"protocol",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        )
        self.session = CollaborativeSession.objects.create(
            project=self.project,
            document_type=CollaborativeSession.DOCUMENT_PROTOCOL,
            started_by=self.ibrahim,
            last_activity_at=timezone.now(),
        )

    def _restore_settings(self):
        self.views.ONLYOFFICE_SETTINGS = self.original_settings

    def test_config_uses_internal_app_base_for_document_and_callback_urls(self):
        request = self.factory.get("/", HTTP_HOST="localhost:8000")
        request.user = self.ibrahim

        config = _build_onlyoffice_config(
            request,
            self.project,
            self.protocol,
            self.session,
            CollaborativeSession.DOCUMENT_PROTOCOL,
        )

        document_url = urlparse(config["document"]["url"])
        callback_url = urlparse(config["editorConfig"]["callbackUrl"])

        self.assertEqual(document_url.scheme, "http")
        self.assertEqual(document_url.netloc, "web:8000")
        self.assertTrue(document_url.path.startswith("/media/"))
        self.assertEqual(callback_url.scheme, "http")
        self.assertEqual(callback_url.netloc, "web:8000")
        self.assertEqual(
            callback_url.path,
            reverse(
                "synopsis:collaborative_edit_callback",
                args=[self.project.id, "protocol", self.session.token],
            ),
        )


class OnlyOfficeExternalAccessTests(TestCase):
    def setUp(self):
        from . import views

        self.views = views
        self.original_settings = views.ONLYOFFICE_SETTINGS
        views.ONLYOFFICE_SETTINGS = {
            "base_url": "http://localhost:8080",
            "internal_url": "http://onlyoffice",
            "app_base_url": "http://web:8000",
            "jwt_secret": "change-me",
            "callback_timeout": 10,
            "trusted_download_urls": [
                "http://localhost:8080",
                "http://onlyoffice",
            ],
        }
        self.addCleanup(self._restore_settings)

        self.media_dir = tempfile.mkdtemp()
        self.addCleanup(lambda: shutil.rmtree(self.media_dir, ignore_errors=True))
        override = override_settings(MEDIA_ROOT=self.media_dir)
        override.enable()
        self.addCleanup(override.disable)

        self.project = Project.objects.create(title="External Collaboration")
        self.author = User.objects.create_user(username="external-author", password="pw")
        UserRole.objects.create(user=self.author, project=self.project, role="author")
        self.member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Asha",
            last_name="Reviewer",
            organisation="CE",
            email="asha@example.com",
            response="Y",
            participation_confirmed=True,
            feedback_on_protocol_deadline=timezone.now() + timedelta(days=7),
        )
        self.protocol = Protocol.objects.create(
            project=self.project,
            document=SimpleUploadedFile(
                "external-protocol.docx",
                b"protocol",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        )
        self.session = CollaborativeSession.objects.create(
            project=self.project,
            document_type=CollaborativeSession.DOCUMENT_PROTOCOL,
            started_by=self.author,
            last_activity_at=timezone.now(),
        )

    def _restore_settings(self):
        self.views.ONLYOFFICE_SETTINGS = self.original_settings

    def _editor_url(self, query):
        return (
            reverse(
                "synopsis:collaborative_edit",
                args=[self.project.id, "protocol", self.session.token],
            )
            + query
        )

    def test_anonymous_reviewer_can_open_editor_with_feedback_token(self):
        feedback = ProtocolFeedback.objects.create(
            project=self.project,
            member=self.member,
            email=self.member.email,
            feedback_deadline_at=self.member.feedback_on_protocol_deadline,
        )

        response = self.client.get(self._editor_url(f"?feedback={feedback.token}"))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.context["editor_config"]["editorConfig"]["user"]["id"],
            f"abm:{self.member.id}",
        )
        self.assertFalse(response.context["editor_config"]["document"]["permissions"]["edit"])
        self.assertTrue(response.context["editor_config"]["document"]["permissions"]["comment"])
        self.assertFalse(response.context["editor_config"]["document"]["permissions"]["review"])
        self.assertContains(response, "Reviewing as")
        self.assertContains(response, "Asha Reviewer")
        self.assertContains(
            response,
            "To comment, highlight text and use the comment button in the toolbar.",
        )
        self.assertContains(
            response,
            "Authors will review your comments and decide whether to apply them.",
        )
        self.assertContains(
            response,
            "Comments save automatically while you work.",
        )
        self.assertContains(
            response,
            "Comments accepted until",
        )
        self.assertContains(
            response,
            _format_deadline(self.member.feedback_on_protocol_deadline),
        )
        self.assertContains(
            response,
            "Comment-only access",
        )
        self.assertContains(response, "Leave review page")
        self.assertContains(response, "reviewer-tab-lock-key")
        self.assertContains(
            response,
            "This review page is already open in another tab. Return to that tab or close it before opening another one.",
        )
        self.assertNotContains(response, "How collaborative editing works")

    def test_anonymous_reviewer_can_open_editor_with_invitation_token(self):
        invitation = AdvisoryBoardInvitation.objects.create(
            project=self.project,
            member=self.member,
            email=self.member.email,
            invited_by=self.author,
            due_date=timezone.localdate() + timedelta(days=7),
        )
        self.session.invitations.add(invitation)

        response = self.client.get(
            self._editor_url(f"?invite={invitation.token}&member={self.member.id}")
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.context["editor_config"]["editorConfig"]["user"]["id"],
            f"abm:{self.member.id}",
        )
        self.assertFalse(response.context["editor_config"]["document"]["permissions"]["edit"])
        self.assertTrue(response.context["editor_config"]["document"]["permissions"]["comment"])
        self.assertFalse(response.context["editor_config"]["document"]["permissions"]["review"])
        self.assertContains(response, "Comments accepted until")
        self.assertContains(response, "Leave review page")
        self.assertNotContains(response, "How collaborative editing works")

    def test_anonymous_reviewer_invitation_link_restarts_when_session_is_closed(self):
        invitation = AdvisoryBoardInvitation.objects.create(
            project=self.project,
            member=self.member,
            email=self.member.email,
            invited_by=self.author,
            due_date=timezone.localdate() + timedelta(days=7),
        )
        self.session.invitations.add(invitation)
        original_token = self.session.token
        self.session.mark_inactive(reason="Closed for restart test")

        response = self.client.get(
            self._editor_url(f"?invite={invitation.token}&member={self.member.id}")
        )

        self.assertEqual(response.status_code, 302)
        self.assertIn(str(invitation.token), response["Location"])
        self.assertNotIn(str(original_token), response["Location"])
        new_session = CollaborativeSession.objects.get(
            project=self.project,
            document_type=CollaborativeSession.DOCUMENT_PROTOCOL,
            is_active=True,
        )
        self.assertNotEqual(new_session.token, original_token)
        self.assertTrue(new_session.invitations.filter(pk=invitation.pk).exists())

    def test_anonymous_reviewer_feedback_link_restarts_when_session_expires(self):
        feedback = ProtocolFeedback.objects.create(
            project=self.project,
            member=self.member,
            email=self.member.email,
            feedback_deadline_at=self.member.feedback_on_protocol_deadline,
        )
        original_token = self.session.token
        self.session.last_activity_at = timezone.now() - timedelta(hours=5)
        self.session.save(update_fields=["last_activity_at"])

        response = self.client.get(self._editor_url(f"?feedback={feedback.token}"))

        self.assertEqual(response.status_code, 302)
        self.assertIn(str(feedback.token), response["Location"])
        self.assertNotIn(str(original_token), response["Location"])
        self.session.refresh_from_db()
        self.assertFalse(self.session.is_active)
        new_session = CollaborativeSession.objects.get(
            project=self.project,
            document_type=CollaborativeSession.DOCUMENT_PROTOCOL,
            is_active=True,
        )
        self.assertNotEqual(new_session.token, original_token)

    def test_member_id_only_link_is_blocked_for_anonymous_users(self):
        response = self.client.get(self._editor_url(f"?member={self.member.id}"))

        self.assertEqual(response.status_code, 403)
        self.assertContains(
            response,
            "This older collaborative link is missing its secure review token.",
            status_code=403,
        )

    def test_project_author_can_open_editor_without_external_token(self):
        self.client.force_login(self.author)

        response = self.client.get(self._editor_url(""))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.context["editor_config"]["editorConfig"]["user"]["id"],
            str(self.author.id),
        )
        self.assertTrue(response.context["editor_config"]["document"]["permissions"]["edit"])
        self.assertTrue(response.context["editor_config"]["document"]["permissions"]["comment"])
        self.assertTrue(response.context["editor_config"]["document"]["permissions"]["review"])
        self.assertContains(response, "Back to protocol page")
        self.assertContains(
            response,
            "To save and close the shared session for everyone, return to the protocol detail page.",
        )
        self.assertContains(response, "Active in this document:")
        self.assertContains(response, "visibilitychange")
        self.assertContains(response, "startPresencePolling")
        self.assertNotContains(response, "reviewer-tab-lock-key")
        self.assertNotContains(response, "How collaborative editing works")
        self.assertNotContains(
            response,
            reverse(
                "synopsis:collaborative_force_end",
                args=[self.project.id, "protocol", self.session.token],
            ),
        )

    @patch(
        "synopsis.views._collaborative_active_participant_names",
        return_value=["Asha Reviewer", "external-author"],
    )
    def test_author_can_fetch_active_collaborative_participants(
        self, mock_active_names
    ):
        self.client.force_login(self.author)

        response = self.client.get(
            reverse(
                "synopsis:collaborative_presence",
                args=[self.project.id, "protocol", self.session.token],
            )
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.json(),
            {"participants": ["Asha Reviewer", "external-author"]},
        )
        mock_active_names.assert_called_once_with(
            self.project,
            CollaborativeSession.DOCUMENT_PROTOCOL,
            self.session,
        )

    def test_anonymous_reviewer_can_leave_editor_without_closing_shared_session(self):
        feedback = ProtocolFeedback.objects.create(
            project=self.project,
            member=self.member,
            email=self.member.email,
            feedback_deadline_at=self.member.feedback_on_protocol_deadline,
        )

        response = self.client.get(
            reverse(
                "synopsis:collaborative_leave",
                args=[self.project.id, "protocol", self.session.token],
            )
            + f"?feedback={feedback.token}"
        )

        self.assertEqual(response.status_code, 200)
        self.session.refresh_from_db()
        self.assertTrue(self.session.is_active)
        self.assertContains(
            response,
            "You left the review page. This did not close the shared session for other participants.",
        )
        self.assertContains(response, "Reviewing as")
        self.assertContains(response, "Asha Reviewer")
        self.assertContains(response, "Comment-only access")
        self.assertContains(response, "Comments accepted until")
        self.assertContains(
            response, _format_deadline(self.member.feedback_on_protocol_deadline)
        )
        self.assertContains(response, "Reopen review page")
        self.assertContains(response, "Close this tab")

class CollaborativeForceSaveCloseTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Will Final Save")
        self.manager = User.objects.create_user(username="ibrahim-manager", password="pw")
        self.manager.is_staff = True
        self.manager.save(update_fields=["is_staff"])
        UserRole.objects.create(user=self.manager, project=self.project, role="manager")
        self.client.force_login(self.manager)
        self.protocol = Protocol.objects.create(
            project=self.project,
            document=SimpleUploadedFile("will-protocol.docx", b"protocol"),
        )
        self.session = CollaborativeSession.objects.create(
            project=self.project,
            document_type=CollaborativeSession.DOCUMENT_PROTOCOL,
            started_by=self.manager,
            last_activity_at=timezone.now(),
        )
        self.url = reverse(
            "synopsis:collaborative_force_end",
            args=[self.project.id, "protocol", self.session.token],
        )

    @patch("synopsis.views._wait_for_collaborative_save", return_value=False)
    @patch(
        "synopsis.views._request_onlyoffice_forcesave",
        return_value=("failed", "Unable to request a final save from OnlyOffice."),
    )
    def test_force_end_keeps_session_open_when_save_request_fails(
        self, mock_request, mock_wait
    ):
        response = self.client.post(self.url, {"reason": "Close from portal"})

        self.assertRedirects(
            response,
            reverse("synopsis:protocol_detail", args=[self.project.id]),
        )
        self.session.refresh_from_db()
        self.assertTrue(self.session.is_active)
        messages_list = [message.message for message in get_messages(response.wsgi_request)]
        self.assertTrue(
            any("session is still open" in message for message in messages_list),
            messages_list,
        )
        mock_request.assert_called_once()
        mock_wait.assert_not_called()

    @patch("synopsis.views._wait_for_collaborative_save", return_value=True)
    @patch(
        "synopsis.views._request_onlyoffice_forcesave",
        return_value=("requested", "Final save requested from OnlyOffice."),
    )
    def test_force_end_reports_success_after_final_save(
        self, mock_request, mock_wait
    ):
        response = self.client.post(self.url, {"reason": "Close from portal"})

        self.assertRedirects(
            response,
            reverse("synopsis:protocol_detail", args=[self.project.id]),
        )
        messages_list = [message.message for message in get_messages(response.wsgi_request)]
        self.assertIn(
            "Protocol saved and collaborative session closed.",
            messages_list,
        )
        mock_request.assert_called_once()
        mock_wait.assert_called_once()

    @patch("synopsis.views._wait_for_collaborative_save")
    @patch(
        "synopsis.views._request_onlyoffice_forcesave",
        return_value=("noop", "No unsaved changes were pending in OnlyOffice."),
    )
    def test_force_end_closes_session_when_no_unsaved_changes(
        self, mock_request, mock_wait
    ):
        response = self.client.post(self.url, {"reason": "Close from portal"})

        self.assertRedirects(
            response,
            reverse("synopsis:protocol_detail", args=[self.project.id]),
        )
        self.session.refresh_from_db()
        self.assertFalse(self.session.is_active)
        self.assertEqual(self.session.ended_by, self.manager)
        self.assertEqual(self.session.end_reason, "Close from portal")
        self.assertEqual(self.session.change_summary, "Close from portal")
        messages_list = [message.message for message in get_messages(response.wsgi_request)]
        self.assertIn(
            "Protocol had no unsaved changes and the session was closed.",
            messages_list,
        )
        mock_request.assert_called_once()
        mock_wait.assert_not_called()


class MediaServingUrlTests(SimpleTestCase):
    def test_media_route_added_when_serve_media_enabled_without_debug(self):
        import ce_portal.urls as project_urls

        with override_settings(DEBUG=False, SERVE_MEDIA=True):
            reloaded = importlib.reload(project_urls)
            try:
                route_texts = [str(pattern.pattern) for pattern in reloaded.urlpatterns]
                self.assertTrue(
                    any(text.startswith("^media/") for text in route_texts),
                    route_texts,
                )
            finally:
                importlib.reload(project_urls)


class CollaborativePanelViewTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Collaborative Panel")
        self.user = User.objects.create_user(username="collab-author", password="pw")
        UserRole.objects.create(user=self.user, project=self.project, role="author")
        self.client.force_login(self.user)

    def test_protocol_panel_disabled_without_document(self):
        response = self.client.get(
            reverse("synopsis:protocol_detail", args=[self.project.id])
        )
        self.assertContains(response, "How this works")
        self.assertContains(
            response, "This page manages the current working protocol for the project."
        )
        self.assertContains(
            response,
            "Sending the protocol for review happens on the Advisory Board page.",
        )
        self.assertContains(
            response,
            "When you are ready to send it to advisory board members, go to the Advisory Board page",
        )
        self.assertContains(response, "Go to Advisory Board")
        self.assertContains(
            response,
            "Sending the protocol to advisory board members is done from the <strong>Advisory Board</strong> page.",
        )
        self.assertContains(response, "How collaborative editing works")
        self.assertContains(
            response,
            "Use this guide for the live OnlyOffice session itself.",
        )
        self.assertContains(
            response, "Upload the protocol before starting a collaborative session."
        )
        self.assertIn('aria-disabled="true"', response.content.decode())

    def test_protocol_panel_enabled_with_document(self):
        Protocol.objects.create(
            project=self.project,
            document=SimpleUploadedFile("protocol.docx", b"protocol"),
        )
        response = self.client.get(
            reverse("synopsis:protocol_detail", args=[self.project.id])
        )
        self.assertNotContains(
            response, "Upload the protocol before starting a collaborative session."
        )

    def test_action_list_panel_disabled_without_document(self):
        response = self.client.get(
            reverse("synopsis:action_list_detail", args=[self.project.id])
        )
        self.assertContains(response, "How this works")
        self.assertContains(
            response,
            "This page manages the current working action list for the project.",
        )
        self.assertContains(
            response,
            "Sending the action list for review happens on the Advisory Board page.",
        )
        self.assertContains(
            response,
            "When you are ready to send it to advisory board members, go to the Advisory Board page",
        )
        self.assertContains(response, "Go to Advisory Board")
        self.assertContains(
            response,
            "Sending the action list to advisory board members is done from the <strong>Advisory Board</strong> page.",
        )
        self.assertContains(response, "How collaborative editing works")
        self.assertContains(
            response,
            "Use this guide for the live OnlyOffice session itself.",
        )
        self.assertContains(
            response, "Upload the action list before starting a collaborative session."
        )
        self.assertIn('aria-disabled="true"', response.content.decode())

    def test_action_list_panel_enabled_with_document(self):
        ActionList.objects.create(
            project=self.project,
            document=SimpleUploadedFile("action-list.docx", b"alist"),
        )
        response = self.client.get(
            reverse("synopsis:action_list_detail", args=[self.project.id])
        )
        self.assertNotContains(
            response, "Upload the action list before starting a collaborative session."
        )

    def test_protocol_panel_active_session_explains_global_close_scope(self):
        from . import views

        original_settings = views.ONLYOFFICE_SETTINGS
        views.ONLYOFFICE_SETTINGS = {
            "base_url": "http://localhost:8080",
            "internal_url": "http://onlyoffice",
            "app_base_url": "http://web:8000",
            "jwt_secret": "change-me",
            "callback_timeout": 10,
            "trusted_download_urls": [
                "http://localhost:8080",
                "http://onlyoffice",
            ],
        }
        self.addCleanup(lambda: setattr(views, "ONLYOFFICE_SETTINGS", original_settings))

        Protocol.objects.create(
            project=self.project,
            document=SimpleUploadedFile("protocol.docx", b"protocol"),
        )
        CollaborativeSession.objects.create(
            project=self.project,
            document_type=CollaborativeSession.DOCUMENT_PROTOCOL,
            started_by=self.user,
            last_activity_at=timezone.now(),
        )

        response = self.client.get(
            reverse("synopsis:protocol_detail", args=[self.project.id])
        )

        self.assertContains(response, "Save and close session for everyone")
        self.assertContains(
            response,
            "Participants can leave the editor without using this button. Use this only when everyone is finished.",
        )

    def test_advisory_board_shows_custom_columns_button_and_not_document_feedback_windows(self):
        response = self.client.get(
            reverse("synopsis:advisory_board_list", args=[self.project.id])
        )
        self.assertContains(response, "Custom columns")
        self.assertContains(response, "Deadlines &amp; reminders")
        self.assertNotContains(response, "Protocol feedback window")
        self.assertNotContains(response, "Action list feedback window")

    def test_document_pages_show_feedback_window_controls(self):
        Protocol.objects.create(
            project=self.project,
            document=SimpleUploadedFile("protocol.docx", b"protocol"),
        )
        ActionList.objects.create(
            project=self.project,
            document=SimpleUploadedFile("action-list.docx", b"alist"),
        )
        AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Pia",
            email="pia@example.com",
            response="Y",
            sent_protocol_at=timezone.now(),
        )
        AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Ali",
            email="ali@example.com",
            response="Y",
            sent_action_list_at=timezone.now(),
        )

        protocol_response = self.client.get(
            reverse("synopsis:protocol_detail", args=[self.project.id])
        )
        self.assertContains(protocol_response, "Protocol feedback window")
        self.assertContains(protocol_response, "Set protocol deadline")
        self.assertContains(protocol_response, 'data-bs-target="#protocolFeedbackWindowCollapse"')
        self.assertContains(protocol_response, "data-collapse-toggle-label")
        self.assertContains(protocol_response, 'data-label-open="Hide"')
        self.assertContains(
            protocol_response,
            "Closing this feedback window will stop advisory members from submitting protocol feedback and will end collaborative editing for this protocol. Are you sure?",
        )

        action_list_response = self.client.get(
            reverse("synopsis:action_list_detail", args=[self.project.id])
        )
        self.assertContains(action_list_response, "Action list feedback window")
        self.assertContains(action_list_response, "Set action list deadline")
        self.assertContains(action_list_response, 'data-bs-target="#actionListFeedbackWindowCollapse"')
        self.assertContains(action_list_response, "data-collapse-toggle-label")
        self.assertContains(action_list_response, 'data-label-open="Hide"')
        self.assertContains(
            action_list_response,
            "Closing this feedback window will stop advisory members from submitting action list feedback and will end collaborative editing for this action list. Are you sure?",
        )

    def test_document_pages_explain_deadlines_require_sent_documents(self):
        Protocol.objects.create(
            project=self.project,
            document=SimpleUploadedFile("protocol.docx", b"protocol"),
        )
        ActionList.objects.create(
            project=self.project,
            document=SimpleUploadedFile("action-list.docx", b"alist"),
        )

        protocol_response = self.client.get(
            reverse("synopsis:protocol_detail", args=[self.project.id])
        )
        self.assertContains(
            protocol_response,
            "Send the protocol from the Advisory Board page first.",
        )
        self.assertContains(
            protocol_response,
            "No protocol feedback deadline can be updated here yet because no accepted advisory board member has been sent the protocol.",
        )
        self.assertNotContains(protocol_response, "Set protocol deadline")

        action_list_response = self.client.get(
            reverse("synopsis:action_list_detail", args=[self.project.id])
        )
        self.assertContains(
            action_list_response,
            "Send the action list from the Advisory Board page first.",
        )
        self.assertContains(
            action_list_response,
            "No action list feedback deadline can be updated here yet because no accepted advisory board member has been sent the action list.",
        )
        self.assertNotContains(action_list_response, "Set action list deadline")

    def test_document_pages_show_reopen_feedback_window_confirmations(self):
        Protocol.objects.create(
            project=self.project,
            document=SimpleUploadedFile("protocol.docx", b"protocol"),
            feedback_closed_at=timezone.now(),
        )
        ActionList.objects.create(
            project=self.project,
            document=SimpleUploadedFile("action-list.docx", b"alist"),
            feedback_closed_at=timezone.now(),
        )

        protocol_response = self.client.get(
            reverse("synopsis:protocol_detail", args=[self.project.id])
        )
        self.assertContains(protocol_response, "Reopen feedback")
        self.assertContains(
            protocol_response,
            "Reopening this feedback window will allow advisory members to submit protocol feedback again and can allow collaborative editing for this protocol again. Are you sure?",
        )

        action_list_response = self.client.get(
            reverse("synopsis:action_list_detail", args=[self.project.id])
        )
        self.assertContains(action_list_response, "Reopen feedback")
        self.assertContains(
            action_list_response,
            "Reopening this feedback window will allow advisory members to submit action list feedback again and can allow collaborative editing for this action list again. Are you sure?",
        )


class AdvisoryBoardCustomColumnsTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Dynamic Columns")
        self.editor = User.objects.create_user(username="editor-secondary")
        self.accepted = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Ann",
            email="ann@example.com",
            response="Y",
        )
        self.pending = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Vanessa",
            email="vanessa@example.com",
            response="",
        )
        self.general_field = AdvisoryBoardCustomField.objects.create(
            project=self.project,
            name="General note",
            data_type=AdvisoryBoardCustomField.TYPE_TEXT,
        )
        self.pending_only_field = AdvisoryBoardCustomField.objects.create(
            project=self.project,
            name="Follow-up date",
            data_type=AdvisoryBoardCustomField.TYPE_DATE,
            sections=[AdvisoryBoardCustomField.SECTION_PENDING],
        )
        self.general_field.set_value_for_member(self.accepted, "Confirmed")
        self.general_field.set_value_for_member(self.pending, "Need reply")
        self.pending_only_field.set_value_for_member(self.pending, date(2025, 5, 1))

    def test_section_fields_match_custom_field_configuration(self):
        context = _advisory_board_context(self.project)
        sections = {section["key"]: section for section in context["member_sections"]}

        accepted_field_ids = [
            field.id
            for field in sections[AdvisoryBoardCustomField.SECTION_ACCEPTED]["fields"]
        ]
        pending_field_ids = [
            field.id
            for field in sections[AdvisoryBoardCustomField.SECTION_PENDING]["fields"]
        ]

        self.assertEqual(accepted_field_ids, [self.general_field.id])
        self.assertEqual(
            pending_field_ids, [self.general_field.id, self.pending_only_field.id]
        )

    def test_member_rows_include_formatted_custom_values(self):
        context = _advisory_board_context(self.project)
        sections = {section["key"]: section for section in context["member_sections"]}

        accepted_member = sections[AdvisoryBoardCustomField.SECTION_ACCEPTED][
            "members"
        ][0]
        pending_member = sections[AdvisoryBoardCustomField.SECTION_PENDING]["members"][
            0
        ]

        self.assertEqual(
            accepted_member.custom_field_values[self.general_field.id], "Confirmed"
        )
        self.assertEqual(
            pending_member.custom_field_values[self.general_field.id], "Need reply"
        )
        self.assertEqual(
            pending_member.custom_field_values[self.pending_only_field.id], "2025-05-01"
        )
        self.assertNotIn(
            self.pending_only_field.id, accepted_member.custom_field_values
        )

    def test_custom_fields_list_exposes_all_configured_fields(self):
        context = _advisory_board_context(self.project)
        custom_field_ids = [field.id for field in context["custom_fields"]]
        self.assertEqual(
            custom_field_ids, [self.general_field.id, self.pending_only_field.id]
        )

    def test_history_records_updates(self):
        base_count = AdvisoryBoardCustomFieldValueHistory.objects.filter(
            field=self.general_field, member=self.accepted
        ).count()

        self.general_field.set_value_for_member(
            self.accepted, "Updated note", changed_by=self.editor
        )
        self.general_field.set_value_for_member(
            self.accepted, "", changed_by=self.editor
        )

        history = AdvisoryBoardCustomFieldValueHistory.objects.filter(
            field=self.general_field, member=self.accepted
        ).order_by("-created_at")

        self.assertEqual(history.count(), base_count + 2)
        latest = history.first()
        self.assertTrue(latest.is_cleared)
        previous = history[1]
        self.assertEqual(previous.value, "Updated note")
        self.assertEqual(previous.changed_by, self.editor)




class UserEditPermissionTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Permissions Project")
        self.manager = User.objects.create_user(username="manager_user")
        self.manager.is_staff = True
        self.manager.save(update_fields=["is_staff"])
        self.author = User.objects.create_user(username="author_user")
        UserRole.objects.create(user=self.author, project=self.project, role="author")
        self.viewer = User.objects.create_user(username="viewer_user")

    def test_manager_can_edit_project(self):
        self.assertTrue(_user_can_edit_project(self.manager, self.project))

    def test_author_can_edit_project(self):
        self.assertTrue(_user_can_edit_project(self.author, self.project))

    def test_other_user_cannot_edit_project(self):
        self.assertFalse(_user_can_edit_project(self.viewer, self.project))


class RevisionDeleteViewTests(TestCase):
    def setUp(self):
        self.media_dir = tempfile.mkdtemp()
        self.addCleanup(lambda: shutil.rmtree(self.media_dir, ignore_errors=True))
        override = override_settings(MEDIA_ROOT=self.media_dir)
        override.enable()
        self.addCleanup(override.disable)

        ensure_global_groups()

        self.factory = RequestFactory()
        self.project = Project.objects.create(title="Revision Project")
        self.author = User.objects.create_user(username="author", password="pwd")
        UserRole.objects.create(user=self.author, project=self.project, role="author")
        self.other_user = User.objects.create_user(username="other", password="pwd")

        # Protocol setup
        base_file = SimpleUploadedFile("protocol_base.docx", b"base")
        self.protocol = Protocol.objects.create(
            project=self.project, document=base_file
        )
        self.rev1 = ProtocolRevision.objects.create(
            protocol=self.protocol,
            file=SimpleUploadedFile("protocol_rev1.docx", b"rev1"),
            stage="draft",
            change_reason="Initial",
        )
        self.protocol.current_revision = self.rev1
        self.protocol.save(update_fields=["current_revision"])
        self.rev2 = ProtocolRevision.objects.create(
            protocol=self.protocol,
            file=SimpleUploadedFile("protocol_rev2.docx", b"rev2"),
            stage="draft",
            change_reason="Second",
        )
        self.protocol.current_revision = self.rev2
        self.protocol.save(update_fields=["current_revision"])

        # Action list setup
        action_file = SimpleUploadedFile("action_base.docx", b"alist")
        self.action_list = ActionList.objects.create(
            project=self.project,
            document=action_file,
        )
        self.al_rev1 = ActionListRevision.objects.create(
            action_list=self.action_list,
            file=SimpleUploadedFile("action_rev1.docx", b"rev1"),
            stage="draft",
            change_reason="Initial",
        )
        self.action_list.current_revision = self.al_rev1
        self.action_list.save(update_fields=["current_revision"])
        self.al_rev2 = ActionListRevision.objects.create(
            action_list=self.action_list,
            file=SimpleUploadedFile("action_rev2.docx", b"rev2"),
            stage="draft",
            change_reason="Second",
        )
        self.action_list.current_revision = self.al_rev2
        self.action_list.save(update_fields=["current_revision"])

    def _add_session_and_messages(self, request):
        SessionMiddleware(lambda req: None).process_request(request)
        request.session.save()
        messages = FallbackStorage(request)
        setattr(request, "_messages", messages)

    def test_author_can_delete_protocol_revision(self):
        request = self.factory.post(
            f"/project/{self.project.id}/protocol/revision/{self.rev2.id}/delete/"
        )
        request.user = self.author
        self._add_session_and_messages(request)

        response = protocol_delete_revision(request, self.project.id, self.rev2.id)
        self.assertEqual(response.status_code, 302)
        self.assertFalse(ProtocolRevision.objects.filter(pk=self.rev2.pk).exists())
        self.protocol.refresh_from_db()
        self.assertEqual(self.protocol.current_revision_id, self.rev1.id)
        self.assertTrue(
            ProjectChangeLog.objects.filter(
                project=self.project, action__icontains="Protocol revision deleted"
            ).exists()
        )

    def test_apply_protocol_revision_does_not_duplicate_upload_prefix(self):
        _apply_revision_to_protocol(self.protocol, self.rev1)
        self.protocol.refresh_from_db()
        self.assertTrue(self.protocol.document.name.startswith("protocols/"))
        self.assertNotIn("protocols/protocols/", self.protocol.document.name)

    def test_non_editor_cannot_delete_protocol_revision(self):
        request = self.factory.post(
            f"/project/{self.project.id}/protocol/revision/{self.rev1.id}/delete/"
        )
        request.user = self.other_user
        self._add_session_and_messages(request)

        response = protocol_delete_revision(request, self.project.id, self.rev1.id)
        self.assertEqual(response.status_code, 302)
        self.assertTrue(ProtocolRevision.objects.filter(pk=self.rev1.pk).exists())

    def test_author_can_delete_action_list_revision(self):
        request = self.factory.post(
            f"/project/{self.project.id}/action-list/revision/{self.al_rev2.id}/delete/"
        )
        request.user = self.author
        self._add_session_and_messages(request)

        response = action_list_delete_revision(
            request, self.project.id, self.al_rev2.id
        )
        self.assertEqual(response.status_code, 302)
        self.assertFalse(ActionListRevision.objects.filter(pk=self.al_rev2.pk).exists())
        self.action_list.refresh_from_db()
        self.assertEqual(self.action_list.current_revision_id, self.al_rev1.id)
        self.assertTrue(
            ProjectChangeLog.objects.filter(
                project=self.project, action__icontains="Action list revision deleted"
            ).exists()
        )

    def test_apply_action_list_revision_does_not_duplicate_upload_prefix(self):
        _apply_revision_to_action_list(self.action_list, self.al_rev1)
        self.action_list.refresh_from_db()
        self.assertTrue(self.action_list.document.name.startswith("action_lists/"))
        self.assertNotIn("action_lists/action_lists/", self.action_list.document.name)

    def test_non_editor_cannot_delete_action_list_revision(self):
        request = self.factory.post(
            f"/project/{self.project.id}/action-list/revision/{self.al_rev1.id}/delete/"
        )
        request.user = self.other_user
        self._add_session_and_messages(request)

        response = action_list_delete_revision(
            request, self.project.id, self.al_rev1.id
        )
        self.assertEqual(response.status_code, 302)
        self.assertTrue(ActionListRevision.objects.filter(pk=self.al_rev1.pk).exists())


class ProjectDeleteFormTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Wetland Recovery")

    def test_requires_matching_title(self):
        form = ProjectDeleteForm(
            data={
                "confirm_title": "Wrong title",
                "acknowledge_irreversible": True,
            },
            project=self.project,
        )
        self.assertFalse(form.is_valid())
        self.assertIn("Title does not match", form.errors["confirm_title"][0])

    def test_requires_acknowledgement(self):
        form = ProjectDeleteForm(
            data={"confirm_title": "Wetland Recovery"}, project=self.project
        )
        self.assertFalse(form.is_valid())
        self.assertIn(
            "This field is required.", form.errors["acknowledge_irreversible"][0]
        )

    def test_valid_when_all_checks_pass(self):
        form = ProjectDeleteForm(
            data={
                "confirm_title": "Wetland Recovery",
                "acknowledge_irreversible": True,
            },
            project=self.project,
        )
        self.assertTrue(form.is_valid())


class ProjectSettingsFormTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Forest Restoration")

    def test_requires_title(self):
        form = ProjectSettingsForm(data={"title": ""}, instance=self.project)
        self.assertFalse(form.is_valid())
        self.assertIn("Enter a title", form.errors["title"][0])

    def test_updates_title(self):
        form = ProjectSettingsForm(
            data={"title": "Forest Recovery"}, instance=self.project
        )
        self.assertTrue(form.is_valid())
        updated = form.save()
        self.assertEqual(updated.title, "Forest Recovery")

    def test_updates_description(self):
        form = ProjectSettingsForm(
            data={
                "title": "Forest Recovery",
                "description": "  A pilot synopsis for forest restoration.  ",
            },
            instance=self.project,
        )
        self.assertTrue(form.is_valid())
        updated = form.save()
        self.assertEqual(updated.description, "A pilot synopsis for forest restoration.")


class ViewHelperTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Helper Project")

    def test_user_is_manager_for_staff(self):
        user = User.objects.create_user(
            username="staffer", password="pw", is_staff=True
        )
        self.assertTrue(_user_is_manager(user))

    def test_user_is_manager_for_group_member(self):
        group, _ = Group.objects.get_or_create(name="manager")
        user = User.objects.create_user(username="manager_user")
        user.groups.add(group)
        self.assertTrue(_user_is_manager(user))

    def test_user_is_manager_false_for_others(self):
        user = User.objects.create_user(username="regular")
        self.assertFalse(_user_is_manager(user))

    def test_log_project_change_records_entry(self):
        user = User.objects.create_user(username="logger")
        _log_project_change(self.project, user, "Edited", "Updated title")
        entry = ProjectChangeLog.objects.get(project=self.project)
        self.assertEqual(entry.action, "Edited")
        self.assertEqual(entry.details, "Updated title")
        self.assertEqual(entry.changed_by, user)

    def test_log_project_change_anonymous(self):
        anonymous = AnonymousUser()
        _log_project_change(self.project, anonymous, "Edited", "Updated title")
        entry = (
            ProjectChangeLog.objects.filter(project=self.project)
            .order_by("-id")
            .first()
        )
        self.assertIsNone(entry.changed_by)

    def test_format_value_handles_none_and_date(self):
        self.assertEqual(_format_value(None), "—")
        today = date(2025, 1, 2)
        self.assertEqual(_format_value(today), "2025-01-02")
        self.assertEqual(_format_value(42), "42")

    def test_funder_contact_label(self):
        self.assertEqual(_funder_contact_label("Ann", "Thornton"), "Ann Thornton")
        self.assertEqual(_funder_contact_label("", ""), "—")

    def test_format_deadline_formats_timezone(self):
        aware = timezone.make_aware(datetime(2025, 7, 1, 15, 0))
        formatted = _format_deadline(aware)
        self.assertEqual(
            formatted, timezone.localtime(aware).strftime("%d %b %Y %H:%M")
        )
        self.assertEqual(_format_deadline(None), "—")

    def test_split_inline_italic_markup_supports_simple_i_tags(self):
        segments = split_inline_italic_markup(
            "Gamma G. (2009) <i>Glipa</i> restoration note."
        )

        self.assertEqual(
            segments,
            [
                ("Gamma G. (2009) ", False),
                ("Glipa", True),
                (" restoration note.", False),
            ],
        )

    def test_user_can_confirm_phase(self):
        staff = User.objects.create_user(username="staff", is_staff=True)
        self.assertTrue(_user_can_confirm_phase(staff, self.project))
        author = User.objects.create_user(username="author")
        UserRole.objects.create(user=author, project=self.project, role="author")
        self.assertTrue(_user_can_confirm_phase(author, self.project))
        manager = User.objects.create_user(username="manager")
        UserRole.objects.create(user=manager, project=self.project, role="manager")
        self.assertTrue(_user_can_confirm_phase(manager, self.project))
        outsider = User.objects.create_user(username="outsider")
        self.assertFalse(_user_can_confirm_phase(outsider, self.project))
        self.assertFalse(_user_can_confirm_phase(AnonymousUser(), self.project))


class CreateProtocolFeedbackTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Protocol Feedback")
        self.protocol = Protocol.objects.create(
            project=self.project,
            document=SimpleUploadedFile("protocol.txt", b"content"),
        )

    def test_feedback_from_member_uses_member_deadline(self):
        member = AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Pat",
            email="pat@example.com",
            response="Y",
            feedback_on_protocol_deadline=timezone.now(),
        )
        feedback = _create_protocol_feedback(self.project, member=member)
        self.assertEqual(feedback.project, self.project)
        self.assertEqual(feedback.member, member)
        self.assertEqual(
            feedback.feedback_deadline_at, member.feedback_on_protocol_deadline
        )
        self.assertEqual(feedback.protocol_stage_snapshot, self.protocol.stage)
        self.assertEqual(
            feedback.protocol_document_last_updated, self.protocol.last_updated
        )
        self.assertEqual(
            feedback.protocol_document_name,
            self.protocol.document.name,
        )

    def test_feedback_from_invitation_sets_end_of_day_deadline(self):
        due = date(2025, 8, 15)
        invitation = AdvisoryBoardInvitation.objects.create(
            project=self.project,
            email="invitee@example.com",
            due_date=due,
        )
        feedback = _create_protocol_feedback(self.project, invitation=invitation)
        self.assertEqual(feedback.invitation, invitation)
        self.assertEqual(feedback.feedback_deadline_at.date(), due)
        self.assertEqual(feedback.feedback_deadline_at.hour, 23)
        self.assertEqual(feedback.feedback_deadline_at.minute, 59)
        self.assertEqual(feedback.protocol_stage_snapshot, self.protocol.stage)


class PlainTextReferenceParserTests(TestCase):
    def test_parses_references_and_extracts_metadata(self):
        payload = textwrap.dedent(
            """
            Angel, D. L.; et al. (2002). "In situ biofiltration: a means to limit the dispersal of effluents from marine finfish cage aquaculture." Hydrobiologia 469(1): 1-10.
            Net pen fish farms generally enrich the surrounding waters and the underlying sediments with nutrients and organic matter.

            This entry is invalid and should be skipped.

            Morgan, W. and Thornton, A. (2018). Another example of parsing. Ecology Letters 11: 9-12.
            Full abstract text including doi:10.5678/example and https://example.com/article for reference.
            """
        ).strip()

        parsed = _parse_plaintext_references(payload)

        self.assertEqual(len(parsed), 2)

        first = parsed[0]
        self.assertEqual(
            first["title"],
            "In situ biofiltration: a means to limit the dispersal of effluents from marine finfish cage aquaculture.",
        )
        self.assertEqual(first["journal_name"], "Hydrobiologia")
        self.assertEqual(first["volume"], "469")
        self.assertEqual(first["issue"], "1")
        self.assertEqual(first["pages"], "1-10")
        self.assertEqual(first["year"], "2002")
        self.assertEqual(first["publication_year"], "2002")
        self.assertEqual(first["authors"], ["Angel, D. L", "et al"])
        self.assertTrue(first["abstract"].startswith("Net pen fish farms"))

        second = parsed[1]
        self.assertEqual(second["title"], "Another example of parsing")
        self.assertEqual(second["journal_name"], "Ecology Letters")
        self.assertEqual(second["volume"], "11")
        self.assertFalse(second["issue"])
        self.assertEqual(second["pages"], "9-12")
        self.assertEqual(second["authors"], ["Morgan, W", "Thornton, A"])
        self.assertEqual(second["doi"], "10.5678/example")
        self.assertEqual(second["url"], "https://example.com/article")

    def test_returns_empty_list_for_blank_payload(self):
        self.assertEqual(_parse_plaintext_references(""), [])
        self.assertEqual(_parse_plaintext_references("   "), [])


class EndNoteXmlParserTests(TestCase):
    def test_parses_endnote_xml_record(self):
        payload = textwrap.dedent(
            """
            <xml>
              <records>
                <record>
                  <titles>
                    <title>Coral restoration methods</title>
                    <secondary-title>Marine Ecology</secondary-title>
                  </titles>
                  <contributors>
                    <authors>
                      <author>Alhas, Ibrahim</author>
                      <author>Morgan, Will</author>
                    </authors>
                  </contributors>
                  <dates>
                    <year>2023</year>
                  </dates>
                  <volume>12</volume>
                  <number>4</number>
                  <pages>101-110</pages>
                  <abstract>Summary text.</abstract>
                  <electronic-resource-num>10.1234/example</electronic-resource-num>
                  <urls>
                    <related-urls>
                      <url>https://example.com/article</url>
                    </related-urls>
                  </urls>
                </record>
              </records>
            </xml>
            """
        ).strip()

        parsed = _parse_endnote_xml(payload)

        self.assertEqual(len(parsed), 1)
        self.assertEqual(parsed[0]["title"], "Coral restoration methods")
        self.assertEqual(parsed[0]["journal_name"], "Marine Ecology")
        self.assertEqual(parsed[0]["authors"], ["Alhas, Ibrahim", "Morgan, Will"])
        self.assertEqual(parsed[0]["publication_year"], "2023")
        self.assertEqual(parsed[0]["doi"], "10.1234/example")
        self.assertEqual(parsed[0]["url"], "https://example.com/article")

    def test_returns_empty_list_for_invalid_xml(self):
        self.assertEqual(_parse_endnote_xml("<xml><records>"), [])


class LibraryReferenceBatchUploadTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Library Upload Project")
        self.user = User.objects.create_user(username="libraryuploader", password="pw")
        UserRole.objects.create(user=self.user, project=self.project, role="author")
        self.client.force_login(self.user)
        self.url = reverse("synopsis:library_reference_batch_upload")

    def _plaintext_payload(self):
        return textwrap.dedent(
            """
            Angel, D. L.; et al. (2002). "In situ biofiltration: a means to limit the dispersal of effluents from marine finfish cage aquaculture." Hydrobiologia 469(1): 1-10.
            Net pen fish farms generally enrich the surrounding waters and the underlying sediments with nutrients and organic matter.

            Morgan, W. and Thornton, A. (2018). Another example of parsing. Ecology Letters 11: 9-12.
            Full abstract text including doi:10.5678/example and https://example.com/article for reference.
            """
        ).strip()

    def test_skips_existing_library_reference_hashes(self):
        existing_hash = reference_hash(
            "In situ biofiltration: a means to limit the dispersal of effluents from marine finfish cage aquaculture.",
            "2002",
            "",
        )
        existing = LibraryReference.objects.create(
            hash_key=existing_hash,
            title="In situ biofiltration: a means to limit the dispersal of effluents from marine finfish cage aquaculture.",
            publication_year=2002,
        )

        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        response = self.client.post(
            self.url,
            {
                "label": "Library batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
            follow=True,
        )

        self.assertEqual(LibraryReference.objects.count(), 2)
        self.assertTrue(LibraryReference.objects.filter(pk=existing.pk).exists())
        batch = LibraryImportBatch.objects.get(label="Library batch")
        self.assertEqual(batch.record_count, 1)
        messages = [str(message) for message in get_messages(response.wsgi_request)]
        self.assertTrue(
            any("Imported 1 reference(s) into 'Library batch'." in message for message in messages)
        )
        self.assertTrue(
            any("Skipped 1 reference(s) already in the library." in message for message in messages)
        )

    def test_imports_ris_like_txt_file_with_utf8_bom(self):
        ris_payload = (
            "\ufeffTY  - JOUR\n"
            "TI  - First RIS entry\n"
            "AU  - Morgan, Will\n"
            "PY  - 2024\n"
            "DO  - 10.1000/first\n"
            "ER  -\n\n"
            "TY  - JOUR\n"
            "TI  - Second RIS entry\n"
            "AU  - Thornton, Ann\n"
            "PY  - 2023\n"
            "DO  - 10.1000/second\n"
            "ER  -\n"
        )
        upload = SimpleUploadedFile(
            "references.txt",
            ris_payload.encode("utf-8"),
            content_type="text/plain",
        )

        response = self.client.post(
            self.url,
            {
                "label": "BOM library batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )

        self.assertRedirects(response, reverse("synopsis:reference_library"))
        self.assertEqual(LibraryReference.objects.count(), 2)
        titles = set(LibraryReference.objects.values_list("title", flat=True))
        self.assertEqual(titles, {"First RIS entry", "Second RIS entry"})
        batch = LibraryImportBatch.objects.get(label="BOM library batch")
        self.assertEqual(batch.record_count, 2)


class ReferenceBatchUploadParsingTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Reference Upload Project")
        self.user = User.objects.create_user(username="uploader", password="pw")
        UserRole.objects.create(user=self.user, project=self.project, role="author")
        self.client.force_login(self.user)
        self.url = reverse(
            "synopsis:reference_batch_upload", args=[self.project.id]
        )

    def _plaintext_payload(self):
        return textwrap.dedent(
            """
            Angel, D. L.; et al. (2002). "In situ biofiltration: a means to limit the dispersal of effluents from marine finfish cage aquaculture." Hydrobiologia 469(1): 1-10.
            Net pen fish farms generally enrich the surrounding waters and the underlying sediments with nutrients and organic matter.

            Morgan, W. and Thornton, A. (2018). Another example of parsing. Ecology Letters 11: 9-12.
            Full abstract text including doi:10.5678/example and https://example.com/article for reference.
            """
        ).strip()

    def test_imports_plaintext_file(self):
        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )

        response = self.client.post(
            self.url,
            {
                "label": "Plain text batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )

        self.assertRedirects(
            response,
            reverse("synopsis:reference_batch_list", args=[self.project.id]),
        )

        self.assertEqual(Reference.objects.filter(project=self.project).count(), 2)
        self.assertEqual(LibraryReference.objects.count(), 2)
        batch = ReferenceSourceBatch.objects.get(project=self.project)
        self.assertEqual(batch.record_count, 2)
        self.assertFalse(
            Reference.objects.filter(project=self.project, library_reference__isnull=True).exists()
        )
        titles = set(
            Reference.objects.filter(project=self.project).values_list(
                "title", flat=True
            )
        )
        self.assertIn(
            "In situ biofiltration: a means to limit the dispersal of effluents from marine finfish cage aquaculture.",
            titles,
        )
        self.assertIn("Another example of parsing", titles)

    def test_rejects_unparseable_plaintext(self):
        upload = SimpleUploadedFile(
            "bad.txt",
            b"not a valid ris record and no parsable citation",
            content_type="text/plain",
        )

        response = self.client.post(
            self.url,
            {
                "label": "Invalid batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )

        self.assertEqual(response.status_code, 200)
        form = response.context["form"]
        self.assertIn("ris_file", form.errors)
        self.assertIn("No references were detected", form.errors["ris_file"][0])
        self.assertEqual(Reference.objects.count(), 0)

    def test_imports_ris_file(self):
        ris_payload = textwrap.dedent(
            """
            TY  - JOUR
            TI  - Example Title
            AU  - Thornton, Ann
            PY  - 2021
            JO  - Marine Science Quarterly
            VL  - 12
            IS  - 3
            SP  - 101
            EP  - 110
            DO  - 10.1000/example
            ER  -
            """
        ).strip()
        upload = SimpleUploadedFile(
            "references.ris",
            ris_payload.encode("utf-8"),
            content_type="application/x-research-info-systems",
        )

        response = self.client.post(
            self.url,
            {
                "label": "RIS batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )

        self.assertRedirects(
            response,
            reverse("synopsis:reference_batch_list", args=[self.project.id]),
        )

        refs = Reference.objects.filter(project=self.project)
        self.assertEqual(refs.count(), 1)
        ref = refs.first()
        self.assertIsNotNone(ref.library_reference)
        self.assertEqual(ref.title, "Example Title")
        self.assertEqual(ref.publication_year, 2021)
        self.assertEqual(ref.journal, "Marine Science Quarterly")
        self.assertEqual(ref.pages, "101-110")
        self.assertEqual(ref.doi, "10.1000/example")

    def test_imports_ris_like_txt_file_with_utf8_bom(self):
        ris_payload = (
            "\ufeffTY  - JOUR\n"
            "TI  - First RIS entry\n"
            "AU  - Morgan, Will\n"
            "PY  - 2024\n"
            "DO  - 10.1000/first\n"
            "ER  -\n\n"
            "TY  - JOUR\n"
            "TI  - Second RIS entry\n"
            "AU  - Thornton, Ann\n"
            "PY  - 2023\n"
            "DO  - 10.1000/second\n"
            "ER  -\n"
        )
        upload = SimpleUploadedFile(
            "references.txt",
            ris_payload.encode("utf-8"),
            content_type="text/plain",
        )

        response = self.client.post(
            self.url,
            {
                "label": "BOM batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )

        self.assertRedirects(
            response,
            reverse("synopsis:reference_batch_list", args=[self.project.id]),
        )
        refs = Reference.objects.filter(project=self.project).order_by("title")
        self.assertEqual(refs.count(), 2)
        self.assertEqual(
            list(refs.values_list("title", flat=True)),
            ["First RIS entry", "Second RIS entry"],
        )
        batch = ReferenceSourceBatch.objects.get(project=self.project, label="BOM batch")
        self.assertEqual(batch.record_count, 2)

    def test_skips_duplicates_within_project(self):
        first_upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Initial batch",
                "source_type": "journal_search",
                "ris_file": first_upload,
            },
        )
        self.assertEqual(Reference.objects.filter(project=self.project).count(), 2)

        duplicate_upload = SimpleUploadedFile(
            "references_again.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        response = self.client.post(
            self.url,
            {
                "label": "Duplicate batch",
                "source_type": "journal_search",
                "ris_file": duplicate_upload,
            },
        )

        self.assertRedirects(
            response,
            reverse("synopsis:reference_batch_list", args=[self.project.id]),
        )
        self.assertEqual(Reference.objects.filter(project=self.project).count(), 2)
        latest_batch = (
            ReferenceSourceBatch.objects.filter(project=self.project)
            .order_by("-id")
            .first()
        )
        self.assertIsNotNone(latest_batch)
        self.assertEqual(latest_batch.record_count, 0)

    def test_reports_invalid_rows_separately_from_duplicates(self):
        existing_hash = reference_hash(
            "In situ biofiltration: a means to limit the dispersal of effluents from marine finfish cage aquaculture.",
            "2002",
            "",
        )
        Reference.objects.create(
            project=self.project,
            batch=ReferenceSourceBatch.objects.create(
                project=self.project,
                label="Existing refs",
                source_type="manual_upload",
                uploaded_by=self.user,
            ),
            hash_key=existing_hash,
            title="In situ biofiltration: a means to limit the dispersal of effluents from marine finfish cage aquaculture.",
            publication_year=2002,
        )

        original_normalise = _normalise_import_record

        def fake_normalise(record):
            if record.get("title") == "Another example of parsing":
                return None
            return original_normalise(record)

        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )

        with patch("synopsis.views._normalise_import_record", side_effect=fake_normalise):
            response = self.client.post(
                self.url,
                {
                    "label": "Mixed skip batch",
                    "source_type": "journal_search",
                    "ris_file": upload,
                },
                follow=True,
            )

        messages = [str(message) for message in get_messages(response.wsgi_request)]
        self.assertTrue(
            any("Imported 0 reference(s) into 'Mixed skip batch'." in message for message in messages)
        )
        self.assertTrue(
            any("Skipped 1 record(s) already present in this project." in message for message in messages)
        )
        self.assertTrue(
            any("Skipped 1 record(s) with no title." in message for message in messages)
        )

    def test_reuses_existing_library_reference_by_hash(self):
        existing_hash = reference_hash(
            "In situ biofiltration: a means to limit the dispersal of effluents from marine finfish cage aquaculture.",
            "2002",
            "",
        )
        existing_library_ref = LibraryReference.objects.create(
            hash_key=existing_hash,
            title="Existing canonical title",
            publication_year=2002,
        )

        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        response = self.client.post(
            self.url,
            {
                "label": "Reuse library ref batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )

        self.assertRedirects(
            response,
            reverse("synopsis:reference_batch_list", args=[self.project.id]),
        )
        self.assertEqual(LibraryReference.objects.count(), 2)
        project_ref = Reference.objects.get(
            project=self.project,
            hash_key=existing_hash,
        )
        self.assertEqual(project_ref.library_reference_id, existing_library_ref.id)

    def test_project_import_prefills_project_reference_from_shared_library_folders(self):
        existing_hash = reference_hash(
            "In situ biofiltration: a means to limit the dispersal of effluents from marine finfish cage aquaculture.",
            "2002",
            "",
        )
        existing_library_ref = LibraryReference.objects.create(
            hash_key=existing_hash,
            title="Existing canonical title",
            publication_year=2002,
            reference_folder=["2", "15"],
        )

        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Reuse shared folders batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )

        project_ref = Reference.objects.get(
            project=self.project,
            hash_key=existing_hash,
        )
        self.assertEqual(project_ref.library_reference_id, existing_library_ref.id)
        self.assertEqual(project_ref.unlinked_reference_folder, [])
        self.assertEqual(project_ref.category_values, ["2", "15"])

    def test_project_reference_uses_shared_categories_as_effective_value(self):
        existing_hash = reference_hash(
            "In situ biofiltration: a means to limit the dispersal of effluents from marine finfish cage aquaculture.",
            "2002",
            "",
        )
        existing_library_ref = LibraryReference.objects.create(
            hash_key=existing_hash,
            title="Existing canonical title",
            publication_year=2002,
            reference_folder=["2", "15"],
        )

        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Effective category batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )

        project_ref = Reference.objects.get(project=self.project, hash_key=existing_hash)
        project_ref.unlinked_reference_folder = ["3a"]
        project_ref.save(update_fields=["unlinked_reference_folder", "updated_at"])

        self.assertEqual(project_ref.unlinked_reference_folder, ["3a"])
        self.assertEqual(project_ref.category_values, ["2", "15"])
        self.assertEqual(project_ref.folder_labels(), ["2. Birds", "15. Forests/Woodland"])

    def test_can_delete_reference_from_batch(self):
        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Delete test batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )
        batch = ReferenceSourceBatch.objects.get(project=self.project)
        self.assertEqual(batch.references.count(), 2)
        ref_to_delete = batch.references.order_by("id").first()

        delete_url = reverse(
            "synopsis:reference_delete",
            args=[self.project.id, ref_to_delete.id],
        )
        response = self.client.post(delete_url, follow=False)
        self.assertRedirects(
            response,
            reverse(
                "synopsis:reference_batch_detail",
                args=[self.project.id, batch.id],
            ),
        )

        self.assertFalse(
            Reference.objects.filter(pk=ref_to_delete.id).exists()
        )
        batch.refresh_from_db()
        self.assertEqual(batch.record_count, batch.references.count())

    def test_delete_reference_requires_edit_permission(self):
        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Permission batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )
        batch = ReferenceSourceBatch.objects.get(project=self.project)
        target = batch.references.first()
        viewer = User.objects.create_user(username="viewer", password="pw")
        delete_url = reverse(
            "synopsis:reference_delete",
            args=[self.project.id, target.id],
        )
        self.client.logout()
        self.client.force_login(viewer)
        response = self.client.post(delete_url)
        self.assertEqual(response.status_code, 403)
        self.assertTrue(Reference.objects.filter(pk=target.id).exists())

    def test_can_delete_batch(self):
        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Batch to delete",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )
        batch = ReferenceSourceBatch.objects.get(project=self.project)
        self.assertEqual(batch.references.count(), 2)
        delete_url = reverse(
            "synopsis:reference_batch_delete",
            args=[self.project.id, batch.id],
        )

        response = self.client.post(delete_url, follow=False)
        self.assertRedirects(
            response,
            reverse("synopsis:reference_batch_list", args=[self.project.id]),
        )
        self.assertFalse(
            ReferenceSourceBatch.objects.filter(pk=batch.id).exists()
        )
        self.assertEqual(Reference.objects.filter(project=self.project).count(), 0)

    def test_bulk_include_selected_references(self):
        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Bulk batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )
        batch = ReferenceSourceBatch.objects.get(project=self.project)
        references = list(batch.references.order_by("id"))
        include_ids = [str(ref.id) for ref in references[:2]]

        detail_url = reverse(
            "synopsis:reference_batch_detail",
            args=[self.project.id, batch.id],
        )
        response = self.client.post(
            detail_url,
            {
                "bulk_action": "include",
                "selected_references": include_ids,
            },
        )

        self.assertRedirects(
            response,
            reverse(
                "synopsis:reference_batch_detail",
                args=[self.project.id, batch.id],
            ),
        )

        refreshed = Reference.objects.filter(pk__in=include_ids)
        self.assertTrue(refreshed.exists())
        for ref in refreshed:
            self.assertEqual(ref.screening_status, "included")
            self.assertEqual(ref.screened_by, self.user)
            self.assertIsNotNone(ref.screening_decision_at)

    def test_bulk_include_can_apply_selected_folders_at_same_time(self):
        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Bulk include folders batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )
        batch = ReferenceSourceBatch.objects.get(project=self.project)
        include_ids = [str(ref.id) for ref in batch.references.order_by("id")[:2]]
        detail_url = reverse(
            "synopsis:reference_batch_detail",
            args=[self.project.id, batch.id],
        )

        response = self.client.post(
            detail_url,
            {
                "bulk_action": "include",
                "selected_references": include_ids,
                "reference_folder": ["3a", "15"],
            },
            follow=True,
        )

        self.assertContains(
            response,
            "Applied the selected categories at the same time.",
        )
        for ref in Reference.objects.filter(pk__in=include_ids):
            self.assertEqual(ref.screening_status, "included")
            self.assertEqual(ref.unlinked_reference_folder, [])
            self.assertEqual(ref.category_values, ["3a", "15"])

        linked_library_refs = list(
            LibraryReference.objects.filter(project_references__id__in=include_ids).distinct()
        )
        self.assertTrue(linked_library_refs)
        for library_ref in linked_library_refs:
            self.assertEqual(library_ref.reference_folder, ["3a", "15"])
        self.assertTrue(
            LibraryReferenceFolderHistory.objects.filter(
                library_reference__in=linked_library_refs
            ).exists()
        )

    def test_screening_page_uses_resizable_folder_select_wrapper(self):
        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Folder width batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )
        batch = ReferenceSourceBatch.objects.get(project=self.project)

        response = self.client.get(
            reverse(
                "synopsis:reference_batch_detail",
                args=[self.project.id, batch.id],
            )
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "folder-select-shell")
        self.assertContains(response, "folder-select")
        self.assertContains(response, "screening-bulk-sticky")
        self.assertContains(response, "Apply categories")
        self.assertContains(response, "Include in synopsis")
        self.assertContains(response, "Exclude from synopsis")
        self.assertContains(response, "Multiple categories are allowed.")
        self.assertContains(response, 'id="reference-batch-page"', html=False)
        self.assertContains(response, "cePreservePageState({", html=False)
        self.assertContains(
            response,
            f'"screening-batch-state-{self.project.id}-{batch.id}"',
            html=False,
        )
        self.assertContains(
            response,
            "Inclusion and exclusion here apply only to this synopsis. Category changes feed back into the shared CE reference library.",
        )
        self.assertContains(
            response,
            "This is the main category-classification step while screening.",
        )
        self.assertContains(
            response,
            "changing categories here updates the shared reference library record and is reflected in linked synopsis copies in other projects",
        )

    def test_screening_page_uses_shared_categories_when_local_fallback_exists(self):
        existing_hash = reference_hash(
            "In situ biofiltration: a means to limit the dispersal of effluents from marine finfish cage aquaculture.",
            "2002",
            "",
        )
        existing_library_ref = LibraryReference.objects.create(
            hash_key=existing_hash,
            title="Existing canonical title",
            publication_year=2002,
            reference_folder=["2", "15"],
        )

        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Shared screening category batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )

        batch = ReferenceSourceBatch.objects.get(project=self.project)
        ref = Reference.objects.get(project=self.project, hash_key=existing_hash)
        ref.unlinked_reference_folder = ["3a"]
        ref.save(update_fields=["unlinked_reference_folder", "updated_at"])

        response = self.client.get(
            reverse(
                "synopsis:reference_batch_detail",
                args=[self.project.id, batch.id],
            )
        )

        self.assertEqual(ref.library_reference_id, existing_library_ref.id)
        self.assertContains(response, 'option value="2" selected')
        self.assertContains(response, 'option value="15" selected')

    def test_focused_screening_shows_fixed_decision_bar_and_current_status(self):
        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Focused screening batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )
        batch = ReferenceSourceBatch.objects.get(project=self.project)
        ref = batch.references.order_by("id").first()
        ref.screening_status = "excluded"
        ref.screening_decision_at = timezone.now()
        ref.screened_by = self.user
        ref.save(
            update_fields=[
                "screening_status",
                "screening_decision_at",
                "screened_by",
                "updated_at",
            ]
        )

        response = self.client.get(
            reverse(
                "synopsis:reference_batch_detail",
                args=[self.project.id, batch.id],
            )
            + f"?focus=1&ref={ref.id}"
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "focus-screening-meta")
        self.assertContains(response, "focus-category-actions")
        self.assertContains(response, "focus-decision-bar")
        self.assertContains(response, "focus-decision-actions")
        self.assertContains(response, "Current synopsis status")
        self.assertNotContains(response, "Current status")
        self.assertContains(response, "Reference notes")
        self.assertContains(
            response,
            "These notes stay on this synopsis copy of the reference. They do not update the shared reference library, and excluding this reference here removes it only from this synopsis.",
        )
        self.assertContains(response, "Save notes")
        self.assertContains(
            response,
            f"Last screening update by {self.user.username} on",
        )
        self.assertNotContains(
            response,
            f'data-bs-target="#refCommentsModal-{ref.id}"',
            html=False,
        )
        self.assertContains(
            response,
            '<span class="focus-status-pill is-active">Excluded</span>',
            html=False,
        )

    def test_focused_save_categories_stays_on_same_reference(self):
        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Focused category save batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )
        batch = ReferenceSourceBatch.objects.get(project=self.project)
        references = list(batch.references.order_by("id"))
        ref = references[0]
        next_ref = references[1]

        response = self.client.post(
            reverse(
                "synopsis:reference_batch_detail",
                args=[self.project.id, batch.id],
            ),
            {
                "action": "save-categories",
                "focus": "1",
                "focus_ref": str(ref.id),
                "reference_id": str(ref.id),
                "screening_status": ref.screening_status,
                "reference_folder": ["3a"],
                "next_ref_id": str(next_ref.id),
            },
        )

        self.assertEqual(
            response["Location"],
            f"{reverse('synopsis:reference_batch_detail', args=[self.project.id, batch.id])}?focus=1&ref={ref.id}",
        )
        ref.refresh_from_db()
        self.assertEqual(ref.category_values, ["3a"])

    def test_focused_save_notes_stays_on_same_reference(self):
        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Focused notes save batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )
        batch = ReferenceSourceBatch.objects.get(project=self.project)
        references = list(batch.references.order_by("id"))
        ref = references[0]
        next_ref = references[1]
        ref.screening_status = "included"
        ref.save(update_fields=["screening_status", "updated_at"])

        response = self.client.post(
            reverse(
                "synopsis:reference_batch_detail",
                args=[self.project.id, batch.id],
            ),
            {
                "action": "save-notes",
                "focus": "1",
                "focus_ref": str(ref.id),
                "reference_id": str(ref.id),
                "screening_status": "included",
                "screening_notes": "Keep for synopsis discussion.",
                "next_ref_id": str(next_ref.id),
            },
        )

        self.assertEqual(
            response["Location"],
            f"{reverse('synopsis:reference_batch_detail', args=[self.project.id, batch.id])}?focus=1&ref={ref.id}",
        )
        ref.refresh_from_db()
        self.assertEqual(ref.screening_status, "included")
        self.assertEqual(ref.screening_notes, "Keep for synopsis discussion.")

    def test_bulk_apply_folders_to_selected_references(self):
        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Bulk folder batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )
        batch = ReferenceSourceBatch.objects.get(project=self.project)
        references = list(batch.references.order_by("id"))
        selected_ids = [str(ref.id) for ref in references[:2]]

        detail_url = reverse(
            "synopsis:reference_batch_detail",
            args=[self.project.id, batch.id],
        )
        response = self.client.post(
            detail_url,
            {
                "bulk_action": "save-folders",
                "selected_references": selected_ids,
                "reference_folder": ["3a", "15"],
            },
        )

        self.assertRedirects(
            response,
            reverse(
                "synopsis:reference_batch_detail",
                args=[self.project.id, batch.id],
            ),
        )

        for ref in Reference.objects.filter(pk__in=selected_ids):
            self.assertEqual(ref.unlinked_reference_folder, [])
            self.assertEqual(ref.category_values, ["3a", "15"])
            self.assertEqual(ref.screened_by, self.user)
            self.assertIsNotNone(ref.screening_decision_at)

    def test_single_screening_update_filters_blank_reference_folder_values(self):
        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Single screening batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )
        batch = ReferenceSourceBatch.objects.get(project=self.project)
        ref = batch.references.order_by("id").first()
        detail_url = reverse(
            "synopsis:reference_batch_detail",
            args=[self.project.id, batch.id],
        )

        response = self.client.post(
            detail_url,
            {
                "reference_id": ref.id,
                "screening_status": "included",
                "screening_notes": "Relevant to the topic.",
                "reference_folder": ["", "3a"],
            },
        )

        self.assertEqual(
            response["Location"],
            f"{reverse('synopsis:reference_batch_detail', args=[self.project.id, batch.id])}#ref-{ref.id}",
        )
        ref.refresh_from_db()
        self.assertEqual(ref.screening_status, "included")
        self.assertEqual(ref.unlinked_reference_folder, [])
        self.assertEqual(ref.category_values, ["3a"])
        self.assertEqual(ref.library_reference.reference_folder, ["3a"])

    def test_save_folders_preserves_existing_screening_notes(self):
        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Folder notes batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )
        batch = ReferenceSourceBatch.objects.get(project=self.project)
        ref = batch.references.order_by("id").first()
        ref.screening_notes = "Keep these notes."
        ref.save(update_fields=["screening_notes", "updated_at"])
        detail_url = reverse(
            "synopsis:reference_batch_detail",
            args=[self.project.id, batch.id],
        )

        response = self.client.post(
            detail_url,
            {
                "reference_id": ref.id,
                "screening_status": ref.screening_status,
                "reference_folder": ["15"],
            },
        )

        self.assertEqual(
            response["Location"],
            f"{reverse('synopsis:reference_batch_detail', args=[self.project.id, batch.id])}#ref-{ref.id}",
        )
        ref.refresh_from_db()
        self.assertEqual(ref.unlinked_reference_folder, [])
        self.assertEqual(ref.category_values, ["15"])
        self.assertEqual(ref.screening_notes, "Keep these notes.")

    def test_single_reference_can_be_reset_to_pending(self):
        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Pending reset batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )
        batch = ReferenceSourceBatch.objects.get(project=self.project)
        ref = batch.references.order_by("id").first()
        ref.screening_status = "included"
        ref.screening_decision_at = timezone.now()
        ref.screened_by = self.user
        ref.save(
            update_fields=["screening_status", "screening_decision_at", "screened_by", "updated_at"]
        )

        detail_url = reverse(
            "synopsis:reference_batch_detail",
            args=[self.project.id, batch.id],
        )
        response = self.client.post(
            detail_url,
            {
                "reference_id": ref.id,
                "screening_status": "pending",
            },
        )

        self.assertEqual(
            response["Location"],
            f"{reverse('synopsis:reference_batch_detail', args=[self.project.id, batch.id])}#ref-{ref.id}",
        )
        ref.refresh_from_db()
        self.assertEqual(ref.screening_status, "pending")

    def test_bulk_action_requires_selection(self):
        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Bulk batch empty",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )
        batch = ReferenceSourceBatch.objects.get(project=self.project)
        detail_url = reverse(
            "synopsis:reference_batch_detail",
            args=[self.project.id, batch.id],
        )

        response = self.client.post(
            detail_url,
            {
                "bulk_action": "exclude",
            },
            follow=True,
        )

        self.assertContains(
            response,
            "Select at least one reference before applying a bulk update.",
        )
        self.assertTrue(
            Reference.objects.filter(project=self.project, screening_status="pending").exists()
        )

    def test_bulk_reset_to_pending(self):
        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Bulk reset batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )
        batch = ReferenceSourceBatch.objects.get(project=self.project)
        references = list(batch.references.order_by("id"))
        include_ids = [str(ref.id) for ref in references[:2]]
        detail_url = reverse(
            "synopsis:reference_batch_detail",
            args=[self.project.id, batch.id],
        )

        self.client.post(
            detail_url,
            {
                "bulk_action": "include",
                "selected_references": include_ids,
            },
        )
        response = self.client.post(
            detail_url,
            {
                "bulk_action": "pending",
                "selected_references": include_ids,
            },
        )
        self.assertRedirects(
            response,
            reverse(
                "synopsis:reference_batch_detail",
                args=[self.project.id, batch.id],
            ),
        )
        for ref in Reference.objects.filter(pk__in=include_ids):
            self.assertEqual(ref.screening_status, "pending")

    def test_update_notes_creates_history(self):
        upload = SimpleUploadedFile(
            "references.txt",
            self._plaintext_payload().encode("utf-8"),
            content_type="text/plain",
        )
        self.client.post(
            self.url,
            {
                "label": "Notes batch",
                "source_type": "journal_search",
                "ris_file": upload,
            },
        )
        batch = ReferenceSourceBatch.objects.get(project=self.project)
        detail_url = reverse(
            "synopsis:reference_batch_detail",
            args=[self.project.id, batch.id],
        )

        response = self.client.post(
            detail_url,
            {
                "action": "update_notes",
                "notes": "Initial notes",
            },
        )
        self.assertRedirects(
            response,
            reverse(
                "synopsis:reference_batch_detail",
                args=[self.project.id, batch.id],
            ),
        )
        batch.refresh_from_db()
        self.assertEqual(batch.notes, "Initial notes")
        history = ReferenceSourceBatchNoteHistory.objects.filter(batch=batch)
        self.assertEqual(history.count(), 1)
        first_entry = history.first()
        self.assertEqual(first_entry.previous_notes, "")
        self.assertEqual(first_entry.new_notes, "Initial notes")
        self.assertEqual(first_entry.changed_by, self.user)

        response = self.client.post(
            detail_url,
            {
                "action": "update_notes",
                "notes": "Updated notes",
            },
        )
        self.assertRedirects(
            response,
            reverse(
                "synopsis:reference_batch_detail",
                args=[self.project.id, batch.id],
            ),
        )
        batch.refresh_from_db()
        self.assertEqual(batch.notes, "Updated notes")
        history = ReferenceSourceBatchNoteHistory.objects.filter(batch=batch).order_by(
            "-changed_at"
        )
        self.assertEqual(history.count(), 2)
        latest = history.first()
        self.assertEqual(latest.previous_notes, "Initial notes")
        self.assertEqual(latest.new_notes, "Updated notes")


class LibraryLinkBatchTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Library Link Project")
        self.user = User.objects.create_user(username="linker", password="pw")

    def test_linking_on_separate_operations_creates_distinct_library_batches(self):
        first_lib_ref = LibraryReference.objects.create(
            title="Library reference one",
            publication_year=2020,
            doi="10.1000/one",
            hash_key="lib-hash-one",
        )
        second_lib_ref = LibraryReference.objects.create(
            title="Library reference two",
            publication_year=2021,
            doi="10.1000/two",
            hash_key="lib-hash-two",
        )

        linked_one, reused_one, batch_one = _link_library_references_to_project(
            self.user,
            self.project,
            [first_lib_ref.id],
            ["15"],
        )
        linked_two, reused_two, batch_two = _link_library_references_to_project(
            self.user,
            self.project,
            [second_lib_ref.id],
            ["15"],
        )

        self.assertEqual((linked_one, reused_one), (1, 0))
        self.assertEqual((linked_two, reused_two), (1, 0))
        self.assertIsNotNone(batch_one)
        self.assertIsNotNone(batch_two)
        self.assertNotEqual(batch_one.id, batch_two.id)
        self.assertNotEqual(batch_one.label, batch_two.label)
        self.assertEqual(batch_one.source_type, "library_link")
        self.assertEqual(batch_two.source_type, "library_link")
        self.assertEqual(batch_one.record_count, 1)
        self.assertEqual(batch_two.record_count, 1)
        self.assertEqual(
            Reference.objects.filter(project=self.project, batch=batch_one).count(), 1
        )
        self.assertEqual(
            Reference.objects.filter(project=self.project, batch=batch_two).count(), 1
        )

    def test_duplicate_only_link_does_not_create_empty_library_batch(self):
        lib_ref = LibraryReference.objects.create(
            title="Duplicate library reference",
            publication_year=2022,
            doi="10.1000/duplicate",
            hash_key="lib-hash-duplicate",
        )
        linked, reused, initial_batch = _link_library_references_to_project(
            self.user,
            self.project,
            [lib_ref.id],
            ["15"],
        )
        self.assertEqual((linked, reused), (1, 0))
        self.assertIsNotNone(initial_batch)

        batch_count_before = ReferenceSourceBatch.objects.filter(
            project=self.project,
            source_type="library_link",
        ).count()

        linked_again, reused_again, duplicate_batch = _link_library_references_to_project(
            self.user,
            self.project,
            [lib_ref.id],
            ["15"],
        )

        self.assertEqual((linked_again, reused_again), (0, 1))
        self.assertIsNone(duplicate_batch)
        self.assertEqual(
            ReferenceSourceBatch.objects.filter(
                project=self.project,
                source_type="library_link",
            ).count(),
            batch_count_before,
        )

    def test_link_uses_existing_shared_library_folders_by_default(self):
        lib_ref = LibraryReference.objects.create(
            title="Shared folders default",
            publication_year=2024,
            doi="10.1000/shared-default",
            hash_key="lib-hash-shared-default",
            reference_folder=["2", "15"],
        )

        linked, reused, batch = _link_library_references_to_project(
            self.user,
            self.project,
            [lib_ref.id],
            [],
        )

        self.assertEqual((linked, reused), (1, 0))
        self.assertIsNotNone(batch)
        project_ref = Reference.objects.get(project=self.project, library_reference=lib_ref)
        self.assertEqual(project_ref.unlinked_reference_folder, [])
        self.assertEqual(project_ref.category_values, ["2", "15"])

    def test_link_folder_override_updates_shared_library_record(self):
        lib_ref = LibraryReference.objects.create(
            title="Shared folders override",
            publication_year=2024,
            doi="10.1000/shared-override",
            hash_key="lib-hash-shared-override",
            reference_folder=["2"],
        )

        linked, reused, _batch = _link_library_references_to_project(
            self.user,
            self.project,
            [lib_ref.id],
            ["15", "2"],
        )

        self.assertEqual((linked, reused), (1, 0))
        lib_ref.refresh_from_db()
        self.assertEqual(lib_ref.reference_folder, ["2", "15"])
        project_ref = Reference.objects.get(project=self.project, library_reference=lib_ref)
        self.assertEqual(project_ref.unlinked_reference_folder, [])
        self.assertEqual(project_ref.category_values, ["2", "15"])
        self.assertTrue(
            LibraryReferenceFolderHistory.objects.filter(
                library_reference=lib_ref,
                new_folders=["2", "15"],
            ).exists()
        )


class LibraryReferenceDetailTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="libeditor", password="pw")
        self.project = Project.objects.create(title="Library Sync Project")
        UserRole.objects.create(user=self.user, project=self.project, role="author")
        self.library_reference = LibraryReference.objects.create(
            title="Library reference",
            publication_year=2024,
            doi="10.1000/library-detail",
            hash_key="lib-detail-hash",
            reference_folder=["2"],
        )
        self.batch = ReferenceSourceBatch.objects.create(
            project=self.project,
            label="Linked batch",
            source_type="library_link",
            uploaded_by=self.user,
        )
        self.project_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            library_reference=self.library_reference,
            hash_key="lib-detail-hash",
            title="Project copy",
        )
        self.client.force_login(self.user)

    def test_library_detail_updates_shared_folders_for_linked_project_copies(self):
        response = self.client.post(
            reverse(
                "synopsis:library_reference_detail",
                args=[self.library_reference.id],
            ),
            {
                "action": "edit",
                "title": "Library reference",
                "authors": "",
                "publication_year": 2024,
                "journal": "",
                "volume": "",
                "issue": "",
                "pages": "",
                "doi": "10.1000/library-detail",
                "url": "",
                "language": "",
                "abstract": "",
                "reference_folder": ["15", "2"],
            },
            follow=True,
        )

        self.library_reference.refresh_from_db()
        self.project_reference.refresh_from_db()
        self.assertEqual(self.library_reference.reference_folder, ["2", "15"])
        self.assertEqual(self.project_reference.unlinked_reference_folder, [])
        self.assertEqual(self.project_reference.category_values, ["2", "15"])
        self.assertContains(
            response,
            "Shared CE subject categories were updated.",
        )
        self.assertContains(
            response,
            "Linked synopsis copies now read those shared categories automatically.",
        )
        self.assertTrue(
            LibraryReferenceFolderHistory.objects.filter(
                library_reference=self.library_reference,
                new_folders=["2", "15"],
                change_source="library_detail",
            ).exists()
        )


class ReferenceSummaryFormTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Summary form project")

    def test_update_form_requires_project(self):
        with self.assertRaises(TypeError):
            ReferenceSummaryUpdateForm()

    def test_habitat_tags_use_detailed_iucn_choices(self):
        values = [value for value, _label in IUCN_HABITAT_CHOICES]
        self.assertEqual(len(values), 63)
        self.assertIn("Marine Coral Reefs", values)
        self.assertIn(
            "Wetlands (inland) - Permanent Freshwater Lakes",
            values,
        )
        self.assertIn(
            "Artificial - Subtropical/Tropical Heavily Degraded Former Forest",
            values,
        )

        form = ReferenceSummaryUpdateForm(
            data={
                "status": ReferenceSummary.STATUS_TODO,
                "habitat_tags": [
                    "Marine Coral Reefs",
                    "Forest - Temperate",
                ],
            },
            project=self.project,
        )
        self.assertTrue(form.is_valid())
        self.assertEqual(
            form.cleaned_data["habitat_tags"],
            [
                "Marine Coral Reefs",
                "Forest - Temperate",
            ],
        )

    def test_action_tags_use_detailed_iucn_choices(self):
        values = [value for value, _label in IUCN_ACTION_CHOICES]
        self.assertEqual(len(values), 37)
        self.assertIn("Land/water protection-Area protection", values)
        self.assertIn(
            "Livelihood, economic & other incentives-Conservation payments", values
        )
        self.assertIn("Research & monitoring-Other", values)

        form = ReferenceSummaryUpdateForm(
            data={
                "status": ReferenceSummary.STATUS_TODO,
                "action_tags": [
                    "Land/water management-Site/area management",
                    "Research & monitoring-Conservation planning",
                ],
            },
            project=self.project,
        )
        self.assertTrue(form.is_valid())
        self.assertEqual(
            form.cleaned_data["action_tags"],
            [
                "Land/water management-Site/area management",
                "Research & monitoring-Conservation planning",
            ],
        )

    def test_threat_tags_use_detailed_iucn_choices(self):
        values = [value for value, _label in IUCN_THREAT_CHOICES]
        self.assertEqual(len(values), 41)
        self.assertIn(
            "Residential & commercial development-Housing/urban areas", values
        )
        self.assertIn("Invasive & other problematic species & genes", values)
        self.assertIn(
            "Climate change & severe weather-Other impacts", values
        )

        form = ReferenceSummaryUpdateForm(
            data={
                "status": ReferenceSummary.STATUS_TODO,
                "threat_tags": [
                    "Residential & commercial development-Housing/urban areas",
                    "Climate change & severe weather-Storms/flooding",
                ],
            },
            project=self.project,
        )
        self.assertTrue(form.is_valid())
        self.assertEqual(
            form.cleaned_data["threat_tags"],
            [
                "Residential & commercial development-Housing/urban areas",
                "Climate change & severe weather-Storms/flooding",
            ],
        )

    def test_research_design_accepts_up_to_four_tags(self):
        form = ReferenceSummaryUpdateForm(
            data={
                "status": ReferenceSummary.STATUS_TODO,
                "research_design": [
                    "Replicated",
                    "Randomized",
                    "Controlled*",
                    "Before-and-after",
                ],
            },
            project=self.project,
        )

        self.assertTrue(form.is_valid())
        self.assertEqual(
            form.cleaned_data["research_design"],
            "Replicated; Randomized; Controlled*; Before-and-after",
        )

    def test_research_design_rejects_more_than_four_tags(self):
        form = ReferenceSummaryUpdateForm(
            data={
                "status": ReferenceSummary.STATUS_TODO,
                "research_design": [
                    "Replicated",
                    "Randomized",
                    "Paired sites",
                    "Controlled*",
                    "Before-and-after",
                ],
            },
            project=self.project,
        )

        self.assertFalse(form.is_valid())
        self.assertIn("Select up to 4 research design tags", str(form.errors))

    def test_research_design_initial_splits_saved_tags(self):
        summary = ReferenceSummary(research_design="Replicated; Controlled*")

        form = ReferenceSummaryUpdateForm(instance=summary, project=self.project)

        self.assertEqual(
            form["research_design"].value(),
            ["Replicated", "Controlled*"],
        )

    def test_blank_study_design_is_built_from_research_design_tags(self):
        project = Project.objects.create(title="Auto design")
        batch = ReferenceSourceBatch.objects.create(
            project=project,
            label="Batch",
            source_type="journal_search",
        )
        reference = Reference.objects.create(
            project=project,
            batch=batch,
            hash_key="d" * 40,
            title="Auto design reference",
        )
        summary = ReferenceSummary.objects.create(
            project=project,
            reference=reference,
            status=ReferenceSummary.STATUS_TODO,
        )

        form = ReferenceSummaryUpdateForm(
            data={
                "status": ReferenceSummary.STATUS_TODO,
                "study_design": "",
                "research_design": [
                    "Replicated",
                    "Randomized",
                    "Controlled*",
                ],
            },
            instance=summary,
            project=project,
        )

        self.assertTrue(form.is_valid(), form.errors)
        saved = form.save()
        self.assertEqual(
            saved.study_design,
            "replicated, randomized, controlled study",
        )

    def test_manual_study_design_overrides_research_design_tags(self):
        project = Project.objects.create(title="Manual design")
        batch = ReferenceSourceBatch.objects.create(
            project=project,
            label="Batch",
            source_type="journal_search",
        )
        reference = Reference.objects.create(
            project=project,
            batch=batch,
            hash_key="e" * 40,
            title="Manual design reference",
        )
        summary = ReferenceSummary.objects.create(
            project=project,
            reference=reference,
            status=ReferenceSummary.STATUS_TODO,
        )

        form = ReferenceSummaryUpdateForm(
            data={
                "status": ReferenceSummary.STATUS_TODO,
                "study_design": "replicated, randomized, controlled, before-and-after study",
                "research_design": [
                    "Replicated",
                    "Randomized",
                    "Controlled*",
                ],
            },
            instance=summary,
            project=project,
        )

        self.assertTrue(form.is_valid(), form.errors)
        saved = form.save()
        self.assertEqual(
            saved.study_design,
            "replicated, randomized, controlled, before-and-after study",
        )

    def test_methods_and_design_initial_merges_existing_fields(self):
        summary = ReferenceSummary(
            action_methods="Used fenced plots and added seed.",
            experimental_design="Compared treated and untreated plots over two years.",
        )

        form = ReferenceSummaryUpdateForm(instance=summary, project=self.project)

        self.assertEqual(
            form["methods_and_design"].value(),
            "Used fenced plots and added seed.\n\nCompared treated and untreated plots over two years.",
        )

    def test_action_dropdown_uses_project_interventions(self):
        project = Project.objects.create(title="Action choices")
        batch = ReferenceSourceBatch.objects.create(
            project=project,
            label="Batch",
            source_type="journal_search",
        )
        reference = Reference.objects.create(
            project=project,
            batch=batch,
            hash_key="a" * 40,
            title="Action reference",
        )
        summary = ReferenceSummary.objects.create(
            project=project,
            reference=reference,
            action_description="Install nest boxes",
        )
        chapter = SynopsisChapter.objects.create(
            project=project,
            title="Evidence",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="General",
            position=1,
        )
        SynopsisIntervention.objects.create(
            subheading=subheading,
            title="Install nest boxes",
            position=1,
        )

        form = ReferenceSummaryUpdateForm(instance=summary, project=project)

        choice_values = [value for value, _label in form.fields["action_choice"].choices]
        self.assertIn("Install nest boxes", choice_values)
        self.assertEqual(form["action_choice"].value(), "Install nest boxes")
        self.assertEqual(form["action_custom"].value(), None)

    def test_action_dropdown_supports_custom_value_when_not_in_structure(self):
        project = Project.objects.create(title="Custom action choice")
        batch = ReferenceSourceBatch.objects.create(
            project=project,
            label="Batch",
            source_type="journal_search",
        )
        reference = Reference.objects.create(
            project=project,
            batch=batch,
            hash_key="b" * 40,
            title="Custom action reference",
        )
        summary = ReferenceSummary.objects.create(
            project=project,
            reference=reference,
            action_description="Reduce ditch dredging",
        )

        form = ReferenceSummaryUpdateForm(instance=summary, project=project)

        self.assertEqual(
            form["action_choice"].value(),
            ReferenceSummaryUpdateForm.ACTION_CUSTOM_VALUE,
        )
        self.assertEqual(form["action_custom"].value(), "Reduce ditch dredging")

    def test_action_dropdown_save_uses_selected_intervention_title(self):
        project = Project.objects.create(title="Dropdown save")
        batch = ReferenceSourceBatch.objects.create(
            project=project,
            label="Batch",
            source_type="journal_search",
        )
        reference = Reference.objects.create(
            project=project,
            batch=batch,
            hash_key="c" * 40,
            title="Save action reference",
        )
        summary = ReferenceSummary.objects.create(
            project=project,
            reference=reference,
            status=ReferenceSummary.STATUS_TODO,
        )
        chapter = SynopsisChapter.objects.create(
            project=project,
            title="Evidence",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="General",
            position=1,
        )
        SynopsisIntervention.objects.create(
            subheading=subheading,
            title="Install nest boxes",
            position=1,
        )

        form = ReferenceSummaryUpdateForm(
            data={
                "status": ReferenceSummary.STATUS_DRAFT,
                "action_choice": "Install nest boxes",
                "action_custom": "",
            },
            instance=summary,
            project=project,
        )

        self.assertTrue(form.is_valid(), form.errors)
        saved = form.save()
        self.assertEqual(saved.action_description, "Install nest boxes")

    def test_methods_and_design_save_flattens_into_single_summary_field(self):
        project = Project.objects.create(title="Methods Merge")
        batch = ReferenceSourceBatch.objects.create(
            project=project,
            label="Batch",
            source_type="journal_search",
        )
        reference = Reference.objects.create(
            project=project,
            batch=batch,
            hash_key="m" * 40,
            title="Methods reference",
        )
        summary = ReferenceSummary.objects.create(
            project=project,
            reference=reference,
            status=ReferenceSummary.STATUS_TODO,
            action_methods="Old methods",
            experimental_design="Old design",
        )

        form = ReferenceSummaryUpdateForm(
            data={
                "status": ReferenceSummary.STATUS_DRAFT,
                "methods_and_design": "Combined methods and design notes.",
            },
            instance=summary,
            project=project,
        )

        self.assertTrue(form.is_valid(), form.errors)
        saved = form.save()

        self.assertEqual(saved.action_methods, "Combined methods and design notes.")
        self.assertEqual(saved.experimental_design, "")

    def test_draft_form_prefills_generated_summary_when_no_saved_draft_exists(self):
        project = Project.objects.create(title="Coral Reefs Synopsis")
        batch = ReferenceSourceBatch.objects.create(
            project=project,
            label="Batch",
            source_type="journal_search",
        )
        reference = Reference.objects.create(
            project=project,
            batch=batch,
            hash_key="c" * 40,
            title="Test reference",
        )
        summary = ReferenceSummary.objects.create(project=project, reference=reference)

        form = ReferenceSummaryDraftForm(
            instance=summary,
            generated_summary="Auto-generated paragraph.",
        )

        self.assertEqual(form["synopsis_draft"].value(), "Auto-generated paragraph.")

    def test_draft_form_prefills_saved_custom_paragraph_when_custom_mode_is_active(self):
        project = Project.objects.create(title="Saved custom paragraph")
        batch = ReferenceSourceBatch.objects.create(
            project=project,
            label="Batch",
            source_type="journal_search",
        )
        reference = Reference.objects.create(
            project=project,
            batch=batch,
            hash_key="f" * 40,
            title="Custom paragraph reference",
        )
        summary = ReferenceSummary.objects.create(
            project=project,
            reference=reference,
            synopsis_draft="Author edited paragraph.",
            use_custom_synopsis_draft=True,
        )

        form = ReferenceSummaryDraftForm(
            instance=summary,
            generated_summary="Auto-generated paragraph.",
        )

        self.assertEqual(form["synopsis_draft"].value(), "Author edited paragraph.")

    def test_citation_field_prefills_with_shared_reference_citation_when_no_local_override_exists(self):
        batch = ReferenceSourceBatch.objects.create(
            project=self.project,
            label="Citation batch",
            source_type="journal_search",
        )
        reference = Reference.objects.create(
            project=self.project,
            batch=batch,
            hash_key="c" * 40,
            title="Corallivorous snail removal",
            authors="Miller M.",
            publication_year=2001,
            journal="Coral Reefs",
            volume="19",
            pages="293-295",
            doi="10.1007/PL00006963",
        )
        summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=reference,
            citation="Miller M. (2001) Corallivorous snail removal",
        )

        form = ReferenceSummaryUpdateForm(instance=summary, project=self.project)

        self.assertEqual(
            form["citation"].value(),
            reference_summary_effective_citation(summary),
        )
        self.assertIn("Coral Reefs, 19, 293-295.", form["citation"].value())

    def test_citation_matching_shared_reference_is_not_saved_as_local_override(self):
        batch = ReferenceSourceBatch.objects.create(
            project=self.project,
            label="Citation save batch",
            source_type="journal_search",
        )
        reference = Reference.objects.create(
            project=self.project,
            batch=batch,
            hash_key="d" * 40,
            title="Corallivorous snail removal",
            authors="Miller M.",
            publication_year=2001,
            journal="Coral Reefs",
            volume="19",
            pages="293-295",
            doi="10.1007/PL00006963",
        )
        summary = ReferenceSummary.objects.create(project=self.project, reference=reference)
        shared_citation = reference_summary_effective_citation(summary)

        form = ReferenceSummaryUpdateForm(
            data={
                "status": ReferenceSummary.STATUS_TODO,
                "citation": shared_citation,
            },
            instance=summary,
            project=self.project,
        )

        self.assertTrue(form.is_valid(), form.errors)
        saved = form.save()
        self.assertEqual(saved.citation, "")

    def test_location_tags_accepts_place_and_coords(self):
        form = ReferenceSummaryUpdateForm(
            data={
                "status": ReferenceSummary.STATUS_TODO,
                "location_tags": "London, UK - 51.50740, -0.12780",
            },
            project=self.project,
        )
        self.assertTrue(form.is_valid())
        self.assertEqual(form.cleaned_data["location_tags"], ["London, UK - 51.50740, -0.12780"])

    def test_location_tags_rejects_out_of_range(self):
        form = ReferenceSummaryUpdateForm(
            data={
                "status": ReferenceSummary.STATUS_TODO,
                "location_tags": "Nowhere - 123.00000, 200.00000",
            },
            project=self.project,
        )
        self.assertFalse(form.is_valid())
        self.assertIn("Coordinates must be valid latitude", str(form.errors))

    def test_outcomes_raw_ignores_empty_rows(self):
        data = {
            "status": ReferenceSummary.STATUS_TODO,
            "outcomes_raw": "Outcome | 1 | treat | 2 | comp | unit | diff | stats | p | notes\n | | | | | | | | | ",
        }
        form = ReferenceSummaryUpdateForm(data=data, project=self.project)
        self.assertTrue(form.is_valid())
        cleaned = form.cleaned_data["outcomes_raw"]
        self.assertEqual(len(cleaned), 1)
        self.assertEqual(cleaned[0]["outcome"], "Outcome")

    def test_outcomes_raw_accepts_free_text_sentence_lines(self):
        data = {
            "status": ReferenceSummary.STATUS_TODO,
            "outcomes_raw": "Species richness increased after scrub removal\nBreeding success stayed similar between treatments.",
        }
        form = ReferenceSummaryUpdateForm(data=data, project=self.project)
        self.assertTrue(form.is_valid(), form.errors)
        cleaned = form.cleaned_data["outcomes_raw"]
        self.assertEqual(
            cleaned,
            [
                {"sentence": "Species richness increased after scrub removal"},
                {"sentence": "Breeding success stayed similar between treatments."},
            ],
        )

    def test_structured_summary_paragraph_uses_free_text_outcome_notes(self):
        summary = ReferenceSummary(
            study_design="replicated study",
            summary_of_results="brush cutting improved habitat condition.",
            outcome_rows=[
                {"sentence": "Species richness increased after scrub removal"},
                {"sentence": "Breeding success stayed similar between treatments."},
            ],
        )

        paragraph = _structured_summary_paragraph(summary)

        self.assertIn("Species richness increased after scrub removal.", paragraph)
        self.assertIn("Breeding success stayed similar between treatments.", paragraph)

    def test_structured_summary_paragraph_excludes_quality_scores_from_text(self):
        summary = ReferenceSummary(
            study_design="replicated study",
            summary_of_results="brush cutting improved habitat condition.",
            benefits_score=80,
            harms_score=5,
            reliability_score=0.7,
            relevance_score=0.9,
        )

        paragraph = _structured_summary_paragraph(summary)

        self.assertNotIn("Benefits:", paragraph)
        self.assertNotIn("Harms:", paragraph)
        self.assertNotIn("Reliability:", paragraph)
        self.assertNotIn("Relevance:", paragraph)

    def test_quality_scores_accept_boundary_values(self):
        form = ReferenceSummaryUpdateForm(
            data={
                "status": ReferenceSummary.STATUS_TODO,
                "benefits_score": "0",
                "harms_score": "100",
                "reliability_score": "0.0",
                "relevance_score": "1.0",
            },
            project=self.project,
        )
        self.assertTrue(form.is_valid())
        self.assertEqual(form.cleaned_data["benefits_score"], 0.0)
        self.assertEqual(form.cleaned_data["harms_score"], 100.0)
        self.assertEqual(form.cleaned_data["reliability_score"], 0.0)
        self.assertEqual(form.cleaned_data["relevance_score"], 1.0)

    def test_quality_scores_reject_values_outside_ranges(self):
        invalid_cases = [
            ("benefits_score", "-0.1"),
            ("benefits_score", "100.1"),
            ("harms_score", "-1"),
            ("harms_score", "101"),
            ("reliability_score", "-0.01"),
            ("reliability_score", "1.01"),
            ("relevance_score", "-0.5"),
            ("relevance_score", "2"),
        ]
        for field_name, value in invalid_cases:
            with self.subTest(field_name=field_name, value=value):
                form = ReferenceSummaryUpdateForm(
                    data={
                        "status": ReferenceSummary.STATUS_TODO,
                        field_name: value,
                    },
                    project=self.project,
                )
                self.assertFalse(form.is_valid())
                self.assertIn(field_name, form.errors)


class ReferenceSummaryDetailViewTests(TestCase):
    def setUp(self):
        cache.clear()
        self.user = User.objects.create_user(username="author", password="pass123")
        self.viewer = User.objects.create_user(username="viewer", password="pass123")
        self.project = Project.objects.create(title="Coral Reefs Synopsis")
        UserRole.objects.create(user=self.user, project=self.project, role="author")
        self.batch = ReferenceSourceBatch.objects.create(
            project=self.project,
            label="Batch",
            source_type="journal_search",
        )
        self.reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            hash_key="a" * 40,
            title="Test reference",
        )
        self.summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=self.reference,
            status=ReferenceSummary.STATUS_TODO,
        )

    def test_detail_page_is_forbidden_for_non_project_editors(self):
        self.client.login(username="viewer", password="pass123")

        response = self.client.get(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            )
        )

        self.assertEqual(response.status_code, 403)

    def test_duplicate_summary_tab_is_forbidden_for_non_project_editors(self):
        self.client.login(username="viewer", password="pass123")

        response = self.client.post(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            ),
            {"action": "duplicate-summary-tab"},
            follow=False,
        )

        self.assertEqual(response.status_code, 403)
        self.assertEqual(
            ReferenceSummary.objects.filter(
                project=self.project,
                reference=self.reference,
            ).count(),
            1,
        )

    def test_detail_page_explains_local_citation_override_behaviour(self):
        self.client.login(username="author", password="pass123")

        response = self.client.get(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            )
        )

        self.assertContains(response, "Citation for synopsis export")
        self.assertContains(response, "Shared reference citation in use")
        self.assertContains(
            response,
            "does not update the shared reference database",
        )
        self.assertContains(
            response,
            "&lt;i&gt;...&lt;/i&gt; or &lt;em&gt;...&lt;/em&gt; for italics.",
            html=False,
        )

    def test_detail_page_shows_project_action_dropdown_options(self):
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="Evidence",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="General",
            position=1,
        )
        SynopsisIntervention.objects.create(
            subheading=subheading,
            title="Install nest boxes",
            position=1,
        )

        self.client.login(username="author", password="pass123")
        response = self.client.get(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            )
        )

        self.assertContains(response, "Choose an action already added to the project intervention list.")
        self.assertContains(response, '<option value="Install nest boxes">Install nest boxes</option>', html=False)

    def test_detail_page_warns_when_another_author_is_active_in_summary(self):
        other_author = User.objects.create_user(
            username="coauthor",
            password="pass123",
            first_name="Co",
            last_name="Author",
        )
        UserRole.objects.create(user=other_author, project=self.project, role="author")

        self.client.login(username="coauthor", password="pass123")
        self.client.post(
            reverse(
                "synopsis:reference_summary_presence",
                args=[self.project.id, self.summary.id],
            )
        )

        self.client.login(username="author", password="pass123")
        response = self.client.get(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            )
        )

        self.assertContains(response, "Active author")
        self.assertContains(response, "You + Co Author")
        self.assertContains(response, "Co Author")
        self.assertContains(
            response,
            reverse(
                "synopsis:reference_summary_presence",
                args=[self.project.id, self.summary.id],
            ),
            html=False,
        )

    def test_summary_presence_endpoint_returns_active_participants(self):
        other_author = User.objects.create_user(
            username="coauthor",
            password="pass123",
            first_name="Co",
            last_name="Author",
        )
        UserRole.objects.create(user=other_author, project=self.project, role="author")

        self.client.login(username="coauthor", password="pass123")
        self.client.post(
            reverse(
                "synopsis:reference_summary_presence",
                args=[self.project.id, self.summary.id],
            )
        )

        self.client.login(username="author", password="pass123")
        response = self.client.post(
            reverse(
                "synopsis:reference_summary_presence",
                args=[self.project.id, self.summary.id],
            )
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload["current_user_active"])
        self.assertIn("Co Author", payload["other_participants"])
        self.assertIn("author", payload["participant_names"])

    def test_summary_presence_endpoint_rejects_get_requests(self):
        self.client.login(username="author", password="pass123")

        response = self.client.get(
            reverse(
                "synopsis:reference_summary_presence",
                args=[self.project.id, self.summary.id],
            )
        )

        self.assertEqual(response.status_code, 400)

    def test_detail_page_explains_optional_fields_and_custom_paragraph_mode(self):
        self.summary.synopsis_draft = "Manual paragraph text."
        self.summary.use_custom_synopsis_draft = True
        self.summary.save(
            update_fields=["synopsis_draft", "use_custom_synopsis_draft", "updated_at"]
        )

        self.client.login(username="author", password="pass123")
        response = self.client.get(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            )
        )

        self.assertContains(response, "These fields do different jobs.")
        self.assertContains(response, "Usually expected")
        self.assertContains(response, "Classification")
        self.assertContains(response, "Writing aid")
        self.assertContains(response, "Internal")
        self.assertContains(
            response,
            "The final compiled text always comes from the summary paragraph below.",
        )
        self.assertContains(response, "Custom paragraph mode is active.")
        self.assertContains(
            response,
            "The summary paragraph is currently the source of truth for compilation and export.",
        )
        self.assertContains(response, "Custom paragraph in use")
        self.assertContains(response, "Save custom paragraph")
        self.assertContains(response, "Switch back to auto-generated")
        self.assertContains(response, "Clear saved custom paragraph")
        self.assertContains(response, "Use these tags to organise, filter and group summaries across the synopsis.")
        self.assertContains(response, "Stored separately for internal use. These scores are not inserted into the generated summary paragraph.")
        self.assertContains(response, "Outcome notes")
        self.assertContains(response, "Main findings summary")
        self.assertContains(response, "More optional detail boxes")

    def test_creating_summary_tab_invalidates_board_presence_summary_id_cache(self):
        self.client.login(username="author", password="pass123")
        initial_ids = _project_reference_summary_ids(self.project.id)

        response = self.client.post(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            ),
            {"action": "create-summary-tab"},
            follow=False,
        )

        self.assertEqual(response.status_code, 302)
        new_summary = (
            ReferenceSummary.objects.filter(project=self.project, reference=self.reference)
            .exclude(pk=self.summary.id)
            .get()
        )
        self.assertNotIn(new_summary.id, initial_ids)
        self.assertIn(new_summary.id, _project_reference_summary_ids(self.project.id))

    def test_save_summary_persists_changes(self):
        self.client.login(username="author", password="pass123")
        url = reverse("synopsis:reference_summary_detail", args=[self.project.id, self.summary.id])
        resp = self.client.post(
            url,
            {
                "action": "save-summary",
                "status": ReferenceSummary.STATUS_DRAFT,
                "habitat_and_sites": "New habitat info",
            },
            follow=True,
        )
        self.summary.refresh_from_db()
        self.assertEqual(self.summary.status, ReferenceSummary.STATUS_DRAFT)
        self.assertEqual(self.summary.habitat_and_sites, "New habitat info")
        messages = list(get_messages(resp.wsgi_request))
        self.assertTrue(any("Summary updated" in str(m) for m in messages))

    def test_save_summary_auto_moves_todo_tab_to_in_progress_when_content_saved(self):
        self.client.login(username="author", password="pass123")
        url = reverse(
            "synopsis:reference_summary_detail", args=[self.project.id, self.summary.id]
        )
        response = self.client.post(
            url,
            {
                "action": "save-summary",
                "status": ReferenceSummary.STATUS_TODO,
                "habitat_and_sites": "New habitat info",
            },
            follow=True,
        )

        self.summary.refresh_from_db()
        self.assertEqual(self.summary.status, ReferenceSummary.STATUS_DRAFT)
        self.assertContains(
            response,
            "Status moved to In progress automatically.",
        )

    def test_save_summary_can_store_selected_project_action(self):
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="Evidence",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="General",
            position=1,
        )
        SynopsisIntervention.objects.create(
            subheading=subheading,
            title="Install nest boxes",
            position=1,
        )

        self.client.login(username="author", password="pass123")
        url = reverse("synopsis:reference_summary_detail", args=[self.project.id, self.summary.id])
        self.client.post(
            url,
            {
                "action": "save-summary",
                "status": ReferenceSummary.STATUS_DRAFT,
                "action_choice": "Install nest boxes",
                "action_custom": "",
            },
            follow=True,
        )

        self.summary.refresh_from_db()
        self.assertEqual(self.summary.action_description, "Install nest boxes")

    def test_summary_status_choices_include_excluded_after_full_text(self):
        labels = dict(ReferenceSummary.STATUS_CHOICES)
        self.assertIn(ReferenceSummary.STATUS_EXCLUDED, labels)
        self.assertEqual(
            labels[ReferenceSummary.STATUS_EXCLUDED],
            "Excluded after full text",
        )

    def test_save_summary_does_not_clear_saved_paragraph_draft(self):
        self.summary.synopsis_draft = "Edited summary paragraph."
        self.summary.use_custom_synopsis_draft = True
        self.summary.save(
            update_fields=["synopsis_draft", "use_custom_synopsis_draft", "updated_at"]
        )

        self.client.login(username="author", password="pass123")
        url = reverse("synopsis:reference_summary_detail", args=[self.project.id, self.summary.id])
        self.client.post(
            url,
            {
                "action": "save-summary",
                "status": ReferenceSummary.STATUS_DRAFT,
                "habitat_and_sites": "Updated habitat info",
            },
            follow=True,
        )

        self.summary.refresh_from_db()
        self.assertEqual(self.summary.synopsis_draft, "Edited summary paragraph.")
        self.assertTrue(self.summary.use_custom_synopsis_draft)

    def test_save_summary_keeps_auto_generated_mode_and_current_paragraph_updates(self):
        self.summary.study_design = "replicated, controlled study"
        self.summary.year_range = "2018-2020"
        self.summary.summary_of_results = "installing nest boxes increased occupancy."
        self.summary.habitat_and_sites = "woodland sites"
        self.summary.country = "UK"
        self.summary.save(
            update_fields=[
                "study_design",
                "year_range",
                "summary_of_results",
                "habitat_and_sites",
                "country",
                "updated_at",
            ]
        )

        self.client.login(username="author", password="pass123")
        url = reverse("synopsis:reference_summary_detail", args=[self.project.id, self.summary.id])
        self.client.post(
            url,
            {
                "action": "save-summary",
                "status": ReferenceSummary.STATUS_DRAFT,
                "study_design": "replicated, controlled study",
                "year_range": "2018-2020",
                "summary_of_results": "installing nest boxes increased occupancy.",
                "habitat_and_sites": "wetland sites",
                "country": "UK",
            },
            follow=True,
        )

        self.summary.refresh_from_db()
        generated_after = _structured_summary_paragraph(self.summary)
        self.assertFalse(self.summary.use_custom_synopsis_draft)
        self.assertEqual(_reference_summary_paragraph(self.summary), generated_after)

    def test_save_summary_paragraph_draft_persists_changes(self):
        self.client.login(username="author", password="pass123")
        url = reverse("synopsis:reference_summary_detail", args=[self.project.id, self.summary.id])
        self.client.post(
            url,
            {
                "action": "save-synopsis-draft",
                "draft_command": "save",
                "synopsis_draft": "A revised summary paragraph written by the author.",
            },
            follow=True,
        )

        self.summary.refresh_from_db()
        self.assertEqual(
            self.summary.synopsis_draft,
            "A revised summary paragraph written by the author.",
        )
        self.assertTrue(self.summary.use_custom_synopsis_draft)

    def test_switching_back_to_auto_generated_clears_custom_paragraph_mode(self):
        self.summary.study_design = "replicated study"
        self.summary.summary_of_results = "occupancy increased."
        self.summary.synopsis_draft = "Custom paragraph text."
        self.summary.use_custom_synopsis_draft = True
        self.summary.save(
            update_fields=[
                "study_design",
                "summary_of_results",
                "synopsis_draft",
                "use_custom_synopsis_draft",
                "updated_at",
            ]
        )

        self.client.login(username="author", password="pass123")
        response = self.client.post(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            ),
            {
                "action": "save-synopsis-draft",
                "draft_command": "use-generated",
            },
            follow=True,
        )

        self.summary.refresh_from_db()
        self.assertFalse(self.summary.use_custom_synopsis_draft)
        self.assertEqual(self.summary.synopsis_draft, "")
        self.assertEqual(
            _reference_summary_paragraph(self.summary),
            _structured_summary_paragraph(self.summary),
        )
        self.assertContains(
            response,
            "Auto-generated paragraph restored.",
        )

    def test_save_summary_paragraph_draft_auto_moves_todo_tab_to_in_progress(self):
        self.client.login(username="author", password="pass123")
        url = reverse(
            "synopsis:reference_summary_detail", args=[self.project.id, self.summary.id]
        )
        response = self.client.post(
            url,
            {
                "action": "save-synopsis-draft",
                "draft_command": "save",
                "synopsis_draft": "A revised summary paragraph written by the author.",
            },
            follow=True,
        )

        self.summary.refresh_from_db()
        self.assertEqual(self.summary.status, ReferenceSummary.STATUS_DRAFT)
        self.assertContains(
            response,
            "Custom paragraph saved and set as the version used for compilation. Status moved to In progress automatically.",
        )

    def test_detail_status_update_requires_reason_for_summary_phase_exclusion(self):
        self.reference.screening_status = "included"
        self.reference.save(update_fields=["screening_status", "updated_at"])
        self.client.login(username="author", password="pass123")

        response = self.client.post(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            ),
            {
                "action": "update-status",
                "status": ReferenceSummary.STATUS_EXCLUDED,
                "needs_help": "",
                "exclusion_reason": "",
            },
            follow=True,
        )

        self.summary.refresh_from_db()
        self.assertEqual(self.summary.status, ReferenceSummary.STATUS_TODO)
        self.assertContains(
            response,
            "Provide a reason before excluding this summary after full-text review.",
        )

    def test_detail_status_exclusion_removes_only_that_summary_from_synopsis(self):
        self.reference.screening_status = "included"
        self.reference.save(update_fields=["screening_status", "updated_at"])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="Evidence",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="General",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="Intervention",
            position=1,
        )
        assignment = SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=self.summary,
            position=1,
        )
        key_message = SynopsisInterventionKeyMessage.objects.create(
            intervention=intervention,
            response_group=SynopsisInterventionKeyMessage.GROUP_POPULATION,
            statement="Supported by this study.",
            position=1,
        )
        key_message.supporting_summaries.set([self.summary])

        self.client.login(username="author", password="pass123")
        response = self.client.post(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            ),
            {
                "action": "update-status",
                "status": ReferenceSummary.STATUS_EXCLUDED,
                "needs_help": "",
                "exclusion_reason": "Full text shows this is not an intervention study.",
            },
            follow=True,
        )

        self.summary.refresh_from_db()
        key_message.refresh_from_db()
        self.assertEqual(self.summary.status, ReferenceSummary.STATUS_EXCLUDED)
        self.assertEqual(
            self.summary.exclusion_reason,
            "Full text shows this is not an intervention study.",
        )
        self.assertFalse(SynopsisAssignment.objects.filter(pk=assignment.id).exists())
        self.assertEqual(key_message.supporting_summaries.count(), 0)
        self.assertContains(
            response,
            "Summary excluded after full-text review.",
        )
        messages = list(get_messages(response.wsgi_request))
        self.assertTrue(
            any("excluded after full-text review" in str(m).lower() for m in messages)
        )

    def test_saved_summary_paragraph_draft_is_used_for_compilation(self):
        self.summary.reference_identifier = "CR1000"
        self.summary.synopsis_draft = "A revised paragraph (CR1000) with edited wording."
        self.summary.use_custom_synopsis_draft = True
        self.summary.save(
            update_fields=[
                "reference_identifier",
                "synopsis_draft",
                "use_custom_synopsis_draft",
                "updated_at",
            ]
        )

        compiled = _reference_summary_paragraph(
            self.summary, reference_identifier_override="2"
        )

        self.assertEqual(
            compiled,
            "A revised paragraph (2) with edited wording.",
        )

    def test_create_summary_tab_adds_second_summary_for_same_reference(self):
        self.summary.assigned_to = self.user
        self.summary.reference_identifier = "manual-ref"
        self.summary.summary_identifier = "manual-summary"
        self.summary.reference_label = "Test reference label"
        self.summary.summary_author = "Existing Author"
        self.summary.citation = "Author (2024)"
        self.summary.save()

        self.client.login(username="author", password="pass123")
        url = reverse("synopsis:reference_summary_detail", args=[self.project.id, self.summary.id])
        resp = self.client.post(
            url,
            {"action": "create-summary-tab"},
            follow=False,
        )

        self.assertEqual(
            ReferenceSummary.objects.filter(
                project=self.project,
                reference=self.reference,
            ).count(),
            2,
        )
        new_summary = (
            ReferenceSummary.objects.filter(
                project=self.project,
                reference=self.reference,
            )
            .exclude(pk=self.summary.id)
            .get()
        )
        self.assertRedirects(
            resp,
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, new_summary.id],
            ),
            fetch_redirect_response=False,
        )
        self.summary.refresh_from_db()
        self.assertEqual(new_summary.assigned_to, self.user)
        self.assertEqual(self.summary.reference_identifier, "manual-ref")
        self.assertEqual(self.summary.summary_identifier, "manual-summary")
        self.assertEqual(new_summary.reference_identifier, "manual-ref")
        self.assertEqual(new_summary.summary_identifier, "manual-ref.a")
        self.assertEqual(new_summary.summary_author, "Existing Author")
        self.assertEqual(new_summary.citation, "Author (2024)")

    def test_duplicate_summary_tab_copies_current_summary_content(self):
        self.summary.assigned_to = self.user
        self.summary.status = ReferenceSummary.STATUS_DONE
        self.summary.reference_identifier = "manual-ref"
        self.summary.summary_identifier = "manual-summary"
        self.summary.reference_label = "Test reference label"
        self.summary.action_description = "Install nest boxes"
        self.summary.study_design = "Replicated study"
        self.summary.summary_of_results = "Occupancy increased."
        self.summary.action_methods = "Installed wooden boxes."
        self.summary.outcome_rows = [{"outcome": "Occupancy", "notes": "Higher"}]
        self.summary.synopsis_draft = "Draft paragraph copied from the first tab."
        self.summary.use_custom_synopsis_draft = True
        self.summary.summary_author = "Existing Author"
        self.summary.keywords = ["boxes", "occupancy"]
        self.summary.action_tags = ["Land/water protection-Area protection"]
        self.summary.research_design = "Replicated; Controlled*"
        self.summary.citation = "Author (2024)"
        self.summary.save()

        self.client.login(username="author", password="pass123")
        url = reverse(
            "synopsis:reference_summary_detail", args=[self.project.id, self.summary.id]
        )
        resp = self.client.post(
            url,
            {"action": "duplicate-summary-tab"},
            follow=False,
        )

        new_summary = (
            ReferenceSummary.objects.filter(
                project=self.project,
                reference=self.reference,
            )
            .exclude(pk=self.summary.id)
            .get()
        )
        self.assertRedirects(
            resp,
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, new_summary.id],
            ),
            fetch_redirect_response=False,
        )
        self.assertEqual(new_summary.assigned_to, self.user)
        self.assertEqual(new_summary.status, ReferenceSummary.STATUS_DRAFT)
        self.assertEqual(new_summary.reference_identifier, "manual-ref")
        self.assertEqual(new_summary.summary_identifier, "manual-ref.a")
        self.assertEqual(new_summary.reference_label, "Test reference label")
        self.assertEqual(new_summary.action_description, "Install nest boxes")
        self.assertEqual(new_summary.study_design, "Replicated study")
        self.assertEqual(new_summary.summary_of_results, "Occupancy increased.")
        self.assertEqual(new_summary.action_methods, "Installed wooden boxes.")
        self.assertEqual(new_summary.outcome_rows, [{"outcome": "Occupancy", "notes": "Higher"}])
        self.assertEqual(
            new_summary.synopsis_draft,
            "Draft paragraph copied from the first tab.",
        )
        self.assertTrue(new_summary.use_custom_synopsis_draft)
        self.assertEqual(new_summary.summary_author, "Existing Author")
        self.assertEqual(new_summary.keywords, ["boxes", "occupancy"])
        self.assertEqual(
            new_summary.action_tags,
            ["Land/water protection-Area protection"],
        )
        self.assertEqual(new_summary.research_design, "Replicated; Controlled*")
        self.assertEqual(new_summary.citation, "Author (2024)")

    def test_duplicate_summary_tab_does_not_copy_comments_assignments_or_exclusion_state(self):
        self.summary.status = ReferenceSummary.STATUS_EXCLUDED
        self.summary.needs_help = True
        self.summary.exclusion_reason = "Not really an intervention."
        self.summary.save(update_fields=["status", "needs_help", "exclusion_reason", "updated_at"])
        ReferenceSummaryComment.objects.create(
            summary=self.summary,
            author=self.user,
            body="Keep this note on the original tab only.",
        )
        ReferenceActionSummary.objects.create(
            reference_summary=self.summary,
            action_name="Install nest boxes",
            summary_text="Action-specific wording.",
            created_by=self.user,
        )
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="Evidence",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="General",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="Intervention",
            position=1,
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=self.summary,
            position=1,
        )

        self.client.login(username="author", password="pass123")
        response = self.client.post(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            ),
            {"action": "duplicate-summary-tab"},
            follow=True,
        )

        duplicated = (
            ReferenceSummary.objects.filter(project=self.project, reference=self.reference)
            .exclude(pk=self.summary.id)
            .get()
        )
        self.assertEqual(duplicated.status, ReferenceSummary.STATUS_DRAFT)
        self.assertFalse(duplicated.needs_help)
        self.assertEqual(duplicated.exclusion_reason, "")
        self.assertEqual(duplicated.comments.count(), 0)
        self.assertEqual(duplicated.action_summaries.count(), 0)
        self.assertEqual(duplicated.synopsis_assignments.count(), 0)
        self.assertContains(
            response,
            "Summary tab duplicated. Review the copied text and adjust it for the new intervention or study summary.",
        )

    def test_delete_summary_tab_removes_extra_tab_and_redirects_to_remaining_tab(self):
        extra_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=self.reference,
            status=ReferenceSummary.STATUS_DRAFT,
        )

        self.client.login(username="author", password="pass123")
        response = self.client.post(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, extra_summary.id],
            ),
            {"action": "delete-summary-tab"},
            follow=False,
        )

        self.assertRedirects(
            response,
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            ),
            fetch_redirect_response=False,
        )
        self.assertFalse(
            ReferenceSummary.objects.filter(pk=extra_summary.id).exists()
        )
        self.summary.refresh_from_db()
        self.assertEqual(self.summary.summary_identifier, "CR1000.a")

    def test_delete_summary_tab_resequences_intervention_assignments(self):
        extra_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=self.reference,
            status=ReferenceSummary.STATUS_DRAFT,
        )
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="Evidence",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="General",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="Intervention",
            position=1,
        )
        retained_assignment = SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=self.summary,
            position=1,
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=extra_summary,
            position=2,
        )

        self.client.login(username="author", password="pass123")
        response = self.client.post(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, extra_summary.id],
            ),
            {"action": "delete-summary-tab"},
            follow=False,
        )

        self.assertRedirects(
            response,
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            ),
            fetch_redirect_response=False,
        )
        retained_assignment.refresh_from_db()
        self.assertEqual(retained_assignment.position, 1)

    def test_delete_summary_tab_is_blocked_for_only_summary(self):
        self.client.login(username="author", password="pass123")
        response = self.client.post(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            ),
            {"action": "delete-summary-tab"},
            follow=True,
        )

        self.assertEqual(
            ReferenceSummary.objects.filter(
                project=self.project,
                reference=self.reference,
            ).count(),
            1,
        )
        messages = list(get_messages(response.wsgi_request))
        self.assertTrue(
            any("only summary tab" in str(message) for message in messages)
        )

    def test_detail_page_shows_generated_identifiers_and_reference_title(self):
        self.client.login(username="author", password="pass123")
        response = self.client.get(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            )
        )

        self.assertContains(response, 'value="CR1000"')
        self.assertContains(response, 'value="CR1000.a"')
        self.assertContains(response, 'value="Test reference"')

    def test_single_summary_tab_defaults_to_generated_summary_identifier_label(self):
        self.client.login(username="author", password="pass123")
        response = self.client.get(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            )
        )

        self.assertEqual(response.status_code, 200)
        tabs = response.context["summary_tabs"]
        self.assertEqual(len(tabs), 1)
        self.assertEqual(tabs[0]["label"], "CR1000.a")

    def test_second_reference_in_project_gets_next_generated_reference_id(self):
        second_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            hash_key="b" * 40,
            title="Second reference",
        )
        second_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=second_reference,
            status=ReferenceSummary.STATUS_TODO,
        )

        self.client.login(username="author", password="pass123")
        response = self.client.get(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, second_summary.id],
            )
        )

        second_summary.refresh_from_db()
        self.assertEqual(second_summary.reference_identifier, "CR1001")
        self.assertEqual(second_summary.summary_identifier, "CR1001.a")
        self.assertContains(response, 'value="CR1001"')
        self.assertContains(response, 'value="CR1001.a"')

    def test_project_title_change_does_not_rewrite_existing_identifiers(self):
        self.client.login(username="author", password="pass123")
        detail_url = reverse(
            "synopsis:reference_summary_detail",
            args=[self.project.id, self.summary.id],
        )

        self.client.get(detail_url)
        self.summary.refresh_from_db()
        self.assertEqual(self.summary.reference_identifier, "CR1000")
        self.assertEqual(self.summary.summary_identifier, "CR1000.a")

        self.project.title = "Marine Restoration Handbook"
        self.project.save(update_fields=["title"])

        self.client.get(detail_url)
        self.summary.refresh_from_db()
        self.assertEqual(self.summary.reference_identifier, "CR1000")
        self.assertEqual(self.summary.summary_identifier, "CR1000.a")

    def test_deleting_earlier_reference_does_not_rewrite_later_reference_identifier(self):
        second_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            hash_key="b" * 40,
            title="Second reference",
        )
        second_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=second_reference,
            status=ReferenceSummary.STATUS_TODO,
        )

        self.client.login(username="author", password="pass123")
        second_detail_url = reverse(
            "synopsis:reference_summary_detail",
            args=[self.project.id, second_summary.id],
        )

        self.client.get(second_detail_url)
        second_summary.refresh_from_db()
        self.assertEqual(second_summary.reference_identifier, "CR1001")
        self.assertEqual(second_summary.summary_identifier, "CR1001.a")

        self.reference.delete()

        self.client.get(second_detail_url)
        second_summary.refresh_from_db()
        self.assertEqual(second_summary.reference_identifier, "CR1001")
        self.assertEqual(second_summary.summary_identifier, "CR1001.a")

    def test_new_reference_after_deleting_earlier_reference_gets_next_unused_id(self):
        second_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            hash_key="b" * 40,
            title="Second reference",
        )
        second_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=second_reference,
            status=ReferenceSummary.STATUS_TODO,
        )

        self.client.login(username="author", password="pass123")
        self.client.get(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, second_summary.id],
            )
        )
        second_summary.refresh_from_db()
        self.assertEqual(second_summary.reference_identifier, "CR1001")

        self.reference.delete()

        third_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            hash_key="c" * 40,
            title="Third reference",
        )
        third_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=third_reference,
            status=ReferenceSummary.STATUS_TODO,
        )

        response = self.client.get(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, third_summary.id],
            )
        )

        third_summary.refresh_from_db()
        self.assertEqual(third_summary.reference_identifier, "CR1002")
        self.assertEqual(third_summary.summary_identifier, "CR1002.a")
        self.assertContains(response, 'value="CR1002"')
        self.assertContains(response, 'value="CR1002.a"')

    def test_new_reference_after_project_rename_keeps_established_prefix(self):
        self.client.login(username="author", password="pass123")
        self.client.get(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            )
        )
        self.summary.refresh_from_db()
        self.assertEqual(self.summary.reference_identifier, "CR1000")

        self.project.title = "Marine Restoration Handbook"
        self.project.save(update_fields=["title"])

        second_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            hash_key="b" * 40,
            title="Second reference",
        )
        second_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=second_reference,
            status=ReferenceSummary.STATUS_TODO,
        )

        response = self.client.get(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, second_summary.id],
            )
        )

        second_summary.refresh_from_db()
        self.assertEqual(second_summary.reference_identifier, "CR1001")
        self.assertEqual(second_summary.summary_identifier, "CR1001.a")
        self.assertContains(response, 'value="CR1001"')
        self.assertContains(response, 'value="CR1001.a"')

    def test_board_still_creates_only_one_default_summary_per_included_reference(self):
        self.reference.screening_status = "included"
        self.reference.save(update_fields=["screening_status"])
        self.client.login(username="author", password="pass123")
        extra_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=self.reference,
            citation="Alt citation",
        )

        resp = self.client.get(
            reverse("synopsis:reference_summary_board", args=[self.project.id])
        )

        self.assertEqual(resp.status_code, 200)
        self.assertEqual(
            ReferenceSummary.objects.filter(
                project=self.project,
                reference=self.reference,
            ).count(),
            2,
        )
        self.assertContains(resp, "of 2 summary tabs for this reference", status_code=200)
        self.assertContains(resp, "CR1000.b")

    def test_board_and_detail_use_library_reference_metadata(self):
        canonical = LibraryReference.objects.create(
            title="Canonical library title",
            authors="Alhas, Ibrahim",
            publication_year=2024,
        )
        self.reference.library_reference = canonical
        self.reference.title = "Project-local title"
        self.reference.screening_status = "included"
        self.reference.save(
            update_fields=["library_reference", "title", "screening_status", "updated_at"]
        )

        self.client.login(username="author", password="pass123")

        board_response = self.client.get(
            reverse("synopsis:reference_summary_board", args=[self.project.id])
        )
        detail_response = self.client.get(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            )
        )

        self.assertContains(board_response, "Canonical library title")
        self.assertContains(board_response, "Alhas, Ibrahim")
        self.assertContains(detail_response, "Canonical library title")
        self.assertContains(detail_response, "Alhas, Ibrahim")

    def test_summary_detail_can_update_reference_classification(self):
        self.reference.screening_status = "included"
        self.reference.save(update_fields=["screening_status", "updated_at"])
        self.client.login(username="author", password="pass123")

        response = self.client.post(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            ),
            {
                "action": "update-classification",
                "classification_command": "save",
                "screening_status": "included",
                "reference_folder": ["3a"],
                "screening_notes": "Freshwater fish evidence.",
            },
            follow=True,
        )

        self.reference.refresh_from_db()
        self.assertEqual(self.reference.screening_status, "included")
        self.assertEqual(self.reference.unlinked_reference_folder, ["3a"])
        self.assertEqual(self.reference.category_values, ["3a"])
        self.assertEqual(self.reference.screening_notes, "Freshwater fish evidence.")
        self.assertContains(response, "Reference classification updated.")

    def test_summary_detail_folder_update_updates_shared_library_reference(self):
        canonical = LibraryReference.objects.create(
            title="Canonical library title",
            authors="Alhas, Ibrahim",
            publication_year=2024,
            hash_key="summary-shared-sync",
            reference_folder=["2"],
        )
        self.reference.library_reference = canonical
        self.reference.hash_key = "summary-shared-sync"
        self.reference.screening_status = "included"
        self.reference.save(
            update_fields=[
                "library_reference",
                "hash_key",
                "screening_status",
                "updated_at",
            ]
        )
        other_project = Project.objects.create(title="Other project")
        other_batch = ReferenceSourceBatch.objects.create(
            project=other_project,
            label="Other batch",
            source_type="library_link",
        )
        other_reference = Reference.objects.create(
            project=other_project,
            batch=other_batch,
            library_reference=canonical,
            hash_key="summary-shared-sync",
            title="Other project copy",
        )

        self.client.login(username="author", password="pass123")
        response = self.client.post(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            ),
            {
                "action": "update-classification",
                "classification_command": "save",
                "screening_status": "included",
                "reference_folder": ["15", "2"],
                "screening_notes": "Freshwater fish evidence.",
            },
            follow=True,
        )

        canonical.refresh_from_db()
        self.reference.refresh_from_db()
        other_reference.refresh_from_db()
        self.assertEqual(canonical.reference_folder, ["2", "15"])
        self.assertEqual(self.reference.unlinked_reference_folder, [])
        self.assertEqual(self.reference.category_values, ["2", "15"])
        self.assertEqual(other_reference.unlinked_reference_folder, [])
        self.assertEqual(other_reference.category_values, ["2", "15"])
        self.assertContains(
            response,
            "Shared CE subject categories were updated for all linked synopsis copies.",
        )
        self.assertTrue(
            LibraryReferenceFolderHistory.objects.filter(
                library_reference=canonical,
                new_folders=["2", "15"],
                source_project=self.project,
                source_reference=self.reference,
                change_source="summary_reference_management",
            ).exists()
        )

    def test_summary_detail_uses_shared_categories_when_local_fallback_exists(self):
        canonical = LibraryReference.objects.create(
            title="Canonical library title",
            authors="Alhas, Ibrahim",
            publication_year=2024,
            hash_key="summary-shared-read",
            reference_folder=["2", "15"],
        )
        self.reference.library_reference = canonical
        self.reference.hash_key = "summary-shared-read"
        self.reference.screening_status = "included"
        self.reference.unlinked_reference_folder = ["3a"]
        self.reference.save(
            update_fields=[
                "library_reference",
                "hash_key",
                "screening_status",
                "unlinked_reference_folder",
                "updated_at",
            ]
        )

        self.client.login(username="author", password="pass123")
        response = self.client.get(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            )
        )

        classification_form = response.context["classification_form"]
        self.assertEqual(
            classification_form.initial["reference_folder"],
            ["2", "15"],
        )
        self.assertContains(response, 'option value="2" selected')
        self.assertContains(response, 'option value="15" selected')

    def test_summary_detail_filters_blank_reference_folder_values(self):
        self.reference.screening_status = "included"
        self.reference.save(update_fields=["screening_status", "updated_at"])
        self.client.login(username="author", password="pass123")

        response = self.client.post(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            ),
            {
                "action": "update-classification",
                "classification_command": "save",
                "screening_status": "included",
                "reference_folder": ["", "3a"],
                "screening_notes": "Freshwater fish evidence.",
            },
            follow=True,
        )

        self.reference.refresh_from_db()
        self.assertEqual(self.reference.unlinked_reference_folder, ["3a"])
        self.assertEqual(self.reference.category_values, ["3a"])
        self.assertContains(response, "Reference classification updated.")

    def test_excluding_reference_requires_reason_on_summary_page(self):
        self.reference.screening_status = "included"
        self.reference.save(update_fields=["screening_status", "updated_at"])
        self.client.login(username="author", password="pass123")

        response = self.client.post(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            ),
            {
                "action": "update-classification",
                "classification_command": "exclude",
                "screening_status": "excluded",
                "reference_folder": [],
                "screening_notes": "",
            },
        )

        self.reference.refresh_from_db()
        self.assertEqual(self.reference.screening_status, "included")
        self.assertContains(
            response,
            "Provide a reason before excluding this reference from the synopsis.",
        )

    def test_excluding_reference_from_summary_removes_synopsis_assignments(self):
        self.reference.screening_status = "included"
        self.reference.save(update_fields=["screening_status", "updated_at"])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="Evidence",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="General",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="Intervention",
            position=1,
        )
        assignment = SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=self.summary,
            position=1,
        )
        key_message = SynopsisInterventionKeyMessage.objects.create(
            intervention=intervention,
            response_group=SynopsisInterventionKeyMessage.GROUP_POPULATION,
            statement="Supported by this study.",
            position=1,
        )
        key_message.supporting_summaries.set([self.summary])

        self.client.login(username="author", password="pass123")
        response = self.client.post(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            ),
            {
                "action": "update-classification",
                "classification_command": "exclude",
                "screening_status": "excluded",
                "reference_folder": ["3a"],
                "screening_notes": "Not relevant to this synopsis.",
            },
            follow=False,
        )

        self.reference.refresh_from_db()
        key_message.refresh_from_db()
        self.assertEqual(self.reference.screening_status, "excluded")
        self.assertEqual(self.reference.screening_notes, "Not relevant to this synopsis.")
        self.assertFalse(SynopsisAssignment.objects.filter(pk=assignment.id).exists())
        self.assertEqual(key_message.supporting_summaries.count(), 0)
        self.assertRedirects(
            response,
            f"{reverse('synopsis:reference_summary_detail', args=[self.project.id, self.summary.id])}?panel=management",
            fetch_redirect_response=False,
        )

    def test_summary_detail_reference_management_panel_reopens_after_classification_update(self):
        self.reference.screening_status = "included"
        self.reference.save(update_fields=["screening_status", "updated_at"])
        self.client.login(username="author", password="pass123")

        response = self.client.post(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            ),
            {
                "action": "update-classification",
                "classification_command": "exclude",
                "screening_status": "excluded",
                "reference_folder": ["3a"],
                "screening_notes": "Not relevant to this synopsis.",
            },
            follow=True,
        )

        self.assertContains(response, "Re-include this reference")
        self.assertTrue(response.context["open_management_panel"])
        self.assertContains(response, "Reference excluded from this synopsis.")

    def test_reference_management_explains_difference_between_summary_and_reference_exclusion(self):
        self.reference.screening_status = "included"
        self.reference.save(update_fields=["screening_status", "updated_at"])
        self.summary.status = ReferenceSummary.STATUS_EXCLUDED
        self.summary.exclusion_reason = "Full text exclusion reason."
        self.summary.save(update_fields=["status", "exclusion_reason", "updated_at"])
        self.client.login(username="author", password="pass123")

        response = self.client.get(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            )
        )

        self.assertContains(
            response,
            "All summary tabs for this reference are excluded after full-text review, but the whole reference is still marked as included for this synopsis.",
        )
        self.assertContains(response, "Exclude whole reference from synopsis too")
        self.assertContains(
            response,
            "Shared CE subject categories are stored on the reference, not on this individual summary tab.",
        )
        self.assertContains(
            response,
            "changing them here updates the shared reference record and is reflected everywhere it is linked",
        )

    def test_summary_detail_renders_restore_state_hooks(self):
        self.client.login(username="author", password="pass123")

        response = self.client.get(
            reverse(
                "synopsis:reference_summary_detail",
                args=[self.project.id, self.summary.id],
            )
        )

        self.assertContains(response, 'id="reference-summary-page"', html=False)
        self.assertContains(response, "cePreservePageState({", html=False)
        self.assertContains(
            response,
            f'"reference-summary-state-{self.project.id}-{self.summary.id}"',
            html=False,
        )
        self.assertContains(response, "managementPanelOpen", html=False)
        self.assertContains(response, "uploadForm.requestSubmit()", html=False)

    def test_board_context_workload_counts_are_aggregated_correctly(self):
        other_author = User.objects.create_user(
            username="coauthor",
            password="pass123",
            first_name="Co",
            last_name="Author",
        )
        UserRole.objects.create(user=other_author, project=self.project, role="author")

        self.reference.screening_status = "included"
        self.reference.save(update_fields=["screening_status"])
        second_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            hash_key="b" * 40,
            title="Second reference",
            screening_status="included",
        )
        third_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            hash_key="c" * 40,
            title="Third reference",
            screening_status="included",
        )
        second_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=second_reference,
            assigned_to=self.user,
            needs_help=True,
            status=ReferenceSummary.STATUS_DONE,
        )
        third_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=third_reference,
            assigned_to=other_author,
            needs_help=False,
            status=ReferenceSummary.STATUS_DRAFT,
        )
        self.summary.assigned_to = self.user
        self.summary.needs_help = False
        self.summary.status = ReferenceSummary.STATUS_DONE
        self.summary.save(
            update_fields=["assigned_to", "needs_help", "status", "updated_at"]
        )

        self.client.login(username="author", password="pass123")
        response = self.client.get(
            reverse("synopsis:reference_summary_board", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        workload = {
            row["author"].id: {
                "assigned": row["assigned"],
                "summarised": row["summarised"],
                "summarised_percent": row["summarised_percent"],
                "needs_help": row["needs_help"],
                "excluded_after_full_text": row["excluded_after_full_text"],
            }
            for row in response.context["workload"]
        }
        self.assertEqual(
            workload[self.user.id],
            {
                "assigned": 2,
                "summarised": 2,
                "summarised_percent": 100,
                "needs_help": 1,
                "excluded_after_full_text": 0,
            },
        )
        self.assertEqual(
            workload[other_author.id],
            {
                "assigned": 1,
                "summarised": 0,
                "summarised_percent": 0,
                "needs_help": 0,
                "excluded_after_full_text": 0,
            },
        )
        self.assertEqual(response.context["unassigned_count"], 0)
        self.assertEqual(response.context["needs_help_count"], 1)

    def test_summary_board_shows_active_editor_badge_for_active_summary(self):
        self.reference.screening_status = "included"
        self.reference.save(update_fields=["screening_status"])

        other_author = User.objects.create_user(
            username="coauthor",
            password="pass123",
            first_name="Co",
            last_name="Author",
        )
        UserRole.objects.create(user=other_author, project=self.project, role="author")

        self.client.login(username="coauthor", password="pass123")
        self.client.post(
            reverse(
                "synopsis:reference_summary_presence",
                args=[self.project.id, self.summary.id],
            )
        )

        self.client.login(username="author", password="pass123")
        response = self.client.get(
            reverse("synopsis:reference_summary_board", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Active now")
        self.assertContains(response, "Co Author")
        self.assertContains(
            response,
            reverse("synopsis:reference_summary_board_presence", args=[self.project.id]),
            html=False,
        )

    def test_summary_board_shows_excluded_column_reason_and_progress_ignores_excluded(self):
        self.reference.screening_status = "included"
        self.reference.save(update_fields=["screening_status"])
        second_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            hash_key="d" * 40,
            title="Excluded summary reference",
            screening_status="included",
        )
        second_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=second_reference,
            assigned_to=self.user,
            status=ReferenceSummary.STATUS_EXCLUDED,
            exclusion_reason="Full text did not test a conservation intervention.",
        )
        self.summary.status = ReferenceSummary.STATUS_DONE
        self.summary.save(update_fields=["status", "updated_at"])

        self.client.login(username="author", password="pass123")
        response = self.client.get(
            reverse("synopsis:reference_summary_board", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Excluded after full text")
        self.assertContains(response, "Jump to excluded after full text")
        self.assertContains(response, 'id="summary-column-excluded"', html=False)
        self.assertContains(
            response,
            "Full text did not test a conservation intervention.",
        )
        self.assertContains(response, "1 excluded after full text")
        self.assertEqual(response.context["excluded_after_full_text_count"], 1)
        self.assertEqual(response.context["summary_count"], 1)
        self.assertEqual(response.context["completed"], 1)
        workload = {row["author"].id: row for row in response.context["workload"]}
        self.assertEqual(workload[self.user.id]["excluded_after_full_text"], 1)


class GlobalReferenceLibraryAccessTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="authorlib", password="pass123")
        self.project = Project.objects.create(title="Coral Project")
        self.other_project = Project.objects.create(title="Unassigned Project")
        UserRole.objects.create(user=self.user, project=self.project, role="author")

    @override_settings(APP_RELEASE_LABEL="pilot-2026-03-29")
    def test_author_sees_global_library_entry_points(self):
        self.client.login(username="authorlib", password="pass123")

        dashboard_response = self.client.get(reverse("synopsis:dashboard"))
        project_response = self.client.get(
            reverse("synopsis:project_hub", args=[self.project.id])
        )

        self.assertContains(dashboard_response, "Deployed version")
        self.assertContains(dashboard_response, "pilot-2026-03-29")
        self.assertContains(dashboard_response, "Shared Reference Library")
        self.assertContains(dashboard_response, "Create New Synopsis")
        self.assertNotContains(dashboard_response, "How this works for authors")
        self.assertNotContains(
            dashboard_response,
            "This sits above individual synopses and can be used to link references into project batches.",
        )
        self.assertContains(dashboard_response, "Coral Project")
        self.assertContains(dashboard_response, "Unassigned Project")
        self.assertNotContains(project_response, "Create New Synopsis")
        self.assertContains(project_response, "Browse Shared Reference Library")
        self.assertContains(
            project_response,
            reverse("synopsis:reference_library") + f"?project={self.project.id}",
            html=False,
        )

    def test_author_can_open_unassigned_synopsis(self):
        self.client.login(username="authorlib", password="pass123")

        response = self.client.get(
            reverse("synopsis:project_hub", args=[self.other_project.id]),
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Unassigned Project")

    def test_library_pages_show_global_workflow_help(self):
        self.client.login(username="authorlib", password="pass123")
        batch = LibraryImportBatch.objects.create(
            label="Library import",
            source_type="journal_search",
            uploaded_by=self.user,
        )

        library_response = self.client.get(reverse("synopsis:reference_library"))
        batch_list_response = self.client.get(reverse("synopsis:library_batch_list"))
        batch_detail_response = self.client.get(
            reverse("synopsis:library_batch_detail", args=[batch.id])
        )

        for response in (library_response, batch_list_response, batch_detail_response):
            self.assertEqual(response.status_code, 200)
            self.assertContains(response, "How this works")
            self.assertContains(response, "How the shared reference library works")
            self.assertContains(response, "shared library of references")
            self.assertContains(response, "It is not yet a full EndNote replacement")


class ProjectReferenceWorkflowHelpUiTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="refauthor", password="pass123")
        self.project = Project.objects.create(title="Reference Help Project")
        UserRole.objects.create(user=self.user, project=self.project, role="author")
        self.batch = ReferenceSourceBatch.objects.create(
            project=self.project,
            label="Search batch",
            source_type="journal_search",
            uploaded_by=self.user,
        )
        self.reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            hash_key="ref-help-1",
            title="Shared library behaviour",
            screening_status="included",
        )
        self.summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=self.reference,
            citation="Shared library behaviour",
        )

    def test_project_reference_pages_show_workflow_help(self):
        self.client.login(username="refauthor", password="pass123")

        responses = [
            self.client.get(reverse("synopsis:reference_batch_list", args=[self.project.id])),
            self.client.get(
                reverse(
                    "synopsis:reference_batch_detail",
                    args=[self.project.id, self.batch.id],
                )
            ),
            self.client.get(
                reverse("synopsis:reference_summary_board", args=[self.project.id])
            ),
        ]

        for response in responses:
            self.assertEqual(response.status_code, 200)
            self.assertContains(response, "How this works")
            self.assertContains(response, "How project references work")
            self.assertContains(response, "Two ways to add references")
            self.assertContains(
                response,
                "It does not copy the current team workflow exactly",
            )


class ExternalAuthorAccessTests(TestCase):
    def setUp(self):
        ensure_global_groups()
        self.media_dir = tempfile.mkdtemp()
        self.addCleanup(lambda: shutil.rmtree(self.media_dir, ignore_errors=True))
        override = override_settings(MEDIA_ROOT=self.media_dir)
        override.enable()
        self.addCleanup(override.disable)
        self.user = User.objects.create_user(
            username="external@example.com",
            email="external@example.com",
            password="pass123",
        )
        self.user.groups.add(Group.objects.get(name="external_collaborator"))
        self.assigned_project = Project.objects.create(title="Assigned Synopsis")
        self.unassigned_project = Project.objects.create(title="Hidden Synopsis")
        UserRole.objects.create(
            user=self.user, project=self.assigned_project, role="author"
        )

    def _docx_upload(self, name, content):
        return SimpleUploadedFile(
            name,
            content,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def test_external_author_dashboard_only_shows_assigned_synopses(self):
        self.client.login(username="external@example.com", password="pass123")

        response = self.client.get(reverse("synopsis:dashboard"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Assigned Synopsis")
        self.assertNotContains(response, "Hidden Synopsis")
        self.assertNotContains(response, "Open Shared Reference Library")
        self.assertNotContains(response, "Shared Reference Library")
        self.assertNotContains(response, "Create New Synopsis")

    def test_external_author_cannot_create_synopsis_or_open_reference_library(self):
        self.client.login(username="external@example.com", password="pass123")

        create_response = self.client.get(reverse("synopsis:project_create"), follow=True)
        library_response = self.client.get(reverse("synopsis:reference_library"))

        self.assertRedirects(create_response, reverse("synopsis:dashboard"))
        self.assertContains(
            create_response,
            "External author accounts cannot create new synopses.",
        )
        self.assertEqual(library_response.status_code, 403)

    def test_external_author_can_open_assigned_synopsis_only(self):
        self.client.login(username="external@example.com", password="pass123")

        assigned_response = self.client.get(
            reverse("synopsis:project_hub", args=[self.assigned_project.id])
        )
        unassigned_response = self.client.get(
            reverse("synopsis:project_hub", args=[self.unassigned_project.id]),
            follow=True,
        )

        self.assertEqual(assigned_response.status_code, 200)
        self.assertNotContains(assigned_response, "Browse Shared Reference Library")
        self.assertNotContains(assigned_response, "Project settings")
        self.assertNotContains(assigned_response, "Manage phase tracker")
        self.assertNotContains(assigned_response, "Move to ")
        self.assertRedirects(unassigned_response, reverse("synopsis:dashboard"))
        self.assertContains(
            unassigned_response,
            "You do not have access to that synopsis.",
        )

    def test_external_author_cannot_open_project_settings(self):
        self.client.login(username="external@example.com", password="pass123")

        response = self.client.get(
            reverse("synopsis:project_settings", args=[self.assigned_project.id]),
            follow=True,
        )

        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.assigned_project.id])
        )
        self.assertContains(
            response,
            "You do not have permission to update project settings for this synopsis.",
        )

    def test_external_author_project_reference_page_hides_library_buttons(self):
        self.client.login(username="external@example.com", password="pass123")

        response = self.client.get(
            reverse("synopsis:reference_batch_list", args=[self.assigned_project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Import RIS")
        self.assertNotContains(response, "Link from library")
        self.assertNotContains(response, "Browse library")

    def test_external_author_cannot_delete_protocol_or_action_list_documents(self):
        protocol = Protocol.objects.create(
            project=self.assigned_project,
            document=self._docx_upload("protocol.docx", b"protocol"),
        )
        action_list = ActionList.objects.create(
            project=self.assigned_project,
            document=self._docx_upload("action-list.docx", b"action-list"),
        )
        self.client.login(username="external@example.com", password="pass123")

        protocol_page = self.client.get(
            reverse("synopsis:protocol_detail", args=[self.assigned_project.id])
        )
        action_list_page = self.client.get(
            reverse("synopsis:action_list_detail", args=[self.assigned_project.id])
        )
        protocol_delete_response = self.client.post(
            reverse("synopsis:protocol_delete_file", args=[self.assigned_project.id]),
            follow=True,
        )
        action_delete_response = self.client.post(
            reverse(
                "synopsis:action_list_delete_file", args=[self.assigned_project.id]
            ),
            follow=True,
        )

        self.assertNotContains(protocol_page, "Danger zone")
        self.assertNotContains(action_list_page, "Danger zone")
        self.assertRedirects(
            protocol_delete_response,
            reverse("synopsis:protocol_detail", args=[self.assigned_project.id]),
        )
        self.assertContains(
            protocol_delete_response,
            "You do not have permission to delete protocol files for this synopsis.",
        )
        self.assertRedirects(
            action_delete_response,
            reverse("synopsis:action_list_detail", args=[self.assigned_project.id]),
        )
        self.assertContains(
            action_delete_response,
            "You do not have permission to delete action list files for this synopsis.",
        )
        protocol.refresh_from_db()
        action_list.refresh_from_db()
        self.assertTrue(protocol.document)
        self.assertTrue(action_list.document)

    def test_external_author_cannot_mark_completed_or_reactivate_completed_synopsis(self):
        self.assigned_project.status = "completed"
        self.assigned_project.save(update_fields=["status"])
        self.client.login(username="external@example.com", password="pass123")

        dashboard_response = self.client.get(reverse("synopsis:dashboard"))
        direct_post_response = self.client.post(
            reverse("synopsis:project_settings", args=[self.assigned_project.id]),
            {"status_action": "reactivate", "return_to": "dashboard"},
            follow=True,
        )

        self.assertContains(dashboard_response, "Assigned Synopsis")
        self.assertNotContains(dashboard_response, "Move to active")
        self.assertRedirects(
            direct_post_response,
            reverse("synopsis:project_hub", args=[self.assigned_project.id]),
        )
        self.assertContains(
            direct_post_response,
            "You do not have permission to update project settings for this synopsis.",
        )
        self.assigned_project.refresh_from_db()
        self.assertEqual(self.assigned_project.status, "completed")


class ProjectAuthorSelectionUiTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="creator", password="pass123")
        self.other_user = User.objects.create_user(
            username="ibrahim",
            password="pass123",
            first_name="Ibrahim",
            last_name="Alhas",
        )
        self.third_user = User.objects.create_user(
            username="will",
            password="pass123",
            first_name="Will",
            last_name="Morgan",
        )

    def test_project_create_uses_readable_author_picker(self):
        self.client.login(username="creator", password="pass123")

        response = self.client.get(reverse("synopsis:project_create"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Description (optional)")
        self.assertContains(response, "Filter authors by name or username")
        self.assertNotContains(response, "Ctrl/Cmd multi-select", html=False)
        self.assertContains(response, "Ibrahim Alhas (ibrahim)")
        self.assertContains(response, "Will Morgan (will)")


class ProjectDescriptionUiTests(TestCase):
    def setUp(self):
        self.manager = User.objects.create_user(
            username="manager",
            password="pass123",
            is_staff=True,
        )
        self.project = Project.objects.create(
            title="Forest Restoration",
            description="A pilot synopsis for forest restoration.",
        )

    def test_project_hub_shows_description_when_present(self):
        self.client.login(username="manager", password="pass123")

        response = self.client.get(reverse("synopsis:project_hub", args=[self.project.id]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Synopsis overview")
        self.assertContains(response, "Synopsis settings")
        self.assertContains(response, "Description")
        self.assertContains(response, "A pilot synopsis for forest restoration.")

    def test_project_settings_shows_description_field(self):
        self.client.login(username="manager", password="pass123")

        response = self.client.get(
            reverse("synopsis:project_settings", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "optional description")
        self.assertContains(response, "Phase tracker")
        self.assertContains(response, "A pilot synopsis for forest restoration.")
        self.assertContains(response, 'value="Forest Restoration"', html=False)

    def test_project_settings_shows_protocol_and_advisory_relevance_fields(self):
        self.client.login(username="manager", password="pass123")

        response = self.client.get(
            reverse("synopsis:project_settings", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Protocol is relevant for this project")
        self.assertContains(response, "Advisory board is relevant for this project")

    def test_project_pages_show_back_to_project_button_in_nav(self):
        self.client.login(username="manager", password="pass123")

        response = self.client.get(
            reverse("synopsis:project_authors_manage", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, ">Back to project</a>", html=False)


class NavbarIdentityUiTests(TestCase):
    def setUp(self):
        ensure_global_groups()
        self.project = Project.objects.create(title="Navbar Synopsis")
        self.manager = User.objects.create_user(
            username="nav-manager",
            password="pass123",
            first_name="Mina",
            last_name="Manager",
            is_staff=True,
        )
        self.author = User.objects.create_user(
            username="nav-author",
            password="pass123",
        )
        self.external = User.objects.create_user(
            username="nav-external@example.com",
            email="nav-external@example.com",
            password="pass123",
            first_name="Eli",
            last_name="External",
        )
        self.external.groups.add(Group.objects.get(name="external_collaborator"))
        UserRole.objects.create(user=self.author, project=self.project, role="author")
        UserRole.objects.create(user=self.external, project=self.project, role="author")

    def test_manager_nav_shows_signed_in_name_and_manager_role(self):
        self.client.login(username="nav-manager", password="pass123")

        response = self.client.get(reverse("synopsis:dashboard"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="nav-user-summary"', html=False)
        self.assertContains(response, "Mina Manager")
        self.assertContains(response, "Manager")

    def test_project_author_nav_uses_project_role_label(self):
        self.client.login(username="nav-author", password="pass123")

        response = self.client.get(
            reverse("synopsis:project_hub", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="nav-user-summary"', html=False)
        self.assertContains(response, "nav-author")
        self.assertContains(response, "Author")

    def test_external_author_nav_prefers_external_author_account_type(self):
        self.client.login(username="nav-external@example.com", password="pass123")

        response = self.client.get(
            reverse("synopsis:project_hub", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="nav-user-summary"', html=False)
        self.assertContains(response, "Eli External")
        self.assertContains(response, "External Author")


class ProjectHomepageStatusUiTests(TestCase):
    def setUp(self):
        self.author = User.objects.create_user(
            username="status-author",
            password="pass123",
        )
        self.project = Project.objects.create(title="Status Managed Synopsis")
        UserRole.objects.create(user=self.author, project=self.project, role="author")
        self.client.login(username="status-author", password="pass123")

    def test_project_settings_shows_homepage_listing_controls(self):
        response = self.client.get(
            reverse("synopsis:project_settings", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Homepage listing")
        self.assertContains(response, "Shown under active synopses")
        self.assertContains(response, "Mark as completed / archived")
        self.assertNotContains(response, "Archive synopsis")

    def test_author_can_move_synopsis_to_completed_section_without_locking_it(self):
        response = self.client.post(
            reverse("synopsis:project_settings", args=[self.project.id]),
            {"status_action": "mark_completed"},
            follow=True,
        )

        self.assertRedirects(
            response, reverse("synopsis:project_settings", args=[self.project.id])
        )
        self.project.refresh_from_db()
        self.assertEqual(self.project.status, "completed")
        self.assertContains(response, "Shown under completed / archived synopses")
        self.assertContains(response, "Move back to active")
        self.assertTrue(
            ProjectChangeLog.objects.filter(
                project=self.project,
                action="Updated project status",
                details="Status: Planning → Completed",
            ).exists()
        )

        dashboard_response = self.client.get(reverse("synopsis:dashboard"))
        self.assertNotIn(self.project, dashboard_response.context["active_projects"])
        self.assertIn(self.project, dashboard_response.context["completed_projects"])
        self.assertContains(
            dashboard_response,
            reverse("synopsis:project_hub", args=[self.project.id]),
            html=False,
        )
        self.assertContains(dashboard_response, "Move to active")

        hub_response = self.client.get(
            reverse("synopsis:project_hub", args=[self.project.id])
        )
        self.assertEqual(hub_response.status_code, 200)
        self.assertContains(hub_response, self.project.title)

    def test_author_can_move_completed_synopsis_back_to_active_section(self):
        self.project.status = "completed"
        self.project.save(update_fields=["status"])

        response = self.client.post(
            reverse("synopsis:project_settings", args=[self.project.id]),
            {"status_action": "reactivate"},
            follow=True,
        )

        self.assertRedirects(
            response, reverse("synopsis:project_settings", args=[self.project.id])
        )
        self.project.refresh_from_db()
        self.assertEqual(self.project.status, "active")
        self.assertContains(response, "Shown under active synopses")
        self.assertTrue(
            ProjectChangeLog.objects.filter(
                project=self.project,
                action="Updated project status",
                details="Status: Completed → Active",
            ).exists()
        )

        dashboard_response = self.client.get(reverse("synopsis:dashboard"))
        self.assertIn(self.project, dashboard_response.context["active_projects"])
        self.assertNotIn(self.project, dashboard_response.context["completed_projects"])

    def test_author_can_move_completed_synopsis_back_to_active_from_dashboard_row(self):
        self.project.status = "completed"
        self.project.save(update_fields=["status"])

        response = self.client.post(
            reverse("synopsis:project_settings", args=[self.project.id]),
            {"status_action": "reactivate", "return_to": "dashboard"},
            follow=True,
        )

        self.assertRedirects(response, reverse("synopsis:dashboard"))
        self.project.refresh_from_db()
        self.assertEqual(self.project.status, "active")
        self.assertIn(self.project, response.context["active_projects"])
        self.assertNotIn(self.project, response.context["completed_projects"])


class ProjectPhaseUiTests(TestCase):
    def setUp(self):
        self.author = User.objects.create_user(
            username="phase-author",
            password="pass123",
        )
        self.project = Project.objects.create(title="Phase Tracker")
        UserRole.objects.create(user=self.author, project=self.project, role="author")
        self.client.login(username="phase-author", password="pass123")

    def test_project_hub_shows_phase_summary_and_shortcut_controls(self):
        response = self.client.get(
            reverse("synopsis:project_hub", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Move to invite advisory board")
        self.assertContains(response, "Manage phase tracker")
        self.assertContains(response, "Default starting phase")
        self.assertContains(response, "Step 1 of 8")
        self.assertNotContains(response, "Set current phase")

    def test_phase_tracker_skips_protocol_and_advisory_when_not_relevant(self):
        self.project.protocol_relevant = False
        self.project.advisory_board_relevant = False
        self.project.save(update_fields=["protocol_relevant", "advisory_board_relevant"])

        response = self.client.get(
            reverse("synopsis:project_hub", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "References screening")
        self.assertContains(response, "Move to summary writing")
        self.assertContains(response, "Step 1 of 6")
        self.assertNotContains(response, "Move to draft protocol")
        self.assertNotContains(response, "Move to invite advisory board")

    def test_protocol_and_advisory_cards_show_not_relevant_state(self):
        self.project.protocol_relevant = False
        self.project.advisory_board_relevant = False
        self.project.save(update_fields=["protocol_relevant", "advisory_board_relevant"])

        response = self.client.get(
            reverse("synopsis:project_hub", args=[self.project.id])
        )

        self.assertContains(response, "Protocol")
        self.assertContains(response, "This synopsis is not using the protocol workflow in the portal.")
        self.assertContains(response, "Advisory Board")
        self.assertContains(response, "This synopsis is not using an advisory board in the portal.")

    def test_project_settings_shows_full_phase_tracker_controls(self):
        response = self.client.get(
            reverse("synopsis:project_settings", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Phase tracker")
        self.assertContains(response, "Set current phase")
        self.assertContains(response, "Default starting phase")

    def test_author_can_set_phase_backwards_or_forwards(self):
        self.project.phase_manual = "summary_writing"
        self.project.phase_manual_updated = timezone.now()
        self.project.save(update_fields=["phase_manual", "phase_manual_updated"])

        response = self.client.post(
            reverse(
                "synopsis:project_phase_confirm",
                args=[self.project.id, "draft_protocol"],
            )
        )

        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.project.id])
        )
        self.project.refresh_from_db()
        self.assertEqual(self.project.phase_manual, "draft_protocol")
        self.assertEqual(self.project.phase, "draft_protocol")
        event = self.project.phase_events.first()
        self.assertIsNotNone(event)
        self.assertEqual(event.phase, "draft_protocol")

    def test_project_settings_shows_phase_history_entries(self):
        self.client.post(
            reverse(
                "synopsis:project_phase_confirm",
                args=[self.project.id, "invite_advisory_board"],
            )
        )

        response = self.client.get(
            reverse("synopsis:project_settings", args=[self.project.id])
        )

        self.assertContains(response, "Phase history")
        self.assertContains(response, "Invite advisory board")
        self.assertContains(
            response,
            "Phase changed from Draft protocol to Invite advisory board.",
        )

    def test_project_hub_shortcut_moves_to_next_phase(self):
        response = self.client.post(
            reverse(
                "synopsis:project_phase_confirm",
                args=[self.project.id, "invite_advisory_board"],
            )
        )

        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.project.id])
        )
        self.project.refresh_from_db()
        self.assertEqual(self.project.phase, "invite_advisory_board")

    def test_phase_updates_are_logged_in_recent_changes(self):
        response = self.client.post(
            reverse(
                "synopsis:project_phase_confirm",
                args=[self.project.id, "invite_advisory_board"],
            )
        )

        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.project.id])
        )
        change = ProjectChangeLog.objects.filter(
            project=self.project,
            action="Updated project phase",
        ).first()
        self.assertIsNotNone(change)
        self.assertIn("Draft protocol", change.details)
        self.assertIn("Invite advisory board", change.details)

    def test_project_hub_shows_revision_history_timeline(self):
        self.client.post(
            reverse(
                "synopsis:project_phase_confirm",
                args=[self.project.id, "invite_advisory_board"],
            )
        )

        response = self.client.get(
            reverse("synopsis:project_hub", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Revision history")
        self.assertContains(response, "Phase confirmed: Invite advisory board")
        self.assertContains(response, "Updated project phase")
        self.assertContains(response, "Draft protocol")
        self.assertContains(response, "phase-author")

    def test_project_hub_normalises_collaborative_history_entries(self):
        ProjectChangeLog.objects.create(
            project=self.project,
            changed_by=self.author,
            action="Protocol collaborative session closed",
            details="Session 123e4567-e89b-12d3-a456-426614174000 closed (status 3).",
        )
        ProjectChangeLog.objects.create(
            project=self.project,
            changed_by=self.author,
            action="Protocol updated via collaborative edit",
            details=(
                "Session: 123e4567-e89b-12d3-a456-426614174000 | "
                "Status: 6 | File: protocol-v2.docx | Users: phase-author | "
                "Size: 24.0 KB | Reason: Updated references section"
            ),
        )

        response = self.client.get(
            reverse("synopsis:project_hub", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response, "Protocol collaborative session ended"
        )
        self.assertContains(
            response, "Closed without additional document changes."
        )
        self.assertContains(
            response, "Protocol revision saved from collaborative editing"
        )
        self.assertContains(response, "Saved file: protocol-v2.docx")
        self.assertContains(response, "Revision note: Updated references section")
        self.assertNotContains(
            response, "123e4567-e89b-12d3-a456-426614174000"
        )
        self.assertNotContains(response, "status 3")

    def test_manager_role_can_update_phase(self):
        manager = User.objects.create_user(username="phase-manager", password="pass123")
        UserRole.objects.create(user=manager, project=self.project, role="manager")
        self.client.login(username="phase-manager", password="pass123")

        response = self.client.post(
            reverse(
                "synopsis:project_phase_confirm",
                args=[self.project.id, "draft_synopsis"],
            )
        )

        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.project.id])
        )
        self.project.refresh_from_db()
        self.assertEqual(self.project.phase, "draft_synopsis")

    def test_cannot_set_phase_to_disabled_step(self):
        self.project.protocol_relevant = False
        self.project.save(update_fields=["protocol_relevant"])

        response = self.client.post(
            reverse(
                "synopsis:project_phase_confirm",
                args=[self.project.id, "draft_protocol"],
            )
        )

        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.project.id])
        )
        messages = [message.message for message in get_messages(response.wsgi_request)]
        self.assertIn("That phase is not available for this project.", messages)


class ProjectWorkflowApplicabilityTests(TestCase):
    def setUp(self):
        self.author = User.objects.create_user(
            username="workflow-author",
            password="pass123",
        )
        self.project = Project.objects.create(title="Workflow flexibility")
        UserRole.objects.create(user=self.author, project=self.project, role="author")
        self.client.login(username="workflow-author", password="pass123")

    def test_project_settings_can_mark_protocol_and_advisory_not_relevant(self):
        response = self.client.post(
            reverse("synopsis:project_settings", args=[self.project.id]),
            {
                "title": self.project.title,
                "description": "",
                "protocol_relevant": "",
                "advisory_board_relevant": "",
            },
        )

        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.project.id])
        )
        self.project.refresh_from_db()
        self.assertFalse(self.project.protocol_relevant)
        self.assertFalse(self.project.advisory_board_relevant)
        change = ProjectChangeLog.objects.filter(
            project=self.project, action="Updated project settings"
        ).first()
        self.assertIsNotNone(change)
        self.assertIn("Protocol: relevant → not relevant", change.details)
        self.assertIn("Advisory board: relevant → not relevant", change.details)

    def test_protocol_page_redirects_when_protocol_not_relevant(self):
        self.project.protocol_relevant = False
        self.project.save(update_fields=["protocol_relevant"])

        response = self.client.get(
            reverse("synopsis:protocol_detail", args=[self.project.id])
        )

        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.project.id])
        )
        messages = [message.message for message in get_messages(response.wsgi_request)]
        self.assertIn(
            "Protocol is marked as not relevant for this project. Update Project settings if you want to use the protocol workflow.",
            messages,
        )

    def test_advisory_board_page_redirects_when_not_relevant(self):
        self.project.advisory_board_relevant = False
        self.project.save(update_fields=["advisory_board_relevant"])

        response = self.client.get(
            reverse("synopsis:advisory_board_list", args=[self.project.id])
        )

        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.project.id])
        )
        messages = [message.message for message in get_messages(response.wsgi_request)]
        self.assertIn(
            "Advisory board is marked as not relevant for this project. Update Project settings if you want to use the advisory board workflow.",
            messages,
        )
