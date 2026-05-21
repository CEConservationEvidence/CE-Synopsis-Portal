import copy
import datetime as dt
import io
import hashlib
import json
import logging
import mimetypes
import os
import re
import random
import uuid
from decimal import Decimal
from urllib.parse import urljoin, urlparse, urlencode, urlunparse

import jwt
import requests
import rispy
import html
import html as html_lib
import time
from defusedxml import ElementTree as ET

from django.conf import settings
from django.contrib import messages
from django.contrib.auth import views as auth_views
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User, Group
from django.contrib.auth.tokens import default_token_generator
from django.core.exceptions import PermissionDenied, ValidationError
from django.core.files.base import ContentFile
from django.core.mail import EmailMultiAlternatives
from django.db import connection, transaction
from collections import Counter, defaultdict

from django.db.models import Count, Max, Prefetch, Q
from django.http import (
    HttpResponseBadRequest,
    Http404,
    JsonResponse,
    HttpResponseNotAllowed,
    HttpResponse,
)
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.template.loader import render_to_string
from django.utils import timezone
from django.utils.encoding import force_bytes
from django.utils.http import url_has_allowed_host_and_scheme
from django.utils.http import urlsafe_base64_encode
from django.utils.text import slugify
from django import forms
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.clickjacking import xframe_options_exempt
from django.http import FileResponse

from .forms import (
    GLOBAL_ROLE_CHOICES,
    ProtocolUpdateForm,
    ActionListUpdateForm,
    CreateUserForm,
    ManagerUserDeleteForm,
    ManagerUserUpdateForm,
    PortalAuthenticationForm,
    PortalPasswordResetForm,
    PortalSetPasswordForm,
    AdvisoryBoardMemberForm,
    AdvisoryInviteForm,
    AssignAuthorsForm,
    FunderForm,
    ProjectDeleteForm,
    ProjectSettingsForm,
    AdvisoryBulkInviteForm,
    ProtocolSendForm,
    ActionListSendForm,
    SynopsisSendForm,
    ReminderScheduleForm,
    ProtocolReminderScheduleForm,
    ActionListReminderScheduleForm,
    SynopsisReminderScheduleForm,
    ParticipationConfirmForm,
    ParticipationDeclineForm,
    ProtocolFeedbackForm,
    ActionListFeedbackForm,
    SynopsisFeedbackForm,
    ProtocolFeedbackCloseForm,
    ActionListFeedbackCloseForm,
    ReferenceBatchUploadForm,
    LibraryReferenceBatchUploadForm,
    ReferenceScreeningForm,
    ReferenceClassificationForm,
    LibraryReferenceUpdateForm,
    CollaborativeUpdateForm,
    AdvisoryCustomFieldForm,
    AdvisoryCustomFieldPlacementForm,
    AdvisoryMemberCustomDataForm,
    ReferenceSummaryAssignmentForm,
    ReferenceSummaryDraftForm,
    ReferenceSummaryUpdateForm,
    ReferenceSummaryCommentForm,
    ReferenceCommentForm,
    ReferenceDocumentForm,
    SynopsisChapterForm,
    SynopsisSubheadingForm,
    SynopsisInterventionForm,
    SynopsisBackgroundForm,
    SynopsisInterventionSynthesisForm,
    SynopsisKeyMessageForm,
    SynopsisAssignmentForm,
    ReferenceActionSummaryForm,
    FunderContactFormSet,
)
from .models import (
    Project,
    Protocol,
    ActionList,
    AdvisoryBoardMember,
    AdvisoryBoardInvitation,
    AdvisoryBoardCustomField,
    AdvisoryBoardCustomFieldValue,
    AdvisoryBoardCustomFieldValueHistory,
    Funder,
    FunderContact,
    UserRole,
    ProjectPhaseEvent,
    ProjectChangeLog,
    ProtocolFeedback,
    ActionListFeedback,
    SynopsisFeedback,
    LibraryImportBatch,
    LibraryReference,
    LibraryReferenceFolderHistory,
    ReferenceSourceBatch,
    ReferenceSourceBatchNoteHistory,
    Reference,
    ReferenceSummary,
    ReferenceSummaryComment,
    ReferenceComment,
    ReferenceActionSummary,
    ProtocolRevision,
    ActionListRevision,
    CollaborativeSession,
    SynopsisChapter,
    SynopsisSubheading,
    SynopsisIntervention,
    SynopsisInterventionKeyMessage,
    SynopsisAssignment,
    IUCNCategory,
    SynopsisExportLog,
    normalize_reference_folder_values,
)
from .presets import PRESETS
from .utils import (
    advisory_member_display_name,
    default_action_list_review_message,
    default_advisory_invitation_message,
    default_protocol_review_message,
    default_synopsis_review_message,
    ensure_global_groups,
    email_subject,
    is_external_author_user,
    minimum_allowed_deadline_date,
    reply_to_list,
    reference_hash,
)


ONLYOFFICE_SETTINGS = getattr(settings, "ONLYOFFICE", {})

logger = logging.getLogger(__name__)
GLOBAL_ROLE_LABELS = dict(GLOBAL_ROLE_CHOICES)


class PortalLoginView(auth_views.LoginView):
    template_name = "registration/login.html"
    authentication_form = PortalAuthenticationForm
    redirect_authenticated_user = True

    def form_valid(self, form):
        response = super().form_valid(form)
        if not form.cleaned_data.get("remember_me"):
            self.request.session.set_expiry(0)
        return response


class PortalLogoutView(auth_views.LogoutView):
    http_method_names = ["post", "options"]


class PortalPasswordResetView(auth_views.PasswordResetView):
    template_name = "registration/password_reset_form.html"
    email_template_name = "registration/password_reset_email.txt"
    subject_template_name = "registration/password_reset_subject.txt"
    form_class = PortalPasswordResetForm


class PortalPasswordResetConfirmView(auth_views.PasswordResetConfirmView):
    template_name = "registration/password_reset_confirm.html"
    form_class = PortalSetPasswordForm


def _send_account_setup_email(user, request):
    # TODO: #100 Move outbound email sending onto Celery so account creation does
    # not block a request thread on SMTP/network latency.
    uid = urlsafe_base64_encode(force_bytes(user.pk))
    token = default_token_generator.make_token(user)
    reset_url = request.build_absolute_uri(
        reverse("synopsis:password_reset_confirm", args=[uid, token])
    )
    login_url = request.build_absolute_uri(reverse("synopsis:login"))
    context = {
        "user": user,
        "reset_url": reset_url,
        "login_url": login_url,
        "request": request,
    }
    subject = render_to_string(
        "registration/account_setup_subject.txt", context
    ).strip()
    body = render_to_string("registration/account_setup_email.txt", context)
    sent = EmailMultiAlternatives(
        subject=subject,
        body=body,
        from_email=settings.DEFAULT_FROM_EMAIL,
        to=[user.email],
    ).send()
    if sent != 1:
        raise RuntimeError("Account setup email was not sent.")


def _send_password_reset_email(user, request):
    # TODO: #100 Move outbound email sending onto Celery so password reset requests
    # stay fast under concurrent usage.
    uid = urlsafe_base64_encode(force_bytes(user.pk))
    token = default_token_generator.make_token(user)
    context = {
        "user": user,
        "uid": uid,
        "token": token,
        "domain": request.get_host(),
        "protocol": "https" if request.is_secure() else "http",
    }
    subject = render_to_string(
        "registration/password_reset_subject.txt", context
    ).strip()
    body = render_to_string("registration/password_reset_email.txt", context)
    sent = EmailMultiAlternatives(
        subject=subject,
        body=body,
        from_email=settings.DEFAULT_FROM_EMAIL,
        to=[user.email],
    ).send()
    if sent != 1:
        raise RuntimeError("Password reset email was not sent.")


def _manager_user_global_role(user):
    if getattr(user, "is_superuser", False):
        return "system_admin"
    user_groups = set(user.groups.values_list("name", flat=True))
    for role_value, _label in GLOBAL_ROLE_CHOICES:
        if role_value in user_groups:
            return role_value
    if user.is_staff:
        return "manager"
    return ""


def _manager_user_global_role_label(user):
    role_value = _manager_user_global_role(user)
    if role_value == "system_admin":
        return "System admin"
    return GLOBAL_ROLE_LABELS.get(role_value, "—")


def _set_manager_user_global_role(user, global_role):
    group_names = [value for value, _label in GLOBAL_ROLE_CHOICES]
    existing_groups = list(Group.objects.filter(name__in=group_names, user__id=user.id))
    if existing_groups:
        user.groups.remove(*existing_groups)
    desired_group = Group.objects.get(name=global_role)
    user.groups.add(desired_group)
    user.is_staff = global_role == "manager"


def _manager_user_entries():
    users = User.objects.prefetch_related("groups").order_by(
        "-is_superuser",
        "-is_staff",
        "username",
    )
    entries = []
    for user in users:
        role_value = _manager_user_global_role(user)
        entries.append(
            {
                "user": user,
                "global_role": role_value,
                "global_role_label": _manager_user_global_role_label(user),
                "password_state": (
                    "Password set" if user.has_usable_password() else "Setup pending"
                ),
                "status_label": "Active" if user.is_active else "Inactive",
                "is_protected": bool(user.is_superuser),
            }
        )
    return entries


def _decode_entities(text):
    if not text:
        return ""
    try:
        text = re.sub(r"\\u([0-9a-fA-F]{4})", lambda m: chr(int(m.group(1), 16)), text)
    except Exception:
        logging.debug("Unexpected error during unicode escape decoding.", exc_info=True)
    try:
        return html.unescape(text)
    except Exception:
        logging.debug("Unexpected error during HTML unescaping.", exc_info=True)
        return text


_COLLAB_INVITE_TABLE_EXISTS = None


def _linked_project_reference_count(library_reference):
    return Reference.objects.filter(library_reference=library_reference).count()


def _update_shared_library_reference_folders(
    library_reference,
    folder_values,
    *,
    changed_by=None,
    source_project=None,
    source_reference=None,
    change_source="",
    previous_folders=None,
):
    shared_folders = normalize_reference_folder_values(folder_values)
    previous_values = (
        normalize_reference_folder_values(previous_folders)
        if previous_folders is not None
        else normalize_reference_folder_values(library_reference.reference_folder)
    )
    changed = previous_values != shared_folders
    if changed:
        library_reference.reference_folder = shared_folders
        library_reference.save(update_fields=["reference_folder", "updated_at"])
        LibraryReferenceFolderHistory.objects.create(
            library_reference=library_reference,
            previous_folders=previous_values,
            new_folders=shared_folders,
            changed_by=changed_by if getattr(changed_by, "is_authenticated", False) else None,
            source_project=source_project,
            source_reference=source_reference,
            change_source=change_source,
        )
    linked_project_refs = _linked_project_reference_count(library_reference)
    return changed, linked_project_refs, previous_values, shared_folders


def _reference_category_values(reference):
    return list(reference.category_values)


def _update_reference_categories(
    reference,
    category_values,
    *,
    changed_by=None,
    source_project=None,
    change_source="",
):
    categories = normalize_reference_folder_values(category_values)
    shared_changed = False
    shared_linked_count = 0
    local_changed = False
    if reference.library_reference_id:
        (
            shared_changed,
            shared_linked_count,
            _previous_shared_folders,
            _new_shared_folders,
        ) = _update_shared_library_reference_folders(
            reference.library_reference,
            categories,
            changed_by=changed_by,
            source_project=source_project,
            source_reference=reference,
            change_source=change_source,
        )
    else:
        local_before = normalize_reference_folder_values(
            reference.unlinked_reference_folder
        )
        local_changed = local_before != categories
        if local_changed:
            reference.unlinked_reference_folder = categories
            reference.save(update_fields=["unlinked_reference_folder", "updated_at"])
    return shared_changed, shared_linked_count, local_changed, categories


def _collaborative_invitation_table_ready():
    global _COLLAB_INVITE_TABLE_EXISTS
    if _COLLAB_INVITE_TABLE_EXISTS is not None:
        return _COLLAB_INVITE_TABLE_EXISTS

    table_name = CollaborativeSession.invitations.through._meta.db_table
    try:
        tables = connection.introspection.table_names()
    except Exception as exc:
        logger.warning(
            "Could not inspect database tables for collaborative invites: %s", exc
        )
        _COLLAB_INVITE_TABLE_EXISTS = False
        return False

    _COLLAB_INVITE_TABLE_EXISTS = table_name in tables
    return _COLLAB_INVITE_TABLE_EXISTS


def _onlyoffice_enabled() -> bool:
    return bool(ONLYOFFICE_SETTINGS.get("base_url"))


def _user_is_manager(user) -> bool:
    if not getattr(user, "is_authenticated", False):
        return False
    if user.is_staff:
        return True
    return user.groups.filter(name="manager").exists()


def _user_can_manage_library(user) -> bool:
    if not getattr(user, "is_authenticated", False):
        return False
    if user.is_staff:
        return True
    if is_external_author_user(user):
        return False
    if user.groups.filter(name__in=["manager", "author"]).exists():
        return True
    return UserRole.objects.filter(user=user, role__in=["manager", "author"]).exists()


def _user_can_view_project(user, project) -> bool:
    if not getattr(user, "is_authenticated", False):
        return False
    if _user_is_manager(user):
        return True
    if is_external_author_user(user):
        try:
            return project.author_users.filter(id=user.id).exists()
        except Exception:
            return False
    return True


def _user_can_manage_project_configuration(user, project) -> bool:
    if not getattr(user, "is_authenticated", False):
        return False
    if _user_is_manager(user):
        return True
    if is_external_author_user(user):
        return False
    return UserRole.objects.filter(
        user=user, project=project, role__in=["author", "manager"]
    ).exists()


def _user_can_edit_project(user, project) -> bool:
    if not getattr(user, "is_authenticated", False):
        return False
    if _user_is_manager(user):
        return True
    try:
        return project.author_users.filter(id=user.id).exists()
    except Exception:
        return False


def _member_section_key(member: AdvisoryBoardMember) -> str:
    response = (member.response or "").upper()
    if response == "Y":
        return AdvisoryBoardCustomField.SECTION_ACCEPTED
    if response == "N":
        return AdvisoryBoardCustomField.SECTION_DECLINED
    return AdvisoryBoardCustomField.SECTION_PENDING


def _localise_datetime(value):
    if not value:
        return None
    if isinstance(value, dt.datetime):
        try:
            return timezone.localtime(value)
        except (ValueError, TypeError, OverflowError):
            return value
    return value


def _status_badge(label, value, badge, *, date=None, date_label=None, note=None):
    payload = {
        "label": label,
        "value": value,
        "badge": badge,
    }
    if date:
        payload["date"] = _localise_datetime(date)
    if date_label:
        payload["date_label"] = date_label
    if note:
        payload["note"] = note
    return payload


def _member_glance_statuses(member: AdvisoryBoardMember) -> list[dict]:
    statuses: list[dict] = []

    if getattr(member, "invite_sent", False):
        statuses.append(
            _status_badge(
                "Invite",
                "Sent",
                "success",
                date=getattr(member, "invite_sent_at", None),
                date_label="Sent",
            )
        )
    else:
        statuses.append(
            _status_badge(
                "Invite",
                "Draft",
                "secondary",
                note="No invite sent yet",
            )
        )

    response = (getattr(member, "response", "") or "").upper()
    if response == "Y":
        note = None
        if getattr(member, "participation_confirmed", False):
            note = "Participation confirmed"
        elif not getattr(member, "participation_confirmed", False):
            note = "Confirmation pending"
        statuses.append(
            _status_badge(
                "Response",
                "Accepted",
                "success",
                date=getattr(member, "response_date", None),
                date_label="Responded",
                note=note,
            )
        )
    elif response == "N":
        statuses.append(
            _status_badge(
                "Response",
                "Declined",
                "danger",
                date=getattr(member, "response_date", None),
                date_label="Responded",
            )
        )
    else:
        statuses.append(
            _status_badge(
                "Response",
                "Pending",
                "warning",
                note="Awaiting response",
            )
        )

    latest_action_feedback = getattr(member, "latest_action_list_feedback_obj", None)
    if latest_action_feedback and getattr(latest_action_feedback, "submitted_at", None):
        statuses.append(
            _status_badge(
                "Action list",
                "Feedback received",
                "success",
                date=latest_action_feedback.submitted_at,
                date_label="Received",
            )
        )
    elif response == "N":
        statuses.append(
            _status_badge(
                "Action list",
                "Not applicable",
                "secondary",
                note="Member declined",
            )
        )
    elif getattr(member, "sent_action_list_at", None):
        if getattr(member, "feedback_on_action_list_deadline", None):
            if getattr(member, "action_list_reminder_sent", False):
                statuses.append(
                    _status_badge(
                        "Action list",
                        "Reminder sent",
                        "info",
                        date=getattr(member, "action_list_reminder_sent_at", None),
                        date_label="Reminded",
                    )
                )
            else:
                statuses.append(
                    _status_badge(
                        "Action list",
                        "Awaiting feedback",
                        "warning",
                        date=getattr(member, "feedback_on_action_list_deadline", None),
                        date_label="Due",
                    )
                )
        else:
            statuses.append(
                _status_badge(
                    "Action list",
                    "Sent",
                    "primary",
                    date=getattr(member, "sent_action_list_at", None),
                    date_label="Sent",
                )
            )
    else:
        statuses.append(
            _status_badge(
                "Action list",
                "Not sent",
                "secondary",
            )
        )

    protocol_feedback_received = getattr(member, "feedback_on_protocol_received", None)
    if protocol_feedback_received:
        statuses.append(
            _status_badge(
                "Protocol",
                "Feedback received",
                "success",
                date=protocol_feedback_received,
                date_label="Received",
            )
        )
    elif response == "N":
        statuses.append(
            _status_badge(
                "Protocol",
                "Not applicable",
                "secondary",
                note="Member declined",
            )
        )
    elif getattr(member, "sent_protocol_at", None):
        if getattr(member, "feedback_on_protocol_deadline", None):
            if getattr(member, "protocol_reminder_sent", False):
                statuses.append(
                    _status_badge(
                        "Protocol",
                        "Reminder sent",
                        "info",
                        date=getattr(member, "protocol_reminder_sent_at", None),
                        date_label="Reminded",
                    )
                )
            else:
                statuses.append(
                    _status_badge(
                        "Protocol",
                        "Awaiting feedback",
                        "warning",
                        date=getattr(member, "feedback_on_protocol_deadline", None),
                        date_label="Due",
                    )
                )
        else:
            statuses.append(
                _status_badge(
                    "Protocol",
                    "Sent",
                    "primary",
                    date=getattr(member, "sent_protocol_at", None),
                    date_label="Sent",
                )
            )
    else:
        statuses.append(
            _status_badge(
                "Protocol",
                "Not sent",
                "secondary",
            )
        )

    synopsis_feedback_received = getattr(member, "feedback_on_synopsis_received", None)
    if synopsis_feedback_received:
        statuses.append(
            _status_badge(
                "Synopsis",
                "Feedback received",
                "success",
                date=synopsis_feedback_received,
                date_label="Received",
            )
        )
    elif response == "N":
        statuses.append(
            _status_badge(
                "Synopsis",
                "Not applicable",
                "secondary",
                note="Member declined",
            )
        )
    elif getattr(member, "sent_synopsis_at", None):
        if getattr(member, "feedback_on_synopsis_deadline", None):
            if getattr(member, "synopsis_reminder_sent", False):
                statuses.append(
                    _status_badge(
                        "Synopsis",
                        "Reminder sent",
                        "info",
                        date=getattr(member, "synopsis_reminder_sent_at", None),
                        date_label="Reminded",
                    )
                )
            else:
                statuses.append(
                    _status_badge(
                        "Synopsis",
                        "Awaiting feedback",
                        "warning",
                        date=getattr(member, "feedback_on_synopsis_deadline", None),
                        date_label="Due",
                    )
                )
        else:
            statuses.append(
                _status_badge(
                    "Synopsis",
                    "Sent",
                    "primary",
                    date=getattr(member, "sent_synopsis_at", None),
                    date_label="Sent",
                )
            )
    else:
        statuses.append(
            _status_badge(
                "Synopsis",
                "Not sent",
                "secondary",
            )
        )

    return statuses


def _user_can_force_end_session(user, project, session) -> bool:
    if not getattr(user, "is_authenticated", False):
        return False
    if _user_is_manager(user):
        return True

    owner_id = session.started_by_id
    user_id = getattr(user, "pk", None)

    if owner_id is None:
        return _user_can_edit_project(user, project)

    if user_id is None:
        return False

    return str(owner_id) == str(user_id)


def _log_project_change(project, user, action: str, details: str = ""):
    changed_by = user if getattr(user, "is_authenticated", False) else None
    ProjectChangeLog.objects.create(
        project=project, changed_by=changed_by, action=action, details=details
    )


def _format_value(value):
    if value in (None, ""):
        return "—"
    if isinstance(value, dt.date):
        return value.strftime("%Y-%m-%d")
    return str(value)


def _current_revision_label(document):
    if not document:
        return ""
    revision = getattr(document, "current_revision", None)
    if not revision:
        return ""
    return revision.version_label or ""


def _user_display(user: User) -> str:
    full = user.get_full_name()
    return full or user.username


def _advisory_member_display(member: AdvisoryBoardMember) -> str:
    if not member:
        return "Advisory board member"
    parts = [member.first_name.strip() if member.first_name else ""]
    if member.last_name:
        parts.append(member.last_name.strip())
    name = " ".join(part for part in parts if part).strip()
    if not name:
        name = member.email or "Advisory board member"
    return name


def _funder_contact_label(contact, last_name=None) -> str:
    """
    Accept either a FunderContact instance or discrete name parts.
    Tests call with first/last strings; runtime calls pass a contact object.
    """
    if last_name is not None and not hasattr(contact, "display_name"):
        first = (contact or "").strip()
        last = (last_name or "").strip()
        label = " ".join(part for part in (first, last) if part).strip()
        return label or "—"

    if not contact:
        return "—"
    label = contact.display_name()
    if getattr(contact, "email", ""):
        return f"{label} ({contact.email})" if label != "—" else contact.email
    return label


def _contact_entries_from_formset(formset):
    contacts = []
    for form in formset:
        data = getattr(form, "cleaned_data", {}) or {}
        if not data or data.get("DELETE"):
            continue
        has_data = getattr(form, "has_contact_data", lambda *_: False)(data)
        if not has_data:
            continue
        contacts.append(
            {
                "title": data.get("title") or "",
                "first_name": data.get("first_name") or "",
                "last_name": data.get("last_name") or "",
                "email": data.get("email") or "",
                "is_primary": bool(data.get("is_primary")),
            }
        )
    contacts.sort(
        key=lambda c: (
            not c["is_primary"],
            (c["last_name"] or "").lower(),
            (c["first_name"] or "").lower(),
        )
    )
    return contacts


def _formset_has_contacts(formset) -> bool:
    return any(_contact_entries_from_formset(formset))


def _contact_summary_text(contacts) -> str:
    contact_list = sorted(
        contacts,
        key=lambda c: (
            not getattr(c, "is_primary", False),
            (getattr(c, "last_name", "") or "").lower(),
            (getattr(c, "first_name", "") or "").lower(),
            getattr(c, "id", 0),
        ),
    )
    if not contact_list:
        return "—"
    return "; ".join(
        f"{'Primary: ' if c.is_primary else ''}{_funder_contact_label(c)}"
        for c in contact_list
    )


def _format_deadline(deadline):
    if not deadline:
        return "—"
    try:
        aware = timezone.localtime(deadline)
    except (ValueError, TypeError):
        aware = deadline
    return aware.strftime("%d %b %Y %H:%M")


def _invite_response_window_days():
    return getattr(settings, "ADVISORY_INVITE_RESPONSE_WINDOW_DAYS", 10)


def _document_feedback_window_days():
    return getattr(settings, "ADVISORY_DOCUMENT_FEEDBACK_WINDOW_DAYS", 10)


def _end_of_day_datetime(date_value):
    combined = dt.datetime.combine(date_value, dt.time(23, 59))
    return timezone.make_aware(combined) if timezone.is_naive(combined) else combined


def _default_invite_due_date():
    return timezone.localdate() + dt.timedelta(days=_invite_response_window_days())


def _default_document_feedback_due_date():
    return timezone.localdate() + dt.timedelta(days=_document_feedback_window_days())


def _default_document_feedback_deadline():
    return _end_of_day_datetime(_default_document_feedback_due_date())


def _resolve_invite_due_date(override=None, member=None):
    return override or getattr(member, "response_date", None) or _default_invite_due_date()


def _resolve_document_feedback_deadline(override_due_date=None, current_deadline=None):
    if override_due_date:
        return _end_of_day_datetime(override_due_date)
    return current_deadline or _default_document_feedback_deadline()


def _normalise_advisory_message(value):
    return (value or "").strip()


def _stored_advisory_invitation_message(value):
    normalised = _normalise_advisory_message(value)
    if normalised == default_advisory_invitation_message():
        return ""
    return normalised


def _project_advisory_invitation_message(project, override=None):
    candidate = _normalise_advisory_message(override)
    if candidate:
        return candidate
    stored = _normalise_advisory_message(
        getattr(project, "advisory_invitation_message", "")
    )
    return stored or default_advisory_invitation_message()


def _default_document_review_message(document_kind):
    if document_kind == "protocol":
        return default_protocol_review_message()
    if document_kind == "action_list":
        return default_action_list_review_message()
    if document_kind == "synopsis":
        return default_synopsis_review_message()
    return ""


def _document_review_message(document_kind, override=None):
    return (
        _normalise_advisory_message(override)
        or _default_document_review_message(document_kind)
    )


def _eligible_advisory_members(project):
    return (
        AdvisoryBoardMember.objects.filter(
            project=project,
            response="Y",
            participation_confirmed=True,
        )
        .exclude(email__isnull=True)
        .exclude(email__exact="")
    )


def _document_preview_recipient_name(project):
    return "advisory board member"


def _invite_preview_recipient_name(project):
    return "advisory board member"


def _html_message_blocks(text):
    blocks = []
    for block in re.split(r"\n\s*\n", _normalise_advisory_message(text)):
        if block.strip():
            escaped = html.escape(block.strip()).replace("\n", "<br>")
            blocks.append(f"<p>{escaped}</p>")
    return "".join(blocks)


def _build_advisory_invitation_email(
    *,
    project,
    recipient_name,
    due_date,
    yes_url,
    no_url,
    standard_message="",
    additional_message="",
    attachment_lines=None,
):
    deadline_txt = due_date.strftime("%d %b %Y") if due_date else "—"
    recipient_label = recipient_name or "advisory board member"
    safe_yes_url = html.escape(yes_url, quote=True)
    safe_no_url = html.escape(no_url, quote=True)
    standard_text = _project_advisory_invitation_message(
        project, override=standard_message
    )
    additional_text = _normalise_advisory_message(additional_message)

    text_parts = [
        f"Dear {recipient_label},",
        "",
        f"You are invited to advise on '{project.title}'.",
        f"Please reply by: {deadline_txt}",
        "",
        standard_text,
    ]
    if additional_text:
        text_parts.extend(["", additional_text])
    text_parts.extend(
        [
            "",
            f"Yes: {yes_url}",
            f"No:  {no_url}",
            "",
            "After clicking Yes you'll be asked to confirm you can actively participate and provide valuable input.",
        ]
    )
    if attachment_lines:
        text_parts.extend(
            [""] + [f"{label}: {url}" for label, url in attachment_lines]
        )
    text_parts.extend(["", "Thank you."])

    html_parts = [
        f"<p>Dear {html.escape(recipient_label)},</p>",
        f"<p>You are invited to advise on '<strong>{html.escape(project.title)}</strong>'.</p>",
        f"<p><strong>Please reply by: {html.escape(deadline_txt)}</strong></p>",
        _html_message_blocks(standard_text),
    ]
    if additional_text:
        html_parts.append(_html_message_blocks(additional_text))
    html_parts.extend(
        [
            (
                f"<p>"
                f"<a href='{safe_yes_url}' style='padding:8px 12px;border:1px solid #0a0;text-decoration:none;'>Yes</a> "
                f"<a href='{safe_no_url}' style='padding:8px 12px;border:1px solid #a00;text-decoration:none;margin-left:8px;'>No</a>"
                f"</p>"
            ),
            "<p><em>After clicking Yes you'll confirm that you will actively participate and provide valuable input.</em></p>",
        ]
    )
    if attachment_lines:
        for label, url in attachment_lines:
            safe_url = html.escape(url, quote=True)
            html_parts.append(
                f"<p><strong>{html.escape(label)}:</strong> "
                f"<a href='{safe_url}'>{html.escape(url)}</a></p>"
            )
    html_parts.append("<p>Thank you.</p>")
    return "\n".join(text_parts), "".join(html_parts)


def _update_project_advisory_invitation_message(project, message, changed_by):
    stored_message = _stored_advisory_invitation_message(message)
    if stored_message == (project.advisory_invitation_message or ""):
        return
    previous_message = bool((project.advisory_invitation_message or "").strip())
    project.advisory_invitation_message = stored_message
    project.save(update_fields=["advisory_invitation_message"])
    if stored_message:
        details = "Saved a custom standard advisory invitation message."
    elif previous_message:
        details = (
            "Cleared the custom advisory invitation message and reverted to the "
            "built-in default."
        )
    else:
        details = "Using the built-in advisory invitation message."
    _log_project_change(
        project,
        changed_by,
        "Updated advisory invitation message",
        details,
    )


def _format_file_size(size_bytes):
    try:
        size = int(size_bytes)
    except (TypeError, ValueError):
        return "—"
    if size < 0:
        return "—"
    units = ["B", "KB", "MB", "GB", "TB"]
    for unit in units:
        if size < 1024 or unit == units[-1]:
            if unit == "B":
                return f"{size} {unit}"
            return f"{size:.1f} {unit}"
        size /= 1024


_UUID_PREFIX_RE = re.compile(
    r"^(?:[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12}_)+"
)


def _normalized_document_filename(filename: str, fallback: str = "document.docx") -> str:
    base_name = os.path.basename(filename or "").strip()
    if not base_name:
        return fallback
    cleaned = base_name
    while True:
        updated = _UUID_PREFIX_RE.sub("", cleaned)
        if updated == cleaned:
            break
        cleaned = updated
    return cleaned or fallback


def _apply_revision_to_protocol(protocol, revision) -> tuple[str, str]:
    try:
        with revision.file.open("rb") as source:
            content = source.read()
    except FileNotFoundError:
        raise FileNotFoundError("Revision file missing")

    if not content:
        raise ValueError("Revision file empty")

    base_name = _normalized_document_filename(
        revision.original_name or os.path.basename(revision.file.name),
        fallback="protocol.docx",
    )
    new_filename = f"{uuid.uuid4()}_{base_name}"
    protocol.document.save(new_filename, ContentFile(content), save=False)
    protocol.current_revision = revision
    protocol.save(update_fields=["document", "current_revision"])
    return base_name, _format_file_size(revision.file_size)


def _apply_revision_to_action_list(action_list, revision) -> tuple[str, str]:
    try:
        with revision.file.open("rb") as source:
            content = source.read()
    except FileNotFoundError:
        raise FileNotFoundError("Revision file missing")

    if not content:
        raise ValueError("Revision file empty")

    base_name = _normalized_document_filename(
        revision.original_name or os.path.basename(revision.file.name),
        fallback="action-list.docx",
    )
    new_filename = f"{uuid.uuid4()}_{base_name}"
    action_list.document.save(new_filename, ContentFile(content), save=False)
    action_list.current_revision = revision
    action_list.save(update_fields=["document", "current_revision"])
    return base_name, _format_file_size(revision.file_size)


COLLAB_SESSION_DURATION = CollaborativeSession.DEFAULT_DURATION


def _normalize_document_type(document_type: str) -> str | None:
    mapping = {
        "protocol": CollaborativeSession.DOCUMENT_PROTOCOL,
        "protocols": CollaborativeSession.DOCUMENT_PROTOCOL,
        "action-list": CollaborativeSession.DOCUMENT_ACTION_LIST,
        "action_list": CollaborativeSession.DOCUMENT_ACTION_LIST,
        "actionlist": CollaborativeSession.DOCUMENT_ACTION_LIST,
    }
    return mapping.get((document_type or "").lower())


def _get_document_for_type(project, document_type):
    if document_type == CollaborativeSession.DOCUMENT_PROTOCOL:
        return getattr(project, "protocol", None)
    if document_type == CollaborativeSession.DOCUMENT_ACTION_LIST:
        return getattr(project, "action_list", None)
    return None


def _get_active_collaborative_session(project, document_type):
    session = (
        CollaborativeSession.objects.filter(
            project=project,
            document_type=document_type,
            is_active=True,
        )
        .order_by("-started_at", "-id")
        .first()
    )
    if session and session.has_expired():
        session.is_active = False
        session.save(update_fields=["is_active"])
        return None
    return session


def _end_active_collaborative_session(
    project, document_type, *, ended_by=None, reason=""
):
    # TODO: #101 Guard collaborative session creation with a distributed lock
    # (Redis) and/or a database uniqueness guarantee so concurrent requests cannot
    # create competing active sessions under load.
    session = _get_active_collaborative_session(project, document_type)
    if not session:
        return None
    session.mark_inactive(
        ended_by=ended_by,
        reason=reason or f"{_document_label(document_type)} collaborative editing closed",
    )
    return session


def _ensure_collaborative_invite_link(
    request,
    project,
    document_type,
    invitation=None,
    *,
    member=None,
    feedback=None,
):
    if not _onlyoffice_enabled():
        return ""

    document = _get_document_for_type(project, document_type)
    if not _document_requires_file(document):
        return ""

    if document_type == CollaborativeSession.DOCUMENT_PROTOCOL and getattr(
        document, "feedback_closed_at", None
    ):
        _end_active_collaborative_session(
            project,
            document_type,
            ended_by=getattr(request, "user", None),
            reason="Protocol feedback window closed",
        )
        return ""

    if document_type == CollaborativeSession.DOCUMENT_ACTION_LIST and getattr(
        document, "feedback_closed_at", None
    ):
        _end_active_collaborative_session(
            project,
            document_type,
            ended_by=getattr(request, "user", None),
            reason="Action list feedback window closed",
        )
        return ""

    session = _get_active_collaborative_session(project, document_type)
    created = False

    if not session:
        initial_revision = getattr(document, "current_revision", None)
        if not initial_revision and hasattr(document, "latest_revision"):
            try:
                initial_revision = document.latest_revision()
            except Exception:
                initial_revision = None

        session_kwargs = {
            "project": project,
            "document_type": document_type,
            "started_by": (
                request.user
                if getattr(request.user, "is_authenticated", False)
                else None
            ),
            "last_activity_at": timezone.now(),
        }

        if document_type == CollaborativeSession.DOCUMENT_PROTOCOL:
            session_kwargs["initial_protocol_revision"] = initial_revision
        else:
            session_kwargs["initial_action_list_revision"] = initial_revision

        session = CollaborativeSession.objects.create(**session_kwargs)
        created = True

    if invitation and session:
        if _collaborative_invitation_table_ready():
            session.invitations.add(invitation)
        else:
            logger.warning(
                "Collaborative invite join table missing; skipping association for session %s",
                session.pk,
            )

    if created:
        _log_project_change(
            project,
            request.user,
            f"{_document_label(document_type)} collaborative editing started",
            f"Session {session.token} shared with advisory invite {invitation.email if invitation else ''}".strip(),
        )

    params = {}
    if invitation:
        params["invite"] = str(invitation.token)
    if member:
        params["member"] = str(member.id)
    if feedback:
        params["feedback"] = str(feedback.token)

    slug = _document_type_slug(document_type)
    path = reverse(
        "synopsis:collaborative_edit", args=[project.id, slug, session.token]
    )
    if params:
        path = f"{path}?{urlencode(params)}"
    return request.build_absolute_uri(path)


def _document_detail_url(project_id, document_type):
    if document_type == CollaborativeSession.DOCUMENT_PROTOCOL:
        return reverse("synopsis:protocol_detail", args=[project_id])
    return reverse("synopsis:action_list_detail", args=[project_id])


def _document_label(document_type):
    if document_type == CollaborativeSession.DOCUMENT_PROTOCOL:
        return "Protocol"
    return "Action list"


def _document_type_slug(document_type):
    if document_type == CollaborativeSession.DOCUMENT_PROTOCOL:
        return "protocol"
    return "action-list"


def _collaborative_session_or_404(project, document_type, token):
    return get_object_or_404(
        CollaborativeSession,
        project=project,
        document_type=document_type,
        token=token,
    )


def _feedback_model_for_document_type(document_type):
    if document_type == CollaborativeSession.DOCUMENT_PROTOCOL:
        return ProtocolFeedback
    if document_type == CollaborativeSession.DOCUMENT_ACTION_LIST:
        return ActionListFeedback
    return None


def _member_feedback_deadline(member, document_type):
    if not member:
        return None
    if document_type == CollaborativeSession.DOCUMENT_PROTOCOL:
        return member.feedback_on_protocol_deadline
    if document_type == CollaborativeSession.DOCUMENT_ACTION_LIST:
        return member.feedback_on_action_list_deadline
    return None


def _collaborative_access_closed_response(
    request,
    project,
    document_label,
    message,
    *,
    status=403,
    detail_url="",
):
    return render(
        request,
        "synopsis/collaborative_editor.html",
        {
            "project": project,
            "document_label": document_label,
            "detail_url": detail_url,
            "window_closed_message": message,
            "can_force_end": False,
            "force_end_url": "",
            "participant_display": "",
        },
        status=status,
    )


def _resolve_external_collaborative_access(request, project, document_type, session):
    feedback_token = request.GET.get("feedback")
    invitation_token = request.GET.get("invite")

    if feedback_token:
        feedback_model = _feedback_model_for_document_type(document_type)
        if not feedback_model:
            return {
                "allowed": False,
                "message": "This collaborative link is not available.",
            }
        try:
            feedback = (
                feedback_model.objects.select_related("member", "invitation")
                .get(token=feedback_token, project=project)
            )
        except (feedback_model.DoesNotExist, TypeError, ValueError, ValidationError):
            return {
                "allowed": False,
                "message": "This collaborative feedback link is not valid.",
            }

        member = feedback.member
        if member and member.response == "N":
            return {
                "allowed": False,
                "message": "This collaborative link is no longer available because the invitation was declined.",
            }

        deadline = _member_feedback_deadline(member, document_type) or feedback.feedback_deadline_at
        if deadline and timezone.now() >= deadline:
            return {
                "allowed": False,
                "message": f"This collaborative feedback link closed on {_format_deadline(deadline)}.",
            }

        if member:
            display = _advisory_member_display(member)
            return {
                "allowed": True,
                "member": member,
                "feedback": feedback,
                "participant_display": display,
                "participant_context": {
                    "id": f"abm:{member.id}",
                    "name": display,
                    "email": member.email,
                },
            }

        display = feedback.email or "Advisory board reviewer"
        return {
            "allowed": True,
            "member": None,
            "feedback": feedback,
            "participant_display": display,
            "participant_context": {
                "id": f"abe:{display.lower()}",
                "name": display,
                "email": feedback.email,
            },
        }

    if invitation_token:
        try:
            invitation = (
                AdvisoryBoardInvitation.objects.select_related("member")
                .get(token=invitation_token, project=project)
            )
        except (
            AdvisoryBoardInvitation.DoesNotExist,
            TypeError,
            ValueError,
            ValidationError,
        ):
            return {
                "allowed": False,
                "message": "This collaborative invitation link is not valid.",
            }

        if (
            _collaborative_invitation_table_ready()
            and session.invitations.exists()
            and not session.invitations.filter(pk=invitation.pk).exists()
        ):
            return {
                "allowed": False,
                "message": "This invitation is not attached to the current collaborative session.",
            }

        if invitation.accepted is False:
            return {
                "allowed": False,
                "message": "This collaborative link is no longer available because the invitation was declined.",
            }

        if invitation.due_date and timezone.localdate() > invitation.due_date:
            return {
                "allowed": False,
                "message": f"This collaborative invitation link closed on {invitation.due_date:%Y-%m-%d}.",
            }

        member = invitation.member
        if member and member.response == "N":
            return {
                "allowed": False,
                "message": "This collaborative link is no longer available because the invitation was declined.",
            }

        if member:
            display = _advisory_member_display(member)
            participant_id = f"abm:{member.id}"
            participant_email = member.email
        else:
            display = invitation.email or "Advisory board invitee"
            participant_id = f"abe:{display.lower()}"
            participant_email = invitation.email

        return {
            "allowed": True,
            "member": member,
            "invitation": invitation,
            "participant_display": display,
            "participant_context": {
                "id": participant_id,
                "name": display,
                "email": participant_email,
            },
        }

    if request.GET.get("member"):
        return {
            "allowed": False,
            "message": "This older collaborative link is missing its secure review token. Ask the authors to resend the link.",
        }

    return {
        "allowed": False,
        "message": "Please sign in as a project author or use the secure collaborative link from your email.",
    }


def _document_requires_file(document) -> bool:
    return bool(document and getattr(document, "document", None))


def _onlyoffice_editor_js_url() -> str:
    base = ONLYOFFICE_SETTINGS.get("base_url", "").rstrip("/")
    if not base:
        return ""
    return f"{base}/web-apps/apps/api/documents/api.js"


def _onlyoffice_service_base_url() -> str:
    return (
        (ONLYOFFICE_SETTINGS.get("internal_url") or "").rstrip("/")
        or (ONLYOFFICE_SETTINGS.get("base_url") or "").rstrip("/")
    )


def _onlyoffice_app_absolute_uri(request, path: str) -> str:
    base = (ONLYOFFICE_SETTINGS.get("app_base_url") or "").rstrip("/")
    if not base:
        return request.build_absolute_uri(path)
    return urljoin(f"{base}/", path.lstrip("/"))


def _document_filetype(file_name: str) -> str:
    ext = (os.path.splitext(file_name)[1] or "").lstrip(".").lower()
    return ext or "docx"


def _collaborative_document_key(project, document_type, session) -> str:
    return f"{project.id}-{document_type}-{session.id}-{int(session.started_at.timestamp())}"[
        -128:
    ]


def _collaborative_query_suffix(querydict) -> str:
    query_string = querydict.urlencode()
    return f"?{query_string}" if query_string else ""


def _restart_external_collaborative_url(
    request, project, document_type, external_access
) -> str:
    if not external_access.get("allowed"):
        return ""
    return _ensure_collaborative_invite_link(
        request,
        project,
        document_type,
        external_access.get("invitation"),
        member=external_access.get("member"),
        feedback=external_access.get("feedback"),
    )


def _build_onlyoffice_config(
    request,
    project,
    document,
    session,
    document_type,
    participant=None,
):
    document_file = getattr(document, "document", None)
    if not document_file:
        raise ValueError("Document has no file attached")

    file_url = _onlyoffice_app_absolute_uri(request, document_file.url)
    file_type = _document_filetype(document_file.name)
    title = os.path.basename(document_file.name) or _document_label(document_type)
    doc_key = _collaborative_document_key(project, document_type, session)

    user = request.user
    user_id = str(getattr(user, "id", "anonymous"))
    user_name = (
        _user_display(user) if getattr(user, "is_authenticated", False) else "Anonymous"
    )
    user_email = (
        getattr(user, "email", "") if getattr(user, "is_authenticated", False) else ""
    )

    if participant:
        user_id = participant.get("id", user_id)
        user_name = participant.get("name", user_name)
        user_email = participant.get("email", user_email)

    callback_url = _onlyoffice_app_absolute_uri(
        request,
        reverse(
            "synopsis:collaborative_edit_callback",
            args=[project.id, _document_type_slug(document_type), session.token],
        )
    )

    config = {
        "document": {
            "fileType": file_type,
            "key": doc_key,
            "title": title,
            "url": file_url,
            "permissions": {
                "edit": True,
                "download": True,
                "print": True,
                "review": True,
            },
        },
        "editorConfig": {
            "callbackUrl": callback_url,
            "user": {
                "id": user_id,
                "name": user_name,
            },
            "mode": "edit",
            "customization": {
                "autosave": True,
                "forcesave": True,
            },
        },
    }

    if user_email:
        config["editorConfig"]["user"]["email"] = user_email

    secret = ONLYOFFICE_SETTINGS.get("jwt_secret")
    if secret:
        token = jwt.encode(config, secret, algorithm="HS256")
        if isinstance(token, bytes):
            token = token.decode("utf-8")
        config["token"] = token

    return config


def _download_onlyoffice_file(file_url: str) -> bytes:
    file_url = _onlyoffice_internal_download_url(file_url)
    raw_entries = [
        ONLYOFFICE_SETTINGS.get("base_url", ""),
        ONLYOFFICE_SETTINGS.get("internal_url", ""),
    ]
    extra = ONLYOFFICE_SETTINGS.get("trusted_download_urls") or []
    if isinstance(extra, (list, tuple)):
        raw_entries.extend(extra)
    else:
        raw_entries.append(extra)

    allowed: list[tuple[str, str, int, str]] = []
    for entry in raw_entries:
        if not entry:
            continue
        try:
            parsed_allowed = urlparse(entry)
        except ValueError:
            continue
        if (
            parsed_allowed.scheme not in {"http", "https"}
            or not parsed_allowed.hostname
        ):
            continue
        allowed.append(
            (
                parsed_allowed.scheme,
                parsed_allowed.hostname.lower(),
                parsed_allowed.port
                or (443 if parsed_allowed.scheme == "https" else 80),
                (parsed_allowed.path or "/").rstrip("/"),
            )
        )

    try:
        parsed = urlparse(file_url)
    except ValueError:
        parsed = None

    if not parsed or parsed.scheme not in {"http", "https"} or not parsed.hostname:
        logger.warning("Blocked OnlyOffice download from invalid URL: %s", file_url)
        raise ValueError("Untrusted OnlyOffice download URL")

    candidate_host = parsed.hostname.lower()
    candidate_port = parsed.port or (443 if parsed.scheme == "https" else 80)
    candidate_path = parsed.path or "/"

    matched = False
    for scheme, host, port, path_prefix in allowed:
        if scheme != parsed.scheme:
            continue
        if candidate_host != host or candidate_port != port:
            continue
        if path_prefix and path_prefix != "/":
            if candidate_path == path_prefix or candidate_path.startswith(
                f"{path_prefix}/"
            ):
                matched = True
                break
        else:
            matched = True
            break

    if not matched:
        logger.warning("Blocked OnlyOffice download from untrusted URL: %s", file_url)
        raise ValueError("Untrusted OnlyOffice download URL")

    timeout = ONLYOFFICE_SETTINGS.get("callback_timeout", 10)
    response = requests.get(file_url, timeout=timeout)
    response.raise_for_status()
    return response.content


def _onlyoffice_secret() -> str:
    return ONLYOFFICE_SETTINGS.get("jwt_secret", "")


def _onlyoffice_internal_download_url(file_url: str) -> str:
    base = (ONLYOFFICE_SETTINGS.get("base_url") or "").rstrip("/")
    internal = (ONLYOFFICE_SETTINGS.get("internal_url") or "").rstrip("/")
    if not base or not internal:
        return file_url

    try:
        parsed_file = urlparse(file_url)
        parsed_base = urlparse(base)
        parsed_internal = urlparse(internal)
    except ValueError:
        return file_url

    if (
        parsed_file.scheme != parsed_base.scheme
        or parsed_file.hostname != parsed_base.hostname
        or (parsed_file.port or (443 if parsed_file.scheme == "https" else 80))
        != (parsed_base.port or (443 if parsed_base.scheme == "https" else 80))
    ):
        return file_url

    replacement_netloc = parsed_internal.netloc or parsed_file.netloc
    return urlunparse(
        (
            parsed_internal.scheme or parsed_file.scheme,
            replacement_netloc,
            parsed_file.path,
            parsed_file.params,
            parsed_file.query,
            parsed_file.fragment,
        )
    )


def _onlyoffice_command_url() -> str:
    base = _onlyoffice_service_base_url()
    if not base:
        return ""
    return f"{base}/coauthoring/CommandService.ashx"


def _onlyoffice_command_headers(payload: dict) -> dict[str, str]:
    headers = {"Content-Type": "application/json"}
    secret = _onlyoffice_secret()
    if not secret:
        return headers
    token = jwt.encode(payload, secret, algorithm="HS256")
    if isinstance(token, bytes):
        token = token.decode("utf-8")
    headers["Authorization"] = f"Bearer {token}"
    return headers


def _request_onlyoffice_forcesave(project, document_type, session) -> tuple[str, str]:
    command_url = _onlyoffice_command_url()
    if not command_url:
        logger.warning(
            "OnlyOffice force-save skipped for session %s because command URL is missing",
            session.pk,
        )
        return "failed", "Collaborative save is not configured correctly."

    payload = {
        "c": "forcesave",
        "key": _collaborative_document_key(project, document_type, session),
    }
    timeout = ONLYOFFICE_SETTINGS.get("callback_timeout", 10)
    try:
        response = requests.post(
            command_url,
            json=payload,
            headers=_onlyoffice_command_headers(payload),
            timeout=timeout,
        )
        response.raise_for_status()
        data = response.json()
    except (requests.RequestException, ValueError) as exc:
        logger.error(
            "OnlyOffice force-save request failed for session %s: %s",
            session.pk,
            exc,
        )
        return "failed", "Unable to request a final save from OnlyOffice."

    error_code = data.get("error", 1) if isinstance(data, dict) else 1
    if error_code in {0, "0"}:
        return "requested", "Final save requested from OnlyOffice."

    if error_code in {4, "4"}:
        logger.info(
            "OnlyOffice force-save reported no pending changes for session %s: %s",
            session.pk,
            data,
        )
        return "noop", "No unsaved changes were pending in OnlyOffice."

    if error_code not in {0, "0"}:
        logger.error(
            "OnlyOffice force-save request returned error for session %s: %s",
            session.pk,
            data,
        )
        return "failed", "OnlyOffice did not accept the final save request."


def _wait_for_collaborative_save(session, document_type, timeout_seconds: int) -> bool:
    # TODO: #102 Reduce blocking waits during collaborative final-save handling,
    # ideally by shifting the long-running part onto callback-driven or background work.
    deadline = time.monotonic() + max(timeout_seconds, 1)
    if document_type == CollaborativeSession.DOCUMENT_PROTOCOL:
        result_field = "result_protocol_revision_id"
    else:
        result_field = "result_action_list_revision_id"

    while time.monotonic() < deadline:
        session.refresh_from_db()
        if getattr(session, result_field):
            return True
        time.sleep(0.5)

    session.refresh_from_db()
    return bool(getattr(session, result_field))


def _extract_onlyoffice_token(request, payload: dict) -> str | None:
    auth_header = request.headers.get("Authorization", "")
    if auth_header.startswith("Bearer "):
        return auth_header.split(" ", 1)[1].strip()
    return payload.get("token") if isinstance(payload, dict) else None


def _parse_onlyoffice_callback(request) -> dict:
    try:
        body = request.body.decode("utf-8") or "{}"
    except UnicodeDecodeError as exc:
        logger.warning("OnlyOffice callback body decode failed: %s", exc)
        raise ValueError("Invalid payload encoding")

    try:
        payload = json.loads(body)
    except json.JSONDecodeError as exc:
        logger.warning("OnlyOffice callback payload JSON error: %s", exc)
        raise ValueError("Invalid JSON payload")

    secret = _onlyoffice_secret()
    if not secret:
        return payload

    token = _extract_onlyoffice_token(request, payload)
    if not token:
        logger.warning("OnlyOffice callback missing JWT token")
        raise PermissionDenied("Missing callback token")

    try:
        decoded = jwt.decode(token, secret, algorithms=["HS256"])
    except jwt.InvalidTokenError as exc:
        logger.warning("OnlyOffice callback token invalid: %s", exc)
        raise PermissionDenied("Invalid callback token")

    if not isinstance(decoded, dict):
        logger.warning("OnlyOffice callback token did not decode to a JSON object")
        raise PermissionDenied("Invalid callback payload")

    nested_payload = decoded.get("payload")
    if isinstance(nested_payload, dict):
        return nested_payload

    return decoded


def _create_protocol_feedback(project, member=None, email=None, invitation=None):
    proto = getattr(project, "protocol", None)
    kwargs = {
        "project": project,
        "member": member,
        "email": email or (member.email if member else ""),
        "invitation": invitation,
    }
    deadline = None
    if member:
        if member.response == "Y" and member.feedback_on_protocol_deadline:
            deadline = member.feedback_on_protocol_deadline
    elif invitation and invitation.due_date:
        combined = dt.datetime.combine(invitation.due_date, dt.time(23, 59))
        deadline = (
            timezone.make_aware(combined) if timezone.is_naive(combined) else combined
        )
    if proto:
        kwargs.update(
            {
                "protocol_document_name": getattr(proto.document, "name", ""),
                "protocol_document_last_updated": proto.last_updated,
                "protocol_stage_snapshot": proto.stage,
            }
        )
    kwargs["feedback_deadline_at"] = deadline
    return ProtocolFeedback.objects.create(**kwargs)


def _create_action_list_feedback(project, member=None, email=None, invitation=None):
    action_list = getattr(project, "action_list", None)
    kwargs = {
        "project": project,
        "action_list": action_list,
        "member": member,
        "email": email or (member.email if member else ""),
        "invitation": invitation,
    }
    deadline = None
    if member:
        if member.response == "Y" and member.feedback_on_action_list_deadline:
            deadline = member.feedback_on_action_list_deadline
    elif invitation and invitation.due_date:
        combined = dt.datetime.combine(invitation.due_date, dt.time(23, 59))
        deadline = (
            timezone.make_aware(combined) if timezone.is_naive(combined) else combined
        )
    if action_list:
        kwargs.update(
            {
                "action_list_document_name": getattr(action_list.document, "name", ""),
                "action_list_document_last_updated": action_list.last_updated,
                "action_list_stage_snapshot": action_list.stage,
            }
        )
    kwargs["feedback_deadline_at"] = deadline
    return ActionListFeedback.objects.create(**kwargs)


def _create_synopsis_feedback(project, member=None, email=None, invitation=None):
    kwargs = {
        "project": project,
        "member": member,
        "email": email or (member.email if member else ""),
        "invitation": invitation,
    }
    deadline = None
    if member:
        if member.response == "Y" and member.feedback_on_synopsis_deadline:
            deadline = member.feedback_on_synopsis_deadline
    elif invitation and invitation.due_date:
        combined = dt.datetime.combine(invitation.due_date, dt.time(23, 59))
        deadline = (
            timezone.make_aware(combined) if timezone.is_naive(combined) else combined
        )
    kwargs["feedback_deadline_at"] = deadline
    return SynopsisFeedback.objects.create(**kwargs)


def _extract_reference_field(record: dict, key: str) -> str:
    value = record.get(key)
    if isinstance(value, list):
        if not value:
            return ""
        value = value[0]
    return value or ""


def _coerce_year(value) -> int | None:
    if isinstance(value, list):
        value = value[0] if value else None
    try:
        text = str(value).strip()
    except Exception:  # pragma: no cover - defensive
        return None
    if not text:
        return None
    for token in (text, text[:4]):
        try:
            return int(token)
        except (ValueError, TypeError):
            continue
    return None


def _combine_pages(record: dict) -> str:
    pages = _extract_reference_field(record, "pages")
    if pages:
        return pages
    start = _extract_reference_field(record, "start_page")
    end = _extract_reference_field(record, "end_page")
    if start and end:
        return f"{start}-{end}".strip("-")
    return start or end or ""


PLAIN_REFERENCE_SPLIT_RE = re.compile(r"\n\s*\n+")
PLAIN_REFERENCE_CITATION_RE = re.compile(
    r"""^(?P<authors>.+?)\s*\((?P<year>\d{4})\)\.\s*(?P<body>.+)$""",
    re.UNICODE,
)
PLAIN_REFERENCE_JOURNAL_RE = re.compile(
    r"^(?P<journal>[^\d:]+?)\s+(?P<volume>\d+)(?:\((?P<issue>[^)]+)\))?\s*:?\s*(?P<pages>[\d\-–]+)?\.?(?:\s|$)",
    re.UNICODE,
)
PLAIN_REFERENCE_DOI_RE = re.compile(r"doi[:\s]+(?P<doi>\S+)", re.IGNORECASE)
PLAIN_REFERENCE_URL_RE = re.compile(r"(https?://\S+)", re.IGNORECASE)


def _decode_reference_upload_text(raw_bytes: bytes) -> str:
    """Decode uploaded reference files and strip a UTF-8 BOM if present."""

    return raw_bytes.decode("utf-8", errors="ignore").lstrip("\ufeff")


def _parse_plaintext_references(payload: str) -> list[dict]:
    """
    Parse a plain-text reference list.

    Each reference is expected to be separated by at least one blank line.
    The first non-empty line should contain the citation in the form:
        Authors (Year). "Title." Journal Volume(Issue): pages.
    Any subsequent non-empty lines are treated as the abstract.
    """

    if not payload or not payload.strip():
        return []

    entries = [
        block.strip()
        for block in PLAIN_REFERENCE_SPLIT_RE.split(payload.strip())
        if block.strip()
    ]
    parsed: list[dict] = []

    for chunk in entries:
        lines = [line.strip() for line in chunk.splitlines() if line.strip()]
        if not lines:
            continue

        citation = lines[0]
        abstract = " ".join(lines[1:]).strip()

        match = PLAIN_REFERENCE_CITATION_RE.match(citation)
        if not match:
            # If parsing fails we skip this chunk rather than guess.
            continue

        authors_part = match.group("authors").strip()
        year = (match.group("year") or "").strip()
        body = (match.group("body") or "").strip()

        title = ""
        remainder = ""

        if body.startswith('"'):
            closing_quote = body.find('"', 1)
            if closing_quote != -1:
                title = body[1:closing_quote].strip()
                remainder = body[closing_quote + 1 :].lstrip(" .")
            else:
                title = body.strip('" ')
        else:
            if ". " in body:
                title_part, remainder_part = body.split(". ", 1)
                title = title_part.strip()
                remainder = remainder_part.strip()
            else:
                dot_index = body.find(".")
                if dot_index != -1:
                    title = body[:dot_index].strip()
                    remainder = body[dot_index + 1 :].lstrip(" .")
                else:
                    title = body.strip()

        if not title:
            title = citation.strip()
        remainder = remainder.lstrip(" .")

        journal = volume = issue = pages = doi = url = ""

        if remainder:
            journal_match = PLAIN_REFERENCE_JOURNAL_RE.match(remainder)
            if journal_match:
                journal = (journal_match.group("journal") or "").strip(" .,;")
                volume = (journal_match.group("volume") or "").strip()
                issue = (journal_match.group("issue") or "").strip()
                pages = (journal_match.group("pages") or "").strip()

            doi_match = PLAIN_REFERENCE_DOI_RE.search(remainder)
            if doi_match:
                doi = doi_match.group("doi").rstrip(".,")

            url_match = PLAIN_REFERENCE_URL_RE.search(remainder)
            if url_match:
                url = url_match.group(1).rstrip(".,)")

        # If DOI or URL appear in the abstract/content include them.
        if not doi:
            doi_match = PLAIN_REFERENCE_DOI_RE.search(abstract)
            if doi_match:
                doi = doi_match.group("doi").rstrip(".,")
        if not url:
            url_match = PLAIN_REFERENCE_URL_RE.search(abstract)
            if url_match:
                url = url_match.group(1).rstrip(".,)")

        # Authors may be separated by semicolons or " and ".
        authors_tokens = []
        for token in re.split(r";|\band\b", authors_part, flags=re.IGNORECASE):
            cleaned = token.strip().strip(".;")
            if cleaned:
                authors_tokens.append(cleaned)

        parsed.append(
            {
                "primary_title": title,
                "title": title,
                "abstract": abstract,
                "authors": authors_tokens or [authors_part],
                "year": year,
                "publication_year": year,
                "journal_name": journal,
                "secondary_title": journal,
                "volume": volume,
                "issue": issue,
                "pages": pages,
                "doi": doi,
                "url": url,
            }
        )

    return parsed


def _strip_xml_namespaces(root):
    for elem in root.iter():
        if "}" in elem.tag:
            elem.tag = elem.tag.split("}", 1)[1]
    return root


def _parse_endnote_xml(payload: str) -> list[dict]:
    if not payload or not payload.strip():
        return []
    try:
        root = ET.fromstring(payload)
    except ET.ParseError:
        return []

    _strip_xml_namespaces(root)
    records = []

    for record in root.findall(".//record"):
        def find_text(*paths: str) -> str:
            for path in paths:
                text = record.findtext(path)
                if text:
                    return text.strip()
            return ""

        authors = [
            a.text.strip()
            for a in record.findall(".//contributors/authors/author")
            if a.text and a.text.strip()
        ]
        keywords = [
            k.text.strip()
            for k in record.findall(".//keywords/keyword")
            if k.text and k.text.strip()
        ]

        record_dict = {
            "title": find_text(".//titles/title", ".//title"),
            "journal_name": find_text(".//titles/secondary-title", ".//secondary-title"),
            "authors": authors,
            "year": find_text(".//dates/year", ".//year"),
            "publication_year": find_text(".//dates/year", ".//year"),
            "volume": find_text(".//volume"),
            "issue": find_text(".//issue", ".//number"),
            "pages": find_text(".//pages"),
            "doi": find_text(".//doi", ".//electronic-resource-num"),
            "url": find_text(".//urls/related-urls/url", ".//url"),
            "abstract": find_text(".//abstract"),
            "language": find_text(".//language"),
            "keywords": keywords,
            "_raw_source": ET.tostring(record, encoding="unicode"),
        }
        records.append(record_dict)

    return records


def _normalise_import_record(record: dict) -> dict | None:
    title = (
        _extract_reference_field(record, "primary_title")
        or _extract_reference_field(record, "title")
        or _extract_reference_field(record, "secondary_title")
    )
    if not title:
        return None

    authors_list = record.get("authors") or record.get("author") or []
    if isinstance(authors_list, str):
        authors_list = [authors_list]
    authors = "; ".join(str(a) for a in authors_list if a)

    year = _extract_reference_field(record, "year") or _extract_reference_field(
        record, "publication_year"
    )
    doi = _extract_reference_field(record, "doi")
    publication_year = _coerce_year(year)
    journal = _extract_reference_field(record, "journal_name") or _extract_reference_field(
        record, "secondary_title"
    )
    pages = _combine_pages(record)

    return {
        "title": title,
        "authors": authors,
        "year": year,
        "publication_year": publication_year,
        "journal": journal,
        "volume": _extract_reference_field(record, "volume"),
        "issue": _extract_reference_field(record, "issue"),
        "pages": pages,
        "doi": doi,
        "url": _extract_reference_field(record, "url"),
        "abstract": _extract_reference_field(record, "abstract"),
        "language": _extract_reference_field(record, "language"),
        "source_identifier": _extract_reference_field(record, "accession_number")
        or _extract_reference_field(record, "id"),
    }


def _advisory_board_context(
    project,
    *,
    user=None,
    member_form=None,
    reminder_form=None,
    protocol_form=None,
    feedback_close_form=None,
    action_list_form=None,
    action_list_feedback_close_form=None,
    synopsis_form=None,
    custom_field_form=None,
):
    members_qs = project.advisory_board_members.prefetch_related(
        "protocol_feedback", "action_list_feedback", "synopsis_feedback", "invitations"
    ).order_by("last_name", "first_name")
    accepted_members = list(members_qs.filter(response="Y"))
    declined_members = list(members_qs.filter(response="N"))
    pending_members = list(members_qs.exclude(response__in=["Y", "N"]))

    for collection in (accepted_members, declined_members, pending_members):
        for member in collection:
            member.latest_feedback = member.latest_protocol_feedback
            latest_action_feedback = member.latest_action_list_feedback
            member.latest_action_list_feedback_obj = latest_action_feedback
            latest_synopsis_feedback = member.latest_synopsis_feedback
            member.latest_synopsis_feedback_obj = latest_synopsis_feedback
            if (
                latest_action_feedback
                and not member.feedback_on_actions_received
                and (
                    latest_action_feedback.content
                    or latest_action_feedback.uploaded_document
                    or latest_action_feedback.submitted_at
                )
            ):
                member.feedback_on_actions_received = True
            invites = list(member.invitations.all())

            def _latest_timestamp(inv_list, accepted_value):
                filtered = [
                    inv
                    for inv in inv_list
                    if inv.accepted is accepted_value
                    and (inv.responded_at or inv.created_at)
                ]
                if not filtered:
                    return None
                filtered.sort(
                    key=lambda inv: inv.responded_at or inv.created_at, reverse=True
                )
                return filtered[0].responded_at or filtered[0].created_at

            member.declined_at = None
            member.accepted_at = None
            response_code = (member.response or "").upper()
            if response_code == "N":
                member.declined_at = _latest_timestamp(invites, False)
            elif response_code == "Y":
                member.accepted_at = _latest_timestamp(invites, True)
                if member.accepted_at is None and member.participation_confirmed_at:
                    member.accepted_at = member.participation_confirmed_at

    all_members = accepted_members + pending_members + declined_members
    declined_with_reason = [
        member
        for member in declined_members
        if (member.participation_statement or "").strip()
    ]
    accepted_with_statement = [
        member
        for member in accepted_members
        if (member.participation_statement or "").strip()
    ]

    custom_fields = list(
        AdvisoryBoardCustomField.objects.filter(project=project)
        .order_by("display_order", "name", "id")
    )
    fields_by_section = {
        key: [f for f in custom_fields if f.applies_to(key)]
        for key, _ in AdvisoryBoardCustomField.SECTION_CHOICES
    }
    grouped_fields_by_section = {
        key: AdvisoryBoardCustomField.group_fields_by_display(fields_by_section.get(key, []))
        for key, _ in AdvisoryBoardCustomField.SECTION_CHOICES
    }
    values_map = {
        (val.member_id, val.field_id): val.value
        for val in AdvisoryBoardCustomFieldValue.objects.filter(
            member__in=all_members, field__project=project
        )
    }

    for member in all_members:
        section_key = _member_section_key(member)
        member.section_key = section_key
        member.is_declined = (
            section_key == AdvisoryBoardCustomField.SECTION_DECLINED
        )
        section_fields = fields_by_section.get(section_key, [])
        member.custom_fields = section_fields
        member.has_custom_fields = bool(section_fields)
        member.custom_display_pairs = []
        member.custom_field_values = {}
        for field in section_fields:
            raw_value = values_map.get((member.id, field.id))
            formatted = field.format_value(raw_value)
            display_value = formatted if formatted not in (None, "") else "—"
            member.custom_display_pairs.append((field, display_value))
            member.custom_field_values[field.id] = display_value
        member.glance_statuses = _member_glance_statuses(member)

    direct_invites = project.invitations.filter(member__isnull=True).order_by(
        "-created_at"
    )

    not_invited_members = project.advisory_board_members.filter(invite_sent=False)
    pending_reminder_dates = [
        d
        for d in not_invited_members.filter(response_date__isnull=False)
        .order_by("response_date")
        .values_list("response_date", flat=True)
    ]
    if reminder_form is None:
        reminder_initial = {}
        if pending_reminder_dates:
            reminder_initial["reminder_date"] = pending_reminder_dates[0]
        else:
            reminder_initial["reminder_date"] = _default_invite_due_date()
        reminder_form = ReminderScheduleForm(initial=reminder_initial)

    protocol_members = project.advisory_board_members.filter(
        sent_protocol_at__isnull=False,
        response="Y",
    )
    protocol_pending_dates = [
        d
        for d in protocol_members.filter(feedback_on_protocol_deadline__isnull=False)
        .order_by("feedback_on_protocol_deadline")
        .values_list("feedback_on_protocol_deadline", flat=True)
    ]
    if protocol_form is None:
        protocol_initial = {}
        if protocol_pending_dates:
            first_deadline = protocol_pending_dates[0]
            try:
                protocol_initial["deadline"] = timezone.localtime(first_deadline)
            except (ValueError, TypeError):
                protocol_initial["deadline"] = first_deadline
        else:
            protocol_initial["deadline"] = timezone.localtime(
                _default_document_feedback_deadline()
            )
        protocol_form = ProtocolReminderScheduleForm(initial=protocol_initial)

    if member_form is None:
        member_form = AdvisoryBoardMemberForm()

    if custom_field_form is None:
        custom_field_form = AdvisoryCustomFieldForm(project)

    protocol_obj = getattr(project, "protocol", None)
    protocol_document_ready = bool(protocol_obj and getattr(protocol_obj, "document", None))
    if feedback_close_form is None:
        initial_close = {}
        if protocol_obj and protocol_obj.feedback_closure_message:
            initial_close["message"] = protocol_obj.feedback_closure_message
        feedback_close_form = ProtocolFeedbackCloseForm(initial=initial_close)
    protocol_feedback_state = {
        "protocol": protocol_obj,
        "is_closed": bool(getattr(protocol_obj, "feedback_closed_at", None)),
        "closed_at": getattr(protocol_obj, "feedback_closed_at", None),
        "closure_message": getattr(protocol_obj, "feedback_closure_message", ""),
        "deadline": protocol_pending_dates[0] if protocol_pending_dates else None,
        "document_ready": protocol_document_ready,
    }

    action_list_obj = getattr(project, "action_list", None)
    action_list_document_ready = bool(
        action_list_obj and getattr(action_list_obj, "document", None)
    )
    action_list_members = project.advisory_board_members.filter(
        sent_action_list_at__isnull=False,
        response="Y",
    )
    action_list_pending_dates = [
        d
        for d in action_list_members.filter(
            feedback_on_action_list_deadline__isnull=False
        )
        .order_by("feedback_on_action_list_deadline")
        .values_list("feedback_on_action_list_deadline", flat=True)
    ]
    if action_list_form is None:
        action_initial = {}
        if action_list_pending_dates:
            first_deadline = action_list_pending_dates[0]
            try:
                action_initial["deadline"] = timezone.localtime(first_deadline)
            except (ValueError, TypeError):
                action_initial["deadline"] = first_deadline
        else:
            action_initial["deadline"] = timezone.localtime(
                _default_document_feedback_deadline()
            )
        action_list_form = ActionListReminderScheduleForm(initial=action_initial)

    if action_list_feedback_close_form is None:
        action_close_initial = {}
        if action_list_obj and action_list_obj.feedback_closure_message:
            action_close_initial["message"] = action_list_obj.feedback_closure_message
        action_list_feedback_close_form = ActionListFeedbackCloseForm(
            initial=action_close_initial
        )
    action_list_feedback_state = {
        "action_list": action_list_obj,
        "is_closed": bool(getattr(action_list_obj, "feedback_closed_at", None)),
        "closed_at": getattr(action_list_obj, "feedback_closed_at", None),
        "closure_message": getattr(action_list_obj, "feedback_closure_message", ""),
        "deadline": action_list_pending_dates[0] if action_list_pending_dates else None,
        "document_ready": action_list_document_ready,
    }

    synopsis_members = project.advisory_board_members.filter(
        sent_synopsis_at__isnull=False,
        response="Y",
    )
    synopsis_pending_dates = [
        d
        for d in synopsis_members.filter(feedback_on_synopsis_deadline__isnull=False)
        .order_by("feedback_on_synopsis_deadline")
        .values_list("feedback_on_synopsis_deadline", flat=True)
    ]
    if synopsis_form is None:
        synopsis_initial = {}
        if synopsis_pending_dates:
            first_deadline = synopsis_pending_dates[0]
            try:
                synopsis_initial["deadline"] = timezone.localtime(first_deadline)
            except (ValueError, TypeError):
                synopsis_initial["deadline"] = first_deadline
        else:
            synopsis_initial["deadline"] = timezone.localtime(
                _default_document_feedback_deadline()
            )
        synopsis_form = SynopsisReminderScheduleForm(initial=synopsis_initial)

    section_palette = {
        AdvisoryBoardCustomField.SECTION_ACCEPTED: {
            "title": "Accepted members",
            "empty": "No accepted members yet.",
            "card": "border-start border-4 border-success",
            "header": "bg-success text-white",
        },
        AdvisoryBoardCustomField.SECTION_PENDING: {
            "title": "Pending members",
            "empty": "No pending members yet.",
            "card": "border-start border-4 border-warning",
            "header": "bg-warning",
        },
        AdvisoryBoardCustomField.SECTION_DECLINED: {
            "title": "Declined members",
            "empty": "No declined members yet.",
            "card": "border-start border-4 border-secondary",
            "header": "bg-secondary text-white",
        },
    }

    member_sections = [
        {
            "key": AdvisoryBoardCustomField.SECTION_ACCEPTED,
            "title": section_palette[AdvisoryBoardCustomField.SECTION_ACCEPTED][
                "title"
            ],
            "members": accepted_members,
            "empty_text": section_palette[AdvisoryBoardCustomField.SECTION_ACCEPTED][
                "empty"
            ],
            "card_class": section_palette[AdvisoryBoardCustomField.SECTION_ACCEPTED][
                "card"
            ],
            "header_class": section_palette[AdvisoryBoardCustomField.SECTION_ACCEPTED][
                "header"
            ],
            "fields": fields_by_section.get(
                AdvisoryBoardCustomField.SECTION_ACCEPTED, []
            ),
            "fields_by_group": grouped_fields_by_section.get(
                AdvisoryBoardCustomField.SECTION_ACCEPTED,
                AdvisoryBoardCustomField.group_fields_by_display([]),
            ),
        },
        {
            "key": AdvisoryBoardCustomField.SECTION_PENDING,
            "title": section_palette[AdvisoryBoardCustomField.SECTION_PENDING]["title"],
            "members": pending_members,
            "empty_text": section_palette[AdvisoryBoardCustomField.SECTION_PENDING][
                "empty"
            ],
            "card_class": section_palette[AdvisoryBoardCustomField.SECTION_PENDING][
                "card"
            ],
            "header_class": section_palette[AdvisoryBoardCustomField.SECTION_PENDING][
                "header"
            ],
            "fields": fields_by_section.get(
                AdvisoryBoardCustomField.SECTION_PENDING, []
            ),
            "fields_by_group": grouped_fields_by_section.get(
                AdvisoryBoardCustomField.SECTION_PENDING,
                AdvisoryBoardCustomField.group_fields_by_display([]),
            ),
        },
        {
            "key": AdvisoryBoardCustomField.SECTION_DECLINED,
            "title": section_palette[AdvisoryBoardCustomField.SECTION_DECLINED][
                "title"
            ],
            "members": declined_members,
            "empty_text": section_palette[AdvisoryBoardCustomField.SECTION_DECLINED][
                "empty"
            ],
            "card_class": section_palette[AdvisoryBoardCustomField.SECTION_DECLINED][
                "card"
            ],
            "header_class": section_palette[AdvisoryBoardCustomField.SECTION_DECLINED][
                "header"
            ],
            "fields": fields_by_section.get(
                AdvisoryBoardCustomField.SECTION_DECLINED, []
            ),
            "fields_by_group": grouped_fields_by_section.get(
                AdvisoryBoardCustomField.SECTION_DECLINED,
                AdvisoryBoardCustomField.group_fields_by_display([]),
            ),
        },
    ]
    for section in member_sections:
        section["has_fields"] = bool(section["fields"])

    group_names = [
        "personal",
        "invitation",
        "action",
        "protocol",
        "synopsis",
        "custom",
    ]
    combined_fields_by_group = {name: [] for name in group_names}
    seen_field_ids = {name: set() for name in group_names}

    for section in member_sections:
        field_ids_by_group = {}
        for group_name in group_names:
            group_fields = section["fields_by_group"].get(group_name, [])
            field_ids = [field.id for field in group_fields]
            field_ids_by_group[group_name] = field_ids
            for field in group_fields:
                if field.id not in seen_field_ids[group_name]:
                    combined_fields_by_group[group_name].append(field)
                    seen_field_ids[group_name].add(field.id)
        section["field_ids_by_group"] = field_ids_by_group

    def safe_count(members):
        count_attr = getattr(members, "count", None)
        if callable(count_attr):
            try:
                return count_attr()
            except TypeError:
                pass
        return len(members)

    total_member_count = (
        safe_count(accepted_members)
        + safe_count(pending_members)
        + safe_count(declined_members)
    )

    can_edit_members = _user_can_edit_project(user, project) if user else False
    protocol_feedback_members = [
        member
        for member in all_members
        if getattr(getattr(member, "latest_feedback", None), "submitted_at", None)
    ]
    action_list_feedback_members = [
        member
        for member in all_members
        if getattr(
            getattr(member, "latest_action_list_feedback_obj", None), "submitted_at", None
        )
    ]
    synopsis_feedback_members = [
        member
        for member in all_members
        if getattr(
            getattr(member, "latest_synopsis_feedback_obj", None), "submitted_at", None
        )
    ]
    status_badges = {
        AdvisoryBoardCustomField.SECTION_ACCEPTED: {
            "label": "Accepted",
            "badge_class": "text-bg-success",
            "row_class": "ab-row-status-accepted",
        },
        AdvisoryBoardCustomField.SECTION_PENDING: {
            "label": "Pending",
            "badge_class": "text-bg-warning",
            "row_class": "ab-row-status-pending",
        },
        AdvisoryBoardCustomField.SECTION_DECLINED: {
            "label": "Declined",
            "badge_class": "text-bg-danger",
            "row_class": "ab-row-status-declined",
        },
    }

    return {
        "project": project,
        "accepted_members": accepted_members,
        "accepted_members_with_statement": accepted_with_statement,
        "declined_members": declined_members,
        "pending_members": pending_members,
        "protocol_feedback_members": protocol_feedback_members,
        "action_list_feedback_members": action_list_feedback_members,
        "synopsis_feedback_members": synopsis_feedback_members,
        "member_sections": member_sections,
        "combined_fields_by_group": combined_fields_by_group,
        "member_status_badges": status_badges,
        "total_member_count": total_member_count,
        "section_fields": fields_by_section,
        "custom_fields": custom_fields,
        "custom_field_form": custom_field_form,
        "direct_invites": direct_invites,
        "form": member_form,
        "reminder_form": reminder_form,
        "pending_reminders": not_invited_members.count(),
        "pending_reminder_dates": pending_reminder_dates,
        "initial_reminder_log": project.change_log.filter(action="Scheduled reminders")
        .order_by("created_at")
        .first(),
        "protocol_reminder_form": protocol_form,
        "protocol_pending_count": protocol_members.count(),
        "protocol_pending_dates": protocol_pending_dates,
        "initial_protocol_reminder_log": project.change_log.filter(
            action="Scheduled protocol reminders"
        )
        .order_by("created_at")
        .first(),
        "protocol_feedback_state": protocol_feedback_state,
        "protocol_feedback_close_form": feedback_close_form,
        "action_list_reminder_form": action_list_form,
        "action_list_pending_count": action_list_members.count(),
        "action_list_pending_dates": action_list_pending_dates,
        "initial_action_list_reminder_log": project.change_log.filter(
            action="Scheduled action list reminders"
        )
        .order_by("created_at")
        .first(),
        "can_edit_members": can_edit_members,
        "action_list_feedback_state": action_list_feedback_state,
        "action_list_feedback_close_form": action_list_feedback_close_form,
        "synopsis_reminder_form": synopsis_form,
        "synopsis_pending_count": synopsis_members.count(),
        "synopsis_pending_dates": synopsis_pending_dates,
        "initial_synopsis_reminder_log": project.change_log.filter(
            action="Scheduled synopsis reminders"
        )
        .order_by("created_at")
        .first(),
        "section_palette": section_palette,
        "custom_field_group_choices": AdvisoryBoardCustomField.DISPLAY_GROUP_CHOICES,
        "declined_members_with_reason": declined_with_reason,
        "minimum_allowed_deadline_date": minimum_allowed_deadline_date(),
    }


# ---------------- Dashboard & Project Hub ----------------


@login_required
def dashboard(request):
    base_qs = Project.objects.prefetch_related("userrole_set__user").order_by(
        "-created_at"
    )
    if is_external_author_user(request.user):
        base_qs = base_qs.filter(userrole__user=request.user, userrole__role="author")
    base_qs = base_qs.distinct()
    completed_statuses = ["completed", "archived"]
    active_projects = list(base_qs.exclude(status__in=completed_statuses))
    completed_projects = list(base_qs.filter(status__in=completed_statuses))

    for proj in active_projects + completed_projects:
        proj.author_list = [
            role.user for role in proj.userrole_set.all() if role.role == "author"
        ]
        proj.can_edit_project = _user_can_edit_project(request.user, proj)
        proj.can_manage_project_configuration = _user_can_manage_project_configuration(
            request.user, proj
        )
    return render(
        request,
        "synopsis/dashboard.html",
        {
            "active_projects": active_projects,
            "completed_projects": completed_projects,
            "can_manage_projects": _user_is_manager(request.user),
            "app_release_label": settings.APP_RELEASE_LABEL,
        },
    )


class ProjectCreateForm(forms.ModelForm):
    start_date = forms.DateField(
        required=False,
        initial=timezone.localdate,
        widget=forms.DateInput(
            attrs={
                "type": "date",
                "class": "form-control bg-light text-muted",
                "readonly": "readonly",
                "tabindex": "-1",
            }
        ),
        disabled=True,
    )

    class Meta:
        model = Project
        fields = ["title", "description", "start_date"]
        widgets = {
            "title": forms.TextInput(attrs={"class": "form-control"}),
            "description": forms.Textarea(
                attrs={
                    "class": "form-control",
                    "rows": 4,
                    "placeholder": "Optional short description of the synopsis",
                }
            ),
        }

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.fields["start_date"].initial = timezone.localdate()

    def clean_description(self):
        return self.cleaned_data.get("description", "").strip()


@login_required
def project_create(request):
    if is_external_author_user(request.user):
        messages.error(
            request,
            "External author accounts cannot create new synopses.",
        )
        return redirect("synopsis:dashboard")

    today = timezone.localdate()
    contact_instance = Funder()
    if request.method == "POST":
        pform = ProjectCreateForm(request.POST)
        aform = AssignAuthorsForm(request.POST)
        fform = FunderForm(request.POST)
        contact_formset = FunderContactFormSet(
            request.POST, prefix="contacts", instance=contact_instance
        )
        if (
            pform.is_valid()
            and aform.is_valid()
            and fform.is_valid()
            and contact_formset.is_valid()
        ):
            if request.POST.get("edit") == "1":
                pform.fields["start_date"].initial = today
                return render(
                    request,
                    "synopsis/project_create.html",
                    {
                        "form": pform,
                        "authors_form": aform,
                        "funder_form": fform,
                        "contact_formset": contact_formset,
                    },
                )

            if request.POST.get("confirm") == "1":
                project = pform.save(commit=False)
                project.start_date = today
                project.save()

                _log_project_change(
                    project,
                    request.user,
                    "Project created",
                    "Title: "
                    f"{project.title}; Description: {_format_value(project.description)}; "
                    f"Start date: {_format_value(project.start_date)}",
                )

                authors = aform.cleaned_data.get("authors") or []
                for user in authors:
                    UserRole.objects.get_or_create(
                        user=user, project=project, role="author"
                    )
                if authors:
                    author_labels = ", ".join(_user_display(user) for user in authors)
                    _log_project_change(
                        project,
                        request.user,
                        "Assigned authors",
                        f"Initial author list: {author_labels}",
                    )

                if fform.has_meaningful_input() or _formset_has_contacts(
                    contact_formset
                ):
                    funder = fform.save(commit=False)
                    funder.project = project
                    funder.save()
                    contact_formset.instance = funder
                    contact_formset.save()
                    funder.update_cached_contact_fields()
                    primary = funder.primary_contact()
                    contact_summary = _contact_summary_text(funder.contacts.all())
                    details = (
                        f"Organisation: {_format_value(funder.organisation)}; "
                        f"Organisation details: {_format_value(funder.organisation_details)}; "
                        f"Primary contact: {_funder_contact_label(primary)}; "
                        f"Contacts: {contact_summary}; "
                        f"Funds allocated: {_format_value(funder.funds_allocated)}; "
                        f"Dates: {_format_value(funder.fund_start_date)} to {_format_value(funder.fund_end_date)}"
                    )
                    _log_project_change(project, request.user, "Added funder", details)

                messages.success(request, "Project created.")
                return redirect("synopsis:project_hub", project_id=project.id)

            hidden_project_form = ProjectCreateForm(request.POST)
            for field in hidden_project_form.fields.values():
                field.widget = forms.HiddenInput()
                field.disabled = False
            hidden_authors_form = AssignAuthorsForm(request.POST)
            hidden_authors_form.fields["authors"].widget = forms.MultipleHiddenInput()

            hidden_funder_form = FunderForm(request.POST)
            for name, field in hidden_funder_form.fields.items():
                field.widget = forms.HiddenInput()

            hidden_contact_formset = FunderContactFormSet(
                request.POST, prefix="contacts", instance=contact_instance
            )
            for form in hidden_contact_formset.forms:
                for field in form.fields.values():
                    field.widget = forms.HiddenInput()

            authors = aform.cleaned_data.get("authors") or []
            author_names = [_user_display(user) for user in authors]

            funder_cleaned = fform.cleaned_data
            contact_entries = _contact_entries_from_formset(contact_formset)
            funder_summary = {
                "organisation": funder_cleaned.get("organisation"),
                "organisation_details": funder_cleaned.get("organisation_details"),
                "funds_allocated": funder_cleaned.get("funds_allocated"),
                "fund_start_date": funder_cleaned.get("fund_start_date"),
                "fund_end_date": funder_cleaned.get("fund_end_date"),
                "contacts": contact_entries,
                "has_details": fform.has_meaningful_input() or bool(contact_entries),
            }

            return render(
                request,
                "synopsis/project_create_confirm.html",
                {
                    "project_form": hidden_project_form,
                    "authors_form_hidden": hidden_authors_form,
                    "funder_form_hidden": hidden_funder_form,
                    "contact_formset_hidden": hidden_contact_formset,
                    "summary": {
                        "title": pform.cleaned_data["title"],
                        "description": pform.cleaned_data["description"],
                        "start_date": today,
                        "authors": author_names,
                        "funder": funder_summary,
                    },
                },
            )
    else:
        pform = ProjectCreateForm(initial={"start_date": today})
        aform = AssignAuthorsForm()
        fform = FunderForm()
        contact_formset = FunderContactFormSet(prefix="contacts", instance=contact_instance)

    return render(
        request,
        "synopsis/project_create.html",
        {
            "form": pform,
            "authors_form": aform,
            "funder_form": fform,
            "contact_formset": contact_formset,
        },
    )


def _project_phase_context(project: Project, user):
    phase_labels = dict(Project.PHASE_CHOICES)
    order = project.available_phase_keys()
    current_phase = project.phase
    try:
        current_index = order.index(current_phase)
    except ValueError:
        current_index = 0
    next_phase = order[current_index + 1] if current_index + 1 < len(order) else None
    last_event = project.phase_events.first()
    can_update_phase = _user_can_confirm_phase(user, project)
    total_phases = len(order)
    current_step = current_index + 1 if total_phases else 0
    phase_progress_percent = (
        int(round((current_step / total_phases) * 100)) if total_phases else 0
    )
    phase_steps = []
    for step_index, key in enumerate(order, start=1):
        phase_steps.append(
            {
                "key": key,
                "label": phase_labels.get(key, key),
                "number": step_index,
                "is_current": key == current_phase,
                "is_manual": key == project.phase_manual,
                "is_complete": step_index < current_step,
                "is_upcoming": step_index > current_step,
                "can_set": can_update_phase and key != current_phase,
            }
        )

    return {
        "phase_labels": phase_labels,
        "current_phase": current_phase,
        "current_phase_label": phase_labels.get(current_phase, current_phase),
        "phase_is_manual": bool(project.phase_manual and project.phase_manual in order),
        "phase_manual_updated": project.phase_manual_updated,
        "phase_steps": phase_steps,
        "phase_progress_percent": phase_progress_percent,
        "phase_progress_text": f"Step {current_step} of {total_phases}",
        "next_phase": next_phase,
        "next_phase_label": phase_labels.get(next_phase) if next_phase else None,
        "last_phase_event": last_event,
        "can_update_phase": can_update_phase,
    }


@login_required
def project_hub(request, project_id):
    project = get_object_or_404(
        Project.objects.prefetch_related(
            "funders",
            Prefetch("funders__contacts"),
            "userrole_set__user",
            "change_log__changed_by",
            "phase_events",
        ),
        pk=project_id,
    )
    if not _user_can_view_project(request.user, project):
        messages.error(request, "You do not have access to that synopsis.")
        return redirect("synopsis:dashboard")
    protocol = getattr(project, "protocol", None)
    action_list = getattr(project, "action_list", None)
    can_manage = _user_is_manager(request.user)

    funders = list(project.funders.all())
    def _funder_sort_key(funder):
        primary = funder.primary_contact()
        last = primary.last_name if primary else ""
        first = primary.first_name if primary else ""
        return (
            funder.fund_start_date or dt.date.max,
            (funder.organisation or last or first or "").lower(),
        )

    funders.sort(key=_funder_sort_key)
    funding_values = [
        f.funds_allocated for f in funders if f.funds_allocated is not None
    ]
    total_funding = sum(funding_values, Decimal("0")) if funding_values else None
    start_dates = [f.fund_start_date for f in funders if f.fund_start_date]
    end_dates = [f.fund_end_date for f in funders if f.fund_end_date]
    funder_summary = {
        "count": len(funders),
        "total": total_funding,
        "earliest_start": min(start_dates) if start_dates else None,
        "latest_end": max(end_dates) if end_dates else None,
    }

    inv_qs = AdvisoryBoardInvitation.objects.filter(project=project)
    members_qs = AdvisoryBoardMember.objects.filter(project=project)
    ab_stats = {
        "members": members_qs.count(),
        "member_invites_sent": members_qs.filter(invite_sent=True).count(),
        "protocol_sent_to_members": members_qs.filter(
            sent_protocol_at__isnull=False
        ).count(),
        "invites_total": inv_qs.count(),
        "invites_member": inv_qs.filter(member__isnull=False).count(),
        "invites_direct": inv_qs.filter(member__isnull=True).count(),
        "accepted": inv_qs.filter(accepted=True).count(),
        "declined": inv_qs.filter(accepted=False).count(),
        "pending": inv_qs.filter(accepted__isnull=True).count(),
    }
    ab_member_updates = list(
        members_qs.order_by("-invite_sent_at", "-response_date", "last_name")[:6]
    )

    latest_batch = project.reference_batches.order_by("-created_at", "-id").first()
    reference_stats = {
        "batches": project.reference_batches.count(),
        "references": project.references.count(),
        "pending": project.references.filter(screening_status="pending").count(),
        "latest_batch": latest_batch,
    }

    summary_qs = ReferenceSummary.objects.filter(project=project)
    summary_total = summary_qs.count()
    summary_counts = {
        row["status"]: row["count"]
        for row in summary_qs.values("status").annotate(count=Count("id"))
    }
    summary_stats = {
        "total": summary_total,
        "todo": summary_counts.get(ReferenceSummary.STATUS_TODO, 0),
        "draft": summary_counts.get(ReferenceSummary.STATUS_DRAFT, 0),
        "review": summary_counts.get(ReferenceSummary.STATUS_REVIEW, 0),
        "done": summary_counts.get(ReferenceSummary.STATUS_DONE, 0),
        "needs_help": summary_qs.filter(needs_help=True).count(),
        "unassigned": summary_qs.filter(assigned_to__isnull=True).count(),
    }

    structure_stats = {
        "chapters": project.synopsis_chapters.count(),
        "subheadings": SynopsisSubheading.objects.filter(
            chapter__project=project
        ).count(),
        "interventions": SynopsisIntervention.objects.filter(
            subheading__chapter__project=project
        ).count(),
        "assignments": SynopsisAssignment.objects.filter(
            intervention__subheading__chapter__project=project
        ).count(),
        "summaries_mapped": SynopsisAssignment.objects.filter(
            intervention__subheading__chapter__project=project
        )
        .values("reference_summary_id")
        .distinct()
        .count(),
    }

    phase_context = _project_phase_context(project, request.user)

    change_log_entries = project.change_log.select_related("changed_by")[:10]

    return render(
        request,
        "synopsis/project_hub.html",
        {
            "project": project,
            "protocol": protocol,
            "action_list": action_list,
            "ab_stats": ab_stats,
            "ab_member_updates": ab_member_updates,
            "reference_stats": reference_stats,
            "summary_stats": summary_stats,
            "structure_stats": structure_stats,
            "authors": list(project.author_users),
            "change_log_entries": change_log_entries,
            "funders": funders,
            "funder_summary": funder_summary,
            "can_manage_project": _user_is_manager(request.user),
            "can_edit_project": _user_can_edit_project(request.user, project),
            "can_manage_project_configuration": _user_can_manage_project_configuration(
                request.user, project
            ),
            **phase_context,
        },
    )


def _user_can_confirm_phase(user, project: Project) -> bool:
    return _user_can_manage_project_configuration(user, project)


@login_required
def project_phase_confirm(request, project_id, phase):
    project = get_object_or_404(Project, pk=project_id)
    if request.method != "POST":
        return HttpResponseBadRequest("POST required")
    if not _user_can_confirm_phase(request.user, project):
        messages.error(
            request,
            "You do not have permission to confirm milestones for this project.",
        )
        return redirect("synopsis:project_hub", project_id=project.id)

    phase_labels = dict(Project.PHASE_CHOICES)
    valid_phases = project.available_phase_keys()
    if phase not in valid_phases:
        messages.error(request, "That phase is not available for this project.")
        return redirect("synopsis:project_hub", project_id=project.id)

    current_phase = project.phase
    if current_phase == phase and project.phase_manual == phase:
        messages.info(
            request,
            f"The current working phase is already set to {phase_labels[phase]}.",
        )
        return redirect("synopsis:project_hub", project_id=project.id)

    project.phase_manual = phase
    project.phase_manual_updated = timezone.now()
    project.save(update_fields=["phase_manual", "phase_manual_updated"])

    note = (request.POST.get("note") or "").strip()
    ProjectPhaseEvent.objects.create(
        project=project,
        phase=phase,
        confirmed_by=request.user,
        note=note
        or (
            f"Phase changed from {phase_labels.get(current_phase, current_phase)} to {phase_labels[phase]}."
        ),
    )
    _log_project_change(
        project,
        request.user,
        "Updated project phase",
        note
        or (
            f"{phase_labels.get(current_phase, current_phase)} → {phase_labels[phase]}"
        ),
    )

    messages.success(request, f"Current phase set to: {phase_labels[phase]}")
    return redirect("synopsis:project_hub", project_id=project.id)


@login_required
def project_authors_manage(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if not _user_is_manager(request.user):
        messages.error(request, "Only managers can update authors for this project.")
        return redirect("synopsis:project_hub", project_id=project.id)

    current_authors_qs = project.author_users
    current_ids = set(current_authors_qs.values_list("id", flat=True))

    if request.method == "POST":
        form = AssignAuthorsForm(request.POST)
        if form.is_valid():
            new_authors = form.cleaned_data.get("authors") or User.objects.none()
            new_ids = set(new_authors.values_list("id", flat=True))
            added_ids = new_ids - current_ids
            removed_ids = current_ids - new_ids

            for user_id in added_ids:
                UserRole.objects.get_or_create(
                    user_id=user_id, project=project, role="author"
                )
            if removed_ids:
                UserRole.objects.filter(
                    project=project, role="author", user_id__in=removed_ids
                ).delete()

            if added_ids or removed_ids:
                added_names = ", ".join(
                    _user_display(user)
                    for user in User.objects.filter(id__in=added_ids)
                )
                removed_names = ", ".join(
                    _user_display(user)
                    for user in User.objects.filter(id__in=removed_ids)
                )
                fragments = []
                if added_names:
                    fragments.append(f"Added: {added_names}")
                if removed_names:
                    fragments.append(f"Removed: {removed_names}")
                _log_project_change(
                    project,
                    request.user,
                    "Updated authors",
                    "; ".join(fragments) or "No membership changes",
                )
                messages.success(request, "Author assignments updated.")
            else:
                messages.info(request, "No changes made to authors.")

            return redirect("synopsis:project_hub", project_id=project.id)
    else:
        form = AssignAuthorsForm(initial={"authors": current_authors_qs})

    return render(
        request,
        "synopsis/project_authors_form.html",
        {
            "project": project,
            "form": form,
            "current_authors": current_authors_qs,
        },
    )


@login_required
def project_funder_add(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if not _user_is_manager(request.user):
        messages.error(request, "Only managers can add funder information.")
        return redirect("synopsis:project_hub", project_id=project.id)

    instance = Funder(project=project)
    funders = list(project.funders.prefetch_related("contacts"))
    def _sort_funder(f):
        primary = f.primary_contact()
        last = primary.last_name if primary else ""
        first = primary.first_name if primary else ""
        return (
            f.fund_start_date or dt.date.max,
            (f.organisation or last or first or "").lower(),
        )
    funders.sort(key=_sort_funder)
    if request.method == "POST":
        form = FunderForm(request.POST, instance=instance)
        contact_formset = FunderContactFormSet(
            request.POST, prefix="contacts", instance=instance
        )
        if form.is_valid() and contact_formset.is_valid():
            if not form.has_meaningful_input() and not _formset_has_contacts(
                contact_formset
            ):
                form.add_error(None, "Enter details before saving a funder.")
            else:
                funder = form.save(commit=False)
                funder.project = project
                funder.save()
                contact_formset.instance = funder
                contact_formset.save()
                funder.update_cached_contact_fields()
                primary = funder.primary_contact()
                contact_summary = _contact_summary_text(funder.contacts.all())
                details = (
                    f"Organisation: {_format_value(funder.organisation)}; "
                    f"Organisation details: {_format_value(funder.organisation_details)}; "
                    f"Primary contact: {_funder_contact_label(primary)}; "
                    f"Contacts: {contact_summary}; "
                    f"Funds allocated: {_format_value(funder.funds_allocated)}; "
                    f"Start date: {_format_value(funder.fund_start_date)}; "
                    f"End date: {_format_value(funder.fund_end_date)}"
                )
                _log_project_change(project, request.user, "Added funder", details)
                messages.success(request, "Funder details added.")
                return redirect("synopsis:project_hub", project_id=project.id)
    else:
        form = FunderForm(instance=instance)
        contact_formset = FunderContactFormSet(prefix="contacts", instance=instance)

    return render(
        request,
        "synopsis/funder_form.html",
        {
            "project": project,
            "form": form,
            "contact_formset": contact_formset,
            "existing_funders": funders,
            "funder": None,
            "mode": "add",
        },
    )


@login_required
def project_funder_edit(request, project_id, funder_id):
    project = get_object_or_404(Project, pk=project_id)
    funder = get_object_or_404(Funder, pk=funder_id, project=project)
    if not _user_is_manager(request.user):
        messages.error(request, "Only managers can edit funder information.")
        return redirect("synopsis:project_hub", project_id=project.id)

    if request.method == "POST":
        form = FunderForm(request.POST, instance=funder)
        contact_formset = FunderContactFormSet(
            request.POST, prefix="contacts", instance=funder
        )
        if form.is_valid() and contact_formset.is_valid():
            old_values = {
                field: getattr(funder, field)
                for field in (
                    "organisation",
                    "organisation_details",
                    "funds_allocated",
                    "fund_start_date",
                    "fund_end_date",
                )
            }
            old_contacts = list(funder.contacts.all())
            updated = form.save(commit=False)
            updated.project = project
            updated.save()
            contact_formset.instance = updated
            contact_formset.save()
            updated.update_cached_contact_fields()
            changes = []
            for field in (
                "organisation",
                "funds_allocated",
                "fund_start_date",
                "fund_end_date",
            ):
                old_value = old_values[field]
                new_value = getattr(updated, field)
                if old_value != new_value:
                    label = field.replace("_", " ").title()
                    changes.append(
                        f"{label}: {_format_value(old_value)} → {_format_value(new_value)}"
                    )
            new_contacts = list(updated.contacts.all())
            if _contact_summary_text(old_contacts) != _contact_summary_text(
                new_contacts
            ):
                changes.append(
                    f"Contacts: {_contact_summary_text(old_contacts)} → {_contact_summary_text(new_contacts)}"
                )
            detail_msg = "; ".join(changes) if changes else "No visible field changes"
            _log_project_change(project, request.user, "Updated funder", detail_msg)
            messages.success(request, "Funder details updated.")
            return redirect("synopsis:project_funder_add", project_id=project.id)
    else:
        form = FunderForm(instance=funder)
        contact_formset = FunderContactFormSet(prefix="contacts", instance=funder)

    return render(
        request,
        "synopsis/funder_form.html",
        {
            "project": project,
            "form": form,
            "contact_formset": contact_formset,
            "funder": funder,
            "mode": "edit",
        },
    )


@login_required
def project_funder_delete(request, project_id, funder_id):
    project = get_object_or_404(Project, pk=project_id)
    funder = get_object_or_404(Funder, pk=funder_id, project=project)
    if not _user_is_manager(request.user):
        messages.error(request, "Only managers can remove funders.")
        return redirect("synopsis:project_hub", project_id=project.id)

    if request.method == "POST":
        detail = "Removed funder " + funder.contact_display_name()
        funder.delete()
        _log_project_change(project, request.user, "Removed funder", detail)
        messages.success(request, "Funder removed.")
        return redirect("synopsis:project_funder_add", project_id=project.id)

    return render(
        request,
        "synopsis/funder_confirm_delete.html",
        {"project": project, "funder": funder},
    )


@login_required
def project_delete(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if not _user_is_manager(request.user):
        messages.error(request, "Only managers can delete projects.")
        return redirect("synopsis:project_hub", project_id=project.id)

    next_url = request.POST.get("next") or request.GET.get("next")
    cancel_url = next_url or reverse("synopsis:manager_dashboard")

    if request.method == "POST":
        form = ProjectDeleteForm(request.POST, project=project)
        if form.is_valid():
            title = project.title
            project.delete()
            messages.success(request, f"Project '{title}' deleted.")
            return redirect(next_url or "synopsis:manager_dashboard")
    else:
        form = ProjectDeleteForm(project=project)

    return render(
        request,
        "synopsis/project_confirm_delete.html",
        {
            "project": project,
            "form": form,
            "next_url": next_url,
            "cancel_url": cancel_url,
        },
    )


@login_required
def protocol_detail(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if not _user_can_view_project(request.user, project):
        messages.error(request, "You do not have access to that synopsis.")
        return redirect("synopsis:dashboard")
    if not project.protocol_relevant:
        messages.info(
            request,
            "Protocol is marked as not relevant for this project. Update Project settings if you want to use the protocol workflow.",
        )
        return redirect("synopsis:project_hub", project_id=project.id)
    protocol = getattr(project, "protocol", None)
    can_manage = _user_is_manager(request.user)
    can_edit_documents = _user_can_edit_project(request.user, project)
    can_delete_documents = _user_can_manage_project_configuration(
        request.user, project
    )
    protocol_history_queryset = (
        project.change_log.filter(action__icontains="protocol")
        .select_related("changed_by")
        .order_by("-created_at", "-id")
    )
    protocol_history_entries = []
    for log in protocol_history_queryset:
        segments = [
            segment.strip() for segment in log.details.split("|") if segment.strip()
        ]
        reason = ""
        changes = []
        for segment in segments:
            if segment.lower().startswith("reason:"):
                reason = segment.split(":", 1)[1].strip()
            else:
                changes.append(segment)
        protocol_history_entries.append(
            {
                "log": log,
                "changes": changes,
                "reason": reason,
                "actor": _user_display(log.changed_by) if log.changed_by else "System",
            }
        )

    revision_entries = []
    if protocol:
        revision_queryset = protocol.revisions.select_related("uploaded_by")
        for revision in revision_queryset:
            file_name = revision.original_name or os.path.basename(revision.file.name)
            try:
                download_url = revision.file.url
            except ValueError:
                download_url = ""
            revision_entries.append(
                {
                    "revision": revision,
                    "is_current": protocol.current_revision_id == revision.id,
                    "file_name": file_name,
                    "file_size": _format_file_size(revision.file_size),
                    "uploaded_by": (
                        _user_display(revision.uploaded_by)
                        if revision.uploaded_by
                        else "—"
                    ),
                    "download_url": download_url,
                    "can_mark_final": can_manage
                    and (
                        protocol.stage != "final"
                        or protocol.current_revision_id != revision.id
                    ),
                    "version_label": revision.version_label or "",
                }
            )

    current_revision_entry = next(
        (entry for entry in revision_entries if entry["is_current"]), None
    )
    current_revision_download_url = ""
    if current_revision_entry:
        try:
            current_revision_download_url = current_revision_entry["revision"].file.url
        except ValueError:
            current_revision_download_url = ""

    protocol_document_ready = bool(protocol and getattr(protocol, "document", None))
    existing_file_name = protocol.document.name if protocol_document_ready else ""
    has_existing_file = bool(existing_file_name or revision_entries)
    first_upload_pending = not has_existing_file
    final_stage_locked = bool(
        protocol and protocol.stage == "final" and has_existing_file
    )

    collaborative_enabled = _onlyoffice_enabled()
    protocol_closed = bool(protocol and protocol.feedback_closed_at)
    collaborative_session = None
    collaborative_resume_url = ""
    collaborative_force_end_url = ""
    collaborative_slug = _document_type_slug(CollaborativeSession.DOCUMENT_PROTOCOL)
    if collaborative_enabled:
        collaborative_session = _get_active_collaborative_session(
            project, CollaborativeSession.DOCUMENT_PROTOCOL
        )
        if collaborative_session and protocol_closed:
            collaborative_session.mark_inactive(
                ended_by=request.user if request.user.is_authenticated else None,
                reason="Protocol feedback window closed",
            )
            collaborative_session = None
        if collaborative_session:
            collaborative_resume_url = reverse(
                "synopsis:collaborative_edit",
                args=[project.id, collaborative_slug, collaborative_session.token],
            )
            collaborative_force_end_url = reverse(
                "synopsis:collaborative_force_end",
                args=[project.id, collaborative_slug, collaborative_session.token],
            )
    if collaborative_session:
        collaborative_can_override = _user_can_force_end_session(
            request.user, project, collaborative_session
        )
    else:
        collaborative_can_override = False
    if protocol_closed:
        collaborative_enabled = False
        collaborative_session = None
        collaborative_resume_url = ""
        collaborative_force_end_url = ""
        collaborative_can_override = False

    if request.method == "POST":
        form = ProtocolUpdateForm(request.POST, request.FILES, instance=protocol)
        if form.is_valid():
            new_stage = form.cleaned_data.get("stage")
            uploaded_file = form.cleaned_data.get("document")
            reason = form.cleaned_data.get("change_reason", "")
            version_label = form.cleaned_data.get("version_label", "")

            is_new_protocol = protocol is None
            stage_changed = bool(protocol) and protocol.stage != new_stage
            replacing_file = bool(uploaded_file)
            previous_label = (
                (protocol.current_revision.version_label or "")
                if protocol and protocol.current_revision
                else ""
            )

            if final_stage_locked and new_stage == "final" and replacing_file:
                form.add_error(
                    "document",
                    "Finalized protocols cannot be replaced. Switch the stage back to Draft to revise the document.",
                )

            active_file_missing = not bool(protocol and protocol.document)
            if active_file_missing and not replacing_file:
                form.add_error(
                    "document",
                    "Choose a protocol file to upload. You can reuse the same filename as a file you deleted.",
                )

            needs_reason = (not is_new_protocol) and (stage_changed or replacing_file)
            if needs_reason and not reason:
                form.add_error(
                    "change_reason",
                    "Please capture the reason for this revision so the team has context.",
                )

            if not form.errors:
                old_stage = protocol.stage if protocol else None
                old_file = (
                    protocol.document.name if protocol and protocol.document else None
                )

                revision_content = None
                revision_filename = ""
                if uploaded_file:
                    uploaded_file.seek(0)
                    revision_content = ContentFile(uploaded_file.read())
                    uploaded_file.seek(0)
                    revision_filename = os.path.basename(
                        uploaded_file.name or "protocol_upload"
                    )

                obj = form.save(commit=False)
                obj.project = project
                obj.save()
                form.save_m2m()

                protocol = obj

                new_file = obj.document.name if obj.document else None
                changes = []

                if is_new_protocol:
                    changes.append("Created protocol record")
                elif old_stage != obj.stage:
                    changes.append(
                        f"Stage: {_format_value(old_stage)} → {_format_value(obj.stage)}"
                    )

                if old_file != new_file:
                    if new_file and old_file:
                        changes.append(f"File replaced: {old_file} → {new_file}")
                    elif new_file:
                        changes.append(f"File uploaded: {new_file}")
                    elif old_file:
                        changes.append(f"File removed: {old_file}")

                revision_instance = None
                if revision_content is not None:
                    revision_instance = ProtocolRevision(
                        protocol=obj,
                        stage=obj.stage,
                        change_reason=(
                            reason
                            if reason
                            else ("Initial upload" if is_new_protocol else "")
                        ),
                        uploaded_by=(
                            request.user if request.user.is_authenticated else None
                        ),
                    )
                    original_name = os.path.basename(
                        uploaded_file.name or revision_filename
                    )
                    revision_instance.original_name = original_name
                    file_size = getattr(uploaded_file, "size", None)
                    if file_size in (None, ""):
                        file_size = getattr(revision_content, "size", None)
                    try:
                        revision_instance.file_size = int(file_size or 0)
                    except (TypeError, ValueError):
                        revision_instance.file_size = 0
                    revision_instance.version_label = version_label
                    revision_instance.file.save(
                        revision_filename, revision_content, save=True
                    )
                    obj.current_revision = revision_instance
                    obj.save(update_fields=["current_revision"])
                elif obj.current_revision:
                    update_fields = []
                    if stage_changed and obj.current_revision.stage != obj.stage:
                        obj.current_revision.stage = obj.stage
                        update_fields.append("stage")
                    current_label = obj.current_revision.version_label or ""
                    if current_label != version_label:
                        obj.current_revision.version_label = version_label
                        update_fields.append("version_label")
                    if reason and (stage_changed or current_label != version_label):
                        obj.current_revision.change_reason = reason
                        update_fields.append("change_reason")
                    if update_fields:
                        obj.current_revision.save(update_fields=update_fields)

                if obj.stage == "final" and obj.current_revision:
                    ProtocolRevision.objects.filter(protocol=obj).exclude(
                        pk=obj.current_revision_id
                    ).update(stage="draft")
                    if obj.current_revision.stage != "final":
                        obj.current_revision.stage = "final"
                        obj.current_revision.save(update_fields=["stage"])

                new_label = (
                    (obj.current_revision.version_label or "")
                    if obj.current_revision
                    else ""
                )
                if previous_label != new_label:
                    changes.append(
                        f"Version label: {_format_value(previous_label)} → {_format_value(new_label)}"
                    )

                detail_parts = []
                if changes:
                    detail_parts.append("; ".join(changes))
                if reason:
                    detail_parts.append(f"Reason: {reason}")
                if not detail_parts:
                    detail_parts.append(
                        "Protocol saved without detectable field changes"
                    )

                _log_project_change(
                    project,
                    request.user,
                    "Protocol updated",
                    " | ".join(detail_parts),
                )
                success_message = "Protocol updated."
                if obj.stage == "final" and (
                    is_new_protocol or stage_changed or replacing_file
                ):
                    success_message = "Protocol updated and marked as final. Switch to Draft before uploading further revisions."
                messages.success(request, success_message)
                return redirect("synopsis:protocol_detail", project_id=project.id)
    else:
        form = ProtocolUpdateForm(instance=protocol)

    if (
        protocol
        and getattr(protocol, "current_revision", None)
        and request.method != "POST"
    ):
        form.fields["version_label"].initial = protocol.current_revision.version_label

    if first_upload_pending:
        form.fields["change_reason"].help_text = (
            "Optional for the first upload. Provide details when you revise an existing protocol."
        )
    else:
        form.fields["change_reason"].help_text = (
            "Required when you replace the file or change the protocol stage."
        )
    if not protocol_document_ready:
        form.fields["document"].widget.attrs["required"] = "required"
        form.fields["document"].help_text = (
            "Upload a PDF or DOCX version of the protocol. You can reuse the same filename after deleting a file."
        )

    protocol_members = project.advisory_board_members.filter(
        sent_protocol_at__isnull=False,
        response="Y",
    )
    protocol_pending_dates = [
        d
        for d in protocol_members.filter(feedback_on_protocol_deadline__isnull=False)
        .order_by("feedback_on_protocol_deadline")
        .values_list("feedback_on_protocol_deadline", flat=True)
    ]
    protocol_reminder_initial = {}
    if protocol_pending_dates:
        first_deadline = protocol_pending_dates[0]
        try:
            protocol_reminder_initial["deadline"] = timezone.localtime(first_deadline)
        except (ValueError, TypeError):
            protocol_reminder_initial["deadline"] = first_deadline
    else:
        protocol_reminder_initial["deadline"] = timezone.localtime(
            _default_document_feedback_deadline()
        )
    protocol_reminder_form = ProtocolReminderScheduleForm(
        initial=protocol_reminder_initial
    )
    protocol_feedback_close_initial = {}
    if protocol and protocol.feedback_closure_message:
        protocol_feedback_close_initial["message"] = protocol.feedback_closure_message
    protocol_feedback_close_form = ProtocolFeedbackCloseForm(
        initial=protocol_feedback_close_initial
    )
    protocol_feedback_state = {
        "protocol": protocol,
        "is_closed": bool(getattr(protocol, "feedback_closed_at", None)),
        "closed_at": getattr(protocol, "feedback_closed_at", None),
        "closure_message": getattr(protocol, "feedback_closure_message", ""),
        "deadline": protocol_pending_dates[0] if protocol_pending_dates else None,
        "document_ready": protocol_document_ready,
    }

    return render(
        request,
        "synopsis/protocol_detail.html",
        {
            "project": project,
            "protocol": protocol,
            "form": form,
            "protocol_history_entries": protocol_history_entries,
            "protocol_revision_entries": revision_entries,
            "current_revision_entry": current_revision_entry,
            "current_revision_download_url": current_revision_download_url,
            "final_stage_locked": final_stage_locked,
            "first_upload_pending": first_upload_pending,
            "can_manage_project": can_manage,
            "can_toggle_stage": can_edit_documents,
            "can_edit_documents": can_edit_documents,
            "can_delete_documents": can_delete_documents,
            "collaborative_enabled": collaborative_enabled,
            "collaborative_session": collaborative_session,
            "collaborative_start_url": reverse(
                "synopsis:collaborative_start", args=[project.id, collaborative_slug]
            ),
            "collaborative_resume_url": collaborative_resume_url,
            "collaborative_force_end_url": collaborative_force_end_url,
            "collaborative_document_ready": protocol_document_ready,
            "collaborative_can_override": collaborative_can_override,
            "protocol_reminder_form": protocol_reminder_form,
            "protocol_pending_count": protocol_members.count(),
            "protocol_pending_dates": protocol_pending_dates,
            "initial_protocol_reminder_log": project.change_log.filter(
                action="Scheduled protocol reminders"
            )
            .order_by("created_at")
            .first(),
            "protocol_feedback_state": protocol_feedback_state,
            "protocol_feedback_close_form": protocol_feedback_close_form,
        },
    )


@login_required
def action_list_detail(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if not _user_can_view_project(request.user, project):
        messages.error(request, "You do not have access to that synopsis.")
        return redirect("synopsis:dashboard")
    action_list = getattr(project, "action_list", None)
    can_manage = _user_is_manager(request.user)
    can_edit_documents = _user_can_edit_project(request.user, project)
    can_delete_documents = _user_can_manage_project_configuration(
        request.user, project
    )
    history_queryset = (
        project.change_log.filter(action__icontains="action list")
        .select_related("changed_by")
        .order_by("-created_at", "-id")
    )
    history_entries = []
    for log in history_queryset:
        segments = [
            segment.strip() for segment in log.details.split("|") if segment.strip()
        ]
        reason = ""
        changes = []
        for segment in segments:
            if segment.lower().startswith("reason:"):
                reason = segment.split(":", 1)[1].strip()
            else:
                changes.append(segment)
        history_entries.append(
            {
                "log": log,
                "changes": changes,
                "reason": reason,
                "actor": _user_display(log.changed_by) if log.changed_by else "System",
            }
        )

    revision_entries = []
    if action_list:
        revision_queryset = action_list.revisions.select_related("uploaded_by")
        for revision in revision_queryset:
            file_name = revision.original_name or os.path.basename(revision.file.name)
            try:
                download_url = revision.file.url
            except ValueError:
                download_url = ""
            revision_entries.append(
                {
                    "revision": revision,
                    "is_current": action_list.current_revision_id == revision.id,
                    "file_name": file_name,
                    "file_size": _format_file_size(revision.file_size),
                    "uploaded_by": (
                        _user_display(revision.uploaded_by)
                        if revision.uploaded_by
                        else "—"
                    ),
                    "download_url": download_url,
                    "can_mark_final": can_manage
                    and (
                        action_list.stage != "final"
                        or action_list.current_revision_id != revision.id
                    ),
                    "version_label": revision.version_label or "",
                }
            )

    current_revision_entry = next(
        (entry for entry in revision_entries if entry["is_current"]), None
    )
    current_revision_download_url = ""
    if current_revision_entry:
        try:
            current_revision_download_url = current_revision_entry["revision"].file.url
        except ValueError:
            current_revision_download_url = ""

    action_document_ready = bool(action_list and getattr(action_list, "document", None))
    existing_file_name = action_list.document.name if action_document_ready else ""
    has_existing_file = bool(existing_file_name or revision_entries)
    first_upload_pending = not has_existing_file
    final_stage_locked = bool(
        action_list and action_list.stage == "final" and has_existing_file
    )

    collaborative_enabled = _onlyoffice_enabled()
    action_closed = bool(action_list and action_list.feedback_closed_at)
    collaborative_session = None
    collaborative_resume_url = ""
    collaborative_force_end_url = ""
    collaborative_slug = _document_type_slug(CollaborativeSession.DOCUMENT_ACTION_LIST)
    if collaborative_enabled:
        collaborative_session = _get_active_collaborative_session(
            project, CollaborativeSession.DOCUMENT_ACTION_LIST
        )
        if collaborative_session and action_closed:
            collaborative_session.mark_inactive(
                ended_by=request.user if request.user.is_authenticated else None,
                reason="Action list feedback window closed",
            )
            collaborative_session = None
        if collaborative_session:
            collaborative_resume_url = reverse(
                "synopsis:collaborative_edit",
                args=[project.id, collaborative_slug, collaborative_session.token],
            )
            collaborative_force_end_url = reverse(
                "synopsis:collaborative_force_end",
                args=[project.id, collaborative_slug, collaborative_session.token],
            )
    if collaborative_session:
        collaborative_can_override = _user_can_force_end_session(
            request.user, project, collaborative_session
        )
    else:
        collaborative_can_override = False
    if action_closed:
        collaborative_enabled = False
        collaborative_session = None
        collaborative_resume_url = ""
        collaborative_force_end_url = ""
        collaborative_can_override = False

    if request.method == "POST":
        form = ActionListUpdateForm(request.POST, request.FILES, instance=action_list)
        if form.is_valid():
            new_stage = form.cleaned_data.get("stage")
            uploaded_file = form.cleaned_data.get("document")
            reason = form.cleaned_data.get("change_reason", "")
            version_label = form.cleaned_data.get("version_label", "")

            is_new_action_list = action_list is None
            stage_changed = bool(action_list) and action_list.stage != new_stage
            replacing_file = bool(uploaded_file)
            previous_label = (
                (action_list.current_revision.version_label or "")
                if action_list and action_list.current_revision
                else ""
            )

            if final_stage_locked and new_stage == "final" and replacing_file:
                form.add_error(
                    "document",
                    "Finalized action lists cannot be replaced. Switch the stage back to Draft to revise the document.",
                )

            active_file_missing = not bool(action_list and action_list.document)
            if active_file_missing and not replacing_file:
                form.add_error(
                    "document",
                    "Choose an action list file to upload. You can reuse the same filename as a file you deleted.",
                )

            needs_reason = (not is_new_action_list) and (
                stage_changed or replacing_file
            )
            if needs_reason and not reason:
                form.add_error(
                    "change_reason",
                    "Please capture the reason for this revision so the team has context.",
                )

            if not form.errors:
                old_stage = action_list.stage if action_list else None
                old_file = (
                    action_list.document.name
                    if action_list and action_list.document
                    else None
                )

                revision_content = None
                revision_filename = ""
                if uploaded_file:
                    uploaded_file.seek(0)
                    revision_content = ContentFile(uploaded_file.read())
                    uploaded_file.seek(0)
                    revision_filename = os.path.basename(
                        uploaded_file.name or "action_list_upload"
                    )

                obj = form.save(commit=False)
                obj.project = project
                obj.save()
                form.save_m2m()

                action_list = obj

                new_file = obj.document.name if obj.document else None
                changes = []

                if is_new_action_list:
                    changes.append("Created action list record")
                elif old_stage != obj.stage:
                    changes.append(
                        f"Stage: {_format_value(old_stage)} → {_format_value(obj.stage)}"
                    )

                if old_file != new_file:
                    if new_file and old_file:
                        changes.append(f"File replaced: {old_file} → {new_file}")
                    elif new_file:
                        changes.append(f"File uploaded: {new_file}")
                    elif old_file:
                        changes.append(f"File removed: {old_file}")

                revision_instance = None
                if revision_content is not None:
                    revision_instance = ActionListRevision(
                        action_list=obj,
                        stage=obj.stage,
                        change_reason=(
                            reason
                            if reason
                            else ("Initial upload" if is_new_action_list else "")
                        ),
                        uploaded_by=(
                            request.user if request.user.is_authenticated else None
                        ),
                    )
                    original_name = os.path.basename(
                        uploaded_file.name or revision_filename
                    )
                    revision_instance.original_name = original_name
                    file_size = getattr(uploaded_file, "size", None)
                    if file_size in (None, ""):
                        file_size = getattr(revision_content, "size", None)
                    try:
                        revision_instance.file_size = int(file_size or 0)
                    except (TypeError, ValueError):
                        revision_instance.file_size = 0
                    revision_instance.version_label = version_label
                    revision_instance.file.save(
                        revision_filename, revision_content, save=True
                    )
                    obj.current_revision = revision_instance
                    obj.save(update_fields=["current_revision"])
                elif obj.current_revision:
                    update_fields = []
                    if stage_changed and obj.current_revision.stage != obj.stage:
                        obj.current_revision.stage = obj.stage
                        update_fields.append("stage")
                    current_label = obj.current_revision.version_label or ""
                    if current_label != version_label:
                        obj.current_revision.version_label = version_label
                        update_fields.append("version_label")
                    if reason and (stage_changed or current_label != version_label):
                        obj.current_revision.change_reason = reason
                        update_fields.append("change_reason")
                    if update_fields:
                        obj.current_revision.save(update_fields=update_fields)

                if obj.stage == "final" and obj.current_revision:
                    ActionListRevision.objects.filter(action_list=obj).exclude(
                        pk=obj.current_revision_id
                    ).update(stage="draft")
                    if obj.current_revision.stage != "final":
                        obj.current_revision.stage = "final"
                        obj.current_revision.save(update_fields=["stage"])

                new_label = (
                    (obj.current_revision.version_label or "")
                    if obj.current_revision
                    else ""
                )
                if previous_label != new_label:
                    changes.append(
                        f"Version label: {_format_value(previous_label)} → {_format_value(new_label)}"
                    )

                detail_parts = []
                if changes:
                    detail_parts.append("; ".join(changes))
                if reason:
                    detail_parts.append(f"Reason: {reason}")
                if not detail_parts:
                    detail_parts.append(
                        "Action list saved without detectable field changes"
                    )

                _log_project_change(
                    project,
                    request.user,
                    "Action list updated",
                    " | ".join(detail_parts),
                )
                success_message = "Action list updated."
                if obj.stage == "final" and (
                    is_new_action_list or stage_changed or replacing_file
                ):
                    success_message = "Action list updated and marked as final. Switch to Draft before uploading further revisions."
                messages.success(request, success_message)
                return redirect("synopsis:action_list_detail", project_id=project.id)
    else:
        form = ActionListUpdateForm(instance=action_list)

    if (
        action_list
        and getattr(action_list, "current_revision", None)
        and request.method != "POST"
    ):
        form.fields["version_label"].initial = action_list.current_revision.version_label

    if first_upload_pending:
        form.fields["change_reason"].help_text = (
            "Optional for the first upload. Provide details when you revise an existing action list."
        )
    else:
        form.fields["change_reason"].help_text = (
            "Required when you replace the file or change the action list stage."
        )
    if not action_document_ready:
        form.fields["document"].widget.attrs["required"] = "required"
        form.fields["document"].help_text = (
            "Upload a PDF or DOCX version of the action list. You can reuse the same filename after deleting a file."
        )

    action_list_members = project.advisory_board_members.filter(
        sent_action_list_at__isnull=False,
        response="Y",
    )
    action_list_pending_dates = [
        d
        for d in action_list_members.filter(
            feedback_on_action_list_deadline__isnull=False
        )
        .order_by("feedback_on_action_list_deadline")
        .values_list("feedback_on_action_list_deadline", flat=True)
    ]
    action_list_reminder_initial = {}
    if action_list_pending_dates:
        first_deadline = action_list_pending_dates[0]
        try:
            action_list_reminder_initial["deadline"] = timezone.localtime(
                first_deadline
            )
        except (ValueError, TypeError):
            action_list_reminder_initial["deadline"] = first_deadline
    else:
        action_list_reminder_initial["deadline"] = timezone.localtime(
            _default_document_feedback_deadline()
        )
    action_list_reminder_form = ActionListReminderScheduleForm(
        initial=action_list_reminder_initial
    )
    action_list_feedback_close_initial = {}
    if action_list and action_list.feedback_closure_message:
        action_list_feedback_close_initial["message"] = (
            action_list.feedback_closure_message
        )
    action_list_feedback_close_form = ActionListFeedbackCloseForm(
        initial=action_list_feedback_close_initial
    )
    action_list_feedback_state = {
        "action_list": action_list,
        "is_closed": bool(getattr(action_list, "feedback_closed_at", None)),
        "closed_at": getattr(action_list, "feedback_closed_at", None),
        "closure_message": getattr(action_list, "feedback_closure_message", ""),
        "deadline": action_list_pending_dates[0] if action_list_pending_dates else None,
        "document_ready": action_document_ready,
    }

    return render(
        request,
        "synopsis/action_list_detail.html",
        {
            "project": project,
            "action_list": action_list,
            "form": form,
            "action_list_history_entries": history_entries,
            "action_list_revision_entries": revision_entries,
            "current_revision_entry": current_revision_entry,
            "current_revision_download_url": current_revision_download_url,
            "final_stage_locked": final_stage_locked,
            "first_upload_pending": first_upload_pending,
            "can_manage_project": can_manage,
            "can_edit_documents": can_edit_documents,
            "can_delete_documents": can_delete_documents,
            "collaborative_enabled": collaborative_enabled,
            "collaborative_session": collaborative_session,
            "collaborative_start_url": reverse(
                "synopsis:collaborative_start", args=[project.id, collaborative_slug]
            ),
            "collaborative_resume_url": collaborative_resume_url,
            "collaborative_force_end_url": collaborative_force_end_url,
            "collaborative_can_override": collaborative_can_override,
            "collaborative_document_ready": action_document_ready,
            "action_list_reminder_form": action_list_reminder_form,
            "action_list_pending_count": action_list_members.count(),
            "action_list_pending_dates": action_list_pending_dates,
            "initial_action_list_reminder_log": project.change_log.filter(
                action="Scheduled action list reminders"
            )
            .order_by("created_at")
            .first(),
            "action_list_feedback_state": action_list_feedback_state,
            "action_list_feedback_close_form": action_list_feedback_close_form,
        },
    )


def _resolve_collaborative_users(user_ids) -> tuple[list[User], list[str]]:
    if not user_ids:
        return [], []

    normalized = [str(uid) for uid in user_ids if uid is not None]
    if not normalized:
        return [], []

    user_id_strings = [key for key in normalized if key.isdigit()]
    member_id_strings = []
    for key in normalized:
        if key.startswith("abm:"):
            try:
                member_id_strings.append(int(key.split(":", 1)[1]))
            except (TypeError, ValueError):
                continue

    users = []
    user_map: dict[str, User] = {}
    if user_id_strings:
        ids = [int(key) for key in user_id_strings]
        for user in User.objects.filter(id__in=ids):
            user_map[str(user.id)] = user

    member_map: dict[str, AdvisoryBoardMember] = {}
    if member_id_strings:
        for member in AdvisoryBoardMember.objects.filter(id__in=member_id_strings):
            member_map[f"abm:{member.id}"] = member

    ordered_users: list[User] = []
    labels: list[str] = []
    seen_users: set[int] = set()

    for key in normalized:
        user = user_map.get(key)
        if user:
            if user.id not in seen_users:
                ordered_users.append(user)
                seen_users.add(user.id)
            labels.append(_user_display(user))
            continue

        member = member_map.get(key)
        if member:
            labels.append(_advisory_member_display(member))
            continue

        if key.startswith("abe:"):
            email_value = key.split(":", 1)[1]
            labels.append(email_value or "Advisory board invitee")
            continue

        if key:
            labels.append(key)

    return ordered_users, labels


def _persist_collaborative_revision(
    document_type,
    document,
    content: bytes,
    original_name: str,
    uploader,
    change_reason: str,
):
    file_length = len(content)
    safe_name = original_name or "document.docx"
    content_file = ContentFile(content)

    if document_type == CollaborativeSession.DOCUMENT_PROTOCOL:
        revision = ProtocolRevision(
            protocol=document,
            stage=document.stage,
            change_reason=change_reason,
            uploaded_by=uploader,
        )
        revision.file.save(safe_name, content_file, save=False)
        revision.original_name = safe_name
        revision.file_size = file_length
        revision.save()
        base_name, size_text = _apply_revision_to_protocol(document, revision)
        if document.stage == "final":
            ProtocolRevision.objects.filter(protocol=document).exclude(
                pk=revision.pk
            ).update(stage="draft")
            if revision.stage != "final":
                revision.stage = "final"
                revision.save(update_fields=["stage"])
        return revision, base_name, size_text

    revision = ActionListRevision(
        action_list=document,
        stage=document.stage,
        change_reason=change_reason,
        uploaded_by=uploader,
    )
    revision.file.save(safe_name, content_file, save=False)
    revision.original_name = safe_name
    revision.file_size = file_length
    revision.save()
    base_name, size_text = _apply_revision_to_action_list(document, revision)
    if document.stage == "final":
        ActionListRevision.objects.filter(action_list=document).exclude(
            pk=revision.pk
        ).update(stage="draft")
        if revision.stage != "final":
            revision.stage = "final"
            revision.save(update_fields=["stage"])
    return revision, base_name, size_text


def _handle_collaborative_save(
    project,
    document_type,
    document,
    session,
    payload: dict,
    status: int,
):
    file_url = payload.get("url")
    if not file_url:
        logger.warning(
            "Collaborative callback missing file URL for session %s", session.pk
        )
        return False

    try:
        content = _download_onlyoffice_file(file_url)
    except requests.RequestException as exc:
        logger.error("Failed to download OnlyOffice file: %s", exc)
        return False

    current_revision = getattr(document, "current_revision", None) or document.latest_revision()
    original_name = _normalized_document_filename(
        payload.get("filename")
        or getattr(current_revision, "original_name", "")
        or getattr(getattr(document, "document", None), "name", ""),
        fallback=f"{_document_type_slug(document_type)}.docx",
    )

    resolved_users, user_labels = _resolve_collaborative_users(payload.get("users", []))
    uploader = resolved_users[0] if resolved_users else session.started_by
    if not user_labels and session.started_by:
        user_labels.append(_user_display(session.started_by))
    if user_labels:
        session.last_participant_name = user_labels[0]

    change_reason = (payload.get("message") or payload.get("comment") or "").strip()
    if not change_reason:
        qualifier = "force save" if status == 6 else "save"
        if user_labels:
            change_reason = (
                f"Collaborative {qualifier} via OnlyOffice by {', '.join(user_labels)}"
            )
        else:
            change_reason = f"Collaborative {qualifier} via OnlyOffice"

    revision, base_name, size_text = _persist_collaborative_revision(
        document_type,
        document,
        content,
        original_name,
        uploader,
        change_reason,
    )

    document_label = _document_label(document_type)
    detail_parts = [
        f"Session: {session.token}",
        f"Status: {status}",
        f"File: {base_name}",
    ]
    if user_labels:
        detail_parts.append(f"Users: {', '.join(user_labels)}")
    if size_text and size_text != "—":
        detail_parts.append(f"Size: {size_text}")
    detail_parts.append(f"Reason: {change_reason}")

    action_label = f"{document_label} updated via collaborative edit"
    _log_project_change(
        project,
        uploader or session.started_by,
        action_label,
        " | ".join(detail_parts),
    )

    session.change_summary = change_reason
    if document_type == CollaborativeSession.DOCUMENT_PROTOCOL:
        session.result_protocol_revision = revision
        extra_updates = ["change_summary", "result_protocol_revision"]
    else:
        session.result_action_list_revision = revision
        extra_updates = ["change_summary", "result_action_list_revision"]

    if session.last_participant_name:
        extra_updates.append("last_participant_name")

    reason_text = (
        "Document saved from OnlyOffice"
        if status in {2, 6}
        else f"Session closed (status {status})"
    )
    session.mark_inactive(
        ended_by=uploader,
        reason=reason_text,
        extra_updates=extra_updates,
    )
    return True


@login_required
def collaborative_start(request, project_id, document_slug):
    if request.method != "POST":
        return HttpResponseBadRequest("POST required")

    project = get_object_or_404(Project, pk=project_id)
    document_type = _normalize_document_type(document_slug)
    if not document_type:
        raise Http404("Unknown document type")

    if not _onlyoffice_enabled():
        messages.error(
            request,
            "Collaborative editing is not configured. Please contact your administrator.",
        )
        return redirect(_document_detail_url(project.id, document_type))

    if not _user_can_edit_project(request.user, project):
        messages.error(
            request, "You do not have permission to start a collaborative session."
        )
        return redirect(_document_detail_url(project.id, document_type))

    document = _get_document_for_type(project, document_type)
    if not _document_requires_file(document):
        messages.error(
            request,
            f"No {_document_label(document_type).lower()} file is available to edit.",
        )
        return redirect(_document_detail_url(project.id, document_type))

    if document_type == CollaborativeSession.DOCUMENT_PROTOCOL and getattr(
        document, "feedback_closed_at", None
    ):
        _end_active_collaborative_session(
            project,
            document_type,
            ended_by=request.user,
            reason="Protocol feedback window closed",
        )
        messages.info(
            request,
            "Collaborative editing is disabled because the protocol feedback window is closed.",
        )
        return redirect(_document_detail_url(project.id, document_type))

    if document_type == CollaborativeSession.DOCUMENT_ACTION_LIST and getattr(
        document, "feedback_closed_at", None
    ):
        _end_active_collaborative_session(
            project,
            document_type,
            ended_by=request.user,
            reason="Action list feedback window closed",
        )
        messages.info(
            request,
            "Collaborative editing is disabled because the action list feedback window is closed.",
        )
        return redirect(_document_detail_url(project.id, document_type))

    # TODO: #17 Make collaborative session creation atomic so concurrent POSTs cannot spawn two active sessions.
    active_session = _get_active_collaborative_session(project, document_type)
    if active_session:
        messages.warning(
            request,
            "A collaborative session is already running. Opening the existing editor instead.",
        )
        return redirect(
            "synopsis:collaborative_edit",
            project_id=project.id,
            document_slug=_document_type_slug(document_type),
            token=active_session.token,
        )

    initial_revision = None
    if document_type == CollaborativeSession.DOCUMENT_PROTOCOL:
        initial_revision = (
            getattr(document, "current_revision", None) or document.latest_revision()
        )
        session = CollaborativeSession.objects.create(
            project=project,
            document_type=document_type,
            started_by=request.user,
            last_activity_at=timezone.now(),
            initial_protocol_revision=initial_revision,
        )
    else:
        initial_revision = (
            getattr(document, "current_revision", None) or document.latest_revision()
        )
        session = CollaborativeSession.objects.create(
            project=project,
            document_type=document_type,
            started_by=request.user,
            last_activity_at=timezone.now(),
            initial_action_list_revision=initial_revision,
        )

    document_label = _document_label(document_type)
    _log_project_change(
        project,
        request.user,
        f"{document_label} collaborative editing started",
        f"Session {session.token} started by {_user_display(request.user)}",
    )

    messages.success(request, f"{document_label} editor is ready.")
    return redirect(
        "synopsis:collaborative_edit",
        project_id=project.id,
        document_slug=_document_type_slug(document_type),
        token=session.token,
    )


def collaborative_edit(request, project_id, document_slug, token):
    project = get_object_or_404(Project, pk=project_id)
    document_type = _normalize_document_type(document_slug)
    if not document_type:
        raise Http404("Unknown document type")

    document_label = _document_label(document_type)
    detail_url = _document_detail_url(project.id, document_type)
    user_can_edit = _user_can_edit_project(request.user, project)
    project_editor_detail_url = detail_url if user_can_edit else ""

    if not _onlyoffice_enabled():
        if user_can_edit:
            messages.error(
                request,
                "Collaborative editing is not configured. Please contact your administrator.",
            )
            return redirect(detail_url)
        return _collaborative_access_closed_response(
            request,
            project,
            document_label,
            "Collaborative editing is not configured. Please contact your administrator.",
            status=503,
        )

    session = _collaborative_session_or_404(project, document_type, token)
    external_access = {"allowed": False}
    if not user_can_edit:
        external_access = _resolve_external_collaborative_access(
            request, project, document_type, session
        )
        if not external_access.get("allowed"):
            return _collaborative_access_closed_response(
                request,
                project,
                document_label,
                external_access.get("message")
                or "You do not have access to this collaborative session.",
            )

    if session.has_expired():
        session.mark_inactive(reason="Session expired")
        if user_can_edit:
            messages.warning(request, "This collaborative session has expired.")
            return redirect(detail_url)
        restart_url = _restart_external_collaborative_url(
            request, project, document_type, external_access
        )
        if restart_url and restart_url != request.build_absolute_uri():
            return redirect(restart_url)
        return _collaborative_access_closed_response(
            request,
            project,
            document_label,
            "This collaborative session has expired. Ask the authors to resend the link.",
        )

    document = _get_document_for_type(project, document_type)
    if not _document_requires_file(document):
        if user_can_edit:
            messages.error(
                request,
                f"No {_document_label(document_type).lower()} file is available to edit.",
            )
            return redirect(detail_url)
        return _collaborative_access_closed_response(
            request,
            project,
            document_label,
            f"No {_document_label(document_type).lower()} file is available to edit.",
        )

    if getattr(document, "feedback_closed_at", None):
        if session.is_active:
            session.mark_inactive(
                ended_by=request.user if user_can_edit else None,
                reason=f"{document_label} feedback window closed",
            )
        closure_message = document.feedback_closure_message or (
            f"Collaborative editing is closed for this {document_label.lower()}."
        )
        return render(
            request,
            "synopsis/collaborative_editor.html",
            {
                "project": project,
                "document_label": document_label,
                "detail_url": project_editor_detail_url,
                "window_closed_message": closure_message,
                "can_force_end": False,
                "force_end_url": "",
                "participant_display": "",
            },
            status=200,
        )

    if not session.is_active:
        if user_can_edit:
            restart_url = _ensure_collaborative_invite_link(
                request, project, document_type
            )
            if restart_url and restart_url != request.build_absolute_uri():
                messages.info(
                    request,
                    "The previous collaborative session ended. A fresh editor has been opened.",
                )
                return redirect(restart_url)
            messages.info(request, "This collaborative session is no longer active.")
            return redirect(detail_url)
        restart_url = _restart_external_collaborative_url(
            request, project, document_type, external_access
        )
        if restart_url and restart_url != request.build_absolute_uri():
            return redirect(restart_url)
        return _collaborative_access_closed_response(
            request,
            project,
            document_label,
            "This collaborative session is no longer active. Ask the authors to resend the link.",
        )

    editor_js_url = _onlyoffice_editor_js_url()
    if not editor_js_url:
        if user_can_edit:
            messages.error(
                request,
                "The OnlyOffice editor script URL is not configured. Please contact your administrator.",
            )
            return redirect(_document_detail_url(project.id, document_type))
        return _collaborative_access_closed_response(
            request,
            project,
            document_label,
            "The OnlyOffice editor script URL is not configured. Please contact the authors.",
            status=503,
        )

    participant_member = None
    participant_feedback = None
    participant_display = ""
    participant_context = None
    if external_access.get("allowed"):
        participant_member = external_access.get("member")
        participant_feedback = external_access.get("feedback")
        participant_display = external_access.get("participant_display", "")
        participant_context = external_access.get("participant_context")

    if not participant_member and not participant_context and user_can_edit:
        member_id = request.GET.get("member")
        if member_id:
            try:
                participant_member = AdvisoryBoardMember.objects.get(
                    pk=member_id, project=project
                )
            except AdvisoryBoardMember.DoesNotExist:
                participant_member = None

    if participant_member:
        participant_display = _advisory_member_display(participant_member)
        participant_context = {
            "id": f"abm:{participant_member.id}",
            "name": participant_display,
            "email": participant_member.email,
        }
    elif participant_feedback and participant_feedback.email and not participant_context:
        participant_display = participant_feedback.email
        participant_context = {
            "id": f"abe:{participant_feedback.email.lower()}",
            "name": participant_display,
            "email": participant_feedback.email,
        }

    try:
        config = _build_onlyoffice_config(
            request,
            project,
            document,
            session,
            document_type,
            participant=participant_context,
        )
    except ValueError as exc:
        if user_can_edit:
            messages.error(request, str(exc))
            return redirect(_document_detail_url(project.id, document_type))
        return _collaborative_access_closed_response(
            request,
            project,
            document_label,
            str(exc),
        )

    session.last_activity_at = timezone.now()
    update_fields = ["last_activity_at"]
    if participant_display:
        session.last_participant_name = participant_display
        update_fields.append("last_participant_name")
    session.save(update_fields=update_fields)

    force_end_url = reverse(
        "synopsis:collaborative_force_end",
        args=[project.id, _document_type_slug(document_type), session.token],
    )
    leave_url = (
        reverse(
            "synopsis:collaborative_leave",
            args=[project.id, _document_type_slug(document_type), session.token],
        )
        + _collaborative_query_suffix(request.GET)
    )
    can_force_end = _user_can_force_end_session(request.user, project, session)

    return render(
        request,
        "synopsis/collaborative_editor.html",
        {
            "project": project,
            "session": session,
            "document_label": document_label,
            "editor_config": config,
            "onlyoffice_js_url": editor_js_url,
            "detail_url": project_editor_detail_url,
            "can_force_end": can_force_end,
            "force_end_url": force_end_url,
            "leave_url": leave_url,
            "participant_display": participant_display,
        },
    )


def collaborative_leave(request, project_id, document_slug, token):
    project = get_object_or_404(Project, pk=project_id)
    document_type = _normalize_document_type(document_slug)
    if not document_type:
        raise Http404("Unknown document type")

    document_label = _document_label(document_type)
    detail_url = _document_detail_url(project.id, document_type)
    user_can_edit = _user_can_edit_project(request.user, project)
    project_editor_detail_url = detail_url if user_can_edit else ""
    session = _collaborative_session_or_404(project, document_type, token)
    external_access = _resolve_external_collaborative_access(
        request, project, document_type, session
    )
    if not user_can_edit and not external_access.get("allowed"):
        return _collaborative_access_closed_response(
            request,
            project,
            document_label,
            external_access.get("message")
            or "You do not have access to this collaborative session.",
        )

    if user_can_edit:
        messages.info(
            request,
            "You left the collaborative editor. The shared session is still open for other participants.",
        )
        return redirect(detail_url)

    reopen_url = ""
    document = _get_document_for_type(project, document_type)
    if (
        session.is_active
        and not session.has_expired()
        and not getattr(document, "feedback_closed_at", None)
    ):
        reopen_url = reverse(
            "synopsis:collaborative_edit",
            args=[project.id, _document_type_slug(document_type), session.token],
        ) + _collaborative_query_suffix(request.GET)

    return render(
        request,
        "synopsis/collaborative_editor.html",
        {
            "project": project,
            "document_label": document_label,
            "detail_url": project_editor_detail_url,
            "leave_message": (
                "You left the collaborative editor. This did not close the shared session for other participants."
            ),
            "reopen_url": reopen_url,
            "can_force_end": False,
            "force_end_url": "",
            "leave_url": "",
            "participant_display": "",
        },
        status=200,
    )


@login_required
def document_view(request, project_id, document_slug):
    project = get_object_or_404(Project, pk=project_id)
    if not _user_can_view_project(request.user, project):
        messages.error(request, "You do not have access to that synopsis.")
        return redirect("synopsis:dashboard")

    document_type = _normalize_document_type(document_slug)
    if not document_type:
        raise Http404("Unknown document type")

    detail_url = _document_detail_url(project.id, document_type)
    document = _get_document_for_type(project, document_type)
    if not _document_requires_file(document):
        messages.error(
            request,
            f"There is no current {_document_label(document_type).lower()} document to view.",
        )
        return redirect(detail_url)

    file_handle = document.document.open("rb")
    filename = document.document.name.rsplit("/", 1)[-1]
    content_type = (
        mimetypes.guess_type(filename)[0] or "application/octet-stream"
    )
    response = FileResponse(file_handle)
    response["Content-Type"] = content_type
    response["Content-Disposition"] = f'inline; filename="{filename}"'
    response["Cache-Control"] = "no-store"
    return response


@login_required
def collaborative_force_end(request, project_id, document_slug, token):
    if request.method != "POST":
        return HttpResponseBadRequest("POST required")

    project = get_object_or_404(Project, pk=project_id)
    document_type = _normalize_document_type(document_slug)
    if not document_type:
        raise Http404("Unknown document type")

    session = _collaborative_session_or_404(project, document_type, token)
    if not session.is_active:
        messages.info(request, "The collaborative session is already closed.")
        return redirect(_document_detail_url(project.id, document_type))

    if not _user_can_force_end_session(request.user, project, session):
        messages.error(request, "You cannot override this collaborative lock.")
        return redirect(_document_detail_url(project.id, document_type))

    reason = (request.POST.get("reason") or "").strip() or "Session ended from portal"
    document_label = _document_label(document_type)
    save_state, save_message = _request_onlyoffice_forcesave(
        project, document_type, session
    )
    if save_state == "noop":
        session.change_summary = session.change_summary or reason
        session.mark_inactive(
            ended_by=request.user,
            reason=reason,
            extra_updates=["change_summary"],
        )
        _log_project_change(
            project,
            request.user,
            f"{document_label} collaborative session closed",
            f"Session {session.token} closed with no unsaved changes ({reason}).",
        )
        messages.success(
            request, f"{document_label} had no unsaved changes and the session was closed."
        )
        return redirect(_document_detail_url(project.id, document_type))

    if save_state == "requested":
        if _wait_for_collaborative_save(
            session,
            document_type,
            timeout_seconds=ONLYOFFICE_SETTINGS.get("callback_timeout", 10),
        ):
            messages.success(
                request, f"{document_label} saved and collaborative session closed."
            )
            return redirect(_document_detail_url(project.id, document_type))
        messages.warning(
            request,
            f"{save_message} The session is still open while OnlyOffice finishes saving.",
        )
    else:
        messages.error(
            request,
            f"{save_message} The session is still open so no edits are discarded.",
        )
    return redirect(_document_detail_url(project.id, document_type))


@csrf_exempt
def collaborative_edit_callback(request, project_id, document_slug, token):
    if request.method != "POST":
        return JsonResponse({"error": 1, "message": "POST required"}, status=405)

    project = get_object_or_404(Project, pk=project_id)
    document_type = _normalize_document_type(document_slug)
    if not document_type:
        return JsonResponse({"error": 1, "message": "Unknown document"}, status=404)

    session = _collaborative_session_or_404(project, document_type, token)

    try:
        payload = _parse_onlyoffice_callback(request)
    except PermissionDenied:
        raise
    except ValueError:
        return JsonResponse({"error": 1, "message": "Invalid payload"}, status=400)

    status = payload.get("status")
    try:
        status = int(status)
    except (TypeError, ValueError):
        status = 0

    session.record_callback(payload)

    document = _get_document_for_type(project, document_type)
    if not _document_requires_file(document):
        session.mark_inactive(reason="Document missing at save time")
        return JsonResponse({"error": 0})

    if status in {2, 6}:
        success = _handle_collaborative_save(
            project, document_type, document, session, payload, status
        )
        if not success:
            return JsonResponse({"error": 1})
        return JsonResponse({"error": 0})

    if status in {3, 4, 7}:
        reason_map = {
            3: "Session closed without changes",
            4: "Session closed by timeout",
            7: "Session closed with errors",
        }
        session.mark_inactive(reason=reason_map.get(status, "Session closed"))
        _log_project_change(
            project,
            session.started_by,
            f"{_document_label(document_type)} collaborative session closed",
            f"Session {session.token} closed (status {status}).",
        )
        return JsonResponse({"error": 0})

    return JsonResponse({"error": 0})


@login_required
def action_list_set_stage(request, project_id):
    if request.method != "POST":
        return HttpResponseBadRequest("Invalid request method.")

    project = get_object_or_404(Project, pk=project_id)
    if not _user_can_edit_project(request.user, project):
        messages.error(
            request, "You do not have permission to update the action list stage."
        )
        return redirect("synopsis:action_list_detail", project_id=project.id)

    action_list = getattr(project, "action_list", None)
    if not action_list:
        messages.error(request, "No action list exists yet.")
        return redirect("synopsis:action_list_detail", project_id=project.id)

    target_stage = (request.POST.get("stage") or "").strip()
    if target_stage not in {"draft", "final"}:
        return HttpResponseBadRequest("Invalid stage value.")

    reason = (request.POST.get("reason") or "").strip()
    old_stage = action_list.stage

    revision = None
    revision_id = request.POST.get("revision_id")
    if revision_id:
        revision = get_object_or_404(
            ActionListRevision, pk=revision_id, action_list=action_list
        )
    elif action_list.current_revision:
        revision = action_list.current_revision
    else:
        revision = action_list.revisions.order_by("-uploaded_at", "-id").first()

    if action_list.stage == target_stage and not (
        target_stage == "final"
        and revision
        and action_list.current_revision
        and revision.id != action_list.current_revision_id
    ):
        messages.info(request, f"Action list is already marked as {target_stage}.")
        return redirect("synopsis:action_list_detail", project_id=project.id)

    if target_stage == "final":
        if not revision:
            messages.error(
                request,
                "No revision found to mark as final. Please upload an action list first.",
            )
            return redirect("synopsis:action_list_detail", project_id=project.id)
        if not reason:
            messages.error(
                request,
                "Please provide a brief reason for marking the action list as final.",
            )
            return redirect("synopsis:action_list_detail", project_id=project.id)

        try:
            base_name, size_text = _apply_revision_to_action_list(action_list, revision)
        except FileNotFoundError:
            messages.error(
                request,
                "The selected revision file could not be found. Please choose another revision or upload a new version.",
            )
            return redirect("synopsis:action_list_detail", project_id=project.id)
        except ValueError:
            messages.error(
                request,
                "The selected revision file is empty. Please choose another revision or upload a new version.",
            )
            return redirect("synopsis:action_list_detail", project_id=project.id)

        action_list.stage = "final"
        action_list.save(update_fields=["stage"])

        ActionListRevision.objects.filter(action_list=action_list).exclude(
            pk=revision.pk
        ).update(stage="draft")
        update_fields = ["stage"]
        if revision.stage != "final":
            revision.stage = "final"
        if reason:
            revision.change_reason = reason
            if "change_reason" not in update_fields:
                update_fields.append("change_reason")
        revision.save(update_fields=update_fields)

        detail_parts = [
            f"Stage: {_format_value(old_stage)} → Final",
            f"File: {base_name}",
        ]
        if size_text != "—":
            detail_parts.append(f"Size: {size_text}")
        if reason:
            detail_parts.append(f"Reason: {reason}")

        _log_project_change(
            project,
            request.user,
            "Action list marked final",
            " | ".join(detail_parts),
        )

        messages.success(
            request,
            "Action list marked as final. Switch back to Draft before uploading new revisions.",
        )
        return redirect("synopsis:action_list_detail", project_id=project.id)

    action_list.stage = "draft"
    action_list.save(update_fields=["stage"])
    if action_list.current_revision and action_list.current_revision.stage != "draft":
        action_list.current_revision.stage = "draft"
        action_list.current_revision.save(update_fields=["stage"])

    detail_parts = [f"Stage: {_format_value(old_stage)} → Draft"]
    if reason:
        detail_parts.append(f"Reason: {reason}")

    _log_project_change(
        project,
        request.user,
        "Action list stage updated",
        " | ".join(detail_parts),
    )

    messages.success(request, "Action list stage set to draft.")
    return redirect("synopsis:action_list_detail", project_id=project.id)


@login_required
def action_list_delete_file(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if not _user_can_manage_project_configuration(request.user, project):
        messages.error(
            request,
            "You do not have permission to delete action list files for this synopsis.",
        )
        return redirect("synopsis:action_list_detail", project_id=project.id)
    action_list = getattr(project, "action_list", None)
    if not action_list or not action_list.document:
        messages.info(request, "No action list file to delete.")
        return redirect("synopsis:action_list_detail", project_id=project.id)

    if request.method == "POST":
        file_name = action_list.document.name
        action_list.document.delete(save=False)
        action_list.document = ""
        action_list.current_revision = None
        action_list.save(update_fields=["document", "current_revision"])
        ended_session = _end_active_collaborative_session(
            project,
            CollaborativeSession.DOCUMENT_ACTION_LIST,
            ended_by=request.user,
            reason="Action list file deleted",
        )
        _log_project_change(
            project,
            request.user,
            "Removed action list file",
            f"File: {file_name}"
            + ("; Collaborative session closed" if ended_session else ""),
        )
        messages.success(
            request,
            "Action list file removed. Upload a replacement from Upload new version; it can use the same filename.",
        )
        return redirect("synopsis:action_list_detail", project_id=project.id)

    return render(
        request,
        "synopsis/action_list_confirm_delete.html",
        {
            "project": project,
            "action_list": action_list,
            "mode": "file",
        },
    )


@login_required
def action_list_restore_revision(request, project_id, revision_id):
    if request.method != "POST":
        return HttpResponseBadRequest("Invalid request method.")

    project = get_object_or_404(Project, pk=project_id)
    action_list = getattr(project, "action_list", None)
    if not action_list:
        messages.error(request, "No action list exists to restore.")
        return redirect("synopsis:action_list_detail", project_id=project.id)

    revision = get_object_or_404(
        ActionListRevision, pk=revision_id, action_list=action_list
    )

    try:
        base_name, size_text = _apply_revision_to_action_list(action_list, revision)
    except FileNotFoundError:
        messages.error(
            request,
            "The selected revision file could not be found. Please upload a new version instead.",
        )
        return redirect("synopsis:action_list_detail", project_id=project.id)
    except ValueError:
        messages.error(
            request,
            "The selected revision file is empty. Please choose another revision or upload a new version.",
        )
        return redirect("synopsis:action_list_detail", project_id=project.id)

    action_list.stage = revision.stage
    action_list.save(update_fields=["stage"])

    restored_at = timezone.localtime(revision.uploaded_at).strftime("%Y-%m-%d %H:%M")
    detail_parts = [
        f"Restored revision uploaded {restored_at}",
        f"Stage reset to {_format_value(revision.stage)}",
        f"File: {base_name}",
    ]
    if size_text != "—":
        detail_parts.append(f"Size: {size_text}")
    if revision.change_reason:
        detail_parts.append(f"Original reason: {revision.change_reason}")

    _log_project_change(
        project,
        request.user,
        "Action list restored",
        " | ".join(detail_parts),
    )

    messages.success(request, "Action list reverted to the selected revision.")
    return redirect("synopsis:action_list_detail", project_id=project.id)


@login_required
def action_list_clear_text(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if not _user_can_manage_project_configuration(request.user, project):
        messages.error(
            request,
            "You do not have permission to clear action list notes for this synopsis.",
        )
        return redirect("synopsis:action_list_detail", project_id=project.id)
    action_list = getattr(project, "action_list", None)
    if not action_list or not (action_list.text_version or "").strip():
        messages.info(request, "No action list notes to clear.")
        return redirect("synopsis:action_list_detail", project_id=project.id)

    if request.method == "POST":
        old_length = len(action_list.text_version or "")
        action_list.text_version = ""
        action_list.save(update_fields=["text_version"])
        _log_project_change(
            project,
            request.user,
            "Cleared action list notes",
            f"Removed rich text content (previous length {old_length} chars)",
        )
        messages.success(request, "Action list notes cleared.")
        return redirect("synopsis:action_list_detail", project_id=project.id)

    return render(
        request,
        "synopsis/action_list_confirm_delete.html",
        {
            "project": project,
            "action_list": action_list,
            "mode": "text",
        },
    )


@login_required
def action_list_delete(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if not _user_can_manage_project_configuration(request.user, project):
        messages.error(
            request,
            "You do not have permission to delete the action list for this synopsis.",
        )
        return redirect("synopsis:action_list_detail", project_id=project.id)
    action_list = getattr(project, "action_list", None)
    if not action_list:
        messages.info(request, "No action list to delete.")
        return redirect("synopsis:action_list_detail", project_id=project.id)

    if request.method == "POST":
        file_name = action_list.document.name if action_list.document else None
        text_len = len(action_list.text_version or "")
        revision_count = action_list.revisions.count()
        if action_list.document:
            action_list.document.delete(save=False)
        ended_session = _end_active_collaborative_session(
            project,
            CollaborativeSession.DOCUMENT_ACTION_LIST,
            ended_by=request.user,
            reason="Action list deleted",
        )
        action_list.delete()
        details = []
        if file_name:
            details.append(f"File: {file_name}")
        details.append(f"Text length removed: {text_len} chars")
        details.append(f"Revisions removed: {revision_count}")
        if ended_session:
            details.append("Collaborative session closed")
        _log_project_change(
            project,
            request.user,
            "Deleted action list",
            "; ".join(details),
        )
        messages.success(request, "Action list deleted.")
        return redirect("synopsis:project_hub", project_id=project.id)

    return render(
        request,
        "synopsis/action_list_confirm_delete.html",
        {
            "project": project,
            "action_list": action_list,
            "mode": "delete",
        },
    )


@login_required
def protocol_set_stage(request, project_id):
    if request.method != "POST":
        return HttpResponseBadRequest("Invalid request method.")

    project = get_object_or_404(Project, pk=project_id)
    if not _user_can_edit_project(request.user, project):
        messages.error(
            request, "You do not have permission to update the protocol stage."
        )
        return redirect("synopsis:protocol_detail", project_id=project.id)
    protocol = getattr(project, "protocol", None)
    if not protocol:
        messages.error(request, "No protocol exists yet.")
        return redirect("synopsis:protocol_detail", project_id=project.id)

    target_stage = (request.POST.get("stage") or "").strip()
    if target_stage not in {"draft", "final"}:
        return HttpResponseBadRequest("Invalid stage value.")

    reason = (request.POST.get("reason") or "").strip()
    old_stage = protocol.stage

    revision = None
    revision_id = request.POST.get("revision_id")
    if revision_id:
        revision = get_object_or_404(
            ProtocolRevision, pk=revision_id, protocol=protocol
        )
    elif protocol.current_revision:
        revision = protocol.current_revision
    else:
        revision = protocol.revisions.order_by("-uploaded_at", "-id").first()

    if protocol.stage == target_stage and not (
        target_stage == "final"
        and revision
        and protocol.current_revision
        and revision.id != protocol.current_revision_id
    ):
        messages.info(request, f"Protocol is already marked as {target_stage}.")
        return redirect("synopsis:protocol_detail", project_id=project.id)

    if target_stage == "final":
        if not revision:
            messages.error(
                request,
                "No revision found to mark as final. Please upload a protocol first.",
            )
            return redirect("synopsis:protocol_detail", project_id=project.id)
        if not reason:
            messages.error(
                request,
                "Please provide a brief reason for marking the protocol as final.",
            )
            return redirect("synopsis:protocol_detail", project_id=project.id)

        try:
            base_name, size_text = _apply_revision_to_protocol(protocol, revision)
        except FileNotFoundError:
            messages.error(
                request,
                "The selected revision file could not be found. Please choose another revision or upload a new version.",
            )
            return redirect("synopsis:protocol_detail", project_id=project.id)
        except ValueError:
            messages.error(
                request,
                "The selected revision file is empty. Please choose another revision or upload a new version.",
            )
            return redirect("synopsis:protocol_detail", project_id=project.id)

        protocol.stage = "final"
        protocol.save(update_fields=["stage"])

        ProtocolRevision.objects.filter(protocol=protocol).exclude(
            pk=revision.pk
        ).update(stage="draft")
        update_fields = ["stage"]
        if revision.stage != "final":
            revision.stage = "final"
        if reason:
            revision.change_reason = reason
            if "change_reason" not in update_fields:
                update_fields.append("change_reason")
        revision.save(update_fields=update_fields)

        detail_parts = [
            f"Stage: {_format_value(old_stage)} → Final",
            f"File: {base_name}",
        ]
        if size_text != "—":
            detail_parts.append(f"Size: {size_text}")
        if reason:
            detail_parts.append(f"Reason: {reason}")

        _log_project_change(
            project,
            request.user,
            "Protocol marked final",
            " | ".join(detail_parts),
        )

        messages.success(
            request,
            "Protocol marked as final. Switch back to Draft before uploading new revisions.",
        )
        return redirect("synopsis:protocol_detail", project_id=project.id)

    # target_stage == "draft"
    old_stage = protocol.stage
    protocol.stage = "draft"
    protocol.save(update_fields=["stage"])
    if protocol.current_revision and protocol.current_revision.stage != "draft":
        protocol.current_revision.stage = "draft"
        protocol.current_revision.save(update_fields=["stage"])

    detail_parts = [f"Stage: {_format_value(old_stage)} → Draft"]
    if reason:
        detail_parts.append(f"Reason: {reason}")

    _log_project_change(
        project,
        request.user,
        "Protocol stage updated",
        " | ".join(detail_parts),
    )

    messages.success(request, "Protocol stage set to Draft.")
    return redirect("synopsis:protocol_detail", project_id=project.id)


@login_required
def protocol_delete_file(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if not _user_can_manage_project_configuration(request.user, project):
        messages.error(
            request,
            "You do not have permission to delete protocol files for this synopsis.",
        )
        return redirect("synopsis:protocol_detail", project_id=project.id)
    protocol = getattr(project, "protocol", None)
    if not protocol or not protocol.document:
        messages.info(request, "No protocol file to delete.")
        return redirect("synopsis:protocol_detail", project_id=project.id)

    if request.method == "POST":
        file_name = protocol.document.name
        protocol.document.delete(save=False)
        protocol.document = ""
        protocol.current_revision = None
        protocol.save(update_fields=["document", "current_revision"])
        ended_session = _end_active_collaborative_session(
            project,
            CollaborativeSession.DOCUMENT_PROTOCOL,
            ended_by=request.user,
            reason="Protocol file deleted",
        )
        _log_project_change(
            project,
            request.user,
            "Removed protocol file",
            f"File: {file_name}"
            + ("; Collaborative session closed" if ended_session else ""),
        )
        messages.success(
            request,
            "Protocol file removed. Upload a replacement from Upload new version; it can use the same filename.",
        )
        return redirect("synopsis:protocol_detail", project_id=project.id)

    return render(
        request,
        "synopsis/protocol_confirm_delete.html",
        {
            "project": project,
            "protocol": protocol,
            "mode": "file",
        },
    )


@login_required
def protocol_delete_revision(request, project_id, revision_id):
    if request.method != "POST":
        return HttpResponseBadRequest("Invalid request method.")

    project = get_object_or_404(Project, pk=project_id)
    if not _user_can_manage_project_configuration(request.user, project):
        messages.error(
            request,
            "You do not have permission to delete protocol revisions for this synopsis.",
        )
        return redirect("synopsis:protocol_detail", project_id=project.id)

    protocol = getattr(project, "protocol", None)
    if not protocol:
        messages.error(request, "No protocol exists for this project.")
        return redirect("synopsis:protocol_detail", project_id=project.id)

    revision = get_object_or_404(ProtocolRevision, pk=revision_id, protocol=protocol)

    was_current = protocol.current_revision_id == revision.id
    file_name = revision.file.name
    revision.delete()

    next_revision = (
        protocol.revisions.exclude(pk=revision_id)
        .order_by("-uploaded_at", "-id")
        .first()
    )

    if was_current:
        if next_revision:
            try:
                base_name, size_text = _apply_revision_to_protocol(
                    protocol, next_revision
                )
            except (FileNotFoundError, ValueError):
                protocol.current_revision = next_revision
                protocol.save(update_fields=["current_revision"])
                base_name = next_revision.original_name or os.path.basename(
                    next_revision.file.name
                )
                size_text = _format_file_size(next_revision.file_size)
        else:
            protocol.current_revision = None
            protocol.save(update_fields=["current_revision"])
            base_name = "none"
            size_text = "—"
    else:
        base_name = revision.original_name or os.path.basename(file_name)
        size_text = _format_file_size(revision.file_size)

    _log_project_change(
        project,
        request.user,
        "Protocol revision deleted",
        f"Revision file: {file_name or 'unknown'}; Remaining current: {base_name} ({size_text})",
    )

    messages.success(request, "Protocol revision deleted.")
    return redirect("synopsis:protocol_detail", project_id=project.id)


@login_required
def protocol_restore_revision(request, project_id, revision_id):
    if request.method != "POST":
        return HttpResponseBadRequest("Invalid request method.")

    project = get_object_or_404(Project, pk=project_id)
    protocol = getattr(project, "protocol", None)
    if not protocol:
        messages.error(request, "No protocol exists to restore.")
        return redirect("synopsis:protocol_detail", project_id=project.id)

    revision = get_object_or_404(ProtocolRevision, pk=revision_id, protocol=protocol)

    try:
        base_name, size_text = _apply_revision_to_protocol(protocol, revision)
    except FileNotFoundError:
        messages.error(
            request,
            "The selected revision file could not be found. Please upload a new version instead.",
        )
        return redirect("synopsis:protocol_detail", project_id=project.id)
    except ValueError:
        messages.error(
            request,
            "The selected revision file is empty. Please choose another revision or upload a new version.",
        )
        return redirect("synopsis:protocol_detail", project_id=project.id)

    protocol.stage = revision.stage
    protocol.save(update_fields=["stage"])

    restored_at = timezone.localtime(revision.uploaded_at).strftime("%Y-%m-%d %H:%M")
    detail_parts = [
        f"Restored revision uploaded {restored_at}",
        f"Stage reset to {_format_value(revision.stage)}",
        f"File: {base_name}",
    ]
    if size_text != "—":
        detail_parts.append(f"Size: {size_text}")
    if revision.change_reason:
        detail_parts.append(f"Original reason: {revision.change_reason}")

    _log_project_change(
        project,
        request.user,
        "Protocol restored",
        " | ".join(detail_parts),
    )

    messages.success(request, "Protocol reverted to the selected revision.")
    return redirect("synopsis:protocol_detail", project_id=project.id)


@login_required
def protocol_clear_text(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if not _user_can_manage_project_configuration(request.user, project):
        messages.error(
            request,
            "You do not have permission to clear protocol text for this synopsis.",
        )
        return redirect("synopsis:protocol_detail", project_id=project.id)
    protocol = getattr(project, "protocol", None)
    if not protocol or not (protocol.text_version or "").strip():
        messages.info(request, "No protocol text to clear.")
        return redirect("synopsis:protocol_detail", project_id=project.id)

    if request.method == "POST":
        old_length = len(protocol.text_version or "")
        protocol.text_version = ""
        protocol.save(update_fields=["text_version"])
        _log_project_change(
            project,
            request.user,
            "Cleared protocol text",
            f"Removed rich text content (previous length {old_length} chars)",
        )
        messages.success(request, "Protocol text cleared.")
        return redirect("synopsis:protocol_detail", project_id=project.id)

    return render(
        request,
        "synopsis/protocol_confirm_delete.html",
        {
            "project": project,
            "protocol": protocol,
            "mode": "text",
        },
    )


@login_required
def protocol_delete(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if not _user_can_manage_project_configuration(request.user, project):
        messages.error(
            request,
            "You do not have permission to delete the protocol for this synopsis.",
        )
        return redirect("synopsis:protocol_detail", project_id=project.id)
    protocol = getattr(project, "protocol", None)
    if not protocol:
        messages.info(request, "No protocol to delete.")
        return redirect("synopsis:protocol_detail", project_id=project.id)

    if request.method == "POST":
        file_name = protocol.document.name if protocol.document else None
        text_len = len(protocol.text_version or "")
        if protocol.document:
            protocol.document.delete(save=False)
        ended_session = _end_active_collaborative_session(
            project,
            CollaborativeSession.DOCUMENT_PROTOCOL,
            ended_by=request.user,
            reason="Protocol deleted",
        )
        protocol.delete()
        details = []
        if file_name:
            details.append(f"File: {file_name}")
        details.append(f"Text length removed: {text_len} chars")
        if ended_session:
            details.append("Collaborative session closed")
        _log_project_change(
            project,
            request.user,
            "Deleted protocol",
            "; ".join(details),
        )
        messages.success(request, "Protocol deleted.")
        return redirect("synopsis:project_hub", project_id=project.id)

    return render(
        request,
        "synopsis/protocol_confirm_delete.html",
        {
            "project": project,
            "protocol": protocol,
            "mode": "all",
        },
    )


@login_required
def manager_dashboard(request):
    if not request.user.is_staff:
        messages.error(request, "Manager access only.")
        return redirect("synopsis:dashboard")

    ensure_global_groups()

    projects = Project.objects.prefetch_related("userrole_set__user").order_by(
        "-created_at", "-id"
    )
    project_entries = []
    for project in projects:
        authors = [
            role.user for role in project.userrole_set.all() if role.role == "author"
        ]
        form = ProjectDeleteForm(project=project, auto_id=f"id_project_{project.id}_%s")
        project_entries.append(
            {
                "project": project,
                "authors": authors,
                "delete_form": form,
            }
        )

    return render(
        request,
        "synopsis/manager_dashboard.html",
        {"user_entries": _manager_user_entries(), "project_entries": project_entries},
    )


@login_required
def manager_user_edit(request, user_id):
    if not request.user.is_staff:
        messages.error(request, "Manager access only.")
        return redirect("synopsis:dashboard")

    ensure_global_groups()
    managed_user = get_object_or_404(User.objects.prefetch_related("groups"), pk=user_id)
    if managed_user.is_superuser:
        messages.info(
            request,
            "System admin accounts are managed outside this screen.",
        )
        return redirect("synopsis:manager_dashboard")

    project_roles = list(
        UserRole.objects.filter(user=managed_user)
        .select_related("project")
        .order_by("project__title", "role")
    )
    existing_author_project_ids = set(
        UserRole.objects.filter(user=managed_user, role="author").values_list(
            "project_id", flat=True
        )
    )

    if request.method == "POST":
        action = request.POST.get("action")
        if action == "update_user":
            form = ManagerUserUpdateForm(request.POST, user=managed_user)
            delete_form = ManagerUserDeleteForm(user=managed_user)
            if form.is_valid():
                if (
                    managed_user == request.user
                    and form.cleaned_data["global_role"] != "manager"
                ):
                    messages.error(request, "You cannot remove your own manager access.")
                elif managed_user == request.user and not form.cleaned_data["is_active"]:
                    messages.error(request, "You cannot deactivate your own account.")
                else:
                    old_first_name = managed_user.first_name
                    old_last_name = managed_user.last_name
                    old_email = managed_user.email
                    old_role = _manager_user_global_role(managed_user)
                    old_active = managed_user.is_active
                    selected_projects = list(form.cleaned_data["assigned_projects"])
                    selected_project_ids = {project.id for project in selected_projects}
                    managed_user.first_name = form.cleaned_data["first_name"].strip()
                    managed_user.last_name = form.cleaned_data["last_name"].strip()
                    managed_user.email = form.cleaned_data["email"]
                    managed_user.username = form.cleaned_data["email"]
                    managed_user.is_active = form.cleaned_data["is_active"]
                    _set_manager_user_global_role(
                        managed_user,
                        form.cleaned_data["global_role"],
                    )
                    managed_user.save()
                    details = []
                    if (
                        old_first_name != managed_user.first_name
                        or old_last_name != managed_user.last_name
                    ):
                        details.append("Name updated")
                    if old_email != managed_user.email:
                        details.append(f"Email/login: {old_email} → {managed_user.email}")
                    if old_role != form.cleaned_data["global_role"]:
                        details.append(
                            "Global role: "
                            f"{GLOBAL_ROLE_LABELS.get(old_role, old_role or '—')} → "
                            f"{GLOBAL_ROLE_LABELS.get(form.cleaned_data['global_role'], form.cleaned_data['global_role'])}"
                        )
                    if old_active != managed_user.is_active:
                        details.append(
                            f"Account {'activated' if managed_user.is_active else 'deactivated'}"
                        )
                    added_project_ids = selected_project_ids - existing_author_project_ids
                    removed_project_ids = (
                        existing_author_project_ids - selected_project_ids
                    )
                    if removed_project_ids:
                        UserRole.objects.filter(
                            user=managed_user,
                            role="author",
                            project_id__in=removed_project_ids,
                        ).delete()
                    for project in selected_projects:
                        UserRole.objects.get_or_create(
                            user=managed_user, project=project, role="author"
                        )
                    if added_project_ids or removed_project_ids:
                        added_titles = list(
                            Project.objects.filter(id__in=added_project_ids)
                            .order_by("title")
                            .values_list("title", flat=True)
                        )
                        removed_titles = list(
                            Project.objects.filter(id__in=removed_project_ids)
                            .order_by("title")
                            .values_list("title", flat=True)
                        )
                        if added_titles:
                            details.append(
                                "Assigned synopses added: "
                                + ", ".join(added_titles)
                            )
                        if removed_titles:
                            details.append(
                                "Assigned synopses removed: "
                                + ", ".join(removed_titles)
                            )
                    if details:
                        messages.success(request, "User account updated.")
                    else:
                        messages.info(request, "No changes saved.")
                    return redirect("synopsis:manager_user_edit", user_id=managed_user.id)
        elif action == "send_access_email":
            form = ManagerUserUpdateForm(
                initial={
                    "first_name": managed_user.first_name,
                    "last_name": managed_user.last_name,
                    "email": managed_user.email,
                    "global_role": _manager_user_global_role(managed_user) or "author",
                    "is_active": managed_user.is_active,
                    "assigned_projects": sorted(existing_author_project_ids),
                },
                user=managed_user,
            )
            delete_form = ManagerUserDeleteForm(user=managed_user)
            if not managed_user.email:
                messages.error(request, "This account does not have an email address.")
            elif not managed_user.is_active:
                messages.error(request, "Reactivate the account before sending access emails.")
            else:
                try:
                    if managed_user.has_usable_password():
                        _send_password_reset_email(managed_user, request)
                        messages.success(
                            request,
                            f"Password reset email sent to {managed_user.email}.",
                        )
                    else:
                        _send_account_setup_email(managed_user, request)
                        messages.success(
                            request,
                            f"Account setup email sent to {managed_user.email}.",
                        )
                except Exception:
                    logger.exception("Failed to send access email for user %s", managed_user.id)
                    messages.error(request, "The access email could not be sent.")
                return redirect("synopsis:manager_user_edit", user_id=managed_user.id)
        elif action == "delete_user":
            form = ManagerUserUpdateForm(
                initial={
                    "first_name": managed_user.first_name,
                    "last_name": managed_user.last_name,
                    "email": managed_user.email,
                    "global_role": _manager_user_global_role(managed_user) or "author",
                    "is_active": managed_user.is_active,
                    "assigned_projects": sorted(existing_author_project_ids),
                },
                user=managed_user,
            )
            delete_form = ManagerUserDeleteForm(request.POST, user=managed_user)
            if managed_user == request.user:
                messages.error(request, "You cannot delete your own account.")
                return redirect("synopsis:manager_user_edit", user_id=managed_user.id)
            elif delete_form.is_valid():
                email = managed_user.email or managed_user.username
                managed_user.delete()
                messages.success(request, f"User {email} deleted.")
                return redirect("synopsis:manager_dashboard")
        else:
            form = ManagerUserUpdateForm(
                initial={
                    "first_name": managed_user.first_name,
                    "last_name": managed_user.last_name,
                    "email": managed_user.email,
                    "global_role": _manager_user_global_role(managed_user) or "author",
                    "is_active": managed_user.is_active,
                    "assigned_projects": sorted(existing_author_project_ids),
                },
                user=managed_user,
            )
            delete_form = ManagerUserDeleteForm(user=managed_user)
    else:
        form = ManagerUserUpdateForm(
            initial={
                "first_name": managed_user.first_name,
                "last_name": managed_user.last_name,
                "email": managed_user.email,
                "global_role": _manager_user_global_role(managed_user) or "author",
                "is_active": managed_user.is_active,
                "assigned_projects": sorted(existing_author_project_ids),
            },
            user=managed_user,
        )
        delete_form = ManagerUserDeleteForm(user=managed_user)

    return render(
        request,
        "synopsis/user_edit.html",
        {
            "managed_user": managed_user,
            "form": form,
            "delete_form": delete_form,
            "project_roles": project_roles,
            "password_state": (
                "Password set" if managed_user.has_usable_password() else "Setup pending"
            ),
            "global_role_label": _manager_user_global_role_label(managed_user),
        },
    )


@login_required
def project_settings(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if not _user_can_manage_project_configuration(request.user, project):
        messages.error(
            request,
            "You do not have permission to update project settings for this synopsis.",
        )
        return redirect("synopsis:project_hub", project_id=project.id)

    if request.method == "POST":
        status_action = request.POST.get("status_action")
        status_targets = {
            "mark_completed": "completed",
            "reactivate": "active",
        }
        if status_action in status_targets:
            old_status = project.status
            new_status = status_targets[status_action]
            status_labels = dict(Project._meta.get_field("status").choices)

            if old_status == new_status:
                messages.info(
                    request,
                    f"Synopsis is already marked as {status_labels.get(new_status, new_status).lower()}.",
                )
            else:
                project.status = new_status
                project.save(update_fields=["status"])
                _log_project_change(
                    project,
                    request.user,
                    "Updated project status",
                    "Status: "
                    f"{status_labels.get(old_status, old_status)} → "
                    f"{status_labels.get(new_status, new_status)}",
                )
                if new_status == "completed":
                    messages.success(
                        request,
                        "Synopsis moved to the completed / archived section on the homepage. It remains fully accessible.",
                    )
                else:
                    messages.success(
                        request,
                        "Synopsis moved back to the active synopses section on the homepage.",
                    )
            if request.POST.get("return_to") == "dashboard":
                return redirect("synopsis:dashboard")
            return redirect("synopsis:project_settings", project_id=project.id)

        original_title = project.title
        original_description = project.description
        original_protocol_relevant = project.protocol_relevant
        original_advisory_board_relevant = project.advisory_board_relevant
        form = ProjectSettingsForm(request.POST, instance=project, project=project)
        if form.is_valid():
            updated_project = form.save()
            changes = []
            if original_title != updated_project.title:
                changes.append(f"Title: {original_title} → {updated_project.title}")
            if original_description != updated_project.description:
                changes.append(
                    "Description: "
                    f"{_format_value(original_description)} → {_format_value(updated_project.description)}"
                )
            if original_protocol_relevant != updated_project.protocol_relevant:
                changes.append(
                    "Protocol: "
                    f"{'relevant' if original_protocol_relevant else 'not relevant'} → "
                    f"{'relevant' if updated_project.protocol_relevant else 'not relevant'}"
                )
            if (
                original_advisory_board_relevant
                != updated_project.advisory_board_relevant
            ):
                changes.append(
                    "Advisory board: "
                    f"{'relevant' if original_advisory_board_relevant else 'not relevant'} → "
                    f"{'relevant' if updated_project.advisory_board_relevant else 'not relevant'}"
                )
            if (
                updated_project.phase_manual
                and updated_project.phase_manual not in updated_project.available_phase_keys()
            ):
                old_phase = updated_project.phase_manual
                new_phase = updated_project.default_phase_key()
                updated_project.phase_manual = new_phase
                updated_project.phase_manual_updated = timezone.now()
                updated_project.save(
                    update_fields=["phase_manual", "phase_manual_updated"]
                )
                ProjectPhaseEvent.objects.create(
                    project=updated_project,
                    phase=new_phase,
                    confirmed_by=request.user,
                    note=(
                        "Phase reset because project settings removed an earlier phase from this workflow."
                    ),
                )
                changes.append(
                    "Phase reset: "
                    f"{dict(Project.PHASE_CHOICES).get(old_phase, old_phase)} → "
                    f"{dict(Project.PHASE_CHOICES).get(new_phase, new_phase)}"
                )
            if changes:
                _log_project_change(
                    updated_project,
                    request.user,
                    "Updated project settings",
                    "; ".join(changes),
                )
                messages.success(request, "Project settings updated.")
            else:
                messages.info(request, "No changes saved.")
            return redirect("synopsis:project_hub", project_id=project.id)
    else:
        form = ProjectSettingsForm(instance=project, project=project)

    previous_titles = []
    seen_titles = set()

    def add_title_entry(title, changed_at, changed_by, *, is_current=False, note=None):
        title = (title or "").strip()
        if not title or title in seen_titles:
            return
        previous_titles.append(
            {
                "title": title,
                "changed_at": changed_at,
                "changed_by": changed_by,
                "is_current": is_current,
                "note": note or "",
            }
        )
        seen_titles.add(title)

    change_logs = project.change_log.filter(
        action="Updated project settings",
        details__icontains="Title:",
    ).order_by("-created_at")

    for log in change_logs:
        actor = _user_display(log.changed_by) if log.changed_by else "System"
        segments = [
            segment.strip() for segment in log.details.split(";") if segment.strip()
        ]
        for segment in segments:
            if segment.startswith("Title:") and "→" in segment:
                old_part, new_part = segment.split("→", 1)
                old_title = old_part.split("Title:", 1)[1].strip()
                new_title = new_part.strip()
                add_title_entry(
                    new_title,
                    log.created_at,
                    actor,
                    is_current=(
                        new_title == project.title and project.title not in seen_titles
                    ),
                )
                add_title_entry(old_title, log.created_at, actor)

    if project.title not in seen_titles:
        created_at = getattr(project, "created_at", None)
        add_title_entry(
            project.title,
            created_at,
            "Created",
            is_current=True,
            note="Current title",
        )
    context = {
        "project": project,
        "form": form,
        "previous_titles": previous_titles,
        "phase_history_entries": project.phase_events.select_related("confirmed_by")[:10],
        **_project_phase_context(project, request.user),
    }

    return render(
        request,
        "synopsis/project_settings_form.html",
        context,
    )


@login_required
def user_create(request):
    if not request.user.is_staff:
        messages.error(request, "Manager access only.")
        return redirect("synopsis:dashboard")

    ensure_global_groups()

    if request.method == "POST":
        form = CreateUserForm(request.POST)
        if form.is_valid():
            email = form.cleaned_data["email"].strip().lower()
            first_name = form.cleaned_data["first_name"].strip()
            last_name = form.cleaned_data["last_name"].strip()
            global_role = form.cleaned_data["global_role"]
            assigned_projects = list(form.cleaned_data["assigned_projects"])

            if User.objects.filter(
                Q(username=email) | Q(email__iexact=email)
            ).exists():
                messages.error(request, "A user with that email already exists.")
            else:
                try:
                    with transaction.atomic():
                        user = User.objects.create_user(
                            username=email,
                            email=email,
                            password=None,
                            first_name=first_name,
                            last_name=last_name,
                        )
                        user.set_unusable_password()
                        if global_role == "manager":
                            user.is_staff = True
                        user.save()

                        group = Group.objects.get(name=global_role)
                        user.groups.add(group)
                        for project in assigned_projects:
                            UserRole.objects.get_or_create(
                                user=user, project=project, role="author"
                            )

                        _send_account_setup_email(user, request)
                except Exception:
                    logger.exception("Failed to create user or send account setup email.")
                    messages.error(
                        request,
                        "The account could not be created because the setup email was not sent.",
                    )
                else:
                    messages.success(
                        request,
                        f"User {email} created as {GLOBAL_ROLE_LABELS.get(global_role, global_role)}. A password setup email has been sent.",
                    )
                    return redirect("synopsis:manager_dashboard")
    else:
        form = CreateUserForm()

    return render(request, "synopsis/user_create.html", {"form": form})


# TODO: #22 Add search, filtering, and pagination to the advisory board list once larger projects need it.
# TODO: #23 Add CSV export for advisory board members and their response state.
# TODO: #25 Finish advisory-board access control and replace any remaining broad file access with scoped links or tokens.
# TODO: #40 Add a resend-invitation action that preserves the original member record and audit history.
# TODO: #39 Add bulk CSV import for advisory board members.

@login_required
def advisory_board_list(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if not project.advisory_board_relevant:
        messages.info(
            request,
            "Advisory board is marked as not relevant for this project. Update Project settings if you want to use the advisory board workflow.",
        )
        return redirect("synopsis:project_hub", project_id=project.id)

    if request.method == "POST":
        action = request.POST.get("action")

        if action == "custom_field_add":
            form = AdvisoryCustomFieldForm(project, request.POST)
            if form.is_valid():
                form.save()
                messages.success(request, "Custom column added to advisory board.")
                return redirect("synopsis:advisory_board_list", project_id=project.id)
            context = _advisory_board_context(
                project,
                user=request.user,
                member_form=AdvisoryBoardMemberForm(),
                custom_field_form=form,
            )
            return render(request, "synopsis/advisory_board_list.html", context)

        if action == "custom_field_delete":
            field_id = request.POST.get("field_id")
            if field_id:
                field = get_object_or_404(
                    AdvisoryBoardCustomField, pk=field_id, project=project
                )
                field.delete()
                messages.info(
                    request,
                    f"Removed custom column '{field.name}'.",
                )
            return redirect("synopsis:advisory_board_list", project_id=project.id)

        if action == "custom_field_move":
            field_id = request.POST.get("field_id")
            field = get_object_or_404(
                AdvisoryBoardCustomField, pk=field_id, project=project
            )
            placement_form = AdvisoryCustomFieldPlacementForm(request.POST)
            if placement_form.is_valid():
                field.display_group = placement_form.cleaned_data["display_group"]
                field.save(update_fields=["display_group"])
                messages.success(
                    request, f"Moved custom column '{field.name}' to a new section."
                )
                return redirect("synopsis:advisory_board_list", project_id=project.id)
            messages.error(request, "Choose a valid section for this column.")
            context = _advisory_board_context(
                project,
                user=request.user,
                member_form=AdvisoryBoardMemberForm(),
                custom_field_form=AdvisoryCustomFieldForm(project),
            )
            return render(request, "synopsis/advisory_board_list.html", context)

        if action == "add_member_confirm":
            form = AdvisoryBoardMemberForm(request.POST)
            if form.is_valid():
                m = form.save(commit=False)
                m.project = project
                m.save()
                messages.success(request, "Advisory Board member added.")
                return redirect("synopsis:advisory_board_list", project_id=project.id)
            context = _advisory_board_context(project, user=request.user, member_form=form)
            return render(request, "synopsis/advisory_board_list.html", context)

        if action == "add_member_back":
            form = AdvisoryBoardMemberForm(request.POST)
            context = _advisory_board_context(project, user=request.user, member_form=form)
            context["open_add_member_modal"] = True
            return render(request, "synopsis/advisory_board_list.html", context)

        if action == "add_member":
            form = AdvisoryBoardMemberForm(request.POST)
            if form.is_valid():
                cleaned = form.cleaned_data
                return render(
                    request,
                    "synopsis/advisory_member_confirm.html",
                    {
                        "project": project,
                        "form": form,
                        "cleaned_data": cleaned,
                    },
                )
            context = _advisory_board_context(project, user=request.user, member_form=form)
            return render(request, "synopsis/advisory_board_list.html", context)

    form = AdvisoryBoardMemberForm()
    context = _advisory_board_context(project, user=request.user, member_form=form)
    return render(request, "synopsis/advisory_board_list.html", context)


@login_required
def advisory_schedule_reminders(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if request.method != "POST":
        return HttpResponseBadRequest("POST required")

    form = ReminderScheduleForm(request.POST)
    pending_members = project.advisory_board_members.filter(invite_sent=False)

    if not form.is_valid():
        context = _advisory_board_context(project, user=request.user, reminder_form=form)
        return render(request, "synopsis/advisory_board_list.html", context)

    reminder_date = form.cleaned_data["reminder_date"]
    updated = 0
    for member in pending_members:
        member.response_date = reminder_date
        member.reminder_sent = False
        member.save(update_fields=["response_date", "reminder_sent"])
        updated += 1

    if updated:
        _log_project_change(
            project,
            request.user,
            "Scheduled reminders",
            f"Date set to {reminder_date} for {updated} pending member(s)",
        )

    messages.success(request, f"Scheduled reminders for {updated} member(s).")
    return redirect("synopsis:advisory_board_list", project_id=project.id)


@login_required
def advisory_schedule_protocol_reminders(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if request.method != "POST":
        return HttpResponseBadRequest("POST required")

    form = ProtocolReminderScheduleForm(request.POST)
    pending_members = project.advisory_board_members.filter(
        sent_protocol_at__isnull=False,
        response="Y",
    )
    if not pending_members.exists():
        messages.warning(
            request,
            "No protocol deadline was updated because no accepted advisory board member has been sent the protocol yet. Send the protocol from the Advisory Board page first, or set the deadline while sending.",
        )
        return redirect("synopsis:protocol_detail", project_id=project.id)

    if not form.is_valid():
        for errors in form.errors.values():
            for error in errors:
                messages.error(request, error)
        return redirect("synopsis:protocol_detail", project_id=project.id)

    deadline = form.cleaned_data["deadline"]
    if timezone.is_naive(deadline):
        deadline = timezone.make_aware(deadline)
    updated = 0
    for member in pending_members:
        member.feedback_on_protocol_deadline = deadline
        member.protocol_reminder_sent = False
        member.protocol_reminder_sent_at = None
        member.save(
            update_fields=[
                "feedback_on_protocol_deadline",
                "protocol_reminder_sent",
                "protocol_reminder_sent_at",
            ]
        )
        ProtocolFeedback.objects.filter(project=project, member=member).update(
            feedback_deadline_at=deadline
        )
        updated += 1

    skipped_members = project.advisory_board_members.filter(
        sent_protocol_at__isnull=False
    ).exclude(response="Y")
    skipped_ids = list(skipped_members.values_list("id", flat=True))
    if skipped_ids:
        project.advisory_board_members.filter(id__in=skipped_ids).update(
            feedback_on_protocol_deadline=None,
            protocol_reminder_sent=False,
            protocol_reminder_sent_at=None,
        )
        ProtocolFeedback.objects.filter(
            project=project, member_id__in=skipped_ids
        ).update(feedback_deadline_at=None)
        messages.info(
            request,
            "Deadline kept unset for members who have not accepted the invitation yet.",
        )

    if updated:
        _log_project_change(
            project,
            request.user,
            "Scheduled protocol reminders",
            f"Protocol deadline {timezone.localtime(deadline).strftime('%Y-%m-%d %H:%M')} for {updated} member(s)",
        )

    messages.success(
        request,
        f"Protocol reminder scheduled for {updated} member(s). Reminder now set as required.",
    )
    return redirect("synopsis:protocol_detail", project_id=project.id)


@login_required
def advisory_schedule_action_list_reminders(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if request.method != "POST":
        return HttpResponseBadRequest("POST required")

    if not getattr(project, "action_list", None):
        messages.error(request, "No action list configured for this project.")
        return redirect("synopsis:action_list_detail", project_id=project.id)

    form = ActionListReminderScheduleForm(request.POST)
    pending_members = project.advisory_board_members.filter(
        sent_action_list_at__isnull=False,
        response="Y",
    )
    if not pending_members.exists():
        messages.warning(
            request,
            "No action list deadline was updated because no accepted advisory board member has been sent the action list yet. Send the action list from the Advisory Board page first, or set the deadline while sending.",
        )
        return redirect("synopsis:action_list_detail", project_id=project.id)

    if not form.is_valid():
        for errors in form.errors.values():
            for error in errors:
                messages.error(request, error)
        return redirect("synopsis:action_list_detail", project_id=project.id)

    deadline = form.cleaned_data["deadline"]
    if timezone.is_naive(deadline):
        deadline = timezone.make_aware(deadline)
    updated = 0
    for member in pending_members:
        member.feedback_on_action_list_deadline = deadline
        member.action_list_reminder_sent = False
        member.action_list_reminder_sent_at = None
        member.save(
            update_fields=[
                "feedback_on_action_list_deadline",
                "action_list_reminder_sent",
                "action_list_reminder_sent_at",
            ]
        )
        ActionListFeedback.objects.filter(project=project, member=member).update(
            feedback_deadline_at=deadline
        )
        updated += 1

    skipped_members = project.advisory_board_members.filter(
        sent_action_list_at__isnull=False
    ).exclude(response="Y")
    skipped_ids = list(skipped_members.values_list("id", flat=True))
    if skipped_ids:
        project.advisory_board_members.filter(id__in=skipped_ids).update(
            feedback_on_action_list_deadline=None,
            action_list_reminder_sent=False,
            action_list_reminder_sent_at=None,
        )
        ActionListFeedback.objects.filter(
            project=project, member_id__in=skipped_ids
        ).update(feedback_deadline_at=None)
        messages.info(
            request,
            "Deadline kept unset for members who have not accepted the invitation yet.",
        )

    if updated:
        _log_project_change(
            project,
            request.user,
            "Scheduled action list reminders",
            f"Action list deadline {timezone.localtime(deadline).strftime('%Y-%m-%d %H:%M')} for {updated} member(s)",
        )

    messages.success(
        request,
        f"Action list reminder scheduled for {updated} member(s). Reminder now set as required.",
    )
    return redirect("synopsis:action_list_detail", project_id=project.id)


@login_required
def advisory_schedule_synopsis_reminders(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if request.method != "POST":
        return HttpResponseBadRequest("POST required")

    form = SynopsisReminderScheduleForm(request.POST)
    pending_members = project.advisory_board_members.filter(
        sent_synopsis_at__isnull=False,
        response="Y",
    )
    if not pending_members.exists():
        messages.warning(
            request,
            "No synopsis deadline was updated because no accepted advisory board member has been sent the synopsis yet. Send the synopsis from the Advisory Board page first, or set the deadline while sending.",
        )
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    if not form.is_valid():
        context = _advisory_board_context(
            project,
            user=request.user,
            synopsis_form=form,
        )
        return render(request, "synopsis/advisory_board_list.html", context)

    deadline = form.cleaned_data["deadline"]
    if timezone.is_naive(deadline):
        deadline = timezone.make_aware(deadline)
    updated = 0
    for member in pending_members:
        member.feedback_on_synopsis_deadline = deadline
        member.synopsis_reminder_sent = False
        member.synopsis_reminder_sent_at = None
        member.save(
            update_fields=[
                "feedback_on_synopsis_deadline",
                "synopsis_reminder_sent",
                "synopsis_reminder_sent_at",
            ]
        )
        SynopsisFeedback.objects.filter(project=project, member=member).update(
            feedback_deadline_at=deadline
        )
        updated += 1

    skipped_members = project.advisory_board_members.filter(
        sent_synopsis_at__isnull=False
    ).exclude(response="Y")
    skipped_ids = list(skipped_members.values_list("id", flat=True))
    if skipped_ids:
        project.advisory_board_members.filter(id__in=skipped_ids).update(
            feedback_on_synopsis_deadline=None,
            synopsis_reminder_sent=False,
            synopsis_reminder_sent_at=None,
        )
        SynopsisFeedback.objects.filter(
            project=project, member_id__in=skipped_ids
        ).update(feedback_deadline_at=None)
        messages.info(
            request,
            "Deadline kept unset for members who have not accepted the invitation yet.",
        )

    if updated:
        _log_project_change(
            project,
            request.user,
            "Scheduled synopsis reminders",
            f"Synopsis deadline {timezone.localtime(deadline).strftime('%Y-%m-%d %H:%M')} for {updated} member(s)",
        )

    messages.success(
        request,
        f"Synopsis reminder scheduled for {updated} member(s). Reminder now set as required.",
    )
    return redirect("synopsis:advisory_board_list", project_id=project.id)


@login_required
def advisory_member_set_deadline(request, project_id, member_id, kind):
    project = get_object_or_404(Project, pk=project_id)
    member = get_object_or_404(AdvisoryBoardMember, pk=member_id, project=project)

    if request.method != "POST":
        return HttpResponseBadRequest("POST required")

    if not _user_can_edit_project(request.user, project):
        messages.error(request, "You do not have permission to update this member.")
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    kind = (kind or "").lower().replace("-", "_")
    form_map = {
        "invite": ReminderScheduleForm,
        "protocol": ProtocolReminderScheduleForm,
        "action_list": ActionListReminderScheduleForm,
        "synopsis": SynopsisReminderScheduleForm,
    }
    if kind not in form_map:
        return HttpResponseBadRequest("Unknown reminder type")

    clearing = request.POST.get("clear_deadline")
    response_code = (member.response or "").upper()

    if kind == "invite" and response_code == "N":
        messages.info(
            request,
            "Reminders are skipped for members who declined the invitation.",
        )
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    if kind == "protocol":
        if member.sent_protocol_at is None or response_code != "Y":
            messages.error(
                request,
                "This member needs an accepted invitation and a sent protocol before setting a protocol deadline.",
            )
            return redirect("synopsis:advisory_board_list", project_id=project.id)
    if kind == "action_list":
        if not getattr(project, "action_list", None):
            messages.error(request, "No action list configured for this project.")
            return redirect("synopsis:advisory_board_list", project_id=project.id)
        if member.sent_action_list_at is None or response_code != "Y":
            messages.error(
                request,
                "This member needs an accepted invitation and the action list before setting an action list deadline.",
            )
            return redirect("synopsis:advisory_board_list", project_id=project.id)
    if kind == "synopsis":
        if member.sent_synopsis_at is None or response_code != "Y":
            messages.error(
                request,
                "This member needs an accepted invitation and a sent synopsis before setting a synopsis deadline.",
            )
            return redirect("synopsis:advisory_board_list", project_id=project.id)
    value = None
    if not clearing:
        form = form_map[kind](request.POST)
        if not form.is_valid():
            error_messages = ", ".join(
                {" ".join(err_list) for err_list in form.errors.values()}
            )
            messages.error(
                request,
                f"Could not update reminder: {error_messages or 'Invalid data.'}",
            )
            return redirect("synopsis:advisory_board_list", project_id=project.id)
        if kind == "invite":
            value = form.cleaned_data["reminder_date"]
        else:
            value = form.cleaned_data["deadline"]
            if timezone.is_naive(value):
                value = timezone.make_aware(value)

    human_name = " ".join(
        part for part in (member.first_name, member.last_name) if part
    ).strip() or member.email

    if kind == "invite":
        member.response_date = value if not clearing else None
        member.reminder_sent = False
        member.reminder_sent_at = None
        member.save(update_fields=["response_date", "reminder_sent", "reminder_sent_at"])
        detail = (
            f"Response deadline {value} for {human_name}"
            if not clearing
            else f"Cleared response deadline for {human_name}"
        )
        _log_project_change(project, request.user, "Updated invite reminder", detail)
        messages.success(
            request,
            "Response deadline updated."
            if not clearing
            else "Response deadline cleared.",
        )
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    if kind == "protocol":
        member.feedback_on_protocol_deadline = value if not clearing else None
        member.protocol_reminder_sent = False
        member.protocol_reminder_sent_at = None
        member.save(
            update_fields=[
                "feedback_on_protocol_deadline",
                "protocol_reminder_sent",
                "protocol_reminder_sent_at",
            ]
        )
        ProtocolFeedback.objects.filter(project=project, member=member).update(
            feedback_deadline_at=value if not clearing else None
        )
        detail_value = (
            timezone.localtime(value).strftime("%Y-%m-%d %H:%M") if value else None
        )
        detail = (
            f"Protocol deadline {detail_value} for {human_name}"
            if not clearing
            else f"Cleared protocol deadline for {human_name}"
        )
        _log_project_change(
            project,
            request.user,
            "Updated protocol reminder",
            detail,
        )
        messages.success(
            request,
            "Protocol deadline updated."
            if not clearing
            else "Protocol deadline cleared.",
        )
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    if kind == "synopsis":
        member.feedback_on_synopsis_deadline = value if not clearing else None
        member.synopsis_reminder_sent = False
        member.synopsis_reminder_sent_at = None
        member.save(
            update_fields=[
                "feedback_on_synopsis_deadline",
                "synopsis_reminder_sent",
                "synopsis_reminder_sent_at",
            ]
        )
        SynopsisFeedback.objects.filter(project=project, member=member).update(
            feedback_deadline_at=value if not clearing else None
        )
        detail_value = (
            timezone.localtime(value).strftime("%Y-%m-%d %H:%M") if value else None
        )
        detail = (
            f"Synopsis deadline {detail_value} for {human_name}"
            if not clearing
            else f"Cleared synopsis deadline for {human_name}"
        )
        _log_project_change(
            project,
            request.user,
            "Updated synopsis reminder",
            detail,
        )
        messages.success(
            request,
            "Synopsis deadline updated."
            if not clearing
            else "Synopsis deadline cleared.",
        )
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    # action list
    member.feedback_on_action_list_deadline = value if not clearing else None
    member.action_list_reminder_sent = False
    member.action_list_reminder_sent_at = None
    member.save(
        update_fields=[
            "feedback_on_action_list_deadline",
            "action_list_reminder_sent",
            "action_list_reminder_sent_at",
        ]
    )
    ActionListFeedback.objects.filter(project=project, member=member).update(
        feedback_deadline_at=value if not clearing else None
    )
    detail_value = (
        timezone.localtime(value).strftime("%Y-%m-%d %H:%M") if value else None
    )
    detail = (
        f"Action list deadline {detail_value} for {human_name}"
        if not clearing
        else f"Cleared action list deadline for {human_name}"
    )
    _log_project_change(
        project,
        request.user,
        "Updated action list reminder",
        detail,
    )
    messages.success(
        request,
        "Action list deadline updated."
        if not clearing
        else "Action list deadline cleared.",
    )
    return redirect("synopsis:advisory_board_list", project_id=project.id)


def _advisory_member_set_tracking_flag(
    request,
    project,
    member,
    *,
    document_name,
    sent_field,
    feedback_received_field,
    latest_feedback,
    flag,
    flag_map,
    log_action,
    success_message,
):
    if request.method != "POST":
        return HttpResponseBadRequest("POST required")

    if not _user_can_edit_project(request.user, project):
        messages.error(request, "You do not have permission to update this member.")
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    if getattr(member, sent_field) is None or (member.response or "").upper() != "Y":
        messages.error(
            request,
            f"This member needs an accepted invitation and sent {document_name} before updating {document_name} tracking.",
        )
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    normalized_flag = (flag or "").lower().replace("-", "_")
    if normalized_flag not in flag_map:
        return HttpResponseBadRequest(f"Unknown {document_name} tracking field")

    has_feedback = bool(
        getattr(member, feedback_received_field)
        or getattr(latest_feedback, "submitted_at", None)
    )
    if not has_feedback:
        messages.error(
            request,
            f"{document_name.capitalize()} tracking can only be updated after feedback has been received.",
        )
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    field_name, label = flag_map[normalized_flag]
    raw_value = request.POST.get("value", "")
    new_value = str(raw_value).lower() in {"1", "true", "on", "yes"}
    current_value = getattr(member, field_name)
    if current_value != new_value:
        setattr(member, field_name, new_value)
        member.save(update_fields=[field_name])
        human_name = " ".join(
            part for part in (member.first_name, member.last_name) if part
        ).strip() or member.email
        state = "Marked" if new_value else "Cleared"
        _log_project_change(
            project,
            request.user,
            log_action,
            f"{state} {label} for {human_name}",
        )

    messages.success(request, success_message)
    return redirect("synopsis:advisory_board_list", project_id=project.id)


@login_required
def advisory_member_set_action_list_flag(request, project_id, member_id, flag):
    project = get_object_or_404(Project, pk=project_id)
    member = get_object_or_404(AdvisoryBoardMember, pk=member_id, project=project)
    return _advisory_member_set_tracking_flag(
        request,
        project,
        member,
        document_name="action list",
        sent_field="sent_action_list_at",
        feedback_received_field="feedback_on_action_list_received",
        latest_feedback=member.latest_action_list_feedback,
        flag=flag,
        flag_map={
            "author_replied": ("wm_replied", "author replied"),
            "added_to_doc": (
                "added_to_action_list_doc",
                "feedback added to the action list document",
            ),
        },
        log_action="Updated action list tracking",
        success_message="Action list tracking updated.",
    )


@login_required
def advisory_member_set_protocol_flag(request, project_id, member_id, flag):
    project = get_object_or_404(Project, pk=project_id)
    member = get_object_or_404(AdvisoryBoardMember, pk=member_id, project=project)
    return _advisory_member_set_tracking_flag(
        request,
        project,
        member,
        document_name="protocol",
        sent_field="sent_protocol_at",
        feedback_received_field="feedback_on_protocol_received",
        latest_feedback=member.latest_protocol_feedback,
        flag=flag,
        flag_map={
            "author_replied": (
                "protocol_author_replied",
                "author replied to protocol feedback",
            ),
            "added_to_doc": (
                "added_to_protocol_doc",
                "feedback added to the protocol document",
            ),
        },
        log_action="Updated protocol tracking",
        success_message="Protocol tracking updated.",
    )


@login_required
def advisory_member_set_synopsis_flag(request, project_id, member_id, flag):
    project = get_object_or_404(Project, pk=project_id)
    member = get_object_or_404(AdvisoryBoardMember, pk=member_id, project=project)
    normalized_flag = (flag or "").lower().replace("-", "_")
    if normalized_flag == "feedback_received":
        if request.method != "POST":
            return HttpResponseBadRequest("POST required")

        if not _user_can_edit_project(request.user, project):
            messages.error(request, "You do not have permission to update this member.")
            return redirect("synopsis:advisory_board_list", project_id=project.id)

        if member.sent_synopsis_at is None or (member.response or "").upper() != "Y":
            messages.error(
                request,
                "This member needs an accepted invitation and sent synopsis before updating synopsis tracking.",
            )
            return redirect("synopsis:advisory_board_list", project_id=project.id)

        raw_value = request.POST.get("value", "")
        new_value = str(raw_value).lower() in {"1", "true", "on", "yes"}
        update_fields = []
        if new_value and not member.feedback_on_synopsis_received:
            member.feedback_on_synopsis_received = timezone.localdate()
            update_fields.append("feedback_on_synopsis_received")
        elif not new_value and member.feedback_on_synopsis_received:
            member.feedback_on_synopsis_received = None
            member.synopsis_author_replied = False
            member.added_to_synopsis_doc = False
            update_fields.extend(
                [
                    "feedback_on_synopsis_received",
                    "synopsis_author_replied",
                    "added_to_synopsis_doc",
                ]
            )

        if update_fields:
            member.save(update_fields=update_fields)
            human_name = " ".join(
                part for part in (member.first_name, member.last_name) if part
            ).strip() or member.email
            state = "Marked feedback received" if new_value else "Cleared feedback received"
            _log_project_change(
                project,
                request.user,
                "Updated synopsis tracking",
                f"{state} for {human_name}",
            )

        messages.success(request, "Synopsis tracking updated.")
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    return _advisory_member_set_tracking_flag(
        request,
        project,
        member,
        document_name="synopsis",
        sent_field="sent_synopsis_at",
        feedback_received_field="feedback_on_synopsis_received",
        latest_feedback=member.latest_synopsis_feedback,
        flag=flag,
        flag_map={
            "author_replied": (
                "synopsis_author_replied",
                "author replied to synopsis feedback",
            ),
            "added_to_doc": (
                "added_to_synopsis_doc",
                "feedback added to the synopsis document",
            ),
        },
        log_action="Updated synopsis tracking",
        success_message="Synopsis tracking updated.",
    )


@login_required
def advisory_member_edit(request, project_id, member_id):
    project = get_object_or_404(Project, pk=project_id)
    member = get_object_or_404(AdvisoryBoardMember, pk=member_id, project=project)

    if not _user_can_edit_project(request.user, project):
        messages.error(request, "You do not have permission to update this member.")
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    if request.method == "POST":
        action = request.POST.get("action")
        if action == "delete_member":
            display_name = _advisory_member_display(member)
            member_email = member.email
            member_id_value = member.id
            with transaction.atomic():
                _log_project_change(
                    project,
                    request.user,
                    "Deleted advisory member",
                    f"Removed {display_name} ({member_email}) from the advisory board; member id {member_id_value}.",
                )
                member.delete()
            messages.success(request, f"Deleted advisory board member {display_name}.")
            return redirect("synopsis:advisory_board_list", project_id=project.id)

        form = AdvisoryBoardMemberForm(request.POST, instance=member)
        if form.is_valid():
            updated_member = form.save()
            display_name = f"{updated_member.first_name} {updated_member.last_name or ''}".strip() or updated_member.email
            _log_project_change(
                project,
                request.user,
                "Updated advisory member",
                f"Edited details for {display_name}",
            )
            messages.success(request, "Member details updated.")
            return redirect("synopsis:advisory_board_list", project_id=project.id)
    else:
        form = AdvisoryBoardMemberForm(instance=member)

    return render(
        request,
        "synopsis/advisory_member_edit.html",
        {
            "project": project,
            "member": member,
            "form": form,
        },
    )


@login_required
def reference_batch_list(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    batches_qs = project.reference_batches.select_related("uploaded_by").order_by(
        "-created_at", "-id"
    )
    batches = list(batches_qs)
    project_references = project.references.select_related("screened_by")

    status_counts = {
        row["screening_status"]: row["count"]
        for row in project_references.values("screening_status").annotate(
            count=Count("id")
        )
    }
    included_count = status_counts.get("included", 0)
    excluded_count = status_counts.get("excluded", 0)
    pending_count = status_counts.get("pending", 0)
    total_references = project_references.count()
    screened_count = included_count + excluded_count
    completion_percent = (
        round((screened_count / total_references) * 100)
        if total_references
        else 0
    )

    latest_screening = (
        project_references.exclude(screening_decision_at__isnull=True)
        .order_by("-screening_decision_at")
        .first()
    )

    batch_stats = defaultdict(
        lambda: {"included": 0, "excluded": 0, "pending": 0, "total": 0}
    )
    for row in project_references.values("batch_id", "screening_status").annotate(
        count=Count("id")
    ):
        batch_id = row["batch_id"]
        if batch_id is None:
            continue
        stats = batch_stats[batch_id]
        stats["total"] += row["count"]
        status = row["screening_status"]
        if status == "included":
            stats["included"] += row["count"]
        elif status == "excluded":
            stats["excluded"] += row["count"]
        else:
            stats["pending"] += row["count"]

    for batch in batches:
        stats = batch_stats.get(batch.id, {"included": 0, "excluded": 0, "pending": 0, "total": 0})
        screened = stats["included"] + stats["excluded"]
        total = stats["total"]
        batch.screened_count = screened
        batch.total_references = total
        batch.pending_count = stats["pending"]
        batch.included_count = stats["included"]
        batch.excluded_count = stats["excluded"]
        batch.progress_percent = (
            round((screened / total) * 100) if total else 0
        )

    next_batch = next((batch for batch in batches if batch.pending_count), None)
    summary = {
        "total": total_references,
        "screened": screened_count,
        "included": included_count,
        "excluded": excluded_count,
        "pending": pending_count,
        "completion_percent": completion_percent,
        "last_screened_at": latest_screening.screening_decision_at
        if latest_screening
        else None,
        "last_screened_by": latest_screening.screened_by
        if latest_screening
        else None,
    }

    return render(
        request,
        "synopsis/reference_batch_list.html",
        {
            "project": project,
            "batches": batches,
            "summary": summary,
            "next_batch": next_batch,
        },
    )


def _link_library_references_to_project(user, target_project, ref_ids, folder):
    folder = normalize_reference_folder_values(folder or [])
    batch = None

    linked = 0
    reused = 0
    for rid in ref_ids:
        lib_ref = LibraryReference.objects.filter(pk=rid).first()
        if not lib_ref:
            continue
        if Reference.objects.filter(
            project=target_project, library_reference=lib_ref
        ).exists():
            reused += 1
            continue

        hash_key = reference_hash(
            lib_ref.title,
            str(lib_ref.publication_year or ""),
            lib_ref.doi,
        )
        if Reference.objects.filter(
            project=target_project, hash_key=hash_key
        ).exists():
            reused += 1
            continue

        if folder:
            _update_shared_library_reference_folders(
                lib_ref,
                folder,
                changed_by=user,
                source_project=target_project,
                change_source="library_link",
            )

        if batch is None:
            now = timezone.now()
            batch = ReferenceSourceBatch.objects.create(
                project=target_project,
                label=f"Library link {now:%Y-%m-%d %H:%M:%S} {uuid.uuid4().hex[:8]}",
                source_type="library_link",
                uploaded_by=user if user.is_authenticated else None,
                record_count=0,
            )

        Reference.objects.create(
            project=target_project,
            batch=batch,
            library_reference=lib_ref,
            hash_key=hash_key,
            source_identifier=lib_ref.source_identifier,
            title=lib_ref.title,
            abstract=lib_ref.abstract,
            authors=lib_ref.authors,
            publication_year=lib_ref.publication_year,
            journal=lib_ref.journal,
            volume=lib_ref.volume,
            issue=lib_ref.issue,
            pages=lib_ref.pages,
            doi=lib_ref.doi,
            url=lib_ref.url,
            language=lib_ref.language,
            raw_ris=lib_ref.raw_ris or {},
            reference_document=lib_ref.reference_document,
            reference_document_uploaded_at=lib_ref.reference_document_uploaded_at,
            screening_status="pending",
        )
        linked += 1

    if batch is not None:
        batch.record_count = batch.references.count()
        batch.save(update_fields=["record_count"])
    return linked, reused, batch


@login_required
def reference_library(request):
    if not _user_can_manage_library(request.user):
        raise PermissionDenied

    q = (request.GET.get("q") or "").strip()
    batch_id = request.GET.get("batch")
    project_id = (request.GET.get("project") or request.POST.get("project") or "").strip()
    next_url = (request.GET.get("next") or request.POST.get("next") or "").strip()
    safe_next_url = ""
    if next_url and url_has_allowed_host_and_scheme(
        next_url,
        allowed_hosts={request.get_host()},
        require_https=request.is_secure(),
    ):
        safe_next_url = next_url
    refs = LibraryReference.objects.select_related("import_batch").order_by("-created_at")
    if batch_id and batch_id.isdigit():
        refs = refs.filter(import_batch_id=batch_id)
    if q:
        refs = refs.filter(
            Q(title__icontains=q)
            | Q(authors__icontains=q)
            | Q(doi__icontains=q)
            | Q(journal__icontains=q)
        )
    refs = refs.annotate(project_count=Count("project_references", distinct=True))

    if request.method == "POST":
        action = request.POST.get("action")
        if action in ["link-to-project", "link-multi"]:
            target_project_id = request.POST.get("target_project")
            target_project = get_object_or_404(Project, pk=target_project_id)
            if not _user_can_edit_project(request.user, target_project):
                raise PermissionDenied

            ref_ids = []
            if action == "link-multi":
                ref_ids = [rid for rid in request.POST.getlist("selected_refs") if rid.isdigit()]
            else:
                single_id = request.POST.get("reference_id")
                if single_id and single_id.isdigit():
                    ref_ids = [single_id]
            if not ref_ids:
                messages.info(request, "Select at least one reference to link.")
                return redirect("synopsis:reference_library")

            folder = [f for f in request.POST.getlist("reference_folder") if f]
            linked, reused, batch = _link_library_references_to_project(
                request.user, target_project, ref_ids, folder
            )

            if linked or reused:
                parts = []
                if linked:
                    parts.append(f"Linked {linked} reference(s)")
                if reused:
                    parts.append(f"Reused {reused} existing reference(s)")
                msg = " and ".join(parts) + f" into {target_project.title}."
                if folder:
                    msg += " Shared CE subject categories were updated before linking."
                messages.success(request, msg)
            else:
                messages.info(request, "No references were linked (possible duplicates).")
            if safe_next_url:
                return redirect(safe_next_url)
            return redirect("synopsis:reference_library")

    project_options = (
        Project.objects.all()
        if request.user.is_staff
        else Project.objects.filter(userrole__user=request.user).distinct()
    )
    selected_project = None
    selected_project_obj = None
    if project_id.isdigit():
        selected_id = int(project_id)
        selected_project_obj = project_options.filter(id=selected_id).first()
        if selected_project_obj:
            selected_project = selected_id
    batch_options = LibraryImportBatch.objects.order_by("-created_at")

    return render(
        request,
        "synopsis/reference_library.html",
        {
            "references": refs,
            "q": q,
            "batch_options": batch_options,
            "selected_batch": int(batch_id) if batch_id and batch_id.isdigit() else None,
            "project_options": project_options,
            "selected_project": selected_project,
            "selected_project_obj": selected_project_obj,
            "next_url": safe_next_url,
            "folder_choices": Reference.FOLDER_CHOICES,
        },
    )


@login_required
def library_batch_list(request):
    if not _user_can_manage_library(request.user):
        raise PermissionDenied

    q = (request.GET.get("q") or "").strip()
    batches = LibraryImportBatch.objects.order_by("-created_at")
    if q:
        batches = batches.filter(
            Q(label__icontains=q) | Q(original_filename__icontains=q)
        )
    batches = batches.annotate(reference_count=Count("references", distinct=True))

    return render(
        request,
        "synopsis/library_batch_list.html",
        {"batches": batches, "q": q},
    )


@login_required
def library_batch_detail(request, batch_id):
    if not _user_can_manage_library(request.user):
        raise PermissionDenied

    batch = get_object_or_404(LibraryImportBatch, pk=batch_id)
    q = (request.GET.get("q") or "").strip()
    refs = (
        batch.references.select_related("import_batch")
        .order_by("-created_at")
    )
    if q:
        refs = refs.filter(
            Q(title__icontains=q)
            | Q(authors__icontains=q)
            | Q(doi__icontains=q)
            | Q(journal__icontains=q)
        )
    refs = refs.annotate(project_count=Count("project_references", distinct=True))

    if request.method == "POST":
        action = request.POST.get("action")
        if action == "link-multi":
            target_project_id = request.POST.get("target_project")
            target_project = get_object_or_404(Project, pk=target_project_id)
            if not _user_can_edit_project(request.user, target_project):
                raise PermissionDenied

            ref_ids = [rid for rid in request.POST.getlist("selected_refs") if rid.isdigit()]
            if not ref_ids:
                messages.info(request, "Select at least one reference to link.")
                return redirect("synopsis:library_batch_detail", batch_id=batch.id)

            valid_ids = list(
                batch.references.filter(id__in=ref_ids).values_list("id", flat=True)
            )
            folder = [f for f in request.POST.getlist("reference_folder") if f]
            linked, reused, _ = _link_library_references_to_project(
                request.user, target_project, valid_ids, folder
            )
            if linked or reused:
                parts = []
                if linked:
                    parts.append(f"Linked {linked} reference(s)")
                if reused:
                    parts.append(f"Reused {reused} existing reference(s)")
                msg = " and ".join(parts) + f" into {target_project.title}."
                if folder:
                    msg += " Shared CE subject categories were updated before linking."
                messages.success(request, msg)
            else:
                messages.info(request, "No references were linked (possible duplicates).")
            return redirect("synopsis:library_batch_detail", batch_id=batch.id)

    project_options = (
        Project.objects.all()
        if request.user.is_staff
        else Project.objects.filter(userrole__user=request.user).distinct()
    )

    return render(
        request,
        "synopsis/library_batch_detail.html",
        {
            "batch": batch,
            "references": refs,
            "q": q,
            "project_options": project_options,
            "folder_choices": Reference.FOLDER_CHOICES,
        },
    )


@login_required
def library_reference_batch_upload(request):
    if not _user_can_manage_library(request.user):
        raise PermissionDenied

    if request.method == "POST":
        form = LibraryReferenceBatchUploadForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = form.cleaned_data["ris_file"]
            raw_bytes = uploaded_file.read()
            if not raw_bytes.strip():
                form.add_error("ris_file", "The uploaded file appears to be empty.")
            else:
                sha1 = hashlib.sha1(raw_bytes).hexdigest()
                text_payload = _decode_reference_upload_text(raw_bytes)
                records = []
                ris_error = None
                plaintext_used = False
                file_ext = os.path.splitext(getattr(uploaded_file, "name", ""))[1].lower()

                if file_ext == ".xml":
                    records = _parse_endnote_xml(text_payload)
                else:
                    try:
                        records = rispy.loads(text_payload)
                    except Exception as exc:  # pragma: no cover - parser errors
                        ris_error = exc

                    if not records:
                        plaintext_records = _parse_plaintext_references(text_payload)
                        if plaintext_records:
                            records = plaintext_records
                            plaintext_used = True

                if not records:
                    if ris_error:
                        form.add_error(
                            "ris_file",
                            f"Could not parse RIS content ({ris_error}).",
                        )
                    else:
                        form.add_error(
                            "ris_file",
                            "No references were detected. Upload a RIS, XML, or plain text file where each entry is separated by a blank line.",
                        )

                if records:
                    with transaction.atomic():
                        batch = LibraryImportBatch.objects.create(
                            label=form.cleaned_data["label"],
                            source_type=form.cleaned_data["source_type"],
                            search_date_start=form.cleaned_data.get(
                                "search_date_start"
                            ),
                            search_date_end=form.cleaned_data.get(
                                "search_date_end"
                            ),
                            uploaded_by=request.user,
                            original_filename=getattr(uploaded_file, "name", ""),
                            record_count=0,
                            ris_sha1=sha1,
                            notes=form.cleaned_data.get("notes", ""),
                        )
                        imported = 0
                        duplicates = 0
                        skipped = 0
                        for record in records:
                            data = _normalise_import_record(record)
                            if not data:
                                skipped += 1
                                continue

                            hash_key = reference_hash(
                                data["title"], data["year"], data["doi"]
                            )
                            raw_source = record.get("_raw_source", "")
                            if file_ext == ".xml":
                                raw_source_format = "endnote_xml"
                            elif plaintext_used or file_ext == ".txt":
                                raw_source_format = "plaintext"
                            else:
                                raw_source_format = "ris"

                            _, created = LibraryReference.objects.get_or_create(
                                hash_key=hash_key,
                                defaults={
                                    "import_batch": batch,
                                    "source_identifier": data["source_identifier"],
                                    "title": data["title"],
                                    "abstract": data["abstract"],
                                    "authors": data["authors"],
                                    "publication_year": data["publication_year"],
                                    "journal": data["journal"],
                                    "volume": data["volume"],
                                    "issue": data["issue"],
                                    "pages": data["pages"],
                                    "doi": data["doi"],
                                    "url": data["url"],
                                    "language": data["language"],
                                    "raw_ris": record,
                                    "raw_source": raw_source,
                                    "raw_source_format": raw_source_format,
                                },
                            )
                            if not created:
                                duplicates += 1
                                continue
                            imported += 1

                        batch.record_count = imported
                        batch.save(update_fields=["record_count", "notes"])

                    messages.success(
                        request,
                        f"Imported {imported} reference(s) into '{batch.label}'.",
                    )
                    if duplicates:
                        messages.info(
                            request,
                            f"Skipped {duplicates} reference(s) already in the library.",
                        )
                    if skipped:
                        messages.info(
                            request,
                            f"Skipped {skipped} record(s) with no title.",
                        )
                    return redirect("synopsis:reference_library")
    else:
        form = LibraryReferenceBatchUploadForm()

    return render(
        request,
        "synopsis/library_reference_batch_upload.html",
        {"form": form},
    )


@login_required
def library_reference_detail(request, reference_id):
    if not _user_can_manage_library(request.user):
        raise PermissionDenied

    library_reference = get_object_or_404(LibraryReference, pk=reference_id)
    project_options = (
        Project.objects.all()
        if request.user.is_staff
        else Project.objects.filter(userrole__user=request.user).distinct()
    )
    link_message = None
    form = None

    if request.method == "POST":
        action = request.POST.get("action")
        if action == "link-to-project":
            target_project_id = request.POST.get("target_project")
            target_project = get_object_or_404(Project, pk=target_project_id)
            if not _user_can_edit_project(request.user, target_project):
                raise PermissionDenied
            folder = [f for f in request.POST.getlist("reference_folder") if f]
            linked, reused, _ = _link_library_references_to_project(
                request.user, target_project, [library_reference.id], folder
            )
            if linked:
                message = f"Linked reference to {target_project.title}."
                if folder:
                    message += " Shared CE subject categories were updated before linking."
                messages.success(request, message)
                return redirect(
                    "synopsis:library_reference_detail",
                    reference_id=library_reference.id,
                )
            if reused:
                link_message = "This reference already exists in that project."
            form = LibraryReferenceUpdateForm(instance=library_reference)
        else:
            previous_folders = normalize_reference_folder_values(
                library_reference.reference_folder
            )
            form = LibraryReferenceUpdateForm(request.POST, instance=library_reference)
            if form.is_valid():
                library_reference = form.save()
                shared_changed, _linked_count, old_folders, new_folders = (
                    False,
                    0,
                    previous_folders,
                    normalize_reference_folder_values(library_reference.reference_folder),
                )
                if previous_folders != new_folders:
                    shared_changed, _linked_count, old_folders, new_folders = (
                        _update_shared_library_reference_folders(
                            library_reference,
                            new_folders,
                            changed_by=request.user,
                            change_source="library_detail",
                            previous_folders=previous_folders,
                        )
                    )
                if shared_changed:
                    message = "Library reference updated. Shared CE subject categories were updated."
                    message += (
                        " Linked synopsis copies now read those shared categories automatically."
                    )
                    messages.success(request, message)
                else:
                    messages.success(request, "Library reference updated.")
                return redirect(
                    "synopsis:library_reference_detail",
                    reference_id=library_reference.id,
                )
            else:
                messages.error(request, "Unable to update the library reference.")
    else:
        form = LibraryReferenceUpdateForm(instance=library_reference)

    if link_message:
        messages.info(request, link_message)

    return render(
        request,
        "synopsis/library_reference_detail.html",
        {
            "reference": library_reference,
            "form": form,
            "project_options": project_options,
            "folder_choices": Reference.FOLDER_CHOICES,
            "folder_history": library_reference.folder_history.select_related(
                "changed_by", "source_project", "source_reference"
            )[:10],
        },
    )


def _reference_summary_citation(reference):
    parts = []
    canonical = reference.canonical if hasattr(reference, "canonical") else reference
    authors = (canonical.authors or "").strip()
    if authors:
        parts.append(authors)
    year = canonical.publication_year
    if year:
        parts.append(f"({year})")
    title = (canonical.title or "").strip()
    if title:
        parts.append(title)
    citation = " ".join(parts).strip()
    if len(citation) > 500:
        citation = citation[:497].rstrip() + "..."
    return citation


def _reference_sort_key(reference):
    canonical = reference.canonical if hasattr(reference, "canonical") else reference
    year = canonical.publication_year if canonical.publication_year is not None else 9999
    title = (canonical.title or "").strip().lower()
    authors = (canonical.authors or "").strip().lower()
    return year, title, authors


def _reference_export_citation(reference):
    canonical = reference.canonical if hasattr(reference, "canonical") else reference

    def _clean(value):
        return (value or "").strip()

    def _doi_url(raw):
        value = _clean(raw)
        if not value:
            return ""
        lowered = value.lower()
        for prefix in ("https://doi.org/", "http://doi.org/", "doi.org/"):
            if lowered.startswith(prefix):
                value = value[len(prefix):]
                break
        if value.lower().startswith("doi:"):
            value = value[4:]
        value = value.strip()
        return f"https://doi.org/{value}" if value else ""

    parts = []
    authors = _clean(canonical.authors)
    year = canonical.publication_year
    if authors and year:
        parts.append(f"{authors} ({year})")
    elif authors:
        parts.append(authors)
    elif year:
        parts.append(f"({year})")

    title = _clean(canonical.title)
    if title:
        parts.append(f"{title.rstrip('.')}.")

    journal = _clean(canonical.journal)
    volume = _clean(canonical.volume)
    issue = _clean(canonical.issue)
    pages = _clean(canonical.pages)
    if journal or volume or issue or pages:
        source_bits = []
        if journal:
            source_bits.append(journal)
        vol_issue = ""
        if volume and issue:
            vol_issue = f"{volume}({issue})"
        elif volume:
            vol_issue = volume
        elif issue:
            vol_issue = f"({issue})"
        if vol_issue:
            source_bits.append(vol_issue)
        if pages:
            source_bits.append(pages)
        if source_bits:
            parts.append(", ".join(source_bits).rstrip(".") + ".")

    doi_url = _doi_url(canonical.doi)
    source_url = _clean(canonical.url)
    if doi_url:
        parts.append(doi_url)
    elif source_url:
        parts.append(source_url)

    return " ".join([part for part in parts if part]).strip()


def _intervention_reference_numbering(assignments):
    ranked_assignments = []
    for assignment in assignments:
        reference = assignment.reference_summary.reference
        ranked_assignments.append(
            (*_reference_sort_key(reference), assignment.position, assignment.id, assignment)
        )
    ranked_assignments.sort()
    ordered_assignments = [entry[-1] for entry in ranked_assignments]
    summary_numbers = {}
    grouped_numbers = defaultdict(list)
    ordered_references = []
    seen_reference_ids = set()
    for idx, assignment in enumerate(ordered_assignments, start=1):
        summary_numbers[assignment.reference_summary_id] = idx
        reference = assignment.reference_summary.reference
        grouped_numbers[reference.id].append(idx)
        if reference.id not in seen_reference_ids:
            ordered_references.append((grouped_numbers[reference.id], reference))
            seen_reference_ids.add(reference.id)
    return ordered_assignments, summary_numbers, ordered_references


def _key_message_supporting_numbers(key_message, summary_numbers):
    numbers = []
    for summary in key_message.supporting_summaries.all():
        number = summary_numbers.get(summary.id)
        if number is not None:
            numbers.append(number)
    return sorted(set(numbers))


def _format_reference_number_list(numbers):
    return "; ".join(str(number) for number in numbers)


def _format_reference_number_ranges(numbers):
    unique_numbers = sorted(set(numbers))
    if not unique_numbers:
        return ""

    ranges = []
    start = unique_numbers[0]
    end = unique_numbers[0]

    for number in unique_numbers[1:]:
        if number == end + 1:
            end = number
            continue
        ranges.append(str(start) if start == end else f"{start}-{end}")
        start = end = number

    ranges.append(str(start) if start == end else f"{start}-{end}")
    return "; ".join(ranges)


def _ensure_reference_summaries(project, references):
    ref_ids = [ref.id for ref in references]
    existing_ref_ids = set(
        ReferenceSummary.objects.filter(reference_id__in=ref_ids).values_list(
            "reference_id", flat=True
        )
    )
    with transaction.atomic():
        for ref in references:
            if ref.id not in existing_ref_ids:
                ReferenceSummary.objects.create(
                    project=project,
                    reference=ref,
                    citation=_reference_summary_citation(ref),
                )
                _sync_reference_summary_identifiers_for_reference(ref, save=True)
                existing_ref_ids.add(ref.id)
    return existing_ref_ids


_REFERENCE_ID_STOP_WORDS = {
    "a",
    "an",
    "and",
    "book",
    "for",
    "in",
    "of",
    "on",
    "project",
    "synopsis",
    "the",
    "to",
}


def _project_reference_prefix(project):
    words = re.findall(r"[A-Za-z0-9]+", project.title or "")
    initials = [
        word[0].upper()
        for word in words
        if word and word.lower() not in _REFERENCE_ID_STOP_WORDS
    ]
    return "".join(initials[:6]) or "REF"


def _prefix_from_identifier(identifier):
    cleaned = _clean_identifier(identifier)
    match = re.match(r"^(.*?)(\d+)$", cleaned)
    if not match:
        return ""
    return match.group(1)


def _sequence_from_identifier(identifier, prefix):
    cleaned = _clean_identifier(identifier)
    if prefix:
        if not cleaned.startswith(prefix):
            return None
        suffix = cleaned[len(prefix) :]
    else:
        suffix = cleaned
    if not suffix.isdigit():
        return None
    return int(suffix)


def _reference_identifier_candidates(project):
    summaries = ReferenceSummary.objects.filter(project=project).only(
        "reference_identifier",
        "summary_identifier"
    )
    for summary in summaries:
        candidate = _clean_identifier(summary.reference_identifier)
        if candidate:
            yield candidate
            continue
        reference_identifier, _suffix = _split_summary_identifier(
            summary.summary_identifier
        )
        if reference_identifier:
            yield reference_identifier


def _stored_project_reference_prefix(project):
    for identifier in _reference_identifier_candidates(project):
        prefix = _prefix_from_identifier(identifier)
        if prefix:
            return prefix
    return ""


def _generated_reference_identifier(reference):
    project = reference.project
    prefix = _stored_project_reference_prefix(project) or _project_reference_prefix(
        project
    )
    count_floor = 1000 + Reference.objects.filter(
        project=project,
        id__lt=reference.id,
    ).count()
    max_sequence = 999
    for identifier in _reference_identifier_candidates(project):
        sequence = _sequence_from_identifier(identifier, prefix)
        if sequence is not None and sequence > max_sequence:
            max_sequence = sequence
    return f"{prefix}{max(count_floor, max_sequence + 1)}"


def _clean_identifier(value):
    return (value or "").strip()


def _split_summary_identifier(summary_identifier):
    cleaned = _clean_identifier(summary_identifier)
    if "." not in cleaned:
        return "", ""
    reference_identifier, suffix = cleaned.rsplit(".", 1)
    if not reference_identifier or not suffix:
        return "", ""
    return reference_identifier, suffix


def _stored_reference_identifier_from_summaries(summaries):
    for summary in summaries:
        reference_identifier = _clean_identifier(summary.reference_identifier)
        if reference_identifier:
            return reference_identifier
    for summary in summaries:
        reference_identifier, _suffix = _split_summary_identifier(
            summary.summary_identifier
        )
        if reference_identifier:
            return reference_identifier
    return ""


def _reference_identifier_for_reference(reference, summaries=None):
    if summaries is None:
        summaries = list(reference.summaries.order_by("created_at", "id"))
    reference_identifier = _stored_reference_identifier_from_summaries(summaries)
    if reference_identifier:
        return reference_identifier
    return _generated_reference_identifier(reference)


def _reference_identifier_for_summary(summary):
    reference_identifier = _clean_identifier(summary.reference_identifier)
    if reference_identifier:
        return reference_identifier
    reference_identifier, _suffix = _split_summary_identifier(
        summary.summary_identifier
    )
    if reference_identifier:
        return reference_identifier
    if summary.reference_id:
        return _reference_identifier_for_reference(summary.reference)
    return ""


def _alphabetical_suffix(index):
    index = max(index, 1)
    letters = []
    while index > 0:
        index -= 1
        index, remainder = divmod(index, 26)
        letters.append(chr(ord("a") + remainder))
    return "".join(reversed(letters))


def _generated_summary_identifier(reference_identifier, index):
    return f"{reference_identifier}.{_alphabetical_suffix(index)}"


def _sync_reference_summary_identifiers_for_reference(reference, *, save=False):
    summaries = list(reference.summaries.order_by("created_at", "id"))
    reference_identifier = _reference_identifier_for_reference(
        reference, summaries=summaries
    )
    used_suffixes_by_reference = defaultdict(set)
    next_index_by_reference = defaultdict(lambda: 1)

    for summary in summaries:
        summary_reference_identifier, suffix = _split_summary_identifier(
            summary.summary_identifier
        )
        if summary_reference_identifier and suffix:
            used_suffixes_by_reference[summary_reference_identifier].add(suffix)

    def _next_unused_summary_identifier(target_reference_identifier):
        next_index = next_index_by_reference[target_reference_identifier]
        while True:
            suffix = _alphabetical_suffix(next_index)
            next_index += 1
            if suffix in used_suffixes_by_reference[target_reference_identifier]:
                continue
            used_suffixes_by_reference[target_reference_identifier].add(suffix)
            next_index_by_reference[target_reference_identifier] = next_index
            return f"{target_reference_identifier}.{suffix}"

    for summary in summaries:
        changed_fields = []
        summary_reference_identifier = _reference_identifier_for_summary(summary)
        if not summary_reference_identifier:
            summary_reference_identifier = reference_identifier
        if _clean_identifier(summary.reference_identifier) != summary_reference_identifier:
            summary.reference_identifier = summary_reference_identifier
            changed_fields.append("reference_identifier")
        if not _clean_identifier(summary.summary_identifier):
            summary.summary_identifier = _next_unused_summary_identifier(
                summary_reference_identifier
            )
            changed_fields.append("summary_identifier")
        if save and changed_fields:
            summary.save(update_fields=changed_fields + ["updated_at"])
    return summaries


def _reference_summary_display_label(summary, index=None):
    label = summary.explicit_label
    if label:
        return label
    if index is not None:
        reference_identifier = _reference_identifier_for_summary(summary)
        return _generated_summary_identifier(reference_identifier, index)
    if summary.reference_id:
        reference_identifier = _reference_identifier_for_summary(summary)
        return _generated_summary_identifier(reference_identifier, 1)
    return summary.display_label


def _reference_summary_workspace_heading(reference):
    canonical = reference.canonical
    author_bits = [canonical.authors or "Unknown authors"]
    if canonical.publication_year:
        author_bits.append(str(canonical.publication_year))
    return " · ".join(author_bits)


def _reference_summary_workspace_context(reference):
    return reference.canonical.title or "Untitled reference"


def _reference_summary_workspace_label(summary, index=None):
    label = _reference_summary_display_label(summary, index)
    summary_identifier = (summary.summary_identifier or "").strip()
    if not summary_identifier and summary.reference_id:
        reference_identifier = _reference_identifier_for_summary(summary)
        if reference_identifier:
            summary_identifier = _generated_summary_identifier(
                reference_identifier, index or 1
            )
    if summary_identifier and label and label != summary_identifier:
        return f"{summary_identifier} — {label}"
    if summary_identifier:
        return summary_identifier
    return _reference_summary_display_label(summary, index)


def _reference_summary_tabs(reference, *, active_summary_id=None):
    summaries = _sync_reference_summary_identifiers_for_reference(reference, save=False)
    tab_count = len(summaries)
    tabs = []
    for index, item in enumerate(summaries, start=1):
        tabs.append(
            {
                "summary": item,
                "label": _reference_summary_display_label(item, index),
                "index": index,
                "is_active": item.id == active_summary_id,
                "tab_count": tab_count,
            }
        )
    return tabs


def _reference_summary_workspace_groups(reference_summaries):
    grouped = defaultdict(list)
    for summary in reference_summaries:
        grouped[summary.reference_id].append(summary)

    groups = []
    summary_meta = {}
    for summaries in grouped.values():
        reference = summaries[0].reference
        canonical = reference.canonical
        reference_heading = _reference_summary_workspace_heading(reference)
        reference_context = _reference_summary_workspace_context(reference)
        paper_title = canonical.title or "Untitled reference"
        meta_parts = []
        if canonical.authors:
            meta_parts.append(canonical.authors)
        if canonical.publication_year:
            meta_parts.append(str(canonical.publication_year))
        paper_meta = " · ".join(meta_parts)

        group_payload = {
            "reference_id": reference.id,
            "reference_heading": reference_heading,
            "reference_context": reference_context,
            "paper_title": paper_title,
            "paper_meta": paper_meta,
            "summaries": [],
        }
        for index, summary in enumerate(summaries, start=1):
            summary_label = _reference_summary_display_label(summary, index)
            summary_display = _reference_summary_workspace_label(summary, index)
            search_text = " ".join(
                bit
                for bit in (
                    reference_heading,
                    reference_context,
                    paper_title,
                    paper_meta,
                    summary_label,
                    summary_display,
                    summary.summary_identifier,
                    summary.action_description,
                )
                if bit
            )
            summary_meta[summary.id] = {
                "reference_heading": reference_heading,
                "reference_context": reference_context,
                "paper_title": paper_title,
                "paper_meta": paper_meta,
                "summary_label": summary_label,
                "summary_display": summary_display,
                "search_text": search_text,
            }
            group_payload["summaries"].append(
                {
                    "id": summary.id,
                    "summary_label": summary_label,
                    "summary_display": summary_display,
                    "search_text": search_text,
                }
            )
        groups.append(group_payload)
    return groups, summary_meta


def _clone_reference_summary(source_summary, user=None):
    summary_author = (
        source_summary.summary_author
        or (user.get_full_name() or user.username if user and user.is_authenticated else "")
    )
    new_summary = ReferenceSummary.objects.create(
        project=source_summary.project,
        reference=source_summary.reference,
        assigned_to=source_summary.assigned_to,
        citation=source_summary.citation or _reference_summary_citation(source_summary.reference),
        summary_author=summary_author,
        source_url=source_summary.source_url,
    )
    synced = _sync_reference_summary_identifiers_for_reference(
        source_summary.reference, save=True
    )
    return next((item for item in synced if item.id == new_summary.id), new_summary)


def _duplicate_reference_summary(source_summary, user=None):
    summary_author = (
        source_summary.summary_author
        or (user.get_full_name() or user.username if user and user.is_authenticated else "")
    )
    duplicate_status = (
        ReferenceSummary.STATUS_TODO
        if source_summary.status == ReferenceSummary.STATUS_TODO
        else ReferenceSummary.STATUS_DRAFT
    )
    new_summary = ReferenceSummary.objects.create(
        project=source_summary.project,
        reference=source_summary.reference,
        assigned_to=source_summary.assigned_to,
        status=duplicate_status,
        needs_help=False,
        reference_label=source_summary.reference_label,
        action_description=source_summary.action_description,
        study_design=source_summary.study_design,
        study_type=source_summary.study_type,
        sites_replications=source_summary.sites_replications,
        year_range=source_summary.year_range,
        habitat_and_sites=source_summary.habitat_and_sites,
        region=source_summary.region,
        country=source_summary.country,
        summary_of_results=source_summary.summary_of_results,
        action_methods=source_summary.action_methods,
        experimental_design=source_summary.experimental_design,
        site_context_details=source_summary.site_context_details,
        sampling_methods_details=source_summary.sampling_methods_details,
        cost_summary=source_summary.cost_summary,
        outcome_rows=copy.deepcopy(source_summary.outcome_rows or []),
        benefits_score=source_summary.benefits_score,
        harms_score=source_summary.harms_score,
        reliability_score=source_summary.reliability_score,
        relevance_score=source_summary.relevance_score,
        summary_text=source_summary.summary_text,
        key_findings=source_summary.key_findings,
        synopsis_draft=source_summary.synopsis_draft,
        summary_author=summary_author,
        broad_category=source_summary.broad_category,
        keywords=copy.deepcopy(source_summary.keywords or []),
        source_url=source_summary.source_url,
        crop_type=source_summary.crop_type,
        action_tags=copy.deepcopy(source_summary.action_tags or []),
        threat_tags=copy.deepcopy(source_summary.threat_tags or []),
        taxon_tags=copy.deepcopy(source_summary.taxon_tags or []),
        habitat_tags=copy.deepcopy(source_summary.habitat_tags or []),
        location_tags=copy.deepcopy(source_summary.location_tags or []),
        research_design=source_summary.research_design,
        citation=source_summary.citation or _reference_summary_citation(source_summary.reference),
    )
    synced = _sync_reference_summary_identifiers_for_reference(
        source_summary.reference, save=True
    )
    return next((item for item in synced if item.id == new_summary.id), new_summary)


def _next_chapter_position(project):
    max_pos = (
        SynopsisChapter.objects.filter(project=project).aggregate(Max("position"))[
            "position__max"
        ]
        or 0
    )
    return max_pos + 1


def _resequence_chapter_positions(project):
    for idx, chapter in enumerate(
        SynopsisChapter.objects.filter(project=project).order_by("position", "id"),
        start=1,
    ):
        if chapter.position != idx:
            chapter.position = idx
            chapter.save(update_fields=["position"])


def _next_subheading_position(chapter):
    max_pos = chapter.subheadings.aggregate(Max("position"))["position__max"] or 0
    return max_pos + 1


def _resequence_subheading_positions(chapter):
    for idx, subheading in enumerate(
        chapter.subheadings.order_by("position", "id"), start=1
    ):
        if subheading.position != idx:
            subheading.position = idx
            subheading.save(update_fields=["position"])


def _next_intervention_position(subheading):
    max_pos = subheading.interventions.aggregate(Max("position"))["position__max"] or 0
    return max_pos + 1


def _resequence_intervention_positions(subheading):
    for idx, intervention in enumerate(
        subheading.interventions.order_by("position", "id"), start=1
    ):
        if intervention.position != idx:
            intervention.position = idx
            intervention.save(update_fields=["position"])


def _next_assignment_position(intervention):
    max_pos = intervention.assignments.aggregate(Max("position"))["position__max"] or 0
    return max_pos + 1


def _resequence_assignment_positions(intervention):
    for idx, assignment in enumerate(
        intervention.assignments.order_by("position", "id"), start=1
    ):
        if assignment.position != idx:
            assignment.position = idx
            assignment.save(update_fields=["position"])


def _next_key_message_position(intervention):
    max_pos = intervention.key_messages.aggregate(Max("position"))["position__max"] or 0
    return max_pos + 1


def _resequence_key_message_positions(intervention):
    for idx, key_message in enumerate(
        intervention.key_messages.order_by("position", "id"), start=1
    ):
        if key_message.position != idx:
            key_message.position = idx
            key_message.save(update_fields=["position"])


def _next_action_summary_order(reference_summary):
    max_pos = (
        reference_summary.action_summaries.aggregate(Max("order"))["order__max"]
        or 0
    )
    return max_pos + 1


def _resequence_action_summaries(reference_summary):
    for idx, action_summary in enumerate(
        reference_summary.action_summaries.order_by("order", "id"),
        start=1,
    ):
        if action_summary.order != idx:
            action_summary.order = idx
            action_summary.save(update_fields=["order"])


def _remove_reference_from_synopsis(reference):
    summaries = list(reference.summaries.all())
    if not summaries:
        return 0

    summary_ids = [summary.id for summary in summaries]
    assignments = list(
        SynopsisAssignment.objects.filter(reference_summary_id__in=summary_ids)
        .select_related("intervention")
    )
    if not assignments:
        return 0

    touched_interventions = {}
    assignment_ids = []
    for assignment in assignments:
        touched_interventions[assignment.intervention_id] = assignment.intervention
        assignment_ids.append(assignment.id)

    supporting_summary_through = (
        SynopsisInterventionKeyMessage.supporting_summaries.through
    )
    with transaction.atomic():
        supporting_summary_through.objects.filter(
            synopsisinterventionkeymessage__intervention_id__in=touched_interventions.keys(),
            referencesummary_id__in=summary_ids,
        ).delete()
        SynopsisAssignment.objects.filter(pk__in=assignment_ids).delete()

    for intervention in touched_interventions.values():
        _resequence_assignment_positions(intervention)

    return len(assignments)


def _remove_summary_from_synopsis(summary):
    assignments = list(
        SynopsisAssignment.objects.filter(reference_summary=summary).select_related(
            "intervention"
        )
    )
    if not assignments:
        return 0

    touched_interventions = {}
    assignment_ids = []
    for assignment in assignments:
        touched_interventions[assignment.intervention_id] = assignment.intervention
        assignment_ids.append(assignment.id)

    supporting_summary_through = (
        SynopsisInterventionKeyMessage.supporting_summaries.through
    )
    with transaction.atomic():
        supporting_summary_through.objects.filter(
            synopsisinterventionkeymessage__intervention_id__in=touched_interventions.keys(),
            referencesummary_id=summary.id,
        ).delete()
        SynopsisAssignment.objects.filter(pk__in=assignment_ids).delete()

    for intervention in touched_interventions.values():
        _resequence_assignment_positions(intervention)

    return len(assignments)


def _structured_summary_paragraph(
    summary: ReferenceSummary, reference_identifier_override: str | None = None
) -> str:
    """Generate a concise paragraph from structured summary fields."""

    def _clean(text):
        return (text or "").strip()

    study_type = _clean(summary.study_type)
    study_design = _clean(summary.study_design)
    year_range = _clean(summary.year_range)
    habitat = _clean(summary.habitat_and_sites)
    location = ", ".join([part for part in [_clean(summary.region), _clean(summary.country)] if part])
    ref_id = _clean(reference_identifier_override)
    if not ref_id:
        ref_id = _reference_identifier_for_summary(summary)
    sites = _clean(summary.sites_replications)
    intro_parts = ["A"]
    intro_parts.append(study_design or study_type or "study")
    if year_range:
        intro_parts.append(f"in {year_range}")
    else:
        intro_parts.append("(year not stated)")
    if habitat:
        intro_parts.append(f"in {habitat}")
    if location:
        intro_parts.append(f"in {location}")
    if sites:
        intro_parts.append(f"({sites})")
    intro_line = " ".join(intro_parts).strip()
    if ref_id:
        intro_line = f"{intro_line} ({ref_id})"
    intro_line = f"{intro_line} found that"

    results = _clean(summary.summary_of_results) or _clean(summary.summary_text)

    methods_parts = []
    if summary.action_methods:
        methods_parts.append(_clean(summary.action_methods))
    if summary.experimental_design:
        methods_parts.append(_clean(summary.experimental_design))
    if summary.site_context_details:
        methods_parts.append(_clean(summary.site_context_details))
    if summary.sampling_methods_details:
        methods_parts.append(_clean(summary.sampling_methods_details))

    outcome_lines = []
    for row in summary.outcome_rows or []:
        outcome = _clean(row.get("outcome", ""))
        difference = _clean(row.get("difference", ""))
        treatment = _clean(row.get("treatment", ""))
        comparator = _clean(row.get("comparator", ""))
        t_val = _clean(row.get("treatment_value", ""))
        c_val = _clean(row.get("comparator_value", ""))
        unit = _clean(row.get("unit", ""))
        notes = _clean(row.get("notes", ""))
        p_val = _clean(row.get("p_value", ""))
        stats = _clean(row.get("stats", ""))

        parts = []
        if outcome:
            parts.append(f"{outcome}:")
        if difference and treatment and comparator:
            parts.append(f"{difference} in {treatment} compared to {comparator}")
        elif difference:
            parts.append(difference)
        if t_val or c_val:
            val_bits = []
            if t_val:
                val_bits.append(t_val)
            if c_val:
                val_bits.append(c_val)
            value_text = " vs ".join(val_bits)
            if unit:
                value_text = f"{value_text} {unit}".strip()
            parts.append(f"({value_text})")
        if stats:
            parts.append(f"Stats: {stats}")
        if p_val:
            parts.append(f"p={p_val}")
        if notes:
            parts.append(notes)
        sentence = " ".join([p for p in parts if p]).strip()
        if sentence:
            outcome_lines.append(sentence if sentence.endswith(".") else f"{sentence}.")

    methods_text = " ".join([part for part in methods_parts if part]).strip()
    scores = []
    if summary.benefits_score is not None:
        scores.append(f"Benefits: {summary.benefits_score}")
    if summary.harms_score is not None:
        scores.append(f"Harms: {summary.harms_score}")
    if summary.reliability_score is not None:
        scores.append(f"Reliability: {summary.reliability_score}")
    if summary.relevance_score is not None:
        scores.append(f"Relevance: {summary.relevance_score}")

    segments = [intro_line]
    if results:
        segments.append(results)
    if outcome_lines:
        segments.append(" ".join(outcome_lines))
    if methods_text:
        segments.append(methods_text)

    paragraph = " ".join([seg.strip() for seg in segments if seg.strip()]).strip()
    if paragraph and not paragraph.endswith("."):
        paragraph = f"{paragraph}."
    if scores:
        paragraph = f"{paragraph}\n\n" + " · ".join(scores)
    return paragraph


def _reference_summary_paragraph(
    summary: ReferenceSummary, reference_identifier_override: str | None = None
) -> str:
    draft = (summary.synopsis_draft or "").strip()
    if draft:
        override = (str(reference_identifier_override).strip() if reference_identifier_override else "")
        if override:
            known_identifiers = {
                identifier.strip()
                for identifier in [
                    summary.reference_identifier,
                    _generated_reference_identifier(summary.reference),
                ]
                if identifier and identifier.strip()
            }
            replaced = False
            for identifier in known_identifiers:
                token = f"({identifier})"
                if token in draft:
                    draft = draft.replace(token, f"({override})")
                    replaced = True
            if not replaced and f"({override})" not in draft:
                draft = f"({override}) {draft}"
        return draft
    return _structured_summary_paragraph(
        summary,
        reference_identifier_override=reference_identifier_override,
    )


def _reference_summary_has_meaningful_progress(summary: ReferenceSummary) -> bool:
    text_fields = [
        "action_description",
        "study_design",
        "study_type",
        "sites_replications",
        "year_range",
        "habitat_and_sites",
        "region",
        "country",
        "summary_of_results",
        "action_methods",
        "experimental_design",
        "site_context_details",
        "sampling_methods_details",
        "cost_summary",
        "summary_text",
        "key_findings",
        "synopsis_draft",
        "broad_category",
        "source_url",
        "crop_type",
        "research_design",
        "citation",
    ]
    list_fields = [
        "outcome_rows",
        "keywords",
        "action_tags",
        "threat_tags",
        "taxon_tags",
        "habitat_tags",
        "location_tags",
    ]
    score_fields = [
        "benefits_score",
        "harms_score",
        "reliability_score",
        "relevance_score",
    ]

    for field_name in text_fields:
        value = getattr(summary, field_name, "")
        if isinstance(value, str) and value.strip():
            return True

    for field_name in list_fields:
        value = getattr(summary, field_name, None)
        if value:
            return True

    for field_name in score_fields:
        if getattr(summary, field_name, None) is not None:
            return True

    return False


def _auto_promote_summary_from_todo(summary: ReferenceSummary, previous_status: str) -> bool:
    if previous_status != ReferenceSummary.STATUS_TODO:
        return False
    if summary.status != ReferenceSummary.STATUS_TODO:
        return False
    if not _reference_summary_has_meaningful_progress(summary):
        return False
    summary.status = ReferenceSummary.STATUS_DRAFT
    summary.save(update_fields=["status", "updated_at"])
    return True


@login_required
def reference_summary_board(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if not _user_can_edit_project(request.user, project):
        raise PermissionDenied

    included_references = list(
        project.references.filter(screening_status="included")
        .select_related("batch")
        .order_by("title")
    )

    _ensure_reference_summaries(project, included_references)

    if request.method == "POST":
        action = request.POST.get("action")
        selected_ids = request.POST.getlist("summary_ids") or []
        if action == "auto-assign":
            authors = list(project.author_users.order_by("id"))
            if not authors:
                messages.error(request, "No authors available to assign.")
                return redirect("synopsis:reference_summary_board", project_id=project.id)
            base_qs = ReferenceSummary.objects.filter(
                project=project,
                reference__screening_status="included",
                assigned_to__isnull=True,
            ).exclude(
                status=ReferenceSummary.STATUS_EXCLUDED,
            )
            if selected_ids:
                base_qs = base_qs.filter(id__in=selected_ids)
            unassigned = list(base_qs.order_by("id"))
            if not unassigned:
                messages.info(request, "No unassigned summaries to distribute.")
                return redirect("synopsis:reference_summary_board", project_id=project.id)
            random.shuffle(unassigned)
            for idx, item in enumerate(unassigned):
                item.assigned_to = authors[idx % len(authors)]
                item.save(update_fields=["assigned_to", "updated_at"])
            messages.success(
                request,
                f"Auto-assigned {len(unassigned)} summaries across {len(authors)} author(s).",
            )
            return redirect("synopsis:reference_summary_board", project_id=project.id)
        elif action == "bulk-assign-author":
            author_id = request.POST.get("bulk_assigned_to")
            if not author_id:
                messages.error(request, "Select an author to assign.")
                return redirect("synopsis:reference_summary_board", project_id=project.id)
            try:
                author = project.author_users.get(pk=author_id)
            except User.DoesNotExist:
                messages.error(request, "Selected author is not part of this project.")
                return redirect("synopsis:reference_summary_board", project_id=project.id)
            qs = ReferenceSummary.objects.filter(
                project=project,
                reference__screening_status="included",
            ).exclude(status=ReferenceSummary.STATUS_EXCLUDED)
            if selected_ids:
                qs = qs.filter(id__in=selected_ids)
            else:
                qs = qs.filter(assigned_to__isnull=True)
            updated = qs.update(assigned_to=author, updated_at=timezone.now())
            if updated:
                messages.success(
                    request,
                    f"Assigned {updated} unassigned summaries to {author.get_full_name() or author.username}.",
                )
            else:
                messages.info(request, "No unassigned summaries to distribute.")
            return redirect("synopsis:reference_summary_board", project_id=project.id)
        elif action == "bulk-unassign":
            qs = ReferenceSummary.objects.filter(
                project=project,
                reference__screening_status="included",
                assigned_to__isnull=False,
            ).exclude(status=ReferenceSummary.STATUS_EXCLUDED)
            if selected_ids:
                qs = qs.filter(id__in=selected_ids)
            updated = qs.update(assigned_to=None, updated_at=timezone.now())
            if updated:
                messages.success(request, f"Unassigned {updated} summaries.")
            else:
                messages.info(request, "No summaries were unassigned.")
            return redirect("synopsis:reference_summary_board", project_id=project.id)

        summary = get_object_or_404(
            ReferenceSummary, pk=request.POST.get("summary_id"), project=project
        )
        if action == "assign":
            form = ReferenceSummaryAssignmentForm(request.POST, project=project)
            if form.is_valid():
                summary.assigned_to = form.cleaned_data["assigned_to"]
                summary.needs_help = form.cleaned_data["needs_help"]
                summary.save(update_fields=["assigned_to", "needs_help", "updated_at"])
                messages.success(request, "Assignment updated.")
            else:
                messages.error(request, "Could not update assignment.")
        elif action == "status":
            status = request.POST.get("status")
            valid_statuses = {choice[0] for choice in ReferenceSummary.STATUS_CHOICES}
            if status in valid_statuses:
                exclusion_reason = (request.POST.get("exclusion_reason") or "").strip()
                if status == ReferenceSummary.STATUS_EXCLUDED and not exclusion_reason:
                    messages.error(
                        request,
                        "Provide a reason before excluding this summary after full-text review.",
                    )
                    return redirect(
                        "synopsis:reference_summary_board", project_id=project.id
                    )
                summary.status = status
                if status == ReferenceSummary.STATUS_EXCLUDED:
                    summary.exclusion_reason = exclusion_reason
                    removed_assignments = _remove_summary_from_synopsis(summary)
                    summary.save(
                        update_fields=["status", "exclusion_reason", "updated_at"]
                    )
                    messages.success(
                        request,
                        "Summary excluded after full-text review."
                        + (
                            f" Removed it from {removed_assignments} intervention assignment(s)."
                            if removed_assignments
                            else ""
                        ),
                    )
                else:
                    summary.exclusion_reason = exclusion_reason if summary.status == ReferenceSummary.STATUS_EXCLUDED else summary.exclusion_reason
                    summary.save(
                        update_fields=["status", "exclusion_reason", "updated_at"]
                    )
                    messages.success(request, "Summary status updated.")
            else:
                messages.error(request, "Invalid summary status selected.")
        return redirect("synopsis:reference_summary_board", project_id=project.id)

    summaries = list(
        ReferenceSummary.objects.filter(
            project=project,
            reference__screening_status="included",
        )
        .select_related("reference", "reference__library_reference", "assigned_to")
        .order_by("reference__title", "created_at", "id")
    )
    summary_groups = defaultdict(list)
    for item in summaries:
        summary_groups[item.reference_id].append(item)
    for group in summary_groups.values():
        group_size = len(group)
        for index, item in enumerate(group, start=1):
            item.variant_label = _reference_summary_display_label(item, index)
            item.variant_count = group_size

    active_summaries = [
        item for item in summaries if item.status != ReferenceSummary.STATUS_EXCLUDED
    ]
    excluded_after_full_text = [
        item for item in summaries if item.status == ReferenceSummary.STATUS_EXCLUDED
    ]

    assigned_counts = Counter()
    needs_help_by_author = Counter()
    summarised_by_author = Counter()
    excluded_by_author = Counter()
    unassigned_count = 0
    needs_help_count = 0
    for item in active_summaries:
        if item.assigned_to_id is None:
            unassigned_count += 1
        else:
            assigned_counts[item.assigned_to_id] += 1
            if item.needs_help:
                needs_help_by_author[item.assigned_to_id] += 1
            if item.status == ReferenceSummary.STATUS_DONE:
                summarised_by_author[item.assigned_to_id] += 1
        if item.needs_help:
            needs_help_count += 1
    for item in excluded_after_full_text:
        if item.assigned_to_id is not None:
            excluded_by_author[item.assigned_to_id] += 1

    author_options = list(project.author_users.order_by("first_name", "last_name"))
    workload = []
    for author in author_options:
        assigned = assigned_counts.get(author.id, 0)
        summarised = summarised_by_author.get(author.id, 0)
        workload.append(
            {
                "author": author,
                "assigned": assigned,
                "summarised": summarised,
                "summarised_percent": int((summarised / assigned) * 100)
                if assigned
                else 0,
                "needs_help": needs_help_by_author.get(author.id, 0),
                "excluded_after_full_text": excluded_by_author.get(author.id, 0),
            }
        )

    status_map = {
        code: {"label": label, "items": []}
        for code, label in ReferenceSummary.STATUS_CHOICES
    }
    for summary in summaries:
        status_map.setdefault(summary.status, {"label": summary.status, "items": []})
        status_map[summary.status]["items"].append(summary)

    columns = [
        {
            "code": code,
            "label": label,
            "items": status_map.get(code, {}).get("items", []),
        }
        for code, label in ReferenceSummary.STATUS_CHOICES
    ]

    total_summaries = len(active_summaries)
    total_included = len(included_references)
    completed = len(status_map.get(ReferenceSummary.STATUS_DONE, {}).get("items", []))
    progress = int((completed / total_summaries) * 100) if total_summaries else 0

    return render(
        request,
        "synopsis/reference_summary_board.html",
        {
            "project": project,
            "columns": columns,
            "total_included": total_summaries,
            "completed": completed,
            "progress": progress,
            "excluded_after_full_text_count": len(excluded_after_full_text),
            "author_options": author_options,
            "status_choices": ReferenceSummary.STATUS_CHOICES,
            "reference_count": total_included,
            "summary_count": total_summaries,
            "workload": workload,
            "unassigned_count": unassigned_count,
            "summaries": sorted(
                summaries,
                key=lambda item: (
                    item.status,
                    (
                        item.assigned_to.first_name.lower()
                        if item.assigned_to and item.assigned_to.first_name
                        else ""
                    ),
                    item.reference.canonical.title.lower(),
                    item.id,
                ),
            ),
            "needs_help_count": needs_help_count,
        },
    )


@login_required
def reference_summary_detail(request, project_id, summary_id):
    project = get_object_or_404(Project, pk=project_id)
    if not _user_can_edit_project(request.user, project):
        raise PermissionDenied

    summary = get_object_or_404(
        ReferenceSummary.objects.select_related(
            "reference", "reference__library_reference", "assigned_to"
        ),
        pk=summary_id,
        project=project,
    )
    synced_summaries = _sync_reference_summary_identifiers_for_reference(
        summary.reference, save=True
    )
    summary = next((item for item in synced_summaries if item.id == summary.id), summary)
    generated_summary = _structured_summary_paragraph(summary)

    active_action = request.POST.get("action") if request.method == "POST" else None
    summary_form = ReferenceSummaryUpdateForm(
        request.POST if active_action == "save-summary" else None,
        instance=summary,
        project=project,
    )
    assignment_initial = {
        "assigned_to": summary.assigned_to_id,
        "needs_help": summary.needs_help,
    }
    assignment_form = ReferenceSummaryAssignmentForm(
        request.POST if request.POST.get("action") == "assign" else None,
        project=project,
        initial=assignment_initial,
    )
    classification_initial = {
        "screening_status": summary.reference.screening_status,
        "reference_folder": _reference_category_values(summary.reference),
        "screening_notes": summary.reference.screening_notes,
    }
    classification_data = None
    if active_action == "update-classification":
        classification_data = request.POST.copy()
        classification_command = (classification_data.get("classification_command") or "").strip()
        if classification_command == "exclude":
            classification_data["screening_status"] = "excluded"
        elif classification_command == "include":
            classification_data["screening_status"] = "included"
    classification_form = ReferenceClassificationForm(
        classification_data,
        initial=classification_initial,
    )
    draft_form = ReferenceSummaryDraftForm(
        request.POST if active_action == "save-synopsis-draft" else None,
        instance=summary,
        generated_summary=generated_summary,
    )
    comment_form = ReferenceSummaryCommentForm()
    document_form = ReferenceDocumentForm()
    action_summary_form = ReferenceActionSummaryForm()

    if request.method == "POST":
        action = active_action
        if action == "create-summary-tab":
            new_summary = _clone_reference_summary(summary, request.user)
            messages.success(
                request,
                "New summary tab created for this reference. Use it for a distinct intervention or study summary.",
            )
            return redirect(
                "synopsis:reference_summary_detail",
                project_id=project.id,
                summary_id=new_summary.id,
            )
        if action == "duplicate-summary-tab":
            new_summary = _duplicate_reference_summary(summary, request.user)
            messages.success(
                request,
                "Summary tab duplicated. Review the copied text and adjust it for the new intervention or study summary.",
            )
            return redirect(
                "synopsis:reference_summary_detail",
                project_id=project.id,
                summary_id=new_summary.id,
            )
        if action == "delete-summary-tab":
            remaining_summaries = list(
                summary.reference.summaries.exclude(pk=summary.pk).order_by("created_at", "id")
            )
            if not remaining_summaries:
                messages.error(
                    request,
                    "You cannot delete the only summary tab for this reference.",
                )
                return redirect(
                    "synopsis:reference_summary_detail",
                    project_id=project.id,
                    summary_id=summary.id,
                )
            next_summary = remaining_summaries[0]
            affected_intervention_ids = list(
                SynopsisAssignment.objects.filter(reference_summary=summary)
                .values_list("intervention_id", flat=True)
                .distinct()
            )
            with transaction.atomic():
                summary.delete()
                if affected_intervention_ids:
                    for intervention in SynopsisIntervention.objects.filter(
                        id__in=affected_intervention_ids
                    ):
                        _resequence_assignment_positions(intervention)
            _sync_reference_summary_identifiers_for_reference(next_summary.reference, save=True)
            messages.success(request, "Summary tab deleted.")
            return redirect(
                "synopsis:reference_summary_detail",
                project_id=project.id,
                summary_id=next_summary.id,
            )
        if action == "save-summary" and summary_form.is_valid():
            previous_status = summary.status
            previous_generated_summary = generated_summary.strip()
            previous_saved_draft = (summary.synopsis_draft or "").strip()
            summary = summary_form.save(commit=False)
            if not summary.summary_author:
                summary.summary_author = (
                    request.user.get_full_name() or request.user.username
                )
            summary.save()
            _sync_reference_summary_identifiers_for_reference(summary.reference, save=True)
            summary_form.save_m2m()
            refreshed_saved_draft = False
            if previous_saved_draft and previous_saved_draft == previous_generated_summary:
                updated_generated_summary = _structured_summary_paragraph(summary).strip()
                summary.synopsis_draft = updated_generated_summary
                summary.save(update_fields=["synopsis_draft", "updated_at"])
                refreshed_saved_draft = True
            auto_promoted = _auto_promote_summary_from_todo(summary, previous_status)
            if refreshed_saved_draft and auto_promoted:
                messages.success(
                    request,
                    "Summary updated. Status moved to In progress automatically. The saved paragraph draft was refreshed automatically.",
                )
            elif refreshed_saved_draft:
                messages.success(
                    request,
                    "Summary updated. The saved paragraph draft was refreshed automatically.",
                )
            elif auto_promoted:
                messages.success(
                    request,
                    "Summary updated. Status moved to In progress automatically.",
                )
            else:
                messages.success(request, "Summary updated.")
            return redirect(
                "synopsis:reference_summary_detail",
                project_id=project.id,
                summary_id=summary.id,
            )
        elif action == "save-summary":
            # Surface validation errors to help users understand why the save failed.
            error_list = []
            for err in summary_form.non_field_errors():
                error_list.append(err)
            for field, errs in summary_form.errors.items():
                for err in errs:
                    label = summary_form.fields.get(field).label if field in summary_form.fields else field
                    error_list.append(f"{label}: {err}")
            if not error_list:
                error_list.append("Unable to save summary. Please review your inputs.")
            messages.error(request, " ".join(error_list))
        if action == "save-synopsis-draft":
            previous_status = summary.status
            draft_command = request.POST.get("draft_command") or "save"
            if draft_command == "use-generated":
                summary.synopsis_draft = generated_summary
                summary.save(update_fields=["synopsis_draft", "updated_at"])
                auto_promoted = _auto_promote_summary_from_todo(summary, previous_status)
                if auto_promoted:
                    messages.success(
                        request,
                        "Generated paragraph copied into the editable draft. Status moved to In progress automatically.",
                    )
                else:
                    messages.success(
                        request,
                        "Generated paragraph copied into the editable draft.",
                    )
                return redirect(
                    "synopsis:reference_summary_detail",
                    project_id=project.id,
                    summary_id=summary.id,
                )
            if draft_command == "clear":
                summary.synopsis_draft = ""
                summary.save(update_fields=["synopsis_draft", "updated_at"])
                messages.success(
                    request,
                    "Saved paragraph draft cleared. Compilation will use the auto-generated paragraph again.",
                )
                return redirect(
                    "synopsis:reference_summary_detail",
                    project_id=project.id,
                    summary_id=summary.id,
                )
            if draft_form.is_valid():
                summary = draft_form.save(commit=False)
                summary.save(update_fields=["synopsis_draft", "updated_at"])
                auto_promoted = _auto_promote_summary_from_todo(summary, previous_status)
                if auto_promoted:
                    messages.success(
                        request,
                        "Summary paragraph draft saved. Status moved to In progress automatically.",
                    )
                else:
                    messages.success(request, "Summary paragraph draft saved.")
                return redirect(
                    "synopsis:reference_summary_detail",
                    project_id=project.id,
                    summary_id=summary.id,
                )
            messages.error(request, "Could not save the summary paragraph draft.")
        if action == "update-classification" and classification_form.is_valid():
            reference = summary.reference
            previous_status = reference.screening_status
            categories = normalize_reference_folder_values(
                classification_form.cleaned_data.get("reference_folder") or []
            )
            reference.screening_status = classification_form.cleaned_data[
                "screening_status"
            ]
            reference.screening_notes = (
                classification_form.cleaned_data.get("screening_notes") or ""
            )
            reference.screening_decision_at = timezone.now()
            reference.screened_by = request.user
            reference.save(
                update_fields=[
                    "screening_status",
                    "screening_notes",
                    "screening_decision_at",
                    "screened_by",
                    "updated_at",
                ]
            )
            (
                shared_folder_changed,
                _shared_linked_count,
                _local_category_changed,
                _saved_categories,
            ) = _update_reference_categories(
                reference,
                categories,
                changed_by=request.user,
                source_project=project,
                change_source="summary_reference_management",
            )

            if reference.screening_status == "excluded":
                removed_assignments = _remove_reference_from_synopsis(reference)
                message = "Reference excluded from this synopsis."
                if removed_assignments:
                    message += f" Removed it from {removed_assignments} intervention assignment(s)."
                if shared_folder_changed:
                    message += (
                        " Shared CE subject categories were updated for all linked synopsis copies."
                    )
                messages.success(request, message)
                return redirect(
                    f"{reverse('synopsis:reference_summary_detail', args=[project.id, summary.id])}?panel=management"
                )

            if previous_status == "excluded":
                message = "Reference re-included in this synopsis. You can now continue summarising it."
                if shared_folder_changed:
                    message += (
                        " Shared CE subject categories were updated for all linked synopsis copies."
                    )
                messages.success(request, message)
            else:
                message = "Reference classification updated."
                if shared_folder_changed:
                    message += (
                        " Shared CE subject categories were updated for all linked synopsis copies."
                    )
                messages.success(request, message)
            return redirect(
                f"{reverse('synopsis:reference_summary_detail', args=[project.id, summary.id])}?panel=management"
            )
        if action == "assign" and assignment_form.is_valid():
            summary.assigned_to = assignment_form.cleaned_data["assigned_to"]
            summary.needs_help = assignment_form.cleaned_data["needs_help"]
            summary.save(update_fields=["assigned_to", "needs_help", "updated_at"])
            messages.success(request, "Assignment updated.")
            return redirect(
                "synopsis:reference_summary_detail",
                project_id=project.id,
                summary_id=summary.id,
            )
        if action == "comment":
            comment_form = ReferenceSummaryCommentForm(request.POST, request.FILES)
            if comment_form.is_valid():
                parent = None
                parent_id = comment_form.cleaned_data.get("parent_id")
                if parent_id:
                    parent = ReferenceSummaryComment.objects.filter(
                        pk=parent_id, summary=summary
                    ).first()
                ReferenceSummaryComment.objects.create(
                    summary=summary,
                    author=request.user,
                    body=comment_form.cleaned_data["body"],
                    parent=parent,
                    attachment=comment_form.cleaned_data.get("attachment"),
                    notify_assignee=comment_form.cleaned_data.get("notify_assignee") or False,
                )
                messages.success(request, "Comment added.")
                return redirect(
                    "synopsis:reference_summary_detail",
                    project_id=project.id,
                    summary_id=summary.id,
                )
            else:
                messages.error(request, "Could not add comment.")
        if action == "update-status":
            previous_status = summary.status
            if request.POST.get("quick_done") == "1":
                summary.status = ReferenceSummary.STATUS_DONE
            elif request.POST.get("status") in dict(ReferenceSummary.STATUS_CHOICES):
                summary.status = request.POST.get("status")
            exclusion_reason = (request.POST.get("exclusion_reason") or "").strip()
            if summary.status == ReferenceSummary.STATUS_EXCLUDED and not exclusion_reason:
                messages.error(
                    request,
                    "Provide a reason before excluding this summary after full-text review.",
                )
                return redirect(
                    "synopsis:reference_summary_detail",
                    project_id=project.id,
                    summary_id=summary.id,
                )
            summary.needs_help = bool(request.POST.get("needs_help"))
            summary.exclusion_reason = (
                exclusion_reason
                if summary.status == ReferenceSummary.STATUS_EXCLUDED
                else summary.exclusion_reason
            )
            if (
                summary.status == ReferenceSummary.STATUS_EXCLUDED
                and previous_status != ReferenceSummary.STATUS_EXCLUDED
            ):
                removed_assignments = _remove_summary_from_synopsis(summary)
            else:
                removed_assignments = 0
            summary.save(
                update_fields=[
                    "status",
                    "needs_help",
                    "exclusion_reason",
                    "updated_at",
                ]
            )
            if summary.status == ReferenceSummary.STATUS_EXCLUDED:
                messages.success(
                    request,
                    "Summary excluded after full-text review."
                    + (
                        f" Removed it from {removed_assignments} intervention assignment(s)."
                        if removed_assignments
                        else ""
                    ),
                )
            else:
                messages.success(request, "Status updated.")
            return redirect(
                "synopsis:reference_summary_detail",
                project_id=project.id,
                summary_id=summary.id,
            )
        if action == "upload-document":
            document_form = ReferenceDocumentForm(request.POST, request.FILES)
            if document_form.is_valid():
                uploaded = document_form.cleaned_data["document"]
                summary.reference.reference_document = uploaded
                summary.reference.reference_document_uploaded_at = timezone.now()
                summary.reference.save(update_fields=["reference_document", "reference_document_uploaded_at"])
                messages.success(request, "PDF uploaded.")
                return redirect(
                    "synopsis:reference_summary_detail",
                    project_id=project.id,
                    summary_id=summary.id,
                )
            else:
                messages.error(request, "Upload failed. Ensure you selected a PDF file.")

        if action == "add-action-summary":
            action_summary_form = ReferenceActionSummaryForm(request.POST)
            if action_summary_form.is_valid():
                action_entry = action_summary_form.save(commit=False)
                action_entry.reference_summary = summary
                action_entry.order = _next_action_summary_order(summary)
                if request.user.is_authenticated:
                    action_entry.created_by = request.user
                action_entry.save()
                messages.success(request, "Action summary added.")
                return redirect(
                    "synopsis:reference_summary_detail",
                    project_id=project.id,
                    summary_id=summary.id,
                )
            else:
                messages.error(request, "Please fix the action summary details.")
        if action == "edit-action-summary":
            target = get_object_or_404(
                ReferenceActionSummary,
                pk=request.POST.get("action_summary_id"),
                reference_summary=summary,
            )
            action_summary_form = ReferenceActionSummaryForm(
                request.POST, instance=target
            )
            if action_summary_form.is_valid():
                action_summary_form.save()
                messages.success(request, "Action summary updated.")
                return redirect(
                    "synopsis:reference_summary_detail",
                    project_id=project.id,
                    summary_id=summary.id,
                )
            else:
                messages.error(request, "Could not update that action summary.")
        if action == "delete-action-summary":
            target = get_object_or_404(
                ReferenceActionSummary,
                pk=request.POST.get("action_summary_id"),
                reference_summary=summary,
            )
            target.delete()
            _resequence_action_summaries(summary)
            messages.success(request, "Action summary removed.")
            return redirect(
                "synopsis:reference_summary_detail",
                project_id=project.id,
                summary_id=summary.id,
            )

    comments = summary.comments.select_related("author")
    action_summaries = summary.action_summaries.order_by("order", "id")
    current_summary_paragraph = _reference_summary_paragraph(summary)
    summary_tabs = _reference_summary_tabs(
        summary.reference, active_summary_id=summary.id
    )
    comment_children = defaultdict(list)
    for c in comments:
        comment_children[c.parent_id].append(c)
    comment_tree = comment_children[None]
    for c in comment_tree:
        c.replies_cached = comment_children.get(c.id, [])

    return render(
        request,
        "synopsis/reference_summary_detail.html",
        {
            "project": project,
            "summary": summary,
            "reference": summary.reference,
            "generated_reference_id": _reference_identifier_for_summary(summary),
            "generated_summary_id": summary.summary_identifier
            or _generated_summary_identifier(
                _reference_identifier_for_summary(summary),
                1,
            ),
            "generated_reference_title": summary.reference.canonical.title
            or summary.reference.title,
            "summary_tabs": summary_tabs,
            "summary_form": summary_form,
            "draft_form": draft_form,
            "assignment_form": assignment_form,
            "classification_form": classification_form,
            "comment_form": comment_form,
            "comments": comments,
            "comment_tree": comment_tree,
            "comment_children": comment_children,
            "document_form": document_form,
            "action_summary_form": action_summary_form,
            "action_summaries": action_summaries,
            "generated_summary": generated_summary,
            "current_summary_paragraph": current_summary_paragraph,
            "status_choices": ReferenceSummary.STATUS_CHOICES,
            "all_summary_tabs_excluded": not summary.reference.summaries.exclude(
                status=ReferenceSummary.STATUS_EXCLUDED
            ).exists(),
            "open_management_panel": (
                request.GET.get("panel") == "management"
                or bool(classification_form.errors)
            ),
        },
    )


@login_required
@xframe_options_exempt
def reference_document_inline(request, project_id, reference_id):
    project = get_object_or_404(Project, pk=project_id)
    reference = get_object_or_404(
        Reference, pk=reference_id, project=project, screening_status="included"
    )
    if not reference.reference_document:
        raise Http404("No document available.")

    file_handle = reference.reference_document.open("rb")
    filename = reference.reference_document.name.rsplit("/", 1)[-1]
    response = FileResponse(file_handle)
    response["Content-Type"] = "application/pdf"
    response["Content-Disposition"] = f'inline; filename="{filename}"'
    response["Cache-Control"] = "no-store"
    response["Content-Security-Policy"] = "default-src 'none'; script-src 'none'; object-src 'none'; base-uri 'none';"
    return response


def _project_synopsis_workspace(
    request,
    project_id,
    *,
    workspace_mode="evidence",
    redirect_name="project_synopsis_structure",
    template_name="synopsis/project_synopsis_structure.html",
):
    project = get_object_or_404(Project, pk=project_id)
    if not _user_can_edit_project(request.user, project):
        raise PermissionDenied
    if workspace_mode not in {"evidence", "narrative"}:
        workspace_mode = "evidence"
    chapter_form = SynopsisChapterForm()
    if workspace_mode == "narrative":
        chapter_form.fields["chapter_type"].initial = SynopsisChapter.TYPE_TEXT
    else:
        chapter_form.fields["chapter_type"].initial = SynopsisChapter.TYPE_EVIDENCE
    subheading_form = SynopsisSubheadingForm()
    intervention_form = SynopsisInterventionForm(project=project)
    intervention_synthesis_form = SynopsisInterventionSynthesisForm()
    key_message_form = SynopsisKeyMessageForm()
    assignment_form = SynopsisAssignmentForm(project=project)
    redirect_url = reverse(
        f"synopsis:{redirect_name}", kwargs={"project_id": project.id}
    )

    interventions_prefetch = Prefetch(
        "interventions",
        queryset=SynopsisIntervention.objects.select_related(
            "iucn_category", "primary_intervention"
        )
        .order_by("position", "id")
        .prefetch_related(
            Prefetch(
                "key_messages",
                queryset=SynopsisInterventionKeyMessage.objects.order_by(
                    "position", "id"
                ).prefetch_related("supporting_summaries"),
            ),
            Prefetch(
                "assignments",
                queryset=SynopsisAssignment.objects.select_related(
                    "reference_summary__reference"
                ).order_by("position", "id"),
            )
        ),
    )
    subheading_prefetch = Prefetch(
        "subheadings",
        queryset=SynopsisSubheading.objects.order_by("position", "id").prefetch_related(
            interventions_prefetch
        ),
    )

    def _chapter_qs():
        return (
            SynopsisChapter.objects.filter(project=project)
            .prefetch_related(subheading_prefetch)
            .order_by("position", "id")
        )

    def _ensure_default_subheading(chapter):
        if not chapter.supports_evidence_structure:
            return
        if chapter.subheadings.exists():
            return
        SynopsisSubheading.objects.create(
            chapter=chapter, title="Interventions", position=1
        )

    if request.method == "POST":
        action = request.POST.get("action")

        def _chapter_from_post():
            return get_object_or_404(
                SynopsisChapter, pk=request.POST.get("chapter_id"), project=project
            )

        def _subheading_from_post():
            subheading = get_object_or_404(
                SynopsisSubheading,
                pk=request.POST.get("subheading_id"),
                chapter__project=project,
            )
            return subheading

        def _intervention_from_post():
            intervention = get_object_or_404(
                SynopsisIntervention,
                pk=request.POST.get("intervention_id"),
                subheading__chapter__project=project,
            )
            return intervention

        def _key_message_from_post():
            key_message = get_object_or_404(
                SynopsisInterventionKeyMessage,
                pk=request.POST.get("key_message_id"),
                intervention__subheading__chapter__project=project,
            )
            return key_message

        if action == "create-chapter":
            chapter_form = SynopsisChapterForm(request.POST)
            if chapter_form.is_valid():
                title = chapter_form.cleaned_data["title"] or "Untitled chapter"
                chapter_type = chapter_form.cleaned_data["chapter_type"]
                SynopsisChapter.objects.create(
                    project=project,
                    title=title,
                    chapter_type=chapter_type,
                    position=_next_chapter_position(project),
                )
                messages.success(request, f"Added chapter “{title}”.")
                return redirect(redirect_url)
            messages.error(request, "Please fix the problems below.")
        elif action == "delete-chapter":
            chapter = _chapter_from_post()
            chapter.delete()
            _resequence_chapter_positions(project)
            messages.success(request, f"Removed chapter “{chapter.title}”.")
            return redirect(redirect_url)
        elif action == "update-chapter-background":
            chapter = _chapter_from_post()
            bg_form = SynopsisBackgroundForm(request.POST)
            if bg_form.is_valid():
                chapter.background_text = bg_form.cleaned_data.get("background_text", "") or ""
                chapter.background_references = (
                    bg_form.cleaned_data.get("background_references", "") or ""
                )
                chapter.save(update_fields=["background_text", "background_references", "updated_at"])
                messages.success(request, "Chapter background saved.")
            else:
                messages.error(request, "Please check the background fields.")
            return redirect(redirect_url)
        elif action == "update-chapter-type":
            chapter = _chapter_from_post()
            chapter_type = (request.POST.get("chapter_type") or "").strip()
            allowed_types = {choice[0] for choice in SynopsisChapter.TYPE_CHOICES}
            if chapter_type not in allowed_types:
                messages.error(request, "Invalid chapter type selected.")
                return redirect(redirect_url)
            if (
                chapter.chapter_type != chapter_type
                and chapter_type != SynopsisChapter.TYPE_EVIDENCE
                and chapter.subheadings.exists()
            ):
                messages.error(
                    request,
                    "Remove subheadings and interventions before changing this chapter to text-only mode.",
                )
                return redirect(redirect_url)
            chapter.chapter_type = chapter_type
            chapter.save(update_fields=["chapter_type", "updated_at"])
            messages.success(request, "Chapter type updated.")
            return redirect(redirect_url)
        elif action == "move-chapter":
            chapter = _chapter_from_post()
            direction = request.POST.get("direction")
            if direction not in {"up", "down"}:
                messages.error(request, "Unknown move direction.")
            else:
                qs = list(_chapter_qs())
                if direction == "up":
                    swap = next((c for c in reversed(qs) if c.position < chapter.position), None)
                else:
                    swap = next((c for c in qs if c.position > chapter.position), None)
                if swap:
                    chapter.position, swap.position = swap.position, chapter.position
                    chapter.save(update_fields=["position"])
                    swap.save(update_fields=["position"])
                    messages.success(request, "Chapter reordered.")
                else:
                    messages.info(request, "Already at the edge.")
            return redirect(redirect_url)
        elif action == "create-subheading":
            chapter = _chapter_from_post()
            if not chapter.supports_evidence_structure:
                messages.error(
                    request, "Subheadings are available only for evidence chapters."
                )
                return redirect(redirect_url)
            subheading_form = SynopsisSubheadingForm(request.POST)
            if subheading_form.is_valid():
                title = subheading_form.cleaned_data["title"] or "Untitled subheading"
                SynopsisSubheading.objects.create(
                    chapter=chapter,
                    title=title,
                    position=_next_subheading_position(chapter),
                )
                messages.success(request, "Subheading added.")
                return redirect(redirect_url)
            messages.error(request, "Could not add the subheading.")
        elif action == "move-subheading":
            subheading = _subheading_from_post()
            direction = request.POST.get("direction")
            siblings = subheading.chapter.subheadings.order_by("position", "id")
            if direction == "up":
                swap = siblings.filter(position__lt=subheading.position).order_by("-position").first()
            elif direction == "down":
                swap = siblings.filter(position__gt=subheading.position).order_by("position").first()
            else:
                swap = None
            if swap:
                subheading.position, swap.position = swap.position, subheading.position
                subheading.save(update_fields=["position"])
                swap.save(update_fields=["position"])
                messages.success(request, "Subheading reordered.")
            else:
                messages.info(request, "Already at the edge.")
            return redirect(redirect_url)
        elif action == "delete-subheading":
            subheading = _subheading_from_post()
            chapter = subheading.chapter
            subheading.delete()
            _resequence_subheading_positions(chapter)
            messages.success(request, "Subheading removed.")
            return redirect(redirect_url)
        elif action == "create-intervention":
            subheading_id = request.POST.get("subheading_id")
            if subheading_id:
                subheading = _subheading_from_post()
                if not subheading.chapter.supports_evidence_structure:
                    messages.error(
                        request,
                        "Interventions are available only for evidence chapters.",
                    )
                    return redirect(redirect_url)
            else:
                chapter = _chapter_from_post()
                if not chapter.supports_evidence_structure:
                    messages.error(
                        request,
                        "Interventions are available only for evidence chapters.",
                    )
                    return redirect(redirect_url)
                _ensure_default_subheading(chapter)
                subheading = chapter.subheadings.first()
            intervention_form = SynopsisInterventionForm(
                request.POST, project=project
            )
            if intervention_form.is_valid():
                title = intervention_form.cleaned_data["title"] or "Untitled intervention"
                SynopsisIntervention.objects.create(
                    subheading=subheading,
                    title=title,
                    iucn_category=intervention_form.cleaned_data.get("iucn_category"),
                    is_cross_reference=intervention_form.cleaned_data.get(
                        "is_cross_reference", False
                    ),
                    primary_intervention=intervention_form.cleaned_data.get(
                        "primary_intervention"
                    ),
                    position=_next_intervention_position(subheading),
                )
                messages.success(request, "Intervention added.")
                return redirect(redirect_url)
            messages.error(request, "Could not add the intervention.")
        elif action == "move-intervention":
            intervention = _intervention_from_post()
            direction = request.POST.get("direction")
            siblings = intervention.subheading.interventions.order_by("position", "id")
            if direction == "up":
                swap = siblings.filter(position__lt=intervention.position).order_by("-position").first()
            elif direction == "down":
                swap = siblings.filter(position__gt=intervention.position).order_by("position").first()
            else:
                swap = None
            if swap:
                intervention.position, swap.position = swap.position, intervention.position
                intervention.save(update_fields=["position"])
                swap.save(update_fields=["position"])
                messages.success(request, "Intervention reordered.")
            else:
                messages.info(request, "Already at the edge.")
            return redirect(redirect_url)
        elif action == "move-intervention-to-subheading":
            intervention = _intervention_from_post()
            old_subheading = intervention.subheading
            target_subheading = get_object_or_404(
                SynopsisSubheading,
                pk=request.POST.get("target_subheading_id"),
                chapter=old_subheading.chapter,
                chapter__project=project,
            )
            if target_subheading.id == old_subheading.id:
                messages.info(request, "Intervention is already in that group.")
                return redirect(redirect_url)

            intervention.subheading = target_subheading
            intervention.position = _next_intervention_position(target_subheading)
            intervention.save(update_fields=["subheading", "position", "updated_at"])
            _resequence_intervention_positions(old_subheading)
            _resequence_intervention_positions(target_subheading)
            messages.success(
                request, f"Moved intervention to “{target_subheading.title}”."
            )
            return redirect(redirect_url)
        elif action == "delete-intervention":
            intervention = _intervention_from_post()
            subheading = intervention.subheading
            intervention.delete()
            _resequence_intervention_positions(subheading)
            messages.success(request, "Intervention removed.")
            return redirect(redirect_url)
        elif action == "update-intervention-background":
            intervention = _intervention_from_post()
            bg_form = SynopsisBackgroundForm(request.POST)
            if bg_form.is_valid():
                intervention.background_text = bg_form.cleaned_data.get("background_text", "") or ""
                intervention.background_references = (
                    bg_form.cleaned_data.get("background_references", "") or ""
                )
                intervention.save(
                    update_fields=["background_text", "background_references", "updated_at"]
                )
                messages.success(request, "Intervention background saved.")
            else:
                messages.error(request, "Please check the background fields.")
            return redirect(redirect_url)
        elif action == "update-intervention-synthesis":
            intervention = _intervention_from_post()
            synthesis_form = SynopsisInterventionSynthesisForm(request.POST)
            if synthesis_form.is_valid():
                intervention.ce_action_url = (
                    synthesis_form.cleaned_data.get("ce_action_url", "") or ""
                )
                intervention.evidence_status = synthesis_form.cleaned_data[
                    "evidence_status"
                ]
                intervention.synthesis_text = (
                    synthesis_form.cleaned_data.get("synthesis_text", "") or ""
                )
                intervention.save(
                    update_fields=[
                        "ce_action_url",
                        "evidence_status",
                        "synthesis_text",
                        "updated_at",
                    ]
                )
                messages.success(request, "Intervention synthesis saved.")
            else:
                messages.error(request, "Please check the synthesis fields.")
            return redirect(redirect_url)
        elif action == "add-key-message":
            intervention = _intervention_from_post()
            key_message_form = SynopsisKeyMessageForm(
                request.POST, intervention=intervention
            )
            if key_message_form.is_valid():
                key_message = SynopsisInterventionKeyMessage.objects.create(
                    intervention=intervention,
                    response_group=key_message_form.cleaned_data["response_group"],
                    outcome_label=key_message_form.cleaned_data.get("outcome_label", "")
                    or "",
                    statement=key_message_form.cleaned_data["statement"],
                    study_count=key_message_form.cleaned_data.get("study_count"),
                    position=_next_key_message_position(intervention),
                )
                key_message.supporting_summaries.set(
                    key_message_form.cleaned_data.get("supporting_summaries")
                )
                messages.success(request, "Key message added.")
            else:
                messages.error(request, "Could not add key message.")
            return redirect(redirect_url)
        elif action == "update-key-message":
            key_message = _key_message_from_post()
            key_message_form = SynopsisKeyMessageForm(
                request.POST, intervention=key_message.intervention
            )
            if key_message_form.is_valid():
                key_message.response_group = key_message_form.cleaned_data[
                    "response_group"
                ]
                key_message.outcome_label = (
                    key_message_form.cleaned_data.get("outcome_label", "") or ""
                )
                key_message.statement = key_message_form.cleaned_data["statement"]
                key_message.study_count = key_message_form.cleaned_data.get(
                    "study_count"
                )
                key_message.save(
                    update_fields=[
                        "response_group",
                        "outcome_label",
                        "statement",
                        "study_count",
                        "updated_at",
                    ]
                )
                key_message.supporting_summaries.set(
                    key_message_form.cleaned_data.get("supporting_summaries")
                )
                messages.success(request, "Key message updated.")
            else:
                messages.error(request, "Could not update key message.")
            return redirect(redirect_url)
        elif action == "move-key-message":
            key_message = _key_message_from_post()
            intervention = key_message.intervention
            direction = request.POST.get("direction")
            siblings = intervention.key_messages.order_by("position", "id")
            if direction == "up":
                swap = (
                    siblings.filter(position__lt=key_message.position)
                    .order_by("-position")
                    .first()
                )
            elif direction == "down":
                swap = (
                    siblings.filter(position__gt=key_message.position)
                    .order_by("position")
                    .first()
                )
            else:
                swap = None
            if swap:
                key_message.position, swap.position = swap.position, key_message.position
                key_message.save(update_fields=["position"])
                swap.save(update_fields=["position"])
                messages.success(request, "Key message reordered.")
            else:
                messages.info(request, "Already at the edge.")
            return redirect(redirect_url)
        elif action == "delete-key-message":
            key_message = _key_message_from_post()
            intervention = key_message.intervention
            key_message.delete()
            _resequence_key_message_positions(intervention)
            messages.success(request, "Key message removed.")
            return redirect(redirect_url)
        elif action == "update-intervention-metadata":
            intervention = _intervention_from_post()
            category = None
            category_id = (request.POST.get("iucn_category") or "").strip()
            if category_id:
                category = IUCNCategory.objects.filter(
                    pk=category_id,
                    kind=IUCNCategory.KIND_ACTION,
                    is_active=True,
                ).first()
                if not category:
                    messages.error(request, "Invalid IUCN category selected.")
                    return redirect(redirect_url)

            primary = None
            primary_id = (request.POST.get("primary_intervention") or "").strip()
            if primary_id:
                primary = SynopsisIntervention.objects.filter(
                    pk=primary_id,
                    subheading__chapter__project=project,
                ).first()
                if not primary:
                    messages.error(request, "Invalid primary intervention selected.")
                    return redirect(redirect_url)
                if primary.id == intervention.id:
                    messages.error(
                        request,
                        "An intervention cannot cross-reference itself.",
                    )
                    return redirect(redirect_url)

            is_cross_reference = bool(request.POST.get("is_cross_reference"))
            if primary and not is_cross_reference:
                is_cross_reference = True
            if is_cross_reference and not primary:
                messages.error(
                    request,
                    "Cross-reference interventions must point to a primary intervention.",
                )
                return redirect(redirect_url)

            intervention.iucn_category = category
            intervention.is_cross_reference = is_cross_reference
            intervention.primary_intervention = primary
            intervention.save(
                update_fields=[
                    "iucn_category",
                    "is_cross_reference",
                    "primary_intervention",
                    "updated_at",
                ]
            )
            messages.success(request, "Intervention metadata saved.")
            return redirect(redirect_url)
        elif action == "add-assignment":
            intervention = _intervention_from_post()
            assignment_form = SynopsisAssignmentForm(request.POST, project=project)
            if assignment_form.is_valid():
                summary = assignment_form.cleaned_data["summary"]
                exists = intervention.assignments.filter(reference_summary=summary).exists()
                if exists:
                    messages.info(request, "That summary is already assigned here.")
                else:
                    SynopsisAssignment.objects.create(
                        intervention=intervention,
                        reference_summary=summary,
                        position=_next_assignment_position(intervention),
                    )
                    messages.success(request, "Summary added to intervention.")
                return redirect(redirect_url)
            messages.error(request, "Select a summary to add.")
        elif action == "delete-assignment":
            assignment = get_object_or_404(
                SynopsisAssignment,
                pk=request.POST.get("assignment_id"),
                intervention__subheading__chapter__project=project,
            )
            intervention = assignment.intervention
            removed_summary = assignment.reference_summary
            assignment.delete()
            for key_message in intervention.key_messages.all():
                key_message.supporting_summaries.remove(removed_summary)
            _resequence_assignment_positions(intervention)
            messages.success(request, "Removed summary from intervention.")
            return redirect(redirect_url)
        elif action == "reorder-assignments" and request.content_type == "application/json":
            try:
                payload = json.loads(request.body.decode("utf-8"))
            except (TypeError, ValueError):
                return JsonResponse({"ok": False, "error": "Invalid payload"}, status=400)
            assignment_ids = payload.get("assignment_ids") or []
            intervention_id = payload.get("intervention_id")
            intervention = get_object_or_404(
                SynopsisIntervention, pk=intervention_id, subheading__chapter__project=project
            )
            assignments = list(intervention.assignments.filter(id__in=assignment_ids))
            id_map = {str(a.id): a for a in assignments}
            for idx, aid in enumerate(assignment_ids, start=1):
                assignment = id_map.get(str(aid))
                if not assignment:
                    continue
                if assignment.position != idx:
                    assignment.position = idx
                    assignment.save(update_fields=["position", "updated_at"])
            return JsonResponse({"ok": True})
        elif action == "apply-preset":
            preset_key = request.POST.get("preset_key")
            preset = PRESETS.get(preset_key)
            if not preset:
                messages.error(request, "Unknown preset.")
                return redirect(redirect_url)
            if SynopsisChapter.objects.filter(project=project).exists():
                messages.error(request, "Presets can only be applied to an empty outline.")
                return redirect(redirect_url)
            allowed_chapter_types = {choice[0] for choice in SynopsisChapter.TYPE_CHOICES}

            chapter_pos = 1
            for chapter_data in preset.chapters:
                chapter_type = (
                    chapter_data.get("chapter_type") or SynopsisChapter.TYPE_EVIDENCE
                )
                if chapter_type not in allowed_chapter_types:
                    chapter_type = SynopsisChapter.TYPE_EVIDENCE
                chapter = SynopsisChapter.objects.create(
                    project=project,
                    title=chapter_data.get("title") or "Untitled chapter",
                    chapter_type=chapter_type,
                    position=chapter_pos,
                )
                chapter_pos += 1
                sub_pos = 1
                for sub_data in chapter_data.get("subheadings", []) or []:
                    sub = SynopsisSubheading.objects.create(
                        chapter=chapter,
                        title=sub_data.get("title") or "Untitled subheading",
                        position=sub_pos,
                    )
                    sub_pos += 1
                    int_pos = 1
                    for int_data in sub_data.get("interventions", []) or []:
                        SynopsisIntervention.objects.create(
                            subheading=sub,
                            title=int_data.get("title") or "Untitled intervention",
                            position=int_pos,
                        )
                        int_pos += 1
            messages.success(request, f"Applied preset: {preset.label}.")
            return redirect(redirect_url)
        elif action == "reset-structure":
            SynopsisAssignment.objects.filter(
                intervention__subheading__chapter__project=project
            ).delete()
            SynopsisIntervention.objects.filter(
                subheading__chapter__project=project
            ).delete()
            SynopsisSubheading.objects.filter(chapter__project=project).delete()
            SynopsisChapter.objects.filter(project=project).delete()
            messages.success(request, "Cleared the outline. You can apply a preset or start fresh.")
            return redirect(redirect_url)

    # Ensure each chapter has at least one subheading so interventions can be added directly.
    for chapter in SynopsisChapter.objects.filter(project=project):
        _ensure_default_subheading(chapter)

    chapters = list(_chapter_qs())
    reference_summaries = (
        project.reference_summaries.select_related("reference")
        .order_by("reference__title", "created_at", "id")
        .all()
    )
    reference_summary_groups, summary_ui_meta = _reference_summary_workspace_groups(
        reference_summaries
    )

    for chapter in chapters:
        chapter.subheading_total = 0
        chapter.intervention_total = 0
        chapter.assignment_total = 0
        if not chapter.supports_evidence_structure:
            continue
        subheadings = list(chapter.subheadings.all())
        chapter.subheading_total = len(subheadings)
        for subheading in subheadings:
            interventions = list(subheading.interventions.all())
            subheading.intervention_total = len(interventions)
            chapter.intervention_total += subheading.intervention_total
            for intervention in interventions:
                assignments = list(intervention.assignments.all())
                (
                    _,
                    summary_numbers,
                    ordered_references,
                ) = _intervention_reference_numbering(assignments)
                reference_counts = Counter(
                    assignment.reference_summary.reference_id for assignment in assignments
                )
                reference_number_labels = {
                    reference.id: _format_reference_number_ranges(numbers)
                    for numbers, reference in ordered_references
                }
                for assignment in assignments:
                    meta = summary_ui_meta.get(assignment.reference_summary_id, {})
                    assignment.ce_reference_number = summary_numbers.get(
                        assignment.reference_summary_id
                    )
                    assignment.paper_title = meta.get(
                        "paper_title",
                        assignment.reference_summary.reference.canonical.title,
                    )
                    assignment.reference_heading = meta.get(
                        "reference_heading",
                        _reference_summary_workspace_heading(
                            assignment.reference_summary.reference
                        ),
                    )
                    assignment.reference_context = meta.get(
                        "reference_context",
                        _reference_summary_workspace_context(
                            assignment.reference_summary.reference
                        ),
                    )
                    assignment.paper_meta = meta.get("paper_meta", "")
                    assignment.summary_label = meta.get(
                        "summary_label", assignment.reference_summary.display_label
                    )
                    assignment.summary_display = meta.get(
                        "summary_display",
                        _reference_summary_workspace_label(assignment.reference_summary),
                    )
                    assignment.duplicate_reference_assignment = (
                        reference_counts.get(
                            assignment.reference_summary.reference_id, 0
                        )
                        > 1
                    )
                    assignment.reference_line_number_label = reference_number_labels.get(
                        assignment.reference_summary.reference_id,
                        (
                            str(assignment.ce_reference_number)
                            if assignment.ce_reference_number is not None
                            else ""
                        ),
                    )
                intervention.supporting_summary_options = [
                    {
                        "id": assignment.reference_summary_id,
                        "number": assignment.ce_reference_number,
                        "reference_heading": assignment.reference_heading,
                        "reference_context": assignment.reference_context,
                        "paper_title": assignment.paper_title,
                        "summary_label": assignment.summary_label,
                        "summary_display": assignment.summary_display,
                        "paragraph": _reference_summary_paragraph(
                            assignment.reference_summary,
                            reference_identifier_override=(
                                str(assignment.ce_reference_number)
                                if assignment.ce_reference_number
                                else None
                            ),
                        ),
                    }
                    for assignment in assignments
                ]
                intervention.compilation_preview = {
                    "study_paragraph_count": len(assignments),
                    "reference_line_count": len(ordered_references),
                    "shared_source_count": sum(
                        1 for numbers, _reference in ordered_references if len(numbers) > 1
                    ),
                    "reference_lines": [
                        {
                            "number_label": _format_reference_number_ranges(numbers),
                            "paper_title": reference.canonical.title,
                            "paper_meta": " · ".join(
                                [
                                    bit
                                    for bit in (
                                        reference.canonical.authors,
                                        (
                                            str(reference.canonical.publication_year)
                                            if reference.canonical.publication_year
                                            else ""
                                        ),
                                    )
                                    if bit
                                ]
                            ),
                            "study_count": len(numbers),
                            "is_shared": len(numbers) > 1,
                        }
                        for numbers, reference in ordered_references
                    ],
                }
                key_messages = list(intervention.key_messages.all())
                for key_message in key_messages:
                    supporting_ids = {
                        summary.id for summary in key_message.supporting_summaries.all()
                    }
                    key_message.supporting_summary_ids = supporting_ids
                    key_message.ce_supporting_numbers = _key_message_supporting_numbers(
                        key_message, summary_numbers
                    )
                intervention.assignment_total = len(assignments)
                intervention.key_message_total = len(key_messages)
                chapter.assignment_total += intervention.assignment_total

    text_chapters = [
        chapter for chapter in chapters if chapter.chapter_type == SynopsisChapter.TYPE_TEXT
    ]
    evidence_chapters = [
        chapter
        for chapter in chapters
        if chapter.chapter_type == SynopsisChapter.TYPE_EVIDENCE
    ]
    appendix_chapters = [
        chapter
        for chapter in chapters
        if chapter.chapter_type == SynopsisChapter.TYPE_APPENDIX
    ]
    evidence_intervention_total = sum(
        chapter.intervention_total for chapter in evidence_chapters
    )
    evidence_assignment_total = sum(
        chapter.assignment_total for chapter in evidence_chapters
    )
    narrative_total = len(text_chapters) + len(appendix_chapters)
    iucn_categories = IUCNCategory.objects.filter(
        kind=IUCNCategory.KIND_ACTION,
        is_active=True,
    ).order_by(
        "position", "name"
    )
    all_interventions = (
        SynopsisIntervention.objects.filter(subheading__chapter__project=project)
        .select_related("subheading__chapter")
        .order_by("title")
    )
    last_export = SynopsisExportLog.objects.filter(project=project).first()

    return render(
        request,
        template_name,
        {
            "project": project,
            "chapters": chapters,
            "chapter_form": chapter_form,
            "subheading_form": subheading_form,
            "intervention_form": intervention_form,
            "intervention_synthesis_form": intervention_synthesis_form,
            "key_message_form": key_message_form,
            "assignment_form": assignment_form,
            "reference_summaries": reference_summaries,
            "reference_summary_groups": reference_summary_groups,
            "text_chapters": text_chapters,
            "evidence_chapters": evidence_chapters,
            "appendix_chapters": appendix_chapters,
            "narrative_total": narrative_total,
            "evidence_intervention_total": evidence_intervention_total,
            "evidence_assignment_total": evidence_assignment_total,
            "iucn_categories": iucn_categories,
            "all_interventions": all_interventions,
            "last_exported": last_export.exported_at if last_export else None,
            "presets": PRESETS.values(),
            "workspace_mode": workspace_mode,
            "workspace_narrative_url": reverse(
                "synopsis:project_synopsis_narrative", kwargs={"project_id": project.id}
            ),
            "workspace_evidence_url": reverse(
                "synopsis:project_synopsis_evidence", kwargs={"project_id": project.id}
            ),
        },
    )


@login_required
def project_synopsis_structure(request, project_id):
    return _project_synopsis_workspace(
        request,
        project_id,
        workspace_mode="evidence",
        redirect_name="project_synopsis_structure",
        template_name="synopsis/project_synopsis_structure.html",
    )


@login_required
def project_synopsis_evidence(request, project_id):
    return _project_synopsis_workspace(
        request,
        project_id,
        workspace_mode="evidence",
        redirect_name="project_synopsis_evidence",
        template_name="synopsis/project_synopsis_structure.html",
    )


@login_required
def project_synopsis_narrative(request, project_id):
    return _project_synopsis_workspace(
        request,
        project_id,
        workspace_mode="narrative",
        redirect_name="project_synopsis_narrative",
        template_name="synopsis/project_synopsis_narrative.html",
    )


def _generate_synopsis_docx(project):
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise ImportError("Install python-docx to enable DOCX export.") from exc

    interventions_prefetch = Prefetch(
        "interventions",
        queryset=SynopsisIntervention.objects.order_by("position", "id").prefetch_related(
            Prefetch(
                "key_messages",
                queryset=SynopsisInterventionKeyMessage.objects.order_by(
                    "position", "id"
                ).prefetch_related("supporting_summaries"),
            ),
            Prefetch(
                "assignments",
                queryset=SynopsisAssignment.objects.select_related(
                    "reference_summary__reference"
                ).order_by("position", "id"),
            ),
        ),
    )
    subheading_prefetch = Prefetch(
        "subheadings",
        queryset=SynopsisSubheading.objects.order_by("position", "id").prefetch_related(
            interventions_prefetch
        ),
    )
    chapters = (
        SynopsisChapter.objects.filter(project=project)
        .prefetch_related(subheading_prefetch)
        .order_by("position", "id")
    )
    doc = Document()
    doc.add_heading(f"{project.title} – Synopsis", 0)
    doc.add_paragraph(f"Generated on {timezone.now().strftime('%Y-%m-%d %H:%M')}")

    def _render_chapter(chapter):
        doc.add_heading(chapter.title or "Untitled chapter", level=1)
        if chapter.background_text:
            doc.add_paragraph(chapter.background_text)
        if chapter.background_references:
            doc.add_paragraph(f"Background references: {chapter.background_references}")
        if not chapter.supports_evidence_structure:
            return
        for subheading in chapter.subheadings.all():
            doc.add_heading(subheading.title or "Untitled subheading", level=2)
            for intervention in subheading.interventions.all():
                doc.add_heading(intervention.title or "Untitled intervention", level=3)
                if intervention.ce_action_url:
                    doc.add_paragraph(f"Conservation Evidence action: {intervention.ce_action_url}")
                if intervention.background_text:
                    doc.add_paragraph(intervention.background_text)
                if intervention.background_references:
                    doc.add_paragraph(
                        f"Background references: {intervention.background_references}"
                    )

                if intervention.is_cross_reference and intervention.primary_intervention:
                    doc.add_paragraph(
                        f"Cross-reference: Evidence is summarized under "
                        f"“{intervention.primary_intervention.title}”."
                    )
                    continue

                assignments = list(intervention.assignments.all())
                (
                    ordered_assignments,
                    summary_numbers,
                    ordered_references,
                ) = _intervention_reference_numbering(assignments)

                key_messages = list(intervention.key_messages.all())
                if key_messages:
                    doc.add_paragraph("Key messages")
                    for message in key_messages:
                        prefix = message.get_response_group_display()
                        study_count = (
                            f" ({message.study_count} studies)"
                            if message.study_count is not None
                            else ""
                        )
                        label = (
                            f"{message.outcome_label}{study_count}: "
                            if message.outcome_label
                            else ""
                        )
                        text = f"{prefix}: {label}{message.statement}".strip()
                        supporting_numbers = _key_message_supporting_numbers(
                            message, summary_numbers
                        )
                        if supporting_numbers:
                            text = (
                                f"{text} ({_format_reference_number_list(supporting_numbers)})"
                            )
                        doc.add_paragraph(text, style="List Bullet")

                if (
                    intervention.evidence_status
                    == SynopsisIntervention.EVIDENCE_STATUS_NO_STUDIES
                ):
                    doc.add_paragraph(
                        "We found no studies that evaluated the effects of this intervention."
                    )

                if intervention.synthesis_text:
                    doc.add_paragraph(intervention.synthesis_text)

                for assignment in ordered_assignments:
                    summary = assignment.reference_summary
                    reference_number = summary_numbers.get(summary.id)
                    paragraph = _reference_summary_paragraph(
                        summary,
                        reference_identifier_override=(
                            str(reference_number) if reference_number else None
                        ),
                    )
                    if paragraph:
                        doc.add_paragraph(paragraph)
                    elif summary.reference.canonical.title:
                        if reference_number:
                            doc.add_paragraph(
                                f"({reference_number}) {summary.reference.canonical.title}."
                            )
                        else:
                            doc.add_paragraph(summary.reference.canonical.title)
                    if summary.cost_summary:
                        doc.add_paragraph(f"Costs: {summary.cost_summary}")

                if ordered_references:
                    doc.add_paragraph("References")
                    for numbers, ref in ordered_references:
                        citation = (
                            _reference_export_citation(ref)
                            or _reference_summary_citation(ref)
                            or (ref.canonical.title if hasattr(ref, "canonical") else ref.title)
                        )
                        number_label = _format_reference_number_ranges(numbers)
                        doc.add_paragraph(f"({number_label}) {citation}")

    for chapter in chapters:
        _render_chapter(chapter)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@login_required
def project_synopsis_export_docx(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    if not _user_can_edit_project(request.user, project):
        raise PermissionDenied
    try:
        payload = _generate_synopsis_docx(project)
    except ImportError as exc:
        messages.error(request, str(exc))
        return redirect("synopsis:project_synopsis_evidence", project_id=project.id)
    filename = slugify(f"{project.title}-synopsis").replace(" ", "-") + ".docx"
    log = SynopsisExportLog.objects.create(
        project=project,
        exported_by=request.user,
        note="Manual export",
    )
    try:
        log.archived_file.save(filename, ContentFile(payload), save=True)
    except Exception:
        # Best-effort logging; continue to serve the file even if archival failed.
        pass
    response = HttpResponse(
        payload,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


@login_required
def reference_batch_detail(request, project_id, batch_id):
    project = get_object_or_404(Project, pk=project_id)
    batch = get_object_or_404(ReferenceSourceBatch, pk=batch_id, project=project)

    references = batch.references.select_related(
        "screened_by", "library_reference"
    ).order_by("library_reference__title", "title")
    status_filter = request.GET.get("status")
    if status_filter in dict(Reference.SCREENING_STATUS_CHOICES):
        references = references.filter(screening_status=status_filter)

    ordered_ids = list(references.values_list("id", flat=True))
    focus_mode = (request.GET.get("focus") or "").strip() == "1"
    focused_reference = None
    focus_prev_id = None
    focus_next_id = None
    focus_index = None
    comment_form = ReferenceCommentForm()
    if focus_mode and ordered_ids:
        focus_id_param = (
            (request.GET.get("ref") or "").strip()
            or (request.POST.get("focus_ref") or "").strip()
            or (request.POST.get("reference_id") or "").strip()
        )
        try:
            focus_id = int(focus_id_param) if focus_id_param else ordered_ids[0]
        except (TypeError, ValueError):
            focus_id = ordered_ids[0]
        if focus_id not in ordered_ids:
            focus_id = ordered_ids[0]
        focus_index = ordered_ids.index(focus_id)
        focused_reference = references.filter(pk=focus_id).select_related(
            "screened_by", "library_reference"
        ).first()
        if focus_index > 0:
            focus_prev_id = ordered_ids[focus_index - 1]
        if focus_index < len(ordered_ids) - 1:
            focus_next_id = ordered_ids[focus_index + 1]

    if request.method == "POST":
        action_type = request.POST.get("action")
        if action_type == "update_notes":
            if not _user_can_edit_project(request.user, project):
                raise PermissionDenied
            new_notes = (request.POST.get("notes") or "").strip()
            current_notes = batch.notes or ""
            status_filter = request.POST.get("status_filter") or status_filter or ""
            if new_notes == current_notes:
                messages.info(request, "Notes unchanged.")
            else:
                ReferenceSourceBatchNoteHistory.objects.create(
                    batch=batch,
                    previous_notes=current_notes,
                    new_notes=new_notes,
                    changed_by=request.user if request.user.is_authenticated else None,
                )
                batch.notes = new_notes
                batch.save(update_fields=["notes"])
                messages.success(request, "Batch notes updated.")

            redirect_url = reverse(
                "synopsis:reference_batch_detail",
                kwargs={"project_id": project.id, "batch_id": batch.id},
            )
            if status_filter in dict(Reference.SCREENING_STATUS_CHOICES):
                redirect_url = f"{redirect_url}?status={status_filter}"
            return redirect(redirect_url)

        if action_type == "add-ref-comment":
            if not _user_can_edit_project(request.user, project):
                raise PermissionDenied
            comment_form = ReferenceCommentForm(request.POST, request.FILES)
            ref_id = request.POST.get("reference_id")
            ref = get_object_or_404(Reference, pk=ref_id, project=project, batch=batch)
            if comment_form.is_valid():
                parent = None
                parent_id = comment_form.cleaned_data.get("parent_id")
                if parent_id:
                    parent = ReferenceComment.objects.filter(
                        pk=parent_id, reference=ref
                    ).first()
                ReferenceComment.objects.create(
                    reference=ref,
                    author=request.user,
                    body=comment_form.cleaned_data["body"],
                    parent=parent,
                    attachment=comment_form.cleaned_data.get("attachment"),
                )
                messages.success(request, "Comment added.")
            else:
                messages.error(request, "Could not add comment.")
            redirect_params = []
            if status_filter:
                redirect_params.append(("status", status_filter))
            if (request.POST.get("focus") or "").strip() == "1":
                redirect_params.append(("focus", "1"))
                redirect_params.append(("ref", ref.id))
            redirect_url = reverse(
                "synopsis:reference_batch_detail",
                kwargs={"project_id": project.id, "batch_id": batch.id},
            )
            if redirect_params:
                redirect_url = f"{redirect_url}?{urlencode(redirect_params)}"
            return redirect(redirect_url)

        bulk_action = request.POST.get("bulk_action")
        if bulk_action:
            if not _user_can_edit_project(request.user, project):
                raise PermissionDenied

            status_choices = dict(Reference.SCREENING_STATUS_CHOICES)
            status_filter = (
                request.POST.get("status_filter")
                or request.GET.get("status")
                or ""
            )
            selected_ids = [
                pk
                for pk in request.POST.getlist("selected_references")
                if pk.isdigit()
            ]

            if not selected_ids:
                messages.warning(
                    request,
                    "Select at least one reference before applying a bulk update.",
                )
                redirect_url = reverse(
                    "synopsis:reference_batch_detail",
                    kwargs={"project_id": project.id, "batch_id": batch.id},
                )
                if status_filter in status_choices:
                    redirect_url = f"{redirect_url}?status={status_filter}"
                return redirect(redirect_url)

            if bulk_action == "save-folders":
                categories = normalize_reference_folder_values(
                    request.POST.getlist("reference_folder")
                )
                updated = 0
                shared_updated = 0
                now = timezone.now()
                for ref in batch.references.select_related("library_reference").filter(
                    pk__in=selected_ids
                ):
                    ref.screening_decision_at = now
                    if request.user.is_authenticated:
                        ref.screened_by = request.user
                    ref.save(
                        update_fields=[
                            "screening_decision_at",
                            "screened_by",
                            "updated_at",
                        ]
                    )
                    changed, _linked_count, _local_changed, _saved_categories = (
                        _update_reference_categories(
                            ref,
                            categories,
                            changed_by=request.user,
                            source_project=project,
                            change_source="screening_bulk_save_folders",
                        )
                    )
                    if changed:
                        shared_updated += 1
                    updated += 1

                message = f"Updated categories for {updated} reference(s)."
                if shared_updated:
                    message += (
                        f" Updated the shared library categories for {shared_updated} linked reference(s)."
                    )
                messages.success(request, message)
                redirect_url = reverse(
                    "synopsis:reference_batch_detail",
                    kwargs={"project_id": project.id, "batch_id": batch.id},
                )
                if status_filter in status_choices:
                    redirect_url = f"{redirect_url}?status={status_filter}"
                return redirect(redirect_url)

            action_map = {
                "include": "included",
                "exclude": "excluded",
                "pending": "pending",
            }
            new_status = action_map.get(bulk_action)
            if not new_status:
                messages.error(
                    request,
                    "Unknown bulk action requested.",
                )
                return redirect(
                    "synopsis:reference_batch_detail",
                    project_id=project.id,
                    batch_id=batch.id,
                )

            updated = 0
            now = timezone.now()
            bulk_folder = normalize_reference_folder_values(
                request.POST.getlist("reference_folder")
            )
            apply_bulk_folder = "reference_folder" in request.POST and bool(bulk_folder)
            shared_updated = 0
            for ref in batch.references.select_related("library_reference").filter(
                pk__in=selected_ids
            ):
                ref.screening_status = new_status
                ref.screening_decision_at = now
                if request.user.is_authenticated:
                    ref.screened_by = request.user
                update_fields = [
                    "screening_status",
                    "screening_decision_at",
                    "screened_by",
                    "updated_at",
                ]
                ref.save(update_fields=update_fields)
                if apply_bulk_folder:
                    changed, _linked_count, _local_changed, _saved_categories = (
                        _update_reference_categories(
                            ref,
                            bulk_folder,
                            changed_by=request.user,
                            source_project=project,
                            change_source=f"screening_bulk_{new_status}",
                        )
                    )
                    if changed:
                        shared_updated += 1
                updated += 1

            if updated:
                status_label = status_choices.get(new_status, new_status.title())
                message = f"Marked {updated} reference(s) as {status_label}."
                if apply_bulk_folder:
                    message += " Applied the selected categories at the same time."
                    if shared_updated:
                        message += (
                            f" Updated the shared library categories for {shared_updated} linked reference(s)."
                        )
                messages.success(request, message)
            else:
                messages.info(request, "No references matched the selection.")

            redirect_url = reverse(
                "synopsis:reference_batch_detail",
                kwargs={"project_id": project.id, "batch_id": batch.id},
            )
            if status_filter in status_choices:
                redirect_url = f"{redirect_url}?status={status_filter}"
            return redirect(redirect_url)

        form = ReferenceScreeningForm(request.POST)
        if form.is_valid():
            is_focus_post = (request.POST.get("focus") or "").strip() == "1"
            focus_stay_put_actions = {"save-categories", "save-notes"}
            is_focus_stay_put_action = (
                is_focus_post
                and request.POST.get("action") in focus_stay_put_actions
            )
            ref_id = form.cleaned_data["reference_id"]
            if is_focus_post:
                focus_ref_override = (request.POST.get("focus_ref") or "").strip()
                if focus_ref_override.isdigit():
                    ref_id = int(focus_ref_override)
            ref = get_object_or_404(
                Reference.objects.select_related("library_reference"),
                pk=ref_id,
                batch=batch,
                project=project,
            )
            status = form.cleaned_data["screening_status"]
            folder = normalize_reference_folder_values(
                form.cleaned_data.get("reference_folder") or []
            )
            ref.screening_status = status
            update_fields = [
                "screening_status",
                "screening_decision_at",
                "screened_by",
                "updated_at",
            ]
            if "screening_notes" in request.POST:
                ref.screening_notes = form.cleaned_data.get("screening_notes") or ""
                update_fields.append("screening_notes")
            ref.screening_decision_at = timezone.now()
            ref.screened_by = request.user
            ref.save(update_fields=update_fields)
            message = f"Updated screening status for '{ref.canonical.title[:80]}'."
            if "reference_folder" in request.POST:
                shared_changed, _linked_count, _local_changed, _saved_categories = (
                    _update_reference_categories(
                        ref,
                        folder,
                        changed_by=request.user,
                        source_project=project,
                        change_source="screening_single",
                    )
                )
                if shared_changed:
                    message += (
                        " Shared CE subject categories were updated for all linked synopsis copies."
                    )
            messages.success(request, message)
            redirect_params = []
            if status_filter:
                redirect_params.append(("status", status_filter))
            if (request.POST.get("focus") or "").strip() == "1":
                redirect_params.append(("focus", "1"))
                next_ref_id = (
                    ref.id
                    if is_focus_stay_put_action
                    else (request.POST.get("next_ref_id") or ref.id)
                )
                redirect_params.append(("ref", next_ref_id))
            redirect_url = reverse(
                "synopsis:reference_batch_detail",
                kwargs={"project_id": project.id, "batch_id": batch.id},
            )
            if redirect_params:
                redirect_url = f"{redirect_url}?{urlencode(redirect_params)}"
            if not is_focus_post:
                redirect_url = f"{redirect_url}#ref-{ref.id}"
            return redirect(redirect_url)
        else:
            messages.error(
                request,
                "Unable to update screening status. Please check the submission.",
            )
    status_counts = {
        row["screening_status"]: row["count"]
        for row in batch.references.values("screening_status").annotate(
            count=Count("id")
        )
    }
    status_summary = [
        {
            "status": value,
            "label": label,
            "count": status_counts.get(value, 0),
        }
        for value, label in Reference.SCREENING_STATUS_CHOICES
    ]
    included_count = status_counts.get("included", 0)
    excluded_count = status_counts.get("excluded", 0)
    pending_count = status_counts.get("pending", 0)
    screened_count = included_count + excluded_count
    total_references = batch.references.count()
    completion_percent = (
        round((screened_count / total_references) * 100) if total_references else 0
    )

    visible_total = references.count()
    visible_screened = (
        references.filter(screening_status__in=["included", "excluded"]).count()
        if visible_total
        else 0
    )
    visible_completion = (
        round((visible_screened / visible_total) * 100) if visible_total else 0
    )

    latest_screening = (
        batch.references.exclude(screening_decision_at__isnull=True)
        .select_related("screened_by")
        .order_by("-screening_decision_at")
        .first()
    )

    summary_stats = {
        "included": included_count,
        "excluded": excluded_count,
        "pending": pending_count,
        "screened": screened_count,
        "total": total_references,
        "completion_percent": completion_percent,
        "visible_total": visible_total,
        "visible_screened": visible_screened,
        "visible_completion": visible_completion,
        "has_filter": bool(status_filter),
        "last_screened_at": latest_screening.screening_decision_at
        if latest_screening
        else None,
        "last_screened_by": latest_screening.screened_by
        if latest_screening
        else None,
    }

    # Decode abstracts for clean display
    for ref in references:
        ref.decoded_abstract = _decode_entities(ref.canonical.abstract)
    if focused_reference:
        focused_reference.decoded_abstract = _decode_entities(
            focused_reference.canonical.abstract
        )

    # Comments/notes per reference (lightweight counts; tree built only when manageable)
    comment_trees = {}
    comment_counts = {}
    target_refs = [focused_reference] if focus_mode and focused_reference else list(
        references
    )
    if target_refs:
        counts = (
            ReferenceComment.objects.filter(reference__in=target_refs)
            .values("reference_id")
            .annotate(count=Count("id"))
        )
        for row in counts:
            comment_counts[row["reference_id"]] = row["count"]

        MAX_REFS_FOR_COMMENT_TREE = 50  # build trees only for small sets to avoid heavy processing
        if focus_mode or len(target_refs) <= MAX_REFS_FOR_COMMENT_TREE:
            comment_qs = (
                ReferenceComment.objects.filter(reference__in=target_refs)
                .select_related("author", "reference")
                .order_by("-created_at", "-id")
            )
            by_ref = defaultdict(list)
            for c in comment_qs:
                by_ref[c.reference_id].append(c)
            for ref_id, items in by_ref.items():
                children = defaultdict(list)
                for c in items:
                    children[c.parent_id].append(c)
                tree = children[None]
                for c in tree:
                    c.replies_cached = children.get(c.id, [])
                comment_trees[ref_id] = tree

    return render(
        request,
        "synopsis/reference_batch_detail.html",
        {
            "project": project,
            "batch": batch,
            "references": references,
            "focus_mode": focus_mode,
            "focused_reference": focused_reference,
            "focus_prev_id": focus_prev_id,
            "focus_next_id": focus_next_id,
            "focus_index": focus_index,
            "focus_total": len(ordered_ids),
            "status_filter": status_filter,
            "status_choices": Reference.SCREENING_STATUS_CHOICES,
            "status_summary": status_summary,
            "summary_stats": summary_stats,
            "note_history": batch.note_history.select_related("changed_by").all(),
            "comment_form": comment_form,
            "comment_trees": comment_trees,
            "comment_counts": comment_counts,
            "folder_choices": Reference.FOLDER_CHOICES,
        },
    )


@login_required
def reference_delete(request, project_id, reference_id):
    project = get_object_or_404(Project, pk=project_id)
    reference = get_object_or_404(
        Reference.objects.select_related("batch"),
        pk=reference_id,
        project=project,
    )

    if request.method != "POST":
        return HttpResponseNotAllowed(["POST"])

    if not _user_can_edit_project(request.user, project):
        raise PermissionDenied

    batch = reference.batch
    status_choices = dict(Reference.SCREENING_STATUS_CHOICES)
    status_filter = (
        request.POST.get("status_filter")
        or request.GET.get("status")
        or ""
    )

    title_fragment = reference.canonical.title[:80]
    reference.delete()

    batch.record_count = batch.references.count()
    batch.save(update_fields=["record_count"])

    messages.success(
        request,
        f"Removed '{title_fragment}' from '{batch.label}'.",
    )

    redirect_url = reverse(
        "synopsis:reference_batch_detail",
        kwargs={"project_id": project.id, "batch_id": batch.id},
    )
    if status_filter in status_choices:
        redirect_url = f"{redirect_url}?status={status_filter}"

    parsed_redirect = urlparse(redirect_url)
    if parsed_redirect.scheme or parsed_redirect.netloc:
        redirect_url = reverse(
            "synopsis:reference_batch_detail",
            kwargs={"project_id": project.id, "batch_id": batch.id},
        )

    return redirect(redirect_url)


@login_required
def reference_batch_delete(request, project_id, batch_id):
    project = get_object_or_404(Project, pk=project_id)
    batch = get_object_or_404(ReferenceSourceBatch, pk=batch_id, project=project)

    if request.method != "POST":
        return HttpResponseNotAllowed(["POST"])

    if not _user_can_edit_project(request.user, project):
        raise PermissionDenied

    label_fragment = batch.label[:80]
    batch.delete()

    messages.success(
        request,
        f"Deleted '{label_fragment}' and its imported references.",
    )

    return redirect(
        "synopsis:reference_batch_list",
        project_id=project.id,
    )


@login_required
def reference_batch_upload(request, project_id):
    project = get_object_or_404(Project, pk=project_id)

    if request.method == "POST":
        form = ReferenceBatchUploadForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = form.cleaned_data["ris_file"]
            raw_bytes = uploaded_file.read()
            if not raw_bytes.strip():
                form.add_error("ris_file", "The uploaded file appears to be empty.")
            else:
                sha1 = hashlib.sha1(raw_bytes).hexdigest()
                text_payload = _decode_reference_upload_text(raw_bytes)
                records = []
                ris_error = None
                try:
                    records = rispy.loads(text_payload)
                except Exception as exc:  # pragma: no cover - parser errors
                    ris_error = exc

                if not records:
                    plaintext_records = _parse_plaintext_references(text_payload)
                    if plaintext_records:
                        records = plaintext_records
                    elif ris_error:
                        form.add_error(
                            "ris_file",
                            f"Could not parse RIS content ({ris_error}).",
                        )
                    else:
                        form.add_error(
                            "ris_file",
                            "No references were detected. Upload a RIS file or a plain text file where each entry is separated by a blank line.",
                        )

                if records:
                    with transaction.atomic():
                        batch = ReferenceSourceBatch.objects.create(
                            project=project,
                            label=form.cleaned_data["label"],
                            source_type=form.cleaned_data["source_type"],
                            search_date_start=form.cleaned_data.get(
                                "search_date_start"
                            ),
                            search_date_end=form.cleaned_data.get(
                                "search_date_end"
                            ),
                            uploaded_by=request.user,
                            original_filename=getattr(uploaded_file, "name", ""),
                            record_count=0,
                            ris_sha1=sha1,
                            notes=form.cleaned_data.get("notes", ""),
                        )
                        imported = 0
                        duplicates = 0
                        skipped = 0
                        for record in records:
                            data = _normalise_import_record(record)
                            if not data:
                                skipped += 1
                                continue

                            hash_key = reference_hash(
                                data["title"], data["year"], data["doi"]
                            )

                            if Reference.objects.filter(
                                project=project, hash_key=hash_key
                            ).exists():
                                duplicates += 1
                                continue

                            library_ref, _ = LibraryReference.objects.get_or_create(
                                hash_key=hash_key,
                                defaults={
                                    "source_identifier": data["source_identifier"],
                                    "title": data["title"],
                                    "abstract": data["abstract"],
                                    "authors": data["authors"],
                                    "publication_year": data["publication_year"],
                                    "journal": data["journal"],
                                    "volume": data["volume"],
                                    "issue": data["issue"],
                                    "pages": data["pages"],
                                    "doi": data["doi"],
                                    "url": data["url"],
                                    "language": data["language"],
                                    "raw_ris": record,
                                },
                            )

                            Reference.objects.create(
                                project=project,
                                batch=batch,
                                library_reference=library_ref,
                                hash_key=hash_key,
                                source_identifier=data["source_identifier"],
                                title=data["title"],
                                abstract=data["abstract"],
                                authors=data["authors"],
                                publication_year=data["publication_year"],
                                journal=data["journal"],
                                volume=data["volume"],
                                issue=data["issue"],
                                pages=data["pages"],
                                doi=data["doi"],
                                url=data["url"],
                                language=data["language"],
                                raw_ris=record,
                            )
                            imported += 1

                        batch.record_count = imported
                        batch.save(update_fields=["record_count", "notes"])

                    messages.success(
                        request,
                        f"Imported {imported} reference(s) into '{batch.label}'.",
                    )
                    if duplicates:
                        messages.info(
                            request,
                            f"Skipped {duplicates} record(s) already present in this project.",
                        )
                    if skipped:
                        messages.info(
                            request,
                            f"Skipped {skipped} record(s) with no title.",
                        )
                    return redirect(
                        "synopsis:reference_batch_list", project_id=project.id
                    )
    else:
        form = ReferenceBatchUploadForm()

    return render(
        request,
        "synopsis/reference_batch_upload.html",
        {"project": project, "form": form},
    )


@login_required
def advisory_protocol_feedback_close(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    proto = getattr(project, "protocol", None)
    if not proto:
        messages.error(request, "No protocol configured for this project.")
        return redirect("synopsis:protocol_detail", project_id=project.id)

    action = request.POST.get("action")
    if action == "reopen":
        if request.method != "POST":
            return HttpResponseBadRequest("POST required")
        proto.feedback_closed_at = None
        proto.feedback_closure_message = ""
        proto.save(update_fields=["feedback_closed_at", "feedback_closure_message"])
        _log_project_change(
            project,
            request.user,
            "Protocol feedback reopened",
        )
        messages.success(request, "Protocol feedback reopened for advisory members.")
        return redirect("synopsis:protocol_detail", project_id=project.id)

    if request.method != "POST":
        return HttpResponseBadRequest("POST required")

    form = ProtocolFeedbackCloseForm(request.POST)
    if not form.is_valid():
        for errors in form.errors.values():
            for error in errors:
                messages.error(request, error)
        return redirect("synopsis:protocol_detail", project_id=project.id)

    message = form.cleaned_data.get("message", "")
    now = timezone.now()
    already_closed = proto.feedback_closed_at is not None
    proto.feedback_closed_at = proto.feedback_closed_at or now
    proto.feedback_closure_message = message
    update_fields = ["feedback_closure_message"]
    if not already_closed:
        update_fields.append("feedback_closed_at")
    proto.save(update_fields=update_fields)

    ended_session = _end_active_collaborative_session(
        project,
        CollaborativeSession.DOCUMENT_PROTOCOL,
        ended_by=request.user if getattr(request.user, "is_authenticated", False) else None,
        reason="Protocol feedback window closed",
    )

    if already_closed:
        _log_project_change(
            project,
            request.user,
            "Protocol feedback closure message updated",
            message,
        )
        messages.info(request, "Closure message updated.")
    else:
        _log_project_change(
            project,
            request.user,
            "Protocol feedback closed",
            message,
        )
        messages.success(request, "Protocol feedback links are now closed.")

    if ended_session:
        _log_project_change(
            project,
            request.user,
            "Protocol collaborative session closed",
            f"Session {ended_session.token} marked inactive.",
        )
    return redirect("synopsis:protocol_detail", project_id=project.id)


@login_required
def advisory_action_list_feedback_close(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    action_list = getattr(project, "action_list", None)
    if not action_list:
        messages.error(request, "No action list configured for this project.")
        return redirect("synopsis:action_list_detail", project_id=project.id)

    action = request.POST.get("action")
    if action == "reopen":
        if request.method != "POST":
            return HttpResponseBadRequest("POST required")
        action_list.feedback_closed_at = None
        action_list.feedback_closure_message = ""
        action_list.save(
            update_fields=["feedback_closed_at", "feedback_closure_message"]
        )
        _log_project_change(
            project,
            request.user,
            "Action list feedback reopened",
        )
        messages.success(request, "Action list feedback reopened for advisory members.")
        return redirect("synopsis:action_list_detail", project_id=project.id)

    if request.method != "POST":
        return HttpResponseBadRequest("POST required")

    form = ActionListFeedbackCloseForm(request.POST)
    if not form.is_valid():
        for errors in form.errors.values():
            for error in errors:
                messages.error(request, error)
        return redirect("synopsis:action_list_detail", project_id=project.id)

    message = form.cleaned_data.get("message", "")
    now = timezone.now()
    already_closed = action_list.feedback_closed_at is not None
    action_list.feedback_closed_at = action_list.feedback_closed_at or now
    action_list.feedback_closure_message = message
    update_fields = ["feedback_closure_message"]
    if not already_closed:
        update_fields.append("feedback_closed_at")
    action_list.save(update_fields=update_fields)

    ended_session = _end_active_collaborative_session(
        project,
        CollaborativeSession.DOCUMENT_ACTION_LIST,
        ended_by=request.user if getattr(request.user, "is_authenticated", False) else None,
        reason="Action list feedback window closed",
    )

    if already_closed:
        _log_project_change(
            project,
            request.user,
            "Action list feedback closure message updated",
            message,
        )
        messages.info(request, "Closure message updated.")
    else:
        _log_project_change(
            project,
            request.user,
            "Action list feedback closed",
            message,
        )
        messages.success(request, "Action list feedback links are now closed.")

    if ended_session:
        _log_project_change(
            project,
            request.user,
            "Action list collaborative session closed",
            f"Session {ended_session.token} marked inactive.",
        )
    return redirect("synopsis:action_list_detail", project_id=project.id)


@login_required
def advisory_invite_create(request, project_id, member_id=None):
    """
    Invite by email. If member_id is provided, the invite links to that member
    and updates their invite flags; otherwise the invite appears in the
    “Invited directly” section.
    """
    project = get_object_or_404(Project, pk=project_id)
    initial = {}
    member = None

    action_list = getattr(project, "action_list", None)
    action_document_available = bool(action_list and getattr(action_list, "document", None))
    collaborative_available = bool(_onlyoffice_enabled() and action_document_available)

    if member_id:
        member = get_object_or_404(AdvisoryBoardMember, pk=member_id, project=project)
        initial["email"] = member.email
        if member.response_date:
            initial["due_date"] = member.response_date
    if "due_date" not in initial:
        initial["due_date"] = _default_invite_due_date()

    form = None
    if request.method == "POST":
        form = AdvisoryInviteForm(request.POST, project=project)
    else:
        form = AdvisoryInviteForm(initial=initial, project=project)
        
    if not action_document_available:
        form.fields["include_action_list"].disabled = True
        form.fields["include_action_list"].help_text = "Upload an action list document to include it here."
        form.fields["include_action_list"].initial = False
    else:
        form.fields["include_action_list"].help_text = "Adds a link to the latest action list file."

    if not collaborative_available:
        form.fields["include_collaborative_link"].disabled = True
        form.fields["include_collaborative_link"].help_text = "Enable the collaborative editor or upload the action list to share this link."
        form.fields["include_collaborative_link"].initial = False
    else:
        form.fields["include_collaborative_link"].help_text = "Shares the live collaborative editor for the action list."

    if request.method == "POST" and form.is_valid():
            email = form.cleaned_data["email"].strip()
            due_date = _resolve_invite_due_date(
                form.cleaned_data.get("due_date"),
                member=member,
            )
            standard_message = form.cleaned_data.get("standard_message") or ""
            message_body = form.cleaned_data.get("message") or ""
            if "standard_message" in request.POST:
                _update_project_advisory_invitation_message(
                    project, standard_message, request.user
                )

            inv = AdvisoryBoardInvitation.objects.create(
                project=project,
                member=member,
                email=email,
                invited_by=request.user,
                due_date=due_date,
            )

            yes_url = request.build_absolute_uri(
                reverse("synopsis:advisory_invite_reply", args=[str(inv.token), "yes"])
            )
            no_url = request.build_absolute_uri(
                reverse("synopsis:advisory_invite_reply", args=[str(inv.token), "no"])
            )
            subject = email_subject("invite", project, due_date)
            # Collect optional attachments/links
            attachment_lines = []

            include_action_list = (
                action_document_available
                and form.cleaned_data.get("include_action_list")
                and action_list
                and getattr(action_list, "document", None)
            )
            if include_action_list:
                action_url = request.build_absolute_uri(action_list.document.url)
                attachment_lines.append(("Action list", action_url))

            include_collaborative_link = (
                collaborative_available
                and form.cleaned_data.get("include_collaborative_link")
            )
            if include_collaborative_link:
                collab_url = _ensure_collaborative_invite_link(
                    request,
                    project,
                    CollaborativeSession.DOCUMENT_ACTION_LIST,
                    inv,
                    member=member,
                )
                if collab_url:
                    attachment_lines.append(("Collaborative editor", collab_url))

            text, html = _build_advisory_invitation_email(
                project=project,
                recipient_name=advisory_member_display_name(member),
                due_date=due_date,
                yes_url=yes_url,
                no_url=no_url,
                standard_message=standard_message,
                additional_message=message_body,
                attachment_lines=attachment_lines,
            )

            msg = EmailMultiAlternatives(
                subject,
                text,
                to=[email],
                reply_to=reply_to_list(getattr(request.user, "email", None)),
            )
            msg.attach_alternative(html, "text/html")
            msg.send()

            if member:
                member.invite_sent = True
                member.invite_sent_at = timezone.now()
                update_fields = {"invite_sent", "invite_sent_at"}
                if member.response_date != due_date:
                    member.response_date = due_date
                    member.reminder_sent = False
                    member.reminder_sent_at = None
                    update_fields.update(
                        {"response_date", "reminder_sent", "reminder_sent_at"}
                    )
                member.save(update_fields=list(update_fields))

            messages.success(request, f"Invitation sent to {email}.")
            return redirect("synopsis:advisory_board_list", project_id=project.id)
    return render(
        request,
        "synopsis/advisory_invite_form.html",
        {
            "project": project,
            "form": form,
            "member": member,
            "action_list_available": action_document_available,
            "collaborative_available": collaborative_available,
            "default_invitation_message": default_advisory_invitation_message(),
            "preview_recipient_name": (
                advisory_member_display_name(member)
            ),
            "preview_is_bulk": False,
        },
    )


@csrf_exempt
def advisory_invite_accept(request, token):
    """
    Kept for compatibility with older links that only had 'accept'.
    New invites should use Yes/No links via advisory_invite_reply.
    """
    inv = get_object_or_404(AdvisoryBoardInvitation, token=token)
    _log_project_change(
        inv.project,
        request.user,
        "Opened advisory invite link",
        f"Source: legacy accept link | Email: {inv.email or '—'}",
    )
    reply_url = reverse("synopsis:advisory_invite_reply", args=[str(inv.token), "yes"])
    return redirect(f"{reply_url}?source=legacy_accept")


def advisory_invite_reply(request, token, choice):
    """
    Public Yes/No link handler from the email.
    - If the invite is tied to a member, update that row.
    - If it’s an email-only invite and they click Yes, create a member.
    """
    inv = get_object_or_404(AdvisoryBoardInvitation, token=token)
    choice = (choice or "").lower()
    if choice not in ("yes", "no"):
        return HttpResponseBadRequest("Invalid choice")

    accepted = choice == "yes"
    member = inv.member
    link_source = (request.GET.get("source") or "reply_link").strip() or "reply_link"

    if accepted:
        if member and member.participation_confirmed:
            if not inv.accepted:
                inv.accepted = True
                inv.responded_at = timezone.now()
                inv.save(update_fields=["accepted", "responded_at"])
            _log_project_change(
                inv.project,
                request.user,
                "Opened advisory invite link",
                f"Source: {link_source} | Choice: yes | Email: {inv.email or '—'} | Already confirmed: yes",
            )
            return render(
                request,
                "synopsis/invite_thanks.html",
                {"member": member, "project": inv.project, "accepted": True},
            )

        form = ParticipationConfirmForm(request.POST or None)
        if request.method != "POST":
            _log_project_change(
                inv.project,
                request.user,
                "Opened advisory invite link",
                f"Source: {link_source} | Choice: yes | Email: {inv.email or '—'} | Showing participation form",
            )
        if request.method == "POST":
            if form.is_valid():
                statement = (form.cleaned_data.get("statement") or "").strip()
                if not statement:
                    statement = "Participation confirmed"
                now = timezone.now()
                today = timezone.localdate()

                if member is None:
                    member = AdvisoryBoardMember.objects.create(
                        project=inv.project,
                        email=inv.email,
                        first_name="",
                        last_name="",
                        organisation="",
                    )
                    inv.member = member

                updates = set()
                if member.response != "Y":
                    member.response = "Y"
                    updates.add("response")
                if not member.response_date:
                    member.response_date = today
                    updates.add("response_date")
                member.participation_confirmed = True
                member.participation_confirmed_at = now
                member.participation_statement = statement
                updates.update(
                    {
                        "participation_confirmed",
                        "participation_confirmed_at",
                        "participation_statement",
                    }
                )
                member.save(update_fields=list(updates))

                inv.accepted = True
                inv.responded_at = now
                inv.save(update_fields=["accepted", "responded_at", "member"])
                _log_project_change(
                    inv.project,
                    request.user,
                    "Confirmed advisory participation",
                    f"Source: {link_source} | Email: {inv.email or '—'}",
                )

                return render(
                    request,
                    "synopsis/invite_thanks.html",
                    {
                        "member": member,
                        "project": inv.project,
                        "accepted": inv.accepted,
                    },
                )

        return render(
            request,
            "synopsis/advisory_participation_confirm.html",
            {
                "project": inv.project,
                "invitation": inv,
                "member": member,
                "form": form,
            },
        )

    decline_form = ParticipationDeclineForm(request.POST or None)
    if request.method != "POST":
        _log_project_change(
            inv.project,
            request.user,
            "Opened advisory invite link",
            f"Source: {link_source} | Choice: no | Email: {inv.email or '—'} | Showing decline form",
        )

    if request.method == "POST" and decline_form.is_valid():
        reason = (decline_form.cleaned_data.get("reason") or "").strip()
        if member:
            updates = {
                "response",
                "response_date",
                "participation_confirmed",
                "participation_confirmed_at",
                "participation_statement",
            }
            member.response = "N"
            member.response_date = timezone.localdate()
            member.participation_confirmed = False
            member.participation_confirmed_at = None
            member.participation_statement = reason
            member.save(update_fields=list(updates))

        inv.accepted = False
        inv.responded_at = timezone.now()
        inv.save(update_fields=["accepted", "responded_at", "member"] if member else ["accepted", "responded_at"])
        _log_project_change(
            inv.project,
            request.user,
            "Declined advisory invitation",
            f"Source: {link_source} | Email: {inv.email or '—'}",
        )

        return render(
            request,
            "synopsis/invite_thanks.html",
            {"member": member, "project": inv.project, "accepted": inv.accepted},
        )

    return render(
        request,
        "synopsis/advisory_participation_decline.html",
        {
            "project": inv.project,
            "invitation": inv,
            "member": member,
            "form": decline_form,
        },
    )


@login_required
def advisory_invite_update_due_date(request, project_id, invitation_id):
    project = get_object_or_404(Project, pk=project_id)
    inv = get_object_or_404(AdvisoryBoardInvitation, pk=invitation_id, project=project)
    if request.method != "POST":
        return HttpResponseBadRequest("POST required")

    date_str = (request.POST.get("due_date") or "").strip()
    new_due_date = dt.date.fromisoformat(date_str) if date_str else None
    if new_due_date and new_due_date < minimum_allowed_deadline_date():
        messages.error(
            request,
            "Response date must be at least one day in the future.",
        )
        return redirect("synopsis:advisory_board_list", project_id=project.id)
    inv.due_date = new_due_date
    inv.save(update_fields=["due_date"])
    messages.success(request, f"Response date updated for {inv.email}.")
    return redirect("synopsis:advisory_board_list", project_id=project.id)


@login_required
def send_advisory_invites(request, project_id):
    project = get_object_or_404(Project, id=project_id)
    ids = request.POST.getlist("member_ids")
    members = AdvisoryBoardMember.objects.filter(project=project, id__in=ids)

    for m in members:
        inv = AdvisoryBoardInvitation.objects.create(
            project=project,
            member=m,
            email=m.email,
            invited_by=request.user,
            due_date=m.response_date,
        )
        yes_url = request.build_absolute_uri(
            reverse("synopsis:advisory_invite_reply", args=[inv.token, "yes"])
        )
        no_url = request.build_absolute_uri(
            reverse("synopsis:advisory_invite_reply", args=[inv.token, "no"])
        )
        subject = email_subject("invite", project, m.response_date)
        text, html = _build_advisory_invitation_email(
            project=project,
            recipient_name="advisory board member",
            due_date=m.response_date,
            yes_url=yes_url,
            no_url=no_url,
        )

        msg = EmailMultiAlternatives(
            subject,
            text,
            to=[m.email],
            reply_to=reply_to_list(getattr(request.user, "email", None)),
        )
        msg.attach_alternative(html, "text/html")
        msg.send()

        m.invite_sent = True
        m.invite_sent_at = timezone.now()
        m.save(update_fields=["invite_sent", "invite_sent_at"])

    messages.success(request, f"Sent {members.count()} invite(s).")
    return redirect("synopsis:advisory_board_list", project_id=project.id)


@login_required
def advisory_send_invites_bulk(request, project_id):
    """Compose and send invitations to all members with an email."""
    project = get_object_or_404(Project, id=project_id)

    action_list = getattr(project, "action_list", None)
    action_document_available = bool(action_list and getattr(action_list, "document", None))
    collaborative_available = bool(_onlyoffice_enabled() and action_document_available)

    if request.method == "POST":
        form = AdvisoryBulkInviteForm(request.POST, project=project)
    else:
        form = AdvisoryBulkInviteForm(
            initial={"due_date": _default_invite_due_date()},
            project=project,
        )

    if not action_document_available:
        form.fields["include_action_list"].disabled = True
        form.fields["include_action_list"].help_text = "Upload an action list document to include it here."
        form.fields["include_action_list"].initial = False
    else:
        form.fields["include_action_list"].help_text = "Adds a link to the latest action list file."

    if not collaborative_available:
        form.fields["include_collaborative_link"].disabled = True
        form.fields["include_collaborative_link"].help_text = "Enable the collaborative editor and upload the action list to share this link."
        form.fields["include_collaborative_link"].initial = False
    else:
        form.fields["include_collaborative_link"].help_text = "Shares the live collaborative editor for the action list."

    if request.method == "POST" and form.is_valid():
        members = (
            AdvisoryBoardMember.objects.filter(project=project, invite_sent=False)
            .exclude(email__isnull=True)
            .exclude(email__exact="")
        )
        if not members:
            messages.info(
                request,
                "All current members have already received an invitation. Add new members before using this bulk send.",
            )
            return redirect("synopsis:advisory_board_list", project_id=project.id)

        standard_message = form.cleaned_data.get("standard_message") or ""
        message_body = form.cleaned_data.get("message") or ""
        bulk_due_date = form.cleaned_data.get("due_date")
        if "standard_message" in request.POST:
            _update_project_advisory_invitation_message(
                project, standard_message, request.user
            )

        include_action_list = action_document_available and form.cleaned_data.get("include_action_list")
        include_collaborative_link = (
            collaborative_available and form.cleaned_data.get("include_collaborative_link")
        )

        action_url = (
            request.build_absolute_uri(action_list.document.url)
            if include_action_list and action_list and getattr(action_list, "document", None)
            else ""
        )

        sent = 0
        for member in members:
            due_date = _resolve_invite_due_date(bulk_due_date, member=member)
            inv = AdvisoryBoardInvitation.objects.create(
                project=project,
                member=member,
                email=member.email,
                invited_by=request.user,
                due_date=due_date,
            )

            yes_url = request.build_absolute_uri(
                reverse("synopsis:advisory_invite_reply", args=[inv.token, "yes"])
            )
            no_url = request.build_absolute_uri(
                reverse("synopsis:advisory_invite_reply", args=[inv.token, "no"])
            )
            subject = email_subject("invite", project, due_date)
            attachment_lines = []
            if action_url:
                attachment_lines.append(("Action list", action_url))

            collab_url = ""
            if include_collaborative_link:
                collab_url = _ensure_collaborative_invite_link(
                    request,
                    project,
                    CollaborativeSession.DOCUMENT_ACTION_LIST,
                    inv,
                    member=member,
                )
                if collab_url:
                    attachment_lines.append(("Collaborative editor", collab_url))

            text, html = _build_advisory_invitation_email(
                project=project,
                recipient_name="advisory board member",
                due_date=due_date,
                yes_url=yes_url,
                no_url=no_url,
                standard_message=standard_message,
                additional_message=message_body,
                attachment_lines=attachment_lines,
            )

            msg = EmailMultiAlternatives(
                subject,
                text,
                to=[member.email],
                reply_to=reply_to_list(getattr(request.user, "email", None)),
            )
            msg.attach_alternative(html, "text/html")
            inviter_email = getattr(request.user, "email", None)
            if inviter_email:
                msg.extra_headers = msg.extra_headers or {}
                msg.extra_headers["List-Unsubscribe"] = f"<mailto:{inviter_email}>"
            msg.send()

            member.invite_sent = True
            member.invite_sent_at = timezone.now()
            update_fields = {"invite_sent", "invite_sent_at"}
            if member.response_date != due_date:
                member.response_date = due_date
                member.reminder_sent = False
                member.reminder_sent_at = None
                update_fields.update({"response_date", "reminder_sent", "reminder_sent_at"})
            member.save(update_fields=list(update_fields))
            sent += 1

        messages.success(request, f"Sent {sent} invite(s).")
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    return render(
        request,
        "synopsis/advisory_invite_compose_all.html",
        {
            "project": project,
            "form": form,
            "action_list_available": action_document_available,
            "collaborative_available": collaborative_available,
            "default_invitation_message": default_advisory_invitation_message(),
            "preview_recipient_name": _invite_preview_recipient_name(project),
            "preview_is_bulk": True,
        },
    )


@login_required
def send_protocol(request, project_id):
    project = get_object_or_404(Project, id=project_id)
    if not hasattr(project, "protocol"):
        messages.error(request, "No protocol uploaded for this project.")
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    ids = request.POST.getlist("member_ids")
    members = AdvisoryBoardMember.objects.filter(project=project, id__in=ids)

    proto_url = request.build_absolute_uri(project.protocol.document.url)
    proto_label = _current_revision_label(project.protocol)
    label_snippet = f" ({proto_label})" if proto_label else ""
    subject = email_subject("protocol_review", project)
    for m in members:
        text = (
            "Dear advisory board member,\n\n"
            f"Please review the protocol{label_snippet} for '{project.title}':\n{proto_url}\n\n"
            f"Deadline for protocol feedback: "
            f"{_format_deadline(m.feedback_on_protocol_deadline)}\n"
        )
        msg = EmailMultiAlternatives(
            subject,
            text,
            to=[m.email],
            reply_to=reply_to_list(getattr(request.user, "email", None)),
        )
        msg.send()
        m.sent_protocol_at = timezone.now()
        m.protocol_reminder_sent = False
        m.protocol_reminder_sent_at = None
        m.save(
            update_fields=[
                "sent_protocol_at",
                "protocol_reminder_sent",
                "protocol_reminder_sent_at",
            ]
        )

    messages.success(request, f"Sent protocol to {members.count()} member(s).")
    return redirect("synopsis:advisory_board_list", project_id=project.id)


@login_required
def advisory_send_protocol_bulk(request, project_id):
    project = get_object_or_404(Project, id=project_id)
    if not hasattr(project, "protocol"):
        messages.error(request, "No protocol uploaded for this project.")
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    members = (
        AdvisoryBoardMember.objects.filter(
            project=project,
            response="Y",
            participation_confirmed=True,
        )
        .exclude(email__isnull=True)
        .exclude(email__exact="")
    )
    if not members:
        messages.info(
            request,
            "No eligible members found. Only members who accepted and confirmed participation can receive the protocol.",
        )
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    proto_url = request.build_absolute_uri(project.protocol.document.url)
    proto_label = _current_revision_label(project.protocol)
    label_snippet = f" ({proto_label})" if proto_label else ""
    subject = email_subject("protocol_review", project)

    sent = 0
    for m in members:
        text = (
            "Dear advisory board member,\n\n"
            f"Please review the protocol{label_snippet} for '{project.title}':\n{proto_url}\n\n"
            f"Deadline for protocol feedback: "
            f"{_format_deadline(m.feedback_on_protocol_deadline)}\n"
        )
        msg = EmailMultiAlternatives(
            subject,
            text,
            to=[m.email],
            reply_to=reply_to_list(getattr(request.user, "email", None)),
        )
        msg.send()
        m.sent_protocol_at = timezone.now()
        m.protocol_reminder_sent = False
        m.protocol_reminder_sent_at = None
        m.save(
            update_fields=[
                "sent_protocol_at",
                "protocol_reminder_sent",
                "protocol_reminder_sent_at",
            ]
        )
        sent += 1

    messages.success(request, f"Sent protocol to {sent} member(s).")
    return redirect("synopsis:advisory_board_list", project_id=project.id)


@login_required
def advisory_send_protocol_member(request, project_id, member_id):
    project = get_object_or_404(Project, id=project_id)
    if not hasattr(project, "protocol"):
        messages.error(request, "No protocol uploaded for this project.")
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    m = get_object_or_404(AdvisoryBoardMember, id=member_id, project=project)
    if not m.email:
        messages.error(request, "This member has no email.")
        return redirect("synopsis:advisory_board_list", project_id=project.id)
    if m.response != "Y" or not m.participation_confirmed:
        messages.error(
            request,
            "This member has not accepted the invitation or has declined participation.",
        )
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    proto_url = request.build_absolute_uri(project.protocol.document.url)
    proto_label = _current_revision_label(project.protocol)
    label_snippet = f" ({proto_label})" if proto_label else ""
    subject = email_subject("protocol_review", project)
    deadline_text = _format_deadline(m.feedback_on_protocol_deadline)
    fb = _create_protocol_feedback(project, member=m, email=m.email)
    feedback_url = request.build_absolute_uri(
        reverse("synopsis:protocol_feedback", args=[str(fb.token)])
    )

    text = (
        f"Dear {advisory_member_display_name(m)},\n\n"
        f"Please review the protocol{label_snippet} for '{project.title}':\n{proto_url}\n\n"
        f"Deadline for protocol feedback: {deadline_text}\n"
        f"Provide feedback: {feedback_url}\n"
    )

    collaborative_url = ""
    protocol_closed = bool(getattr(project.protocol, "feedback_closed_at", None))
    if (
        _onlyoffice_enabled()
        and _document_requires_file(project.protocol)
        and not protocol_closed
    ):
        collaborative_url = _ensure_collaborative_invite_link(
            request,
            project,
            CollaborativeSession.DOCUMENT_PROTOCOL,
            None,
            member=m,
            feedback=fb,
        )
        if collaborative_url:
            text += f"Collaborative editor: {collaborative_url}\n"

    msg = EmailMultiAlternatives(
        subject,
        text,
        to=[m.email],
        reply_to=reply_to_list(getattr(request.user, "email", None)),
    )
    html = (
        f"<p>Dear {html_lib.escape(advisory_member_display_name(m))},</p>"
        f"<p>Please review the protocol{label_snippet} for '<strong>{project.title}</strong>': "
        f"<a href='{proto_url}'>View document</a></p>"
        f"<p>Deadline for protocol feedback: {deadline_text}</p>"
        f"<p><a href='{feedback_url}'>Provide feedback</a></p>"
    )
    if collaborative_url:
        html += (
            "<p><strong>Collaborative editor:</strong> "
            f"<a href='{collaborative_url}'>Open live editor</a></p>"
        )
    msg.attach_alternative(html, "text/html")
    msg.send()

    m.sent_protocol_at = timezone.now()
    m.protocol_reminder_sent = False
    m.protocol_reminder_sent_at = None
    m.save(
        update_fields=[
            "sent_protocol_at",
            "protocol_reminder_sent",
            "protocol_reminder_sent_at",
        ]
    )

    messages.success(request, f"Sent protocol to {m.email}.")
    return redirect("synopsis:advisory_board_list", project_id=project.id)


@login_required
def advisory_send_protocol_compose_all(request, project_id):
    project = get_object_or_404(Project, id=project_id)
    proto = getattr(project, "protocol", None)
    if not proto:
        messages.error(request, "No protocol configured for this project.")
        return redirect("synopsis:advisory_board_list", project_id=project.id)
    proto_doc = getattr(proto, "document", None)
    protocol_document_available = bool(proto_doc)
    protocol_closed = bool(getattr(proto, "feedback_closed_at", None))
    collaborative_enabled = (
        _onlyoffice_enabled()
        and _document_requires_file(proto)
        and not protocol_closed
    )
    if request.method == "POST":
        form = ProtocolSendForm(
            request.POST,
            collaborative_enabled=collaborative_enabled,
            document_available=protocol_document_available,
        )
        if form.is_valid():
            members = _eligible_advisory_members(project)
            if not members:
                messages.info(
                    request,
                    "No eligible members found. Only members who accepted and confirmed participation can receive the protocol.",
                )
                return redirect("synopsis:advisory_board_list", project_id=project.id)
            standard_message = _document_review_message(
                "protocol", form.cleaned_data.get("standard_message")
            )
            message_body = form.cleaned_data.get("message") or ""
            include_collab = collaborative_enabled and form.cleaned_data.get(
                "include_collaborative_link"
            )
            include_document = protocol_document_available and form.cleaned_data.get(
                "include_protocol_document"
            )
            due_date = form.cleaned_data.get("due_date")
            proto_url = (
                request.build_absolute_uri(proto_doc.url)
                if include_document and proto_doc
                else ""
            )
            proto_text = (proto.text_version or "").strip()
            proto_label = _current_revision_label(proto)
            label_snippet = f" ({proto_label})" if proto_label else ""
            sent = 0
            for m in members:
                member_deadline = m.feedback_on_protocol_deadline
                deadline_changed = False
                resolved_deadline = _resolve_document_feedback_deadline(
                    due_date,
                    current_deadline=member_deadline,
                )
                if member_deadline != resolved_deadline:
                    member_deadline = resolved_deadline
                    m.feedback_on_protocol_deadline = resolved_deadline
                    m.protocol_reminder_sent = False
                    m.protocol_reminder_sent_at = None
                    deadline_changed = True

                fb = _create_protocol_feedback(project, member=m, email=m.email)
                feedback_url = request.build_absolute_uri(
                    reverse("synopsis:protocol_feedback", args=[str(fb.token)])
                )
                subject = email_subject("protocol_review", project)
                recipient_name = "advisory board member"
                text = f"Dear {recipient_name},\n\n"
                html = f"<p>Dear {html_lib.escape(recipient_name)},</p>"
                if standard_message:
                    text += f"{standard_message}\n\n"
                    html += _html_message_blocks(standard_message)
                if message_body:
                    text += f"{message_body}\n\n"
                    html += _html_message_blocks(message_body)
                if proto_url:
                    text += f"Protocol document{label_snippet}: {proto_url}\n\n"
                    html += (
                        "<p>Protocol document"
                        f"{label_snippet}: <a href='{proto_url}'>View document</a></p>"
                    )
                elif proto_text:
                    html += "<hr>" + proto_text
                deadline_text = _format_deadline(member_deadline)
                if deadline_text:
                    text += f"Deadline for protocol feedback: {deadline_text}\n"
                    html += f"<p>Deadline for protocol feedback: {deadline_text}</p>"
                text += f"Provide feedback: {feedback_url}\n"
                html += f"<p><a href='{feedback_url}'>Provide feedback</a></p>"
                if include_collab:
                    collaborative_url = _ensure_collaborative_invite_link(
                        request,
                        project,
                        CollaborativeSession.DOCUMENT_PROTOCOL,
                        None,
                        member=m,
                        feedback=fb,
                    )
                    if collaborative_url:
                        text += f"Collaborative editor: {collaborative_url}\n"
                        html += (
                            "<p><strong>Collaborative editor:</strong> "
                            f"<a href='{collaborative_url}'>Open live editor</a></p>"
                        )

                msg = EmailMultiAlternatives(
                    subject,
                    text,
                    to=[m.email],
                    reply_to=reply_to_list(getattr(request.user, "email", None)),
                )
                msg.attach_alternative(html, "text/html")
                inviter_email = getattr(request.user, "email", None)
                if inviter_email:
                    msg.extra_headers = msg.extra_headers or {}
                    msg.extra_headers["List-Unsubscribe"] = f"<mailto:{inviter_email}>"
                msg.send()
                m.sent_protocol_at = timezone.now()
                m.protocol_reminder_sent = False
                m.protocol_reminder_sent_at = None
                update_fields = [
                    "sent_protocol_at",
                    "protocol_reminder_sent",
                    "protocol_reminder_sent_at",
                ]
                if deadline_changed:
                    update_fields.append("feedback_on_protocol_deadline")
                m.save(update_fields=update_fields)
                sent += 1
            messages.success(request, f"Sent protocol to {sent} member(s).")
            return redirect("synopsis:advisory_board_list", project_id=project.id)
    else:
        form = ProtocolSendForm(
            initial={"due_date": _default_document_feedback_due_date()},
            collaborative_enabled=collaborative_enabled,
            document_available=protocol_document_available,
        )
    return render(
        request,
        "synopsis/protocol_send_compose.html",
        {
            "project": project,
            "form": form,
            "scope": "all",
            "collaborative_available": collaborative_enabled,
            "preview_recipient_name": _document_preview_recipient_name(project),
            "preview_is_bulk": True,
        },
    )


@login_required
def advisory_send_protocol_compose_member(request, project_id, member_id):
    project = get_object_or_404(Project, id=project_id)
    proto = getattr(project, "protocol", None)
    if not proto:
        messages.error(request, "No protocol configured for this project.")
        return redirect("synopsis:advisory_board_list", project_id=project.id)
    proto_doc = getattr(proto, "document", None)
    protocol_document_available = bool(proto_doc)
    m = get_object_or_404(AdvisoryBoardMember, id=member_id, project=project)
    if m.response != "Y" or not m.participation_confirmed:
        messages.error(
            request,
            "This member has not accepted the invitation or has declined participation.",
        )
        return redirect("synopsis:advisory_board_list", project_id=project.id)
    protocol_closed = bool(getattr(proto, "feedback_closed_at", None))
    collaborative_enabled = (
        _onlyoffice_enabled()
        and _document_requires_file(proto)
        and not protocol_closed
    )
    if request.method == "POST":
        form = ProtocolSendForm(
            request.POST,
            collaborative_enabled=collaborative_enabled,
            document_available=protocol_document_available,
        )
        if form.is_valid():
            standard_message = _document_review_message(
                "protocol", form.cleaned_data.get("standard_message")
            )
            message_body = form.cleaned_data.get("message") or ""
            include_document = protocol_document_available and form.cleaned_data.get(
                "include_protocol_document"
            )
            member_deadline = m.feedback_on_protocol_deadline
            deadline_changed = False
            resolved_deadline = _resolve_document_feedback_deadline(
                form.cleaned_data.get("due_date"),
                current_deadline=member_deadline,
            )
            if member_deadline != resolved_deadline:
                member_deadline = resolved_deadline
                m.feedback_on_protocol_deadline = resolved_deadline
                m.protocol_reminder_sent = False
                m.protocol_reminder_sent_at = None
                deadline_changed = True
            fb = _create_protocol_feedback(project, member=m, email=m.email)
            feedback_url = request.build_absolute_uri(
                reverse("synopsis:protocol_feedback", args=[str(fb.token)])
            )
            subject = email_subject("protocol_review", project)
            recipient_name = advisory_member_display_name(m)
            text = f"Dear {recipient_name},\n\n"
            html = f"<p>Dear {html_lib.escape(recipient_name)},</p>"
            if standard_message:
                text += f"{standard_message}\n\n"
                html += _html_message_blocks(standard_message)
            if message_body:
                text += f"{message_body}\n\n"
                html += _html_message_blocks(message_body)
            proto_url = (
                request.build_absolute_uri(proto_doc.url)
                if include_document and proto_doc
                else ""
            )
            proto_text = (proto.text_version or "").strip()
            proto_label = _current_revision_label(proto)
            label_snippet = f" ({proto_label})" if proto_label else ""
            if proto_url:
                text += f"Protocol document{label_snippet}: {proto_url}\n\n"
                html += (
                    "<p>Protocol document"
                    f"{label_snippet}: <a href='{proto_url}'>View document</a></p>"
                )
            elif proto_text:
                html += "<hr>" + proto_text
            deadline_text = _format_deadline(member_deadline)
            if deadline_text:
                text += f"Deadline for protocol feedback: {deadline_text}\n"
                html += f"<p>Deadline for protocol feedback: {deadline_text}</p>"
            text += f"Provide feedback: {feedback_url}\n"
            html += f"<p><a href='{feedback_url}'>Provide feedback</a></p>"
            if collaborative_enabled and form.cleaned_data.get(
                "include_collaborative_link"
            ):
                collaborative_url = _ensure_collaborative_invite_link(
                    request,
                    project,
                    CollaborativeSession.DOCUMENT_PROTOCOL,
                    None,
                    member=m,
                    feedback=fb,
                )
                if collaborative_url:
                    text += f"Collaborative editor: {collaborative_url}\n"
                    html += (
                        "<p><strong>Collaborative editor:</strong> "
                        f"<a href='{collaborative_url}'>Open live editor</a></p>"
                    )

            msg = EmailMultiAlternatives(
                subject,
                text,
                to=[m.email],
                reply_to=reply_to_list(getattr(request.user, "email", None)),
            )
            msg.attach_alternative(html, "text/html")
            msg.send()

            m.sent_protocol_at = timezone.now()
            m.protocol_reminder_sent = False
            m.protocol_reminder_sent_at = None
            update_fields = [
                "sent_protocol_at",
                "protocol_reminder_sent",
                "protocol_reminder_sent_at",
            ]
            if deadline_changed:
                update_fields.append("feedback_on_protocol_deadline")
            m.save(update_fields=update_fields)
            messages.success(request, f"Sent protocol to {m.email}.")
            return redirect("synopsis:advisory_board_list", project_id=project.id)
    else:
        deadline_initial = None
        if m.feedback_on_protocol_deadline:
            local_deadline = timezone.localtime(m.feedback_on_protocol_deadline)
            deadline_initial = local_deadline.date()
        else:
            deadline_initial = _default_document_feedback_due_date()
        form = ProtocolSendForm(
            initial={
                "due_date": deadline_initial,
            },
            collaborative_enabled=collaborative_enabled,
            document_available=protocol_document_available,
        )
    return render(
        request,
        "synopsis/protocol_send_compose.html",
        {
            "project": project,
            "form": form,
            "scope": "member",
            "member": m,
            "collaborative_available": collaborative_enabled,
            "preview_recipient_name": advisory_member_display_name(m),
            "preview_is_bulk": False,
        },
    )


@login_required
def advisory_send_action_list_compose_all(request, project_id):
    project = get_object_or_404(Project, id=project_id)
    action_list = getattr(project, "action_list", None)
    if not action_list:
        messages.error(request, "No action list configured for this project.")
        return redirect("synopsis:advisory_board_list", project_id=project.id)
    action_document_available = bool(getattr(action_list, "document", None))
    action_closed = bool(getattr(action_list, "feedback_closed_at", None))
    collaborative_enabled = (
        _onlyoffice_enabled()
        and _document_requires_file(action_list)
        and not action_closed
    )
    if request.method == "POST":
        form = ActionListSendForm(
            request.POST,
            collaborative_enabled=collaborative_enabled,
            document_available=action_document_available,
        )
        if form.is_valid():
            members = _eligible_advisory_members(project)
            if not members:
                messages.info(
                    request,
                    "No eligible members found. Only members who accepted and confirmed participation can receive the action list.",
                )
                return redirect("synopsis:advisory_board_list", project_id=project.id)
            standard_message = _document_review_message(
                "action_list", form.cleaned_data.get("standard_message")
            )
            message_body = form.cleaned_data.get("message") or ""
            include_document = action_document_available and form.cleaned_data.get(
                "include_action_list_document"
            )
            include_collab = collaborative_enabled and form.cleaned_data.get(
                "include_collaborative_link"
            )
            doc_url = (
                request.build_absolute_uri(action_list.document.url)
                if include_document and action_document_available
                else ""
            )
            text_version = (action_list.text_version or "").strip()
            action_label = _current_revision_label(action_list)
            label_snippet = f" ({action_label})" if action_label else ""
            sent = 0
            for m in members:
                member_deadline = m.feedback_on_action_list_deadline
                deadline_changed = False
                resolved_deadline = _resolve_document_feedback_deadline(
                    form.cleaned_data.get("due_date"),
                    current_deadline=member_deadline,
                )
                if member_deadline != resolved_deadline:
                    member_deadline = resolved_deadline
                    m.feedback_on_action_list_deadline = resolved_deadline
                    m.action_list_reminder_sent = False
                    m.action_list_reminder_sent_at = None
                    deadline_changed = True
                fb = _create_action_list_feedback(project, member=m, email=m.email)
                feedback_url = request.build_absolute_uri(
                    reverse("synopsis:action_list_feedback", args=[str(fb.token)])
                )
                subject = email_subject("action_list_review", project)
                recipient_name = "advisory board member"
                text = f"Dear {recipient_name},\n\n"
                html = f"<p>Dear {html_lib.escape(recipient_name)},</p>"
                if standard_message:
                    text += f"{standard_message}\n\n"
                    html += _html_message_blocks(standard_message)
                if message_body:
                    text += f"{message_body}\n\n"
                    html += _html_message_blocks(message_body)
                if doc_url:
                    text += f"Action list document{label_snippet}: {doc_url}\n\n"
                    html += (
                        "<p>Action list document"
                        f"{label_snippet}: <a href='{doc_url}'>View document</a></p>"
                    )
                elif text_version:
                    html += "<hr>" + text_version
                deadline_text = _format_deadline(member_deadline)
                if deadline_text:
                    text += f"Deadline for action list feedback: {deadline_text}\n"
                    html += f"<p>Deadline for action list feedback: {deadline_text}</p>"
                text += f"Provide feedback: {feedback_url}\n"
                html += f"<p><a href='{feedback_url}'>Provide feedback</a></p>"
                if include_collab:
                    collaborative_url = _ensure_collaborative_invite_link(
                        request,
                        project,
                        CollaborativeSession.DOCUMENT_ACTION_LIST,
                        None,
                        member=m,
                        feedback=fb,
                    )
                    if collaborative_url:
                        text += f"Collaborative editor: {collaborative_url}\n"
                        html += (
                            "<p><strong>Collaborative editor:</strong> "
                            f"<a href='{collaborative_url}'>Open live editor</a></p>"
                        )

                msg = EmailMultiAlternatives(
                    subject,
                    text,
                    to=[m.email],
                    reply_to=reply_to_list(getattr(request.user, "email", None)),
                )
                msg.attach_alternative(html, "text/html")
                inviter_email = getattr(request.user, "email", None)
                if inviter_email:
                    msg.extra_headers = msg.extra_headers or {}
                    msg.extra_headers["List-Unsubscribe"] = f"<mailto:{inviter_email}>"
                msg.send()
                m.sent_action_list_at = timezone.now()
                m.action_list_reminder_sent = False
                m.action_list_reminder_sent_at = None
                update_fields = [
                    "sent_action_list_at",
                    "action_list_reminder_sent",
                    "action_list_reminder_sent_at",
                ]
                if deadline_changed:
                    update_fields.append("feedback_on_action_list_deadline")
                m.save(update_fields=update_fields)
                sent += 1
            messages.success(request, f"Sent action list to {sent} member(s).")
            return redirect("synopsis:advisory_board_list", project_id=project.id)
    else:
        form = ActionListSendForm(
            initial={"due_date": _default_document_feedback_due_date()},
            collaborative_enabled=collaborative_enabled,
            document_available=action_document_available,
        )
    return render(
        request,
        "synopsis/action_list_send_compose.html",
        {
            "project": project,
            "form": form,
            "scope": "all",
            "member": None,
            "collaborative_available": collaborative_enabled,
            "preview_recipient_name": _document_preview_recipient_name(project),
            "preview_is_bulk": True,
        },
    )


@login_required
def advisory_send_action_list_compose_member(request, project_id, member_id):
    project = get_object_or_404(Project, id=project_id)
    action_list = getattr(project, "action_list", None)
    if not action_list:
        messages.error(request, "No action list configured for this project.")
        return redirect("synopsis:advisory_board_list", project_id=project.id)
    action_document_available = bool(getattr(action_list, "document", None))
    member = get_object_or_404(AdvisoryBoardMember, id=member_id, project=project)
    if not member.email:
        messages.error(request, "This member has no email.")
        return redirect("synopsis:advisory_board_list", project_id=project.id)
    if member.response != "Y" or not member.participation_confirmed:
        messages.error(
            request,
            "This member has not accepted the invitation or has declined participation.",
        )
        return redirect("synopsis:advisory_board_list", project_id=project.id)
    action_closed = bool(getattr(action_list, "feedback_closed_at", None))
    collaborative_enabled = (
        _onlyoffice_enabled()
        and _document_requires_file(action_list)
        and not action_closed
    )
    if request.method == "POST":
        form = ActionListSendForm(
            request.POST,
            collaborative_enabled=collaborative_enabled,
            document_available=action_document_available,
        )
        if form.is_valid():
            standard_message = _document_review_message(
                "action_list", form.cleaned_data.get("standard_message")
            )
            message_body = form.cleaned_data.get("message") or ""
            include_document = action_document_available and form.cleaned_data.get(
                "include_action_list_document"
            )
            include_collab = collaborative_enabled and form.cleaned_data.get(
                "include_collaborative_link"
            )
            member_deadline = member.feedback_on_action_list_deadline
            deadline_changed = False
            resolved_deadline = _resolve_document_feedback_deadline(
                form.cleaned_data.get("due_date"),
                current_deadline=member_deadline,
            )
            if member_deadline != resolved_deadline:
                member_deadline = resolved_deadline
                member.feedback_on_action_list_deadline = resolved_deadline
                member.action_list_reminder_sent = False
                member.action_list_reminder_sent_at = None
                deadline_changed = True
            fb = _create_action_list_feedback(
                project, member=member, email=member.email
            )
            feedback_url = request.build_absolute_uri(
                reverse("synopsis:action_list_feedback", args=[str(fb.token)])
            )
            subject = email_subject("action_list_review", project)
            recipient_name = advisory_member_display_name(member)
            text = f"Dear {recipient_name},\n\n"
            html = f"<p>Dear {html_lib.escape(recipient_name)},</p>"
            if standard_message:
                text += f"{standard_message}\n\n"
                html += _html_message_blocks(standard_message)
            if message_body:
                text += f"{message_body}\n\n"
                html += _html_message_blocks(message_body)
            doc_url = (
                request.build_absolute_uri(action_list.document.url)
                if include_document and action_document_available
                else ""
            )
            text_version = (action_list.text_version or "").strip()
            action_label = _current_revision_label(action_list)
            label_snippet = f" ({action_label})" if action_label else ""
            if doc_url:
                text += f"Action list document{label_snippet}: {doc_url}\n\n"
                html += (
                    "<p>Action list document"
                    f"{label_snippet}: <a href='{doc_url}'>View document</a></p>"
                )
            elif text_version:
                html += "<hr>" + text_version
            deadline_text = _format_deadline(member_deadline)
            if deadline_text:
                text += f"Deadline for action list feedback: {deadline_text}\n"
                html += f"<p>Deadline for action list feedback: {deadline_text}</p>"
            text += f"Provide feedback: {feedback_url}\n"
            html += f"<p><a href='{feedback_url}'>Provide feedback</a></p>"
            if include_collab:
                collaborative_url = _ensure_collaborative_invite_link(
                    request,
                    project,
                    CollaborativeSession.DOCUMENT_ACTION_LIST,
                    None,
                    member=member,
                    feedback=fb,
                )
                if collaborative_url:
                    text += f"Collaborative editor: {collaborative_url}\n"
                    html += (
                        "<p><strong>Collaborative editor:</strong> "
                        f"<a href='{collaborative_url}'>Open live editor</a></p>"
                    )

            msg = EmailMultiAlternatives(
                subject,
                text,
                to=[member.email],
                reply_to=reply_to_list(getattr(request.user, "email", None)),
            )
            msg.attach_alternative(html, "text/html")
            msg.send()

            member.sent_action_list_at = timezone.now()
            member.action_list_reminder_sent = False
            member.action_list_reminder_sent_at = None
            update_fields = [
                "sent_action_list_at",
                "action_list_reminder_sent",
                "action_list_reminder_sent_at",
            ]
            if deadline_changed:
                update_fields.append("feedback_on_action_list_deadline")
            member.save(update_fields=update_fields)
            messages.success(request, f"Sent action list to {member.email}.")
            return redirect("synopsis:advisory_board_list", project_id=project.id)
    else:
        deadline_initial = None
        if member.feedback_on_action_list_deadline:
            local_deadline = timezone.localtime(member.feedback_on_action_list_deadline)
            deadline_initial = local_deadline.date()
        else:
            deadline_initial = _default_document_feedback_due_date()
        form = ActionListSendForm(
            initial={
                "due_date": deadline_initial,
            },
            collaborative_enabled=collaborative_enabled,
            document_available=action_document_available,
        )
    return render(
        request,
        "synopsis/action_list_send_compose.html",
        {
            "project": project,
            "form": form,
            "scope": "member",
            "member": member,
            "collaborative_available": collaborative_enabled,
            "preview_recipient_name": advisory_member_display_name(member),
            "preview_is_bulk": False,
        },
    )


def _synopsis_export_attachment(project):
    payload = _generate_synopsis_docx(project)
    filename = slugify(f"{project.title}-synopsis").replace(" ", "-") + ".docx"
    return (
        filename,
        payload,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _synopsis_send_attachment(form, project):
    uploaded = form.cleaned_data.get("synopsis_document")
    if uploaded:
        uploaded.seek(0)
        return (
            os.path.basename(uploaded.name),
            uploaded.read(),
            uploaded.content_type or "application/octet-stream",
        )
    return _synopsis_export_attachment(project)


def _send_synopsis_review_email(
    request,
    project,
    member,
    *,
    standard_message="",
    message_body="",
    deadline=None,
    attachment_filename=None,
    attachment_payload=None,
    attachment_content_type=None,
    feedback_url="",
    recipient_name=None,
):
    subject = email_subject("synopsis_review", project)
    deadline_text = _format_deadline(deadline)
    recipient_name = recipient_name or advisory_member_display_name(member)
    text = f"Dear {recipient_name},\n\n"
    html = f"<p>Dear {html_lib.escape(recipient_name)},</p>"
    if standard_message:
        text += f"{standard_message}\n\n"
        html += _html_message_blocks(standard_message)
    if message_body:
        text += f"{message_body}\n\n"
        html += _html_message_blocks(message_body)
    text += f"Synopsis document for '{project.title}' is attached to this email.\n\n"
    html += (
        "<p>Synopsis document for "
        f"'<strong>{project.title}</strong>' is attached to this email.</p>"
    )
    if deadline_text:
        text += f"Deadline for synopsis feedback: {deadline_text}\n"
        html += f"<p>Deadline for synopsis feedback: {deadline_text}</p>"
    if feedback_url:
        text += f"Provide feedback: {feedback_url}\n"
        html += f"<p><a href='{feedback_url}'>Provide feedback</a></p>"
    else:
        text += "\nPlease send your feedback to the synopsis author team.\n"
        html += "<p>Please send your feedback to the synopsis author team.</p>"

    msg = EmailMultiAlternatives(
        subject,
        text,
        to=[member.email],
        reply_to=reply_to_list(getattr(request.user, "email", None)),
    )
    msg.attach_alternative(html, "text/html")
    if attachment_filename and attachment_payload:
        msg.attach(
            attachment_filename,
            attachment_payload,
            attachment_content_type or "application/octet-stream",
        )
    inviter_email = getattr(request.user, "email", None)
    if inviter_email:
        msg.extra_headers = msg.extra_headers or {}
        msg.extra_headers["List-Unsubscribe"] = f"<mailto:{inviter_email}>"
    msg.send()


@login_required
def advisory_send_synopsis_compose_all(request, project_id):
    project = get_object_or_404(Project, id=project_id)
    if request.method == "POST":
        form = SynopsisSendForm(request.POST, request.FILES)
        if form.is_valid():
            members = _eligible_advisory_members(project)
            if not members:
                messages.info(
                    request,
                    "No eligible members found. Only members who accepted and confirmed participation can receive the synopsis.",
                )
                return redirect("synopsis:advisory_board_list", project_id=project.id)
            try:
                (
                    attachment_filename,
                    attachment_payload,
                    attachment_content_type,
                ) = _synopsis_send_attachment(form, project)
            except ImportError as exc:
                messages.error(request, str(exc))
                return redirect("synopsis:advisory_board_list", project_id=project.id)

            standard_message = _document_review_message(
                "synopsis", form.cleaned_data.get("standard_message")
            )
            message_body = form.cleaned_data.get("message") or ""
            sent = 0
            for member in members:
                member_deadline = member.feedback_on_synopsis_deadline
                deadline_changed = False
                resolved_deadline = _resolve_document_feedback_deadline(
                    form.cleaned_data.get("due_date"),
                    current_deadline=member_deadline,
                )
                if member_deadline != resolved_deadline:
                    member_deadline = resolved_deadline
                    member.feedback_on_synopsis_deadline = resolved_deadline
                    member.synopsis_reminder_sent = False
                    member.synopsis_reminder_sent_at = None
                    deadline_changed = True
                fb = _create_synopsis_feedback(project, member=member, email=member.email)
                feedback_url = request.build_absolute_uri(
                    reverse("synopsis:synopsis_feedback", args=[str(fb.token)])
                )

                _send_synopsis_review_email(
                    request,
                    project,
                    member,
                    standard_message=standard_message,
                    message_body=message_body,
                    deadline=member_deadline,
                    attachment_filename=attachment_filename,
                    attachment_payload=attachment_payload,
                    attachment_content_type=attachment_content_type,
                    feedback_url=feedback_url,
                    recipient_name="advisory board member",
                )
                member.sent_synopsis_at = timezone.now()
                member.synopsis_reminder_sent = False
                member.synopsis_reminder_sent_at = None
                update_fields = [
                    "sent_synopsis_at",
                    "synopsis_reminder_sent",
                    "synopsis_reminder_sent_at",
                ]
                if deadline_changed:
                    update_fields.append("feedback_on_synopsis_deadline")
                member.save(update_fields=update_fields)
                sent += 1
            messages.success(request, f"Sent synopsis to {sent} member(s).")
            return redirect("synopsis:advisory_board_list", project_id=project.id)
    else:
        form = SynopsisSendForm(
            initial={"due_date": _default_document_feedback_due_date()}
        )
    return render(
        request,
        "synopsis/synopsis_send_compose.html",
        {
            "project": project,
            "form": form,
            "scope": "all",
            "member": None,
            "preview_recipient_name": _document_preview_recipient_name(project),
            "preview_is_bulk": True,
        },
    )


@login_required
def advisory_send_synopsis_compose_member(request, project_id, member_id):
    project = get_object_or_404(Project, id=project_id)
    member = get_object_or_404(AdvisoryBoardMember, id=member_id, project=project)
    if not member.email:
        messages.error(request, "This member has no email.")
        return redirect("synopsis:advisory_board_list", project_id=project.id)
    if member.response != "Y" or not member.participation_confirmed:
        messages.error(
            request,
            "This member has not accepted the invitation or has declined participation.",
        )
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    if request.method == "POST":
        form = SynopsisSendForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                (
                    attachment_filename,
                    attachment_payload,
                    attachment_content_type,
                ) = _synopsis_send_attachment(form, project)
            except ImportError as exc:
                messages.error(request, str(exc))
                return redirect("synopsis:advisory_board_list", project_id=project.id)

            member_deadline = member.feedback_on_synopsis_deadline
            deadline_changed = False
            resolved_deadline = _resolve_document_feedback_deadline(
                form.cleaned_data.get("due_date"),
                current_deadline=member_deadline,
            )
            if member_deadline != resolved_deadline:
                member_deadline = resolved_deadline
                member.feedback_on_synopsis_deadline = resolved_deadline
                member.synopsis_reminder_sent = False
                member.synopsis_reminder_sent_at = None
                deadline_changed = True
            fb = _create_synopsis_feedback(project, member=member, email=member.email)
            feedback_url = request.build_absolute_uri(
                reverse("synopsis:synopsis_feedback", args=[str(fb.token)])
            )

            _send_synopsis_review_email(
                request,
                project,
                member,
                standard_message=_document_review_message(
                    "synopsis", form.cleaned_data.get("standard_message")
                ),
                message_body=form.cleaned_data.get("message") or "",
                deadline=member_deadline,
                attachment_filename=attachment_filename,
                attachment_payload=attachment_payload,
                attachment_content_type=attachment_content_type,
                feedback_url=feedback_url,
            )
            member.sent_synopsis_at = timezone.now()
            member.synopsis_reminder_sent = False
            member.synopsis_reminder_sent_at = None
            update_fields = [
                "sent_synopsis_at",
                "synopsis_reminder_sent",
                "synopsis_reminder_sent_at",
            ]
            if deadline_changed:
                update_fields.append("feedback_on_synopsis_deadline")
            member.save(update_fields=update_fields)
            messages.success(request, f"Sent synopsis to {member.email}.")
            return redirect("synopsis:advisory_board_list", project_id=project.id)
    else:
        if member.feedback_on_synopsis_deadline:
            local_deadline = timezone.localtime(member.feedback_on_synopsis_deadline)
            deadline_initial = local_deadline.date()
        else:
            deadline_initial = _default_document_feedback_due_date()
        form = SynopsisSendForm(initial={"due_date": deadline_initial})
    return render(
        request,
        "synopsis/synopsis_send_compose.html",
        {
            "project": project,
            "form": form,
            "scope": "member",
            "member": member,
            "preview_recipient_name": advisory_member_display_name(member),
            "preview_is_bulk": False,
        },
    )


@login_required
def advisory_member_custom_data(request, project_id, member_id):
    project = get_object_or_404(Project, pk=project_id)
    member = get_object_or_404(AdvisoryBoardMember, pk=member_id, project=project)

    if not _user_can_edit_project(request.user, project):
        messages.error(request, "You do not have permission to update this member.")
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    status_key = _member_section_key(member)
    all_fields = list(
        AdvisoryBoardCustomField.objects.filter(project=project).order_by(
            "display_order", "name", "id"
        )
    )
    applicable_fields = [field for field in all_fields if field.applies_to(status_key)]

    focus_field_id = request.GET.get("field") or request.POST.get("focused_field")
    focused_field = None
    if focus_field_id:
        try:
            focused_field = next(
                field for field in applicable_fields if str(field.id) == str(focus_field_id)
            )
            applicable_fields = [focused_field]
        except StopIteration:
            focused_field = None

    existing_values = {
        value.field_id: value.value
        for value in AdvisoryBoardCustomFieldValue.objects.filter(
            member=member, field__project=project
        )
    }

    if request.method == "POST":
        form = AdvisoryMemberCustomDataForm(
            applicable_fields,
            status_key,
            existing_values,
            request.POST,
            form_id=f"member-form-{member.id}",
        )
        if form.is_valid():
            actor = request.user if getattr(request.user, "is_authenticated", False) else None
            for field in applicable_fields:
                cleaned_value = form.cleaned_value(field)
                field.set_value_for_member(
                    member,
                    cleaned_value,
                    changed_by=actor,
                )
            messages.success(
                request,
                f"Updated custom data for {member.first_name or member.email}.",
            )
            return redirect("synopsis:advisory_board_list", project_id=project.id)
    else:
        form = AdvisoryMemberCustomDataForm(
            applicable_fields, status_key, existing_values
        )

    form.apply_widget_configuration()

    history_map = {}
    if applicable_fields:
        history_qs = (
            AdvisoryBoardCustomFieldValueHistory.objects.filter(
                member=member, field__in=applicable_fields
            )
            .select_related("field", "changed_by")
            .order_by("-created_at", "-id")
        )
        for entry in history_qs:
            bucket = history_map.setdefault(entry.field_id, [])
            if len(bucket) >= 10:
                continue
            bucket.append(entry)

    return render(
        request,
        "synopsis/advisory_member_custom_data.html",
        {
            "project": project,
            "member": member,
            "form": form,
            "fields": applicable_fields,
            "status_key": status_key,
            "history_map": history_map,
            "focused_field": focused_field,
        },
    )


def protocol_feedback(request, token):
    fb = get_object_or_404(ProtocolFeedback, token=token)
    member = fb.member
    project = fb.project
    proto = getattr(project, "protocol", None)

    deadline = fb.feedback_deadline_at
    if member and member.feedback_on_protocol_deadline:
        deadline = member.feedback_on_protocol_deadline
        if fb.feedback_deadline_at != deadline:
            fb.feedback_deadline_at = deadline
            fb.save(update_fields=["feedback_deadline_at"])
    if member and member.response != "Y":
        if fb.feedback_deadline_at:
            fb.feedback_deadline_at = None
            fb.save(update_fields=["feedback_deadline_at"])
        deadline = None
    closure_message = None
    now = timezone.now()

    if member and member.response == "N":
        return render(
            request,
            "synopsis/protocol_feedback_thanks.html",
            {
                "project": project,
                "error": "This link is no longer available because you declined the invitation.",
            },
        )

    if proto and proto.feedback_closed_at:
        closure_message = proto.feedback_closure_message or (
            "The authors have closed feedback for this protocol."
        )
        return render(
            request,
            "synopsis/protocol_feedback_thanks.html",
            {
                "project": project,
                "feedback": fb,
                "closed_message": closure_message,
                "deadline": deadline,
                "closed": True,
            },
        )

    if deadline and now >= deadline:
        closure_message = (
            "The feedback deadline has passed (" f"{_format_deadline(deadline)})."
        )
        return render(
            request,
            "synopsis/protocol_feedback_thanks.html",
            {
                "project": project,
                "feedback": fb,
                "closed_message": closure_message,
                "deadline": deadline,
                "closed": True,
            },
        )

    if request.method == "POST":
        form = ProtocolFeedbackForm(request.POST, request.FILES)
        if form.is_valid():
            content = form.cleaned_data["content"].strip()
            uploaded_doc = form.cleaned_data["uploaded_document"]
            updates = []
            if content:
                fb.content = content
                updates.append("content")
            if uploaded_doc:
                fb.uploaded_document = uploaded_doc
                updates.append("uploaded_document")
            if updates:
                fb.submitted_at = timezone.now()
                updates.append("submitted_at")
                fb.save(update_fields=updates)

                if member:
                    member_updates = set()
                    today = timezone.localdate()
                    if member.feedback_on_protocol_received != today:
                        member.feedback_on_protocol_received = today
                        member_updates.add("feedback_on_protocol_received")
                    if member_updates:
                        member.save(update_fields=list(member_updates))

                details = []
                if uploaded_doc:
                    details.append(
                        f"Document uploaded: {fb.latest_document_label or uploaded_doc.name}"
                    )
                if content:
                    snippet = (content[:97] + "…") if len(content) > 100 else content
                    details.append(f"Comments provided: {snippet}")
                if details:
                    _log_project_change(
                        project,
                        request.user,
                        "Protocol feedback submitted",
                        " | ".join(details),
                    )
            return render(
                request,
                "synopsis/protocol_feedback_thanks.html",
                {
                    "project": project,
                    "feedback": fb,
                    "deadline": deadline,
                },
            )
    else:
        form = ProtocolFeedbackForm(initial={"content": fb.content})

    return render(
        request,
        "synopsis/protocol_feedback_form.html",
        {
            "project": project,
            "token": fb.token,
            "feedback": fb,
            "deadline": deadline,
            "form": form,
        },
    )


def action_list_feedback(request, token):
    fb = get_object_or_404(ActionListFeedback, token=token)
    member = fb.member
    project = fb.project
    action_list = fb.action_list or getattr(project, "action_list", None)

    deadline = fb.feedback_deadline_at
    if member and member.feedback_on_action_list_deadline:
        deadline = member.feedback_on_action_list_deadline
        if fb.feedback_deadline_at != deadline:
            fb.feedback_deadline_at = deadline
            fb.save(update_fields=["feedback_deadline_at"])
    if member and member.response != "Y":
        if fb.feedback_deadline_at:
            fb.feedback_deadline_at = None
            fb.save(update_fields=["feedback_deadline_at"])
        deadline = None
    closure_message = None
    now = timezone.now()

    if member and member.response == "N":
        return render(
            request,
            "synopsis/action_list_feedback_thanks.html",
            {
                "project": project,
                "error": "This link is no longer available because you declined the invitation.",
            },
        )

    if action_list and action_list.feedback_closed_at:
        closure_message = action_list.feedback_closure_message or (
            "The authors have closed feedback for this action list."
        )
        return render(
            request,
            "synopsis/action_list_feedback_thanks.html",
            {
                "project": project,
                "feedback": fb,
                "closed_message": closure_message,
                "deadline": deadline,
                "closed": True,
            },
        )

    if deadline and now >= deadline:
        closure_message = (
            "The feedback deadline has passed (" f"{_format_deadline(deadline)})."
        )
        return render(
            request,
            "synopsis/action_list_feedback_thanks.html",
            {
                "project": project,
                "feedback": fb,
                "closed_message": closure_message,
                "deadline": deadline,
                "closed": True,
            },
        )

    if request.method == "POST":
        form = ActionListFeedbackForm(request.POST, request.FILES)
        if form.is_valid():
            content = form.cleaned_data["content"].strip()
            uploaded_doc = form.cleaned_data["uploaded_document"]
            updates = []
            if content:
                fb.content = content
                updates.append("content")
            if uploaded_doc:
                fb.uploaded_document = uploaded_doc
                updates.append("uploaded_document")
            if updates:
                fb.submitted_at = timezone.now()
                updates.append("submitted_at")
                fb.save(update_fields=updates)

                if member:
                    member_updates = set()
                    if not member.feedback_on_actions_received:
                        member.feedback_on_actions_received = True
                        member_updates.add("feedback_on_actions_received")
                    today = timezone.localdate()
                    if member.feedback_on_action_list_received != today:
                        member.feedback_on_action_list_received = today
                        member_updates.add("feedback_on_action_list_received")
                    if member_updates:
                        member.save(update_fields=list(member_updates))

                details = []
                if uploaded_doc:
                    details.append(
                        f"Document uploaded: {fb.latest_document_label or uploaded_doc.name}"
                    )
                if content:
                    snippet = (content[:97] + "…") if len(content) > 100 else content
                    details.append(f"Comments provided: {snippet}")
                if details:
                    _log_project_change(
                        project,
                        request.user,
                        "Action list feedback submitted",
                        " | ".join(details),
                    )

            return render(
                request,
                "synopsis/action_list_feedback_thanks.html",
                {
                    "project": project,
                    "feedback": fb,
                    "deadline": deadline,
                },
            )
    else:
        form = ActionListFeedbackForm(initial={"content": fb.content})

    return render(
        request,
        "synopsis/action_list_feedback_form.html",
        {
            "project": project,
            "token": fb.token,
            "feedback": fb,
            "deadline": deadline,
            "action_list": action_list,
            "form": form,
        },
    )


def synopsis_feedback(request, token):
    fb = get_object_or_404(SynopsisFeedback, token=token)
    member = fb.member
    project = fb.project

    deadline = fb.feedback_deadline_at
    if member and member.feedback_on_synopsis_deadline:
        deadline = member.feedback_on_synopsis_deadline
        if fb.feedback_deadline_at != deadline:
            fb.feedback_deadline_at = deadline
            fb.save(update_fields=["feedback_deadline_at"])
    if member and member.response != "Y":
        if fb.feedback_deadline_at:
            fb.feedback_deadline_at = None
            fb.save(update_fields=["feedback_deadline_at"])
        deadline = None
    now = timezone.now()

    if member and member.response == "N":
        return render(
            request,
            "synopsis/synopsis_feedback_thanks.html",
            {
                "project": project,
                "error": "This link is no longer available because you declined the invitation.",
            },
        )

    if deadline and now >= deadline:
        closure_message = (
            "The feedback deadline has passed (" f"{_format_deadline(deadline)})."
        )
        return render(
            request,
            "synopsis/synopsis_feedback_thanks.html",
            {
                "project": project,
                "feedback": fb,
                "closed_message": closure_message,
                "deadline": deadline,
                "closed": True,
            },
        )

    if request.method == "POST":
        form = SynopsisFeedbackForm(request.POST, request.FILES)
        if form.is_valid():
            content = form.cleaned_data["content"].strip()
            uploaded_doc = form.cleaned_data["uploaded_document"]
            updates = []
            if content:
                fb.content = content
                updates.append("content")
            if uploaded_doc:
                fb.uploaded_document = uploaded_doc
                updates.append("uploaded_document")
            if updates:
                fb.submitted_at = timezone.now()
                updates.append("submitted_at")
                fb.save(update_fields=updates)

                if member:
                    member_updates = set()
                    today = timezone.localdate()
                    if member.feedback_on_synopsis_received != today:
                        member.feedback_on_synopsis_received = today
                        member_updates.add("feedback_on_synopsis_received")
                    if member_updates:
                        member.save(update_fields=list(member_updates))

                details = []
                if uploaded_doc:
                    details.append(
                        f"Document uploaded: {fb.latest_document_label or uploaded_doc.name}"
                    )
                if content:
                    snippet = (content[:97] + "…") if len(content) > 100 else content
                    details.append(f"Comments provided: {snippet}")
                if details:
                    _log_project_change(
                        project,
                        request.user,
                        "Synopsis feedback submitted",
                        " | ".join(details),
                    )
            return render(
                request,
                "synopsis/synopsis_feedback_thanks.html",
                {
                    "project": project,
                    "feedback": fb,
                    "deadline": deadline,
                },
            )
    else:
        form = SynopsisFeedbackForm(initial={"content": fb.content})

    return render(
        request,
        "synopsis/synopsis_feedback_form.html",
        {
            "project": project,
            "token": fb.token,
            "feedback": fb,
            "deadline": deadline,
            "form": form,
        },
    )


@login_required
def action_list_delete_revision(request, project_id, revision_id):
    if request.method != "POST":
        return HttpResponseBadRequest("Invalid request method.")

    project = get_object_or_404(Project, pk=project_id)
    if not _user_can_manage_project_configuration(request.user, project):
        messages.error(
            request,
            "You do not have permission to delete action list revisions for this synopsis.",
        )
        return redirect("synopsis:action_list_detail", project_id=project.id)

    action_list = getattr(project, "action_list", None)
    if not action_list:
        messages.error(request, "No action list exists for this project.")
        return redirect("synopsis:action_list_detail", project_id=project.id)

    revision = get_object_or_404(
        ActionListRevision, pk=revision_id, action_list=action_list
    )

    was_current = action_list.current_revision_id == revision.id
    file_name = revision.file.name
    revision.delete()

    next_revision = (
        action_list.revisions.exclude(pk=revision_id)
        .order_by("-uploaded_at", "-id")
        .first()
    )

    if was_current:
        if next_revision:
            try:
                base_name, size_text = _apply_revision_to_action_list(
                    action_list, next_revision
                )
            except (FileNotFoundError, ValueError):
                action_list.current_revision = next_revision
                action_list.save(update_fields=["current_revision"])
                base_name = next_revision.original_name or os.path.basename(
                    next_revision.file.name
                )
                size_text = _format_file_size(next_revision.file_size)
        else:
            action_list.current_revision = None
            action_list.save(update_fields=["current_revision"])
            base_name = "none"
            size_text = "—"
    else:
        base_name = revision.original_name or os.path.basename(file_name)
        size_text = _format_file_size(revision.file_size)

    _log_project_change(
        project,
        request.user,
        "Action list revision deleted",
        f"Revision file: {file_name or 'unknown'}; Remaining current: {base_name} ({size_text})",
    )

    messages.success(request, "Action list revision deleted.")
    return redirect("synopsis:action_list_detail", project_id=project.id)


@login_required
def advisory_send_invite_member(request, project_id, member_id):
    project = get_object_or_404(Project, id=project_id)
    m = get_object_or_404(AdvisoryBoardMember, id=member_id, project=project)
    if request.method != "POST":
        return HttpResponseBadRequest("POST required")
    if not m.email:
        messages.error(request, "This member has no email.")
        return redirect("synopsis:advisory_board_list", project_id=project.id)

    due_date = _resolve_invite_due_date(member=m)
    inv = AdvisoryBoardInvitation.objects.create(
        project=project,
        member=m,
        email=m.email,
        invited_by=request.user,
        due_date=due_date,
    )
    yes_url = request.build_absolute_uri(
        reverse("synopsis:advisory_invite_reply", args=[inv.token, "yes"])
    )
    no_url = request.build_absolute_uri(
        reverse("synopsis:advisory_invite_reply", args=[inv.token, "no"])
    )
    subject = email_subject("invite", project, due_date)
    text, html = _build_advisory_invitation_email(
        project=project,
        recipient_name=advisory_member_display_name(m),
        due_date=due_date,
        yes_url=yes_url,
        no_url=no_url,
    )

    msg = EmailMultiAlternatives(
        subject,
        text,
        to=[m.email],
        reply_to=reply_to_list(getattr(request.user, "email", None)),
    )
    msg.attach_alternative(html, "text/html")
    msg.send()

    m.invite_sent = True
    m.invite_sent_at = timezone.now()
    if m.response_date != due_date:
        m.response_date = due_date
        m.reminder_sent = False
        m.reminder_sent_at = None
        m.save(
            update_fields=[
                "invite_sent",
                "invite_sent_at",
                "response_date",
                "reminder_sent",
                "reminder_sent_at",
            ]
        )
    else:
        m.save(update_fields=["invite_sent", "invite_sent_at"])

    messages.success(request, f"Invitation sent to {m.email}.")
    return redirect("synopsis:advisory_board_list", project_id=project.id)
