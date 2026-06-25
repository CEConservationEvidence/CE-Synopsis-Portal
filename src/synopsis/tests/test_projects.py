"""Project setup, phase, settings, and synopsis structure tests."""

from .common import *


class ProjectPhaseTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Marine Study")

    def test_defaults_to_draft_protocol_without_manual_phase(self):
        self.assertEqual(self.project.phase, "draft_protocol")
        self.assertEqual(self.project.get_phase_display(), "Draft protocol")

    def test_manual_phase_does_not_regress(self):
        self.project.phase_manual = "draft_protocol"
        self.project.save(update_fields=["phase_manual"])
        self.assertEqual(self.project.phase, "draft_protocol")

    def test_manual_phase_can_advance(self):
        self.project.phase_manual = "summary_writing"
        self.project.save(update_fields=["phase_manual"])
        self.assertEqual(self.project.phase, "summary_writing")
        self.assertEqual(
            self.project.get_phase_display(),
            "Summary writing",
        )

    def test_computed_phase_advances_after_protocol_upload(self):
        Protocol.objects.create(
            project=self.project,
            document=SimpleUploadedFile("protocol.docx", b"protocol"),
        )

        self.assertEqual(self.project.phase, "invite_advisory_board")

    def test_computed_phase_does_not_regress_existing_project_without_manual_phase(self):
        Protocol.objects.create(
            project=self.project,
            document=SimpleUploadedFile("protocol.docx", b"protocol"),
        )
        AdvisoryBoardInvitation.objects.create(
            project=self.project,
            email="advisor@example.com",
            accepted=True,
        )
        AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Advisor",
            email="advisor@example.com",
            response="Y",
            feedback_on_protocol_received=timezone.localdate(),
        )

        self.assertEqual(self.project.phase_manual, None)
        self.assertEqual(self.project.phase, "draft_chapters")

    def test_manual_phase_cannot_regress_below_computed_progress(self):
        Protocol.objects.create(
            project=self.project,
            document=SimpleUploadedFile("protocol.docx", b"protocol"),
        )
        AdvisoryBoardInvitation.objects.create(
            project=self.project,
            email="advisor@example.com",
            accepted=True,
        )
        AdvisoryBoardMember.objects.create(
            project=self.project,
            first_name="Advisor",
            email="advisor@example.com",
            response="Y",
            feedback_on_protocol_received=timezone.localdate(),
        )
        self.project.phase_manual = "draft_protocol"
        self.project.save(update_fields=["phase_manual"])

        self.assertEqual(self.project.phase, "draft_chapters")

    def test_computed_phase_respects_disabled_protocol_and_advisory_steps(self):
        self.project.protocol_relevant = False
        self.project.advisory_board_relevant = False
        self.project.save(update_fields=["protocol_relevant", "advisory_board_relevant"])

        self.assertEqual(self.project.phase, "references_screening")


class ProjectAuthorUsersTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Author Demo")
        self.author = User.objects.create_user(username="zoe")
        self.coauthor = User.objects.create_user(username="adam")
        self.manager = User.objects.create_user(username="manager")

        UserRole.objects.create(user=self.author, project=self.project, role="author")
        UserRole.objects.create(user=self.coauthor, project=self.project, role="author")
        UserRole.objects.create(user=self.manager, project=self.project, role="manager")

    def test_returns_authors_sorted_by_username(self):
        usernames = list(self.project.author_users.values_list("username", flat=True))
        self.assertEqual(usernames, ["adam", "zoe"])


class SynopsisStructureTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Synopsis Demo")
        self.user = User.objects.create_user(username="manager", password="pw", is_staff=True)
        UserRole.objects.create(user=self.user, project=self.project, role="manager")
        self.client.force_login(self.user)
        # Attach a reference/summary for assignment flows
        self.batch = ReferenceSourceBatch.objects.create(
            project=self.project,
            label="Batch 1",
            source_type="manual_upload",
            uploaded_by=self.user,
        )
        self.reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Test ref",
            hash_key="hash-test-ref",
            screening_status="included",
        )
        self.summary = ReferenceSummary.objects.create(
            project=self.project, reference=self.reference
        )

    def test_apply_preset_creates_chapters_once(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        response = self.client.post(
            url,
            {"action": "apply-preset", "preset_key": "standard_ce_toc"},
        )
        self.assertEqual(response.status_code, 302)
        self.assertTrue(SynopsisChapter.objects.filter(project=self.project).exists())
        front_chapter = SynopsisChapter.objects.get(
            project=self.project, title="Advisory Board"
        )
        self.assertEqual(front_chapter.chapter_type, SynopsisChapter.TYPE_TEXT)
        evidence_chapter = SynopsisChapter.objects.get(
            project=self.project,
            title="2. Threat: Residential and commercial development",
        )
        self.assertEqual(
            evidence_chapter.chapter_type, SynopsisChapter.TYPE_EVIDENCE
        )
        change = ProjectChangeLog.objects.filter(
            project=self.project,
            action="Synopsis outline preset applied",
        ).first()
        self.assertIsNotNone(change)
        self.assertIn(
            "Preset: Standard CE synopsis (full ToC, chapters only)",
            change.details,
        )
        self.assertIn("Chapters created:", change.details)
        count_after_first = SynopsisChapter.objects.filter(project=self.project).count()
        # Second apply should be blocked because outline is not empty
        response = self.client.post(
            url,
            {"action": "apply-preset", "preset_key": "standard_ce_toc"},
        )
        self.assertEqual(response.status_code, 302)
        count_after_second = SynopsisChapter.objects.filter(project=self.project).count()
        self.assertEqual(count_after_first, count_after_second)

    def test_reset_structure_clears_outline(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        SynopsisChapter.objects.create(project=self.project, title="Tmp", position=1)
        response = self.client.post(url, {"action": "reset-structure"})
        self.assertEqual(response.status_code, 302)
        self.assertFalse(SynopsisChapter.objects.filter(project=self.project).exists())

    def test_create_intervention_without_subheading_creates_default(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(project=self.project, title="Ch", position=1)
        response = self.client.post(
            url,
            {
                "action": "create-intervention",
                "chapter_id": chapter.id,
                "title": "Intervention A",
            },
        )
        self.assertEqual(response.status_code, 302)
        subheadings = SynopsisSubheading.objects.filter(chapter=chapter)
        self.assertEqual(subheadings.count(), 1)
        intervention_titles = list(
            SynopsisIntervention.objects.filter(subheading__chapter=chapter).values_list("title", flat=True)
        )
        self.assertIn("Intervention A", intervention_titles)
        change = ProjectChangeLog.objects.filter(
            project=self.project,
            action="Intervention added",
        ).first()
        self.assertIsNotNone(change)
        self.assertIn("Chapter: Ch", change.details)
        self.assertIn("Intervention: Intervention A", change.details)

    def test_structure_workspace_updates_chapter_title(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "update-chapter-title",
                "chapter_id": chapter.id,
                "title": "99. Threat: Renamed",
            },
        )

        self.assertEqual(response.status_code, 302)
        chapter.refresh_from_db()
        self.assertEqual(chapter.title, "2. Threat: Renamed")
        change = ProjectChangeLog.objects.filter(
            project=self.project,
            action="Chapter title updated",
        ).first()
        self.assertIsNotNone(change)
        self.assertIn("Chapter: 2. Threat: Renamed", change.details)
        self.assertIn("Title: 2. Threat: Demo → 2. Threat: Renamed", change.details)

    def test_structure_workspace_rejects_blank_chapter_title_update(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "update-chapter-title",
                "chapter_id": chapter.id,
                "title": "   ",
            },
            follow=True,
        )

        self.assertEqual(response.status_code, 200)
        chapter.refresh_from_db()
        self.assertEqual(chapter.title, "2. Threat: Demo")
        self.assertContains(
            response,
            "Enter a chapter title of 255 characters or fewer.",
        )

    def test_structure_workspace_updates_subheading_title(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="General",
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "update-subheading-title",
                "subheading_id": subheading.id,
                "title": "General farmland",
            },
        )

        self.assertEqual(response.status_code, 302)
        subheading.refresh_from_db()
        self.assertEqual(subheading.title, "General farmland")
        change = ProjectChangeLog.objects.filter(
            project=self.project,
            action="Intervention group title updated",
        ).first()
        self.assertIsNotNone(change)
        self.assertIn("Chapter: 2. Threat: Demo", change.details)
        self.assertIn("Intervention group: General farmland", change.details)
        self.assertIn("Title: General → General farmland", change.details)

    def test_evidence_workspace_creates_evidence_chapter_without_visible_type_controls(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        response = self.client.post(
            url,
            {
                "action": "create-chapter",
                "title": "2. Threat: Demo",
            },
            follow=True,
        )

        self.assertEqual(response.status_code, 200)
        chapter = SynopsisChapter.objects.get(project=self.project, title="1. Threat: Demo")
        self.assertEqual(chapter.chapter_type, SynopsisChapter.TYPE_EVIDENCE)
        self.assertContains(response, 'name="chapter_type" value="evidence"', html=False)
        self.assertContains(response, "This workspace creates evidence chapters only.")
        self.assertNotContains(response, 'value="update-chapter-type"', html=False)
        self.assertNotContains(response, ">Save type<", html=False)

    def test_evidence_workspace_create_chapter_uses_next_numbered_title(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        SynopsisChapter.objects.create(
            project=self.project,
            title="5. Threat: Existing",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        appendix = SynopsisChapter.objects.create(
            project=self.project,
            title="Appendix 1: Sources",
            chapter_type=SynopsisChapter.TYPE_APPENDIX,
            position=2,
        )

        response = self.client.post(
            url,
            {
                "action": "create-chapter",
                "title": "Threat: New",
            },
            follow=True,
        )

        self.assertEqual(response.status_code, 200)
        chapter = SynopsisChapter.objects.get(project=self.project, title="6. Threat: New")
        self.assertEqual(chapter.chapter_type, SynopsisChapter.TYPE_EVIDENCE)
        appendix.refresh_from_db()
        self.assertLess(chapter.position, appendix.position)
        self.assertContains(response, "Added chapter")
        self.assertContains(response, "6. Threat: New")

    def test_evidence_workspace_create_chapter_uses_highest_existing_evidence_number(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        SynopsisChapter.objects.create(
            project=self.project,
            title="1. About this book",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=1,
        )
        SynopsisChapter.objects.create(
            project=self.project,
            title="15. Threat: Existing later",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=2,
        )
        SynopsisChapter.objects.create(
            project=self.project,
            title="13. Threat: Existing earlier",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=3,
        )

        response = self.client.post(
            url,
            {
                "action": "create-chapter",
                "title": "Threat: New",
            },
            follow=True,
        )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(
            SynopsisChapter.objects.filter(
                project=self.project,
                title="16. Threat: New",
                chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            ).exists()
        )
        self.assertContains(response, "16. Threat: New")

    def test_evidence_workspace_ignores_non_about_narrative_numbers(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        SynopsisChapter.objects.create(
            project=self.project,
            title="99. Advisory Board",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=1,
        )
        SynopsisChapter.objects.create(
            project=self.project,
            title="5. Threat: Existing",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=2,
        )

        response = self.client.post(
            url,
            {
                "action": "create-chapter",
                "title": "Threat: New",
            },
            follow=True,
        )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(
            SynopsisChapter.objects.filter(
                project=self.project,
                title="6. Threat: New",
                chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            ).exists()
        )
        self.assertContains(response, "6. Threat: New")

    def test_evidence_workspace_does_not_reserve_number_for_non_intro_text_chapter(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        SynopsisChapter.objects.create(
            project=self.project,
            title="1. Advisory Board",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "create-chapter",
                "title": "Threat: New",
            },
            follow=True,
        )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(
            SynopsisChapter.objects.filter(
                project=self.project,
                title="1. Threat: New",
                chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            ).exists()
        )
        self.assertContains(response, "1. Threat: New")

    def test_new_evidence_chapter_can_move_up_when_appendices_exist(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        first = SynopsisChapter.objects.create(
            project=self.project,
            title="5. Threat: Existing",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        SynopsisChapter.objects.create(
            project=self.project,
            title="Appendix 1: Sources",
            chapter_type=SynopsisChapter.TYPE_APPENDIX,
            position=2,
        )

        create_response = self.client.post(
            url,
            {
                "action": "create-chapter",
                "title": "Threat: New",
            },
        )
        self.assertEqual(create_response.status_code, 302)

        new_chapter = SynopsisChapter.objects.get(project=self.project, title="6. Threat: New")
        move_response = self.client.post(
            url,
            {
                "action": "move-chapter",
                "chapter_id": new_chapter.id,
                "direction": "up",
            },
        )

        self.assertEqual(move_response.status_code, 302)
        first.refresh_from_db()
        new_chapter.refresh_from_db()
        self.assertEqual(new_chapter.position, 1)
        self.assertEqual(new_chapter.title, "1. Threat: New")
        self.assertEqual(first.position, 2)
        self.assertEqual(first.title, "2. Threat: Existing")

    def test_structure_workspace_requires_confirmation_flag_to_delete_chapter(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "delete-chapter",
                "chapter_id": chapter.id,
            },
            follow=True,
        )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(SynopsisChapter.objects.filter(pk=chapter.id).exists())
        self.assertContains(
            response,
            "Please confirm chapter deletion before removing it.",
        )

    def test_structure_workspace_deletes_chapter_after_confirmation(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "delete-chapter",
                "chapter_id": chapter.id,
                "confirm_delete_chapter": "1",
            },
        )

        self.assertEqual(response.status_code, 302)
        self.assertFalse(SynopsisChapter.objects.filter(pk=chapter.id).exists())

    def test_move_chapter_renumbers_evidence_titles_only(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        front_matter = SynopsisChapter.objects.create(
            project=self.project,
            title="Advisory Board",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=1,
        )
        about = SynopsisChapter.objects.create(
            project=self.project,
            title="1. About this book",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=2,
        )
        threat = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=3,
        )
        next_threat = SynopsisChapter.objects.create(
            project=self.project,
            title="3. Threat: Next",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=4,
        )

        response = self.client.post(
            url,
            {
                "action": "move-chapter",
                "chapter_id": next_threat.id,
                "direction": "up",
            },
        )

        self.assertEqual(response.status_code, 302)
        front_matter.refresh_from_db()
        about.refresh_from_db()
        threat.refresh_from_db()
        next_threat.refresh_from_db()
        self.assertEqual(front_matter.title, "Advisory Board")
        self.assertEqual(about.position, 2)
        self.assertEqual(about.title, "1. About this book")
        self.assertEqual(next_threat.position, 3)
        self.assertEqual(next_threat.title, "2. Threat: Next")
        self.assertEqual(threat.position, 4)
        self.assertEqual(threat.title, "3. Threat: Demo")

    def test_structure_page_renders_saved_action_name_suggestions(self):
        self.project.saved_action_names = "Install nest boxes\nReduce grazing"
        self.project.save(update_fields=["saved_action_names"])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="Evidence",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        SynopsisSubheading.objects.create(
            chapter=chapter,
            title="General",
            position=1,
        )

        response = self.client.get(
            reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        )

        self.assertContains(
            response,
            'id="project-action-name-suggestions"',
            html=False,
        )
        self.assertContains(
            response,
            '<option value="Install nest boxes"></option>',
            html=False,
        )
        self.assertContains(
            response,
            "Saved action names from the Action List page appear as suggestions",
        )

    def test_move_intervention_to_another_subheading(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        general = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="General",
            position=1,
        )
        arable = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Arable",
            position=2,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=general,
            title="Mow more frequently",
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "move-intervention-to-subheading",
                "intervention_id": intervention.id,
                "target_subheading_id": arable.id,
            },
        )

        self.assertEqual(response.status_code, 302)
        intervention.refresh_from_db()
        self.assertEqual(intervention.subheading, arable)
        self.assertEqual(intervention.position, 1)
        self.assertEqual(general.interventions.count(), 0)
        self.assertEqual(arable.interventions.count(), 1)

    def test_structure_page_explains_intervention_group_linking(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Arable",
            position=1,
        )
        SynopsisIntervention.objects.create(
            subheading=subheading,
            title="Mow more frequently",
            position=1,
        )
        self.summary.synopsis_draft = (
            "A replicated study found that mowing more frequently increased arable plant richness."
        )
        self.summary.use_custom_synopsis_draft = True
        self.summary.save(update_fields=["synopsis_draft", "use_custom_synopsis_draft"])
        SynopsisAssignment.objects.create(
            intervention=SynopsisIntervention.objects.get(title="Mow more frequently"),
            reference_summary=self.summary,
            position=1,
        )

        response = self.client.get(url)

        self.assertContains(response, "Add intervention to Arable")
        self.assertContains(response, "Intervention group")
        self.assertContains(response, "Move to group")
        self.assertContains(response, "Edit metadata, background and key messages")
        self.assertContains(response, "Metadata")
        self.assertContains(response, "Background")
        self.assertContains(response, "Key messages")
        self.assertContains(response, "Assigned summaries")
        self.assertContains(response, "Intervention evidence details")
        self.assertContains(
            response,
            "These fields are kept only as intervention metadata and do not add extra narrative text to the compiled synopsis.",
        )
        self.assertContains(response, "Assigned study summaries to review")
        self.assertContains(response, "review the assigned summaries here first")
        self.assertContains(
            response,
            "mowing more frequently increased arable plant richness",
        )
        self.assertContains(response, "Vegetation Community")
        self.assertContains(response, "Vegetation Abundance")
        self.assertContains(response, "Vegetation Structure")
        self.assertContains(response, "Other")

    def test_structure_page_places_no_studies_sentence_before_background_preview(self):
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
            background_text="Chapter background text.",
            background_references="Chapter Background Ref",
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
            evidence_status=SynopsisIntervention.EVIDENCE_STATUS_NO_STUDIES,
            background_text="Intervention background text.",
            background_references="Intervention Background Ref",
        )

        response = self.client.get(
            reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        )

        self.assertContains(
            response,
            "We found no studies that evaluated the effects of this intervention.",
        )
        self.assertContains(response, "Background references: Chapter Background Ref")
        self.assertContains(
            response,
            "Background references: Intervention Background Ref",
        )

        content = response.content.decode()
        title_index = content.find("2.1 Demo intervention")
        no_studies_index = content.find(
            "We found no studies that evaluated the effects of this intervention."
        )
        intervention_background_index = content.find("Intervention background text.")

        self.assertGreaterEqual(title_index, 0)
        self.assertGreater(no_studies_index, title_index)
        self.assertGreater(intervention_background_index, no_studies_index)

    def test_structure_page_renders_inline_markup_for_backgrounds_key_messages_and_summaries(self):
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
            background_text="Forest <i>edge</i> with CO<sub>2</sub> buildup.",
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
            background_text="Reached the 10<sup>th</sup> plot.",
        )
        self.summary.synopsis_draft = (
            "A replicated study found that <i>Festuca</i> cover increased at plot 10<sup>th</sup>."
        )
        self.summary.use_custom_synopsis_draft = True
        self.summary.save(update_fields=["synopsis_draft", "use_custom_synopsis_draft"])
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=self.summary,
            position=1,
        )
        SynopsisInterventionKeyMessage.objects.create(
            intervention=intervention,
            response_group=SynopsisInterventionKeyMessage.GROUP_RESPONSE,
            statement="CO<sub>2</sub> uptake improved in the 10<sup>th</sup> plot.",
            position=1,
        )

        response = self.client.get(
            reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        )

        self.assertContains(response, "<i>edge</i>", html=False)
        self.assertContains(response, "CO<sub>2</sub> buildup.", html=False)
        self.assertContains(response, "10<sup>th</sup> plot.", html=False)
        self.assertContains(response, "<i>Festuca</i>", html=False)
        self.assertContains(response, "CO<sub>2</sub> uptake improved", html=False)
        self.assertContains(response, 'data-inline-markup="true"', html=False)

    def test_structure_page_includes_distinct_subscript_and_superscript_toolbar_contract(self):
        response = self.client.get(
            reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        )

        self.assertContains(response, "x₂", html=False)
        self.assertContains(response, "x²", html=False)
        self.assertContains(response, "bottom: -0.3em;", html=False)
        self.assertContains(response, "top: -0.55em;", html=False)

    def test_structure_page_explains_export_order_and_chapter_delete_impact(self):
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )

        response = self.client.get(
            reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        )

        self.assertContains(response, "How This Compiles to DOCX")
        self.assertContains(
            response,
            "The DOCX follows this page order.",
        )
        self.assertContains(
            response,
            "Each intervention exports in this order:",
        )
        self.assertContains(
            response,
            "This removes the chapter and everything inside it from the synopsis outline",
        )
        self.assertContains(
            response,
            'id="delete-evidence-chapter-modal"',
            html=False,
        )
        self.assertContains(
            response,
            'data-delete-chapter-trigger',
            html=False,
        )
        self.assertContains(
            response,
            'name="confirm_delete_chapter" value="1"',
            html=False,
        )
        self.assertContains(response, f'href="#chapter-{chapter.id}"', html=False)

    def test_structure_page_hides_cross_reference_controls(self):
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        primary = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Primary intervention",
            position=1,
        )
        SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.2 Secondary intervention",
            position=2,
            is_cross_reference=True,
            primary_intervention=primary,
        )

        response = self.client.get(
            reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        )

        self.assertNotContains(response, 'name="primary_intervention"', html=False)
        self.assertNotContains(response, 'name="is_cross_reference"', html=False)
        self.assertNotContains(response, ">Cross-ref<", html=False)
        self.assertNotContains(response, "Cross-reference:", html=False)

    def test_synopsis_ris_export_includes_unique_summarised_references(self):
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        self.reference.title = "Assigned done ref"
        self.reference.authors = "Alpha A.; Beta B."
        self.reference.publication_year = 2004
        self.reference.journal = "Journal One"
        self.reference.doi = "10.1234/example-one"
        self.reference.url = "https://example.com/one"
        self.reference.save(
            update_fields=[
                "title",
                "authors",
                "publication_year",
                "journal",
                "doi",
                "url",
                "updated_at",
            ]
        )
        self.summary.status = ReferenceSummary.STATUS_DONE
        self.summary.save(update_fields=["status", "updated_at"])
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=self.summary,
            position=1,
        )

        duplicate_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=self.reference,
            status=ReferenceSummary.STATUS_DONE,
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=duplicate_summary,
            position=2,
        )

        second_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Second done ref",
            authors="Gamma G.",
            publication_year=2010,
            journal="Journal Two",
            hash_key="hash-second-done-ref",
            screening_status="included",
            raw_ris={
                "type_of_reference": "JOUR",
                "title": "Second done ref",
                "authors": ["Gamma G."],
                "year": "2010",
                "journal_name": "Journal Two",
            },
        )
        second_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=second_reference,
            status=ReferenceSummary.STATUS_DONE,
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=second_summary,
            position=3,
        )

        draft_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Assigned draft ref",
            hash_key="hash-assigned-draft-ref",
            screening_status="included",
        )
        draft_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=draft_reference,
            status=ReferenceSummary.STATUS_DRAFT,
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=draft_summary,
            position=4,
        )

        ReferenceSummary.objects.create(
            project=self.project,
            reference=Reference.objects.create(
                project=self.project,
                batch=self.batch,
                title="Unassigned done ref",
                authors="Delta D.",
                publication_year=2015,
                hash_key="hash-unassigned-done-ref",
                screening_status="included",
            ),
            status=ReferenceSummary.STATUS_DONE,
        )

        structure_response = self.client.get(
            reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        )
        self.assertContains(
            structure_response,
            reverse("synopsis:project_synopsis_export_ris", args=[self.project.id]),
            html=False,
        )

        response = self.client.get(
            reverse("synopsis:project_synopsis_export_ris", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertIn("application/x-research-info-systems", response["Content-Type"])
        self.assertIn(".ris", response["Content-Disposition"])
        records = rispy.loads(response.content.decode("utf-8"))
        self.assertEqual(len(records), 3)
        by_title = {record["title"]: record for record in records}
        self.assertEqual(
            sorted(by_title.keys()),
            ["Assigned done ref", "Second done ref", "Unassigned done ref"],
        )
        self.assertEqual(by_title["Assigned done ref"]["doi"], "10.1234/example-one")
        self.assertEqual(by_title["Assigned done ref"]["year"], "2004")
        self.assertEqual(by_title["Second done ref"]["journal_name"], "Journal Two")
        self.assertEqual(by_title["Unassigned done ref"]["year"], "2015")
        self.assertTrue(
            SynopsisExportLog.objects.filter(
                project=self.project,
                note="Manual RIS export",
            ).exists()
        )

    def test_synopsis_ris_export_preserves_ambiguous_pages_as_note(self):
        self.reference.title = "Ambiguous pages ref"
        self.reference.pages = "12-15; 20"
        self.reference.save(update_fields=["title", "pages", "updated_at"])
        self.summary.status = ReferenceSummary.STATUS_DONE
        self.summary.save(update_fields=["status", "updated_at"])

        response = self.client.get(
            reverse("synopsis:project_synopsis_export_ris", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        records = rispy.loads(response.content.decode("utf-8"))
        self.assertEqual(len(records), 1)
        record = records[0]
        self.assertEqual(record["title"], "Ambiguous pages ref")
        self.assertEqual(record.get("notes"), ["Pages: 12-15; 20"])
        self.assertNotIn("start_page", record)
        self.assertNotIn("end_page", record)

    def test_synopsis_structure_csv_export_includes_paragraphs_tags_and_structure(self):
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="Install nest boxes",
            position=1,
        )
        self.reference.title = "Assigned structure ref"
        self.reference.authors = "Alpha A."
        self.reference.publication_year = 2004
        self.reference.doi = "10.1234/structure"
        self.reference.unlinked_reference_folder = ["1", "10"]
        self.reference.screening_status = "included"
        self.reference.save(
            update_fields=[
                "title",
                "authors",
                "publication_year",
                "doi",
                "unlinked_reference_folder",
                "screening_status",
                "updated_at",
            ]
        )
        self.summary.status = ReferenceSummary.STATUS_DONE
        self.summary.action_description = "Install nest boxes"
        self.summary.study_design = "replicated study"
        self.summary.research_design = "Before-and-after"
        self.summary.broad_category = "Birds"
        self.summary.action_tags = ["Install nest boxes"]
        self.summary.threat_tags = ["Residential & commercial development-Housing & urban areas"]
        self.summary.habitat_tags = ["5. Forests"]
        self.summary.taxon_tags = ["Birds"]
        self.summary.location_tags = ["United Kingdom"]
        self.summary.synopsis_draft = "A replicated study (CR1000) found increased bird occupancy in nest boxes."
        self.summary.use_custom_synopsis_draft = True
        self.summary.save(
            update_fields=[
                "status",
                "action_description",
                "study_design",
                "research_design",
                "broad_category",
                "action_tags",
                "threat_tags",
                "habitat_tags",
                "taxon_tags",
                "location_tags",
                "synopsis_draft",
                "use_custom_synopsis_draft",
                "updated_at",
            ]
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=self.summary,
            position=1,
        )

        unassigned_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Unassigned included ref",
            authors="Beta B.",
            publication_year=2015,
            hash_key="hash-unassigned-structure",
            screening_status="included",
        )
        ReferenceSummary.objects.create(
            project=self.project,
            reference=unassigned_reference,
            status=ReferenceSummary.STATUS_TODO,
        )

        excluded_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Excluded summary ref",
            hash_key="hash-excluded-structure",
            screening_status="included",
        )
        ReferenceSummary.objects.create(
            project=self.project,
            reference=excluded_reference,
            status=ReferenceSummary.STATUS_EXCLUDED,
            exclusion_reason="Not relevant after full text.",
        )

        structure_response = self.client.get(
            reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        )
        self.assertContains(
            structure_response,
            reverse(
                "synopsis:project_synopsis_export_structure_csv",
                args=[self.project.id],
            ),
            html=False,
        )

        response = self.client.get(
            reverse(
                "synopsis:project_synopsis_export_structure_csv",
                args=[self.project.id],
            )
        )

        self.assertEqual(response.status_code, 200)
        self.assertIn("text/csv", response["Content-Type"])
        self.assertIn(".csv", response["Content-Disposition"])
        rows = list(csv.DictReader(io.StringIO(response.content.decode("utf-8"))))
        self.assertEqual(len(rows), 1)
        rows_by_title = {row["paper_title"]: row for row in rows}
        assigned_row = rows_by_title["Assigned structure ref"]
        self.assertTrue(assigned_row["reference_identifier"])
        self.assertTrue(assigned_row["summary_identifier"])
        self.assertEqual(assigned_row["summary_status"], "Summarised")
        self.assertEqual(assigned_row["action"], "Install nest boxes")
        self.assertEqual(assigned_row["study_design"], "replicated study")
        self.assertEqual(assigned_row["research_design"], "Before-and-after")
        self.assertEqual(assigned_row["broad_category"], "Birds")
        self.assertEqual(assigned_row["action_tags"], "Install nest boxes")
        self.assertIn("Housing & urban areas", assigned_row["threat_tags"])
        self.assertEqual(assigned_row["habitat_tags"], "5. Forests")
        self.assertEqual(assigned_row["taxon_tags"], "Birds")
        self.assertEqual(assigned_row["location_tags"], "United Kingdom")
        self.assertEqual(assigned_row["assignment_count"], "1")
        self.assertEqual(assigned_row["chapters"], "2. Threat: Demo")
        self.assertEqual(assigned_row["intervention_groups"], "Interventions")
        self.assertEqual(assigned_row["interventions"], "Install nest boxes")
        self.assertIn(
            "2. Threat: Demo > Interventions > Install nest boxes",
            assigned_row["structure_locations"],
        )
        self.assertEqual(assigned_row["paragraph_mode"], "custom")
        self.assertIn("increased bird occupancy in nest boxes", assigned_row["summary_paragraph"])
        self.assertIn("Assigned structure ref", assigned_row["citation_for_synopsis_export"])
        self.assertIn("Amphibians", assigned_row["reference_categories"])
        self.assertIn("Plants/algae ex situ", assigned_row["reference_categories"])

        self.assertNotIn("Unassigned included ref", rows_by_title)
        self.assertNotIn("Excluded summary ref", rows_by_title)
        self.assertTrue(
            SynopsisExportLog.objects.filter(
                project=self.project,
                note="Manual structure CSV export",
            ).exists()
        )

    def test_synopsis_structure_csv_export_sanitizes_formula_like_cells(self):
        self.reference.title = "=Dangerous title"
        self.reference.authors = "@Author"
        self.reference.screening_status = "included"
        self.reference.save(
            update_fields=["title", "authors", "screening_status", "updated_at"]
        )
        self.summary.status = ReferenceSummary.STATUS_DONE
        self.summary.synopsis_draft = "+Dangerous paragraph"
        self.summary.use_custom_synopsis_draft = True
        self.summary.save(
            update_fields=[
                "status",
                "synopsis_draft",
                "use_custom_synopsis_draft",
                "updated_at",
            ]
        )

        response = self.client.get(
            reverse(
                "synopsis:project_synopsis_export_structure_csv",
                args=[self.project.id],
            )
        )

        self.assertEqual(response.status_code, 200)
        rows = list(csv.DictReader(io.StringIO(response.content.decode("utf-8"))))
        self.assertEqual(len(rows), 1)
        row = rows[0]
        self.assertEqual(row["paper_title"], "'=Dangerous title")
        self.assertEqual(row["authors"], "'@Author")
        self.assertEqual(row["summary_paragraph"], "'+Dangerous paragraph")

    def test_structure_page_background_reference_guidance_is_not_limited_by_search_end_date(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )

        response = self.client.get(url)

        self.assertContains(
            response,
            "Optional contextual references. These do not need to be published before the search end date.",
        )
        self.assertNotContains(response, "published before search end date")

    def test_structure_page_renders_restore_state_hooks(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Arable",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="Mow more frequently",
            position=1,
        )

        response = self.client.get(url)

        self.assertContains(response, 'id="synopsis-structure-page"', html=False)
        self.assertContains(
            response,
            f'id="subheading-{subheading.id}"',
            html=False,
        )
        self.assertContains(
            response,
            f'id="intervention-{intervention.id}"',
            html=False,
        )
        self.assertContains(
            response,
            f'id="intervention-editor-{intervention.id}"',
            html=False,
        )
        self.assertContains(response, "cePreservePageState({", html=False)
        self.assertContains(
            response,
            f'"synopsis-structure-state-evidence-{self.project.id}"',
            html=False,
        )
        self.assertContains(response, 'closest("details[id]")', html=False)

    def test_text_chapter_blocks_subheading_and_intervention(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        text_chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="Advisory Board",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=1,
        )
        response = self.client.post(
            url,
            {
                "action": "create-subheading",
                "chapter_id": text_chapter.id,
                "title": "Should fail",
            },
        )
        self.assertEqual(response.status_code, 302)
        self.assertFalse(SynopsisSubheading.objects.filter(chapter=text_chapter).exists())
        response = self.client.post(
            url,
            {
                "action": "create-intervention",
                "chapter_id": text_chapter.id,
                "title": "Should also fail",
            },
        )
        self.assertEqual(response.status_code, 302)
        self.assertFalse(
            SynopsisIntervention.objects.filter(subheading__chapter=text_chapter).exists()
        )

    def test_structure_workspace_rejects_invalid_inline_markup_in_background(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "update-chapter-background",
                "chapter_id": chapter.id,
                "background_text": "CO<sup>2",
                "background_references": "",
            },
            follow=True,
        )

        self.assertEqual(response.status_code, 200)
        chapter.refresh_from_db()
        self.assertEqual(chapter.background_text, "")
        self.assertContains(
            response,
            "Please check the background fields.",
        )
        self.assertContains(response, "matching closing tag")

    def test_structure_workspace_rejects_chapter_type_change_outside_workspace_choices(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "update-chapter-type",
                "chapter_id": chapter.id,
                "chapter_type": SynopsisChapter.TYPE_TEXT,
            },
            follow=True,
        )

        self.assertEqual(response.status_code, 200)
        chapter.refresh_from_db()
        self.assertEqual(chapter.chapter_type, SynopsisChapter.TYPE_EVIDENCE)
        self.assertContains(
            response,
            "Selected chapter type is not available in this workspace.",
        )

    def test_update_intervention_details_fields(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "update-intervention-details",
                "intervention_id": intervention.id,
                "ce_action_url": "https://www.conservationevidence.com/actions/4018",
                "evidence_status": SynopsisIntervention.EVIDENCE_STATUS_NO_STUDIES,
            },
        )
        self.assertEqual(response.status_code, 302)
        intervention.refresh_from_db()
        self.assertEqual(
            intervention.ce_action_url,
            "https://www.conservationevidence.com/actions/4018",
        )
        self.assertEqual(
            intervention.evidence_status,
            SynopsisIntervention.EVIDENCE_STATUS_NO_STUDIES,
        )

    def test_add_and_update_key_message(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        second_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Second test ref",
            hash_key="hash-second-test-ref",
            screening_status="included",
        )
        second_summary = ReferenceSummary.objects.create(
            project=self.project, reference=second_reference
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=self.summary,
            position=1,
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=second_summary,
            position=2,
        )

        response = self.client.post(
            url,
            {
                "action": "add-key-message",
                "intervention_id": intervention.id,
                "response_group": SynopsisInterventionKeyMessage.GROUP_POPULATION,
                "outcome_label": "Abundance/Cover",
                "study_count": 3,
                "statement": "Three studies found increased coral cover after intervention.",
                "supporting_summaries": [str(self.summary.id)],
            },
        )
        self.assertEqual(response.status_code, 302)
        key_message = intervention.key_messages.first()
        self.assertIsNotNone(key_message)
        self.assertEqual(key_message.study_count, 3)
        self.assertEqual(
            list(
                key_message.supporting_summaries.order_by("id").values_list("id", flat=True)
            ),
            [self.summary.id],
        )
        created_change = ProjectChangeLog.objects.filter(
            project=self.project,
            action="Key message added",
        ).first()
        self.assertIsNotNone(created_change)
        self.assertIn("Key message: Abundance/Cover", created_change.details)
        self.assertIn("Supporting summaries: 1", created_change.details)

        response = self.client.post(
            url,
            {
                "action": "update-key-message",
                "intervention_id": intervention.id,
                "key_message_id": key_message.id,
                "response_group": SynopsisInterventionKeyMessage.GROUP_COMMUNITY,
                "outcome_label": "Richness/diversity",
                "study_count": 5,
                "statement": "Five studies found no clear community-level change.",
                "supporting_summaries": [str(second_summary.id), str(self.summary.id)],
            },
        )
        self.assertEqual(response.status_code, 302)
        key_message.refresh_from_db()
        self.assertEqual(
            key_message.response_group,
            SynopsisInterventionKeyMessage.GROUP_COMMUNITY,
        )
        self.assertEqual(key_message.outcome_label, "Richness/diversity")
        self.assertEqual(key_message.study_count, 5)
        self.assertEqual(
            list(
                key_message.supporting_summaries.order_by("id").values_list("id", flat=True)
            ),
            sorted([self.summary.id, second_summary.id]),
        )
        updated_change = ProjectChangeLog.objects.filter(
            project=self.project,
            action="Key message updated",
        ).first()
        self.assertIsNotNone(updated_change)
        self.assertIn("Key message: Richness/diversity", updated_change.details)
        self.assertIn("Supporting summaries: 2", updated_change.details)

    def test_intervention_reference_numbering_uses_oldest_first_order(self):
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        older_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Older study",
            authors="Alpha A.",
            publication_year=2001,
            hash_key="hash-old-study",
            screening_status="included",
        )
        newer_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Newer study",
            authors="Beta B.",
            publication_year=2018,
            hash_key="hash-new-study",
            screening_status="included",
        )
        older_summary = ReferenceSummary.objects.create(
            project=self.project, reference=older_reference
        )
        newer_summary = ReferenceSummary.objects.create(
            project=self.project, reference=newer_reference
        )
        SynopsisAssignment.objects.create(
            intervention=intervention, reference_summary=newer_summary, position=1
        )
        SynopsisAssignment.objects.create(
            intervention=intervention, reference_summary=older_summary, position=2
        )

        assignments = list(intervention.assignments.all())
        ordered_assignments, summary_numbers, ordered_references = (
            _intervention_reference_numbering(assignments)
        )

        self.assertEqual(
            [assignment.reference_summary_id for assignment in ordered_assignments],
            [older_summary.id, newer_summary.id],
        )
        self.assertEqual(summary_numbers[older_summary.id], 1)
        self.assertEqual(summary_numbers[newer_summary.id], 2)
        self.assertEqual(
            [reference.id for _, reference in ordered_references],
            [older_reference.id, newer_reference.id],
        )
        self.assertEqual(
            [numbers for numbers, _ in ordered_references],
            [[1], [2]],
        )

    def test_intervention_reference_numbering_groups_duplicate_references(self):
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Group duplicate paper studies",
            position=2,
        )
        shared_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Shared study paper",
            authors="Gamma G.",
            publication_year=2009,
            hash_key="hash-shared-study-paper",
            screening_status="included",
        )
        first_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=shared_reference,
            action_description="Study A",
        )
        second_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=shared_reference,
            action_description="Study B",
        )
        SynopsisAssignment.objects.create(
            intervention=intervention, reference_summary=first_summary, position=1
        )
        SynopsisAssignment.objects.create(
            intervention=intervention, reference_summary=second_summary, position=2
        )

        assignments = list(intervention.assignments.all())
        ordered_assignments, summary_numbers, ordered_references = (
            _intervention_reference_numbering(assignments)
        )

        self.assertEqual(
            [assignment.reference_summary_id for assignment in ordered_assignments],
            [first_summary.id, second_summary.id],
        )
        self.assertEqual(summary_numbers[first_summary.id], 1)
        self.assertEqual(summary_numbers[second_summary.id], 2)
        self.assertEqual(len(ordered_references), 1)
        self.assertEqual(ordered_references[0][1].id, shared_reference.id)
        self.assertEqual(ordered_references[0][0], [1, 2])
        self.assertEqual(_format_reference_number_ranges([1, 2]), "1-2")

    def test_key_message_rejects_unassigned_supporting_summary(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=self.summary,
            position=1,
        )
        outside_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Outside intervention summary",
            hash_key="hash-outside-summary",
            screening_status="included",
        )
        outside_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=outside_reference,
        )

        response = self.client.post(
            url,
            {
                "action": "add-key-message",
                "intervention_id": intervention.id,
                "response_group": SynopsisInterventionKeyMessage.GROUP_RESPONSE,
                "statement": "Message with invalid supporting study link.",
                "supporting_summaries": [str(outside_summary.id)],
            },
        )
        self.assertEqual(response.status_code, 302)
        self.assertFalse(intervention.key_messages.exists())

    def test_evidence_page_groups_summary_tabs_under_reference(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        self.reference.authors = "Rebecca Smith"
        self.reference.publication_year = 2024
        self.reference.save(update_fields=["authors", "publication_year", "updated_at"])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        alt_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=self.reference,
            action_description="Study B",
        )
        self.summary.action_description = "Study A"
        self.summary.save(update_fields=["action_description"])
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=self.summary,
            position=1,
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=alt_summary,
            position=2,
        )

        response = self.client.get(url)

        self.assertEqual(response.status_code, 200)
        grouped = response.context["reference_summary_groups"]
        self.assertEqual(len(grouped), 1)
        self.assertEqual(grouped[0]["reference_heading"], "Rebecca Smith · 2024")
        self.assertEqual(grouped[0]["reference_context"], "Test ref")
        self.assertEqual(
            [item["summary_display"] for item in grouped[0]["summaries"]],
            ["D1000.a — Study A", "D1000.b — Study B"],
        )
        self.assertContains(response, "Assign summary tabs, not whole papers.")
        self.assertContains(response, "Grouped by reference author.")
        self.assertContains(
            response,
            "Choose a summary tab to preview its reference and tab label.",
        )
        self.assertContains(response, "Rebecca Smith · 2024")
        self.assertContains(response, "D1000.a — Study A")
        self.assertContains(response, "D1000.b — Study B")
        self.assertContains(response, "Same source paper")
        self.assertContains(response, "shared reference line (1-2)")
        self.assertContains(response, "Compilation preview")
        self.assertContains(response, "study paragraphs")
        self.assertContains(response, "source paper")
        self.assertContains(response, "2 summary tabs from the same paper")

    def test_evidence_workspace_uses_action_only_iucn_categories(self):
        response = self.client.get(
            reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        )

        context_categories = list(response.context["iucn_categories"])
        self.assertTrue(context_categories)
        self.assertTrue(
            all(category.kind == IUCNCategory.KIND_ACTION for category in context_categories)
        )
        self.assertIn(
            "Land/water protection-Area protection",
            [category.name for category in context_categories],
        )
        self.assertNotIn(
            "Residential & commercial development",
            [category.name for category in context_categories],
        )

        form_categories = list(
            response.context["intervention_form"].fields["iucn_actions"].queryset
        )
        self.assertTrue(form_categories)
        self.assertTrue(
            all(category.kind == IUCNCategory.KIND_ACTION for category in form_categories)
        )

    def test_update_intervention_metadata_rejects_threat_category(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        threat_category = IUCNCategory.objects.filter(
            kind=IUCNCategory.KIND_THREAT,
            is_active=True,
        ).first()

        self.assertIsNotNone(threat_category)

        response = self.client.post(
            url,
            {
                "action": "update-intervention-metadata",
                "intervention_id": intervention.id,
                "iucn_actions": [str(threat_category.id)],
            },
        )

        self.assertEqual(response.status_code, 302)
        intervention.refresh_from_db()
        self.assertFalse(intervention.iucn_actions.exists())

    def test_create_intervention_allows_multiple_iucn_actions(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(project=self.project, title="Ch", position=1)
        action_categories = list(
            IUCNCategory.objects.filter(
                kind=IUCNCategory.KIND_ACTION,
                is_active=True,
            ).order_by("position", "name")[:2]
        )

        self.assertEqual(len(action_categories), 2)

        response = self.client.post(
            url,
            {
                "action": "create-intervention",
                "chapter_id": chapter.id,
                "title": "Intervention with multiple actions",
                "iucn_actions": [str(category.id) for category in action_categories],
            },
        )

        self.assertEqual(response.status_code, 302)
        intervention = SynopsisIntervention.objects.get(
            subheading__chapter=chapter,
            title="Intervention with multiple actions",
        )
        self.assertEqual(
            list(intervention.iucn_actions.order_by("position", "name").values_list("id", flat=True)),
            [category.id for category in action_categories],
        )

    def test_update_intervention_metadata_allows_multiple_iucn_actions(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        action_categories = list(
            IUCNCategory.objects.filter(
                kind=IUCNCategory.KIND_ACTION,
                is_active=True,
            ).order_by("position", "name")[:2]
        )

        self.assertEqual(len(action_categories), 2)

        response = self.client.post(
            url,
            {
                "action": "update-intervention-metadata",
                "intervention_id": intervention.id,
                "iucn_actions": [str(category.id) for category in action_categories],
            },
        )

        self.assertEqual(response.status_code, 302)
        intervention.refresh_from_db()
        self.assertEqual(
            list(intervention.iucn_actions.order_by("position", "name").values_list("id", flat=True)),
            [category.id for category in action_categories],
        )

    def test_update_intervention_metadata_allows_title_rename(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="Old action title",
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "update-intervention-metadata",
                "intervention_id": intervention.id,
                "title": "Install nest boxes",
            },
        )

        self.assertEqual(response.status_code, 302)
        intervention.refresh_from_db()
        self.assertEqual(intervention.title, "Install nest boxes")

    def test_update_intervention_metadata_rejects_overlong_title(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="Original title",
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "update-intervention-metadata",
                "intervention_id": intervention.id,
                "title": "x" * 256,
            },
            follow=True,
        )

        self.assertContains(
            response,
            "Intervention title must be 255 characters or fewer.",
        )
        intervention.refresh_from_db()
        self.assertEqual(intervention.title, "Original title")

    def test_delete_assignment_removes_supporting_links_from_key_messages(self):
        url = reverse("synopsis:project_synopsis_structure", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        assignment = SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=self.summary,
            position=1,
        )
        key_message = SynopsisInterventionKeyMessage.objects.create(
            intervention=intervention,
            response_group=SynopsisInterventionKeyMessage.GROUP_RESPONSE,
            statement="Linked message",
            position=1,
        )
        key_message.supporting_summaries.add(self.summary)

        response = self.client.post(
            url,
            {
                "action": "delete-assignment",
                "assignment_id": assignment.id,
            },
        )
        self.assertEqual(response.status_code, 302)
        key_message.refresh_from_db()
        self.assertFalse(key_message.supporting_summaries.exists())
        change = ProjectChangeLog.objects.filter(
            project=self.project,
            action="Removed study summary from intervention",
        ).first()
        self.assertIsNotNone(change)
        self.assertIn("Key messages updated: 1", change.details)

    def test_project_hub_shows_structure_audit_entries(self):
        url = reverse("synopsis:project_synopsis_narrative", args=[self.project.id])

        self.client.post(
            url,
            {
                "action": "create-chapter",
                "title": "Executive summary",
                "chapter_type": SynopsisChapter.TYPE_TEXT,
            },
            follow=True,
        )

        response = self.client.get(
            reverse("synopsis:project_hub", args=[self.project.id])
        )

        self.assertContains(response, "Chapter added")
        self.assertContains(response, "Chapter: Executive summary")

    def test_export_citation_and_reference_identifier_override(self):
        reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Corallivorous snail removal",
            authors="Miller M.",
            publication_year=2001,
            journal="Coral Reefs",
            volume="19",
            pages="293-295",
            doi="10.1007/PL00006963",
            hash_key="hash-citation-study",
            screening_status="included",
        )
        summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=reference,
            study_design="replicated, controlled study",
            year_range="1999",
            summary_of_results="snail removal reduced live tissue loss.",
            reference_identifier="X",
        )

        citation = _reference_export_citation(reference)
        self.assertIn("Miller M. (2001)", citation)
        self.assertIn("Corallivorous snail removal.", citation)
        self.assertIn("Coral Reefs, 19, 293-295.", citation)
        self.assertIn("https://doi.org/10.1007/PL00006963", citation)

        paragraph = _structured_summary_paragraph(
            summary, reference_identifier_override="3"
        )
        self.assertIn("(3)", paragraph)
        self.assertNotIn("(X)", paragraph)

    def test_generate_docx_collapses_duplicate_reference_lines_but_keeps_study_paragraphs(self):
        from docx import Document

        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        shared_reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Shared study paper",
            authors="Gamma G.",
            publication_year=2009,
            hash_key="hash-docx-shared-paper",
            screening_status="included",
        )
        first_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=shared_reference,
            study_design="replicated, controlled study",
            year_range="2009",
            summary_of_results="first finding improved coral cover.",
        )
        second_summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=shared_reference,
            study_design="replicated, controlled study",
            year_range="2010",
            summary_of_results="second finding improved coral recruitment.",
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=first_summary,
            position=1,
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=second_summary,
            position=2,
        )

        payload = _generate_synopsis_docx(self.project)
        document = Document(io.BytesIO(payload))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

        self.assertTrue(
            any("(1)" in paragraph and "first finding improved coral cover" in paragraph for paragraph in paragraphs)
        )
        self.assertTrue(
            any("(2)" in paragraph and "second finding improved coral recruitment" in paragraph for paragraph in paragraphs)
        )
        collapsed_reference_lines = [
            paragraph for paragraph in paragraphs if paragraph.startswith("(1-2) ")
        ]
        self.assertEqual(len(collapsed_reference_lines), 1)
        self.assertIn("Shared study paper.", collapsed_reference_lines[0])
        self.assertFalse(
            any(
                paragraph.startswith("(1) Gamma G. (2009) Shared study paper.")
                or paragraph.startswith("(2) Gamma G. (2009) Shared study paper.")
                for paragraph in paragraphs
            )
        )

    def test_generate_docx_uses_summary_citation_override_with_italics(self):
        from docx import Document

        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Shared study paper",
            authors="Gamma G.",
            publication_year=2009,
            journal="Journal of Marine Trials",
            volume="12",
            pages="34-40",
            hash_key="hash-docx-citation-override",
            screening_status="included",
        )
        summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=reference,
            study_design="replicated study",
            year_range="2009",
            summary_of_results="kelp cover improved.",
            citation="Gamma G. (2009) <i>Glipa</i> restoration note.",
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=summary,
            position=1,
        )

        payload = _generate_synopsis_docx(self.project)
        document = Document(io.BytesIO(payload))
        reference_paragraph = next(
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text.startswith("(1) Gamma G. (2009) Glipa restoration note.")
        )

        self.assertEqual(
            reference_paragraph.text,
            "(1) Gamma G. (2009) Glipa restoration note.",
        )
        self.assertTrue(any(run.text == "Glipa" and run.italic for run in reference_paragraph.runs))

    def test_generate_docx_places_key_message_citations_before_full_stop(self):
        from docx import Document

        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Supporting paper",
            authors="Gamma G.",
            publication_year=2009,
            hash_key="hash-key-message-citation",
            screening_status="included",
        )
        summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=reference,
            study_design="replicated study",
            year_range="2009",
            summary_of_results="kelp cover improved.",
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=summary,
            position=1,
        )
        key_message = SynopsisInterventionKeyMessage.objects.create(
            intervention=intervention,
            response_group=SynopsisInterventionKeyMessage.GROUP_RESPONSE,
            statement="Kelp cover improved.",
            position=1,
        )
        key_message.supporting_summaries.set([summary])

        payload = _generate_synopsis_docx(self.project)
        document = Document(io.BytesIO(payload))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

        self.assertIn(
            "General response: Kelp cover improved (1).",
            paragraphs,
        )

    def test_generate_docx_places_formatted_key_message_citation_before_full_stop(self):
        from docx import Document

        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
        )
        reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Supporting paper",
            authors="Gamma G.",
            publication_year=2009,
            hash_key="hash-key-message-markup-citation",
            screening_status="included",
        )
        summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=reference,
            study_design="replicated study",
            year_range="2009",
            summary_of_results="kelp cover improved.",
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=summary,
            position=1,
        )
        key_message = SynopsisInterventionKeyMessage.objects.create(
            intervention=intervention,
            response_group=SynopsisInterventionKeyMessage.GROUP_RESPONSE,
            statement="<i>Festuca.</i>",
            position=1,
        )
        key_message.supporting_summaries.set([summary])

        payload = _generate_synopsis_docx(self.project)
        document = Document(io.BytesIO(payload))
        paragraph = next(
            item
            for item in document.paragraphs
            if item.text == "General response: Festuca (1)."
        )

        self.assertTrue(any(run.text == "Festuca" and run.italic for run in paragraph.runs))
        self.assertFalse(
            any(run.text == "." and run.italic for run in paragraph.runs)
        )

    def test_generate_docx_places_no_studies_before_background_and_labels_background(self):
        from docx import Document

        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
            background_text="Chapter background text.",
            background_references="Chapter Background Ref",
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
            evidence_status=SynopsisIntervention.EVIDENCE_STATUS_NO_STUDIES,
            background_text="Intervention background text.",
            background_references="Intervention Background Ref",
        )

        payload = _generate_synopsis_docx(self.project)
        document = Document(io.BytesIO(payload))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

        self.assertGreaterEqual(paragraphs.count("Background"), 2)

        chapter_title_index = paragraphs.index("2. Threat: Demo")
        first_background_index = paragraphs.index("Background", chapter_title_index + 1)
        self.assertEqual(paragraphs[first_background_index + 1], "Chapter background text.")
        self.assertEqual(
            paragraphs[first_background_index + 2],
            "Background references: Chapter Background Ref",
        )

        intervention_title_index = paragraphs.index("2.1 Demo intervention")
        no_studies_index = paragraphs.index(
            "We found no studies that evaluated the effects of this intervention."
        )
        intervention_background_index = paragraphs.index(
            "Background", no_studies_index + 1
        )
        self.assertLess(intervention_title_index, no_studies_index)
        self.assertLess(no_studies_index, intervention_background_index)
        self.assertEqual(
            paragraphs[intervention_background_index + 1],
            "Intervention background text.",
        )
        self.assertEqual(
            paragraphs[intervention_background_index + 2],
            "Background references: Intervention Background Ref",
        )

    def test_generate_docx_does_not_label_narrative_chapter_text_as_background(self):
        from docx import Document

        SynopsisChapter.objects.create(
            project=self.project,
            title="1. Introduction",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=1,
            background_text="This is the narrative introduction.",
        )

        payload = _generate_synopsis_docx(self.project)
        document = Document(io.BytesIO(payload))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

        self.assertIn("1. Introduction", paragraphs)
        self.assertIn("This is the narrative introduction.", paragraphs)
        self.assertNotIn("Background", paragraphs)

    def test_generate_docx_preserves_inline_markup_in_backgrounds_key_messages_and_summary_paragraphs(self):
        from docx import Document

        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
            background_text="Forest <i>edge</i> with CO<sub>2</sub> buildup.",
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Demo intervention",
            position=1,
            background_text="Reached the 10<sup>th</sup> plot.",
        )
        reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Supporting paper",
            authors="Gamma G.",
            publication_year=2009,
            hash_key="hash-inline-markup-export",
            screening_status="included",
        )
        summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=reference,
            synopsis_draft=(
                "A replicated study found that <i>Festuca</i> cover increased at plot 10<sup>th</sup>."
            ),
            use_custom_synopsis_draft=True,
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=summary,
            position=1,
        )
        SynopsisInterventionKeyMessage.objects.create(
            intervention=intervention,
            response_group=SynopsisInterventionKeyMessage.GROUP_RESPONSE,
            statement="CO<sub>2</sub> uptake improved in the 10<sup>th</sup> plot.",
            position=1,
        )

        payload = _generate_synopsis_docx(self.project)
        document = Document(io.BytesIO(payload))

        chapter_background = next(
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text == "Forest edge with CO2 buildup."
        )
        intervention_background = next(
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text == "Reached the 10th plot."
        )
        key_message_paragraph = next(
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text == "General response: CO2 uptake improved in the 10th plot."
        )
        summary_paragraph = next(
            paragraph
            for paragraph in document.paragraphs
            if "Festuca cover increased at plot 10th." in paragraph.text
        )

        self.assertTrue(any(run.text == "edge" and run.italic for run in chapter_background.runs))
        self.assertTrue(
            any(run.text == "2" and run.font.subscript for run in chapter_background.runs)
        )
        self.assertTrue(
            any(run.text == "th" and run.font.superscript for run in intervention_background.runs)
        )
        self.assertTrue(
            any(run.text == "2" and run.font.subscript for run in key_message_paragraph.runs)
        )
        self.assertTrue(
            any(run.text == "th" and run.font.superscript for run in key_message_paragraph.runs)
        )
        self.assertTrue(any(run.text == "Festuca" and run.italic for run in summary_paragraph.runs))
        self.assertTrue(
            any(run.text == "th" and run.font.superscript for run in summary_paragraph.runs)
        )

    def test_generate_docx_ignores_legacy_cross_reference_flags(self):
        from docx import Document

        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=1,
        )
        subheading = SynopsisSubheading.objects.create(
            chapter=chapter,
            title="Interventions",
            position=1,
        )
        primary = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.1 Primary intervention",
            position=1,
        )
        intervention = SynopsisIntervention.objects.create(
            subheading=subheading,
            title="2.2 Secondary intervention",
            position=2,
            is_cross_reference=True,
            primary_intervention=primary,
        )
        reference = Reference.objects.create(
            project=self.project,
            batch=self.batch,
            title="Legacy cross reference study",
            authors="Delta D.",
            publication_year=2012,
            hash_key="hash-legacy-cross-ref",
            screening_status="included",
        )
        summary = ReferenceSummary.objects.create(
            project=self.project,
            reference=reference,
            study_design="replicated study",
            year_range="2012",
            summary_of_results="evidence still exports under the intervention.",
        )
        SynopsisAssignment.objects.create(
            intervention=intervention,
            reference_summary=summary,
            position=1,
        )

        payload = _generate_synopsis_docx(self.project)
        document = Document(io.BytesIO(payload))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

        self.assertIn("2.2 Secondary intervention", paragraphs)
        self.assertTrue(
            any(
                "evidence still exports under the intervention." in paragraph
                for paragraph in paragraphs
            )
        )
        self.assertFalse(
            any(paragraph.startswith("Cross-reference: Evidence is summarized under") for paragraph in paragraphs)
        )

    def test_workspace_routes_load(self):
        narrative_url = reverse(
            "synopsis:project_synopsis_narrative", args=[self.project.id]
        )
        evidence_url = reverse(
            "synopsis:project_synopsis_evidence", args=[self.project.id]
        )
        structure_url = reverse(
            "synopsis:project_synopsis_structure", args=[self.project.id]
        )
        narrative_response = self.client.get(narrative_url)
        evidence_response = self.client.get(evidence_url)
        structure_response = self.client.get(structure_url)

        self.assertEqual(narrative_response.status_code, 200)
        self.assertEqual(evidence_response.status_code, 200)
        self.assertEqual(structure_response.status_code, 200)
        self.assertContains(
            narrative_response, "synopsis-workspace-container", html=False
        )
        self.assertContains(narrative_response, "Back to summary workspace")
        self.assertContains(
            evidence_response, "synopsis-workspace-container", html=False
        )
        self.assertContains(
            structure_response, "synopsis-workspace-container", html=False
        )

    def test_narrative_workspace_post_redirects_to_narrative_route(self):
        url = reverse("synopsis:project_synopsis_narrative", args=[self.project.id])
        response = self.client.post(
            url,
            {
                "action": "create-chapter",
                "title": "Advisory Board",
                "chapter_type": SynopsisChapter.TYPE_TEXT,
            },
        )
        self.assertEqual(response.status_code, 302)
        self.assertEqual(
            urlparse(response["Location"]).path,
            urlparse(url).path,
        )

    def test_narrative_workspace_create_chapter_does_not_auto_number(self):
        url = reverse("synopsis:project_synopsis_narrative", args=[self.project.id])
        SynopsisChapter.objects.create(
            project=self.project,
            title="5. Existing chapter",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "create-chapter",
                "title": "Executive summary",
                "chapter_type": SynopsisChapter.TYPE_TEXT,
            },
            follow=True,
        )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(
            SynopsisChapter.objects.filter(
                project=self.project,
                title="Executive summary",
                chapter_type=SynopsisChapter.TYPE_TEXT,
            ).exists()
        )
        self.assertContains(response, "Executive summary")

    def test_narrative_workspace_limits_chapter_type_choices_to_narrative_and_appendix(self):
        url = reverse("synopsis:project_synopsis_narrative", args=[self.project.id])
        response = self.client.get(url)

        self.assertContains(response, "Narrative chapter")
        self.assertContains(response, "Appendix / back matter")
        self.assertContains(response, "Evidence chapters are created in the Evidence workspace.")
        self.assertNotContains(response, ">Evidence chapter<", html=False)

        invalid_response = self.client.post(
            url,
            {
                "action": "create-chapter",
                "title": "Should fail",
                "chapter_type": SynopsisChapter.TYPE_EVIDENCE,
            },
        )

        self.assertEqual(invalid_response.status_code, 200)
        self.assertFalse(
            SynopsisChapter.objects.filter(project=self.project, title="Should fail").exists()
        )

    def test_narrative_workspace_renders_restore_state_hooks(self):
        url = reverse("synopsis:project_synopsis_narrative", args=[self.project.id])
        SynopsisChapter.objects.create(
            project=self.project,
            title="Advisory Board",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=1,
        )

        response = self.client.get(url)

        self.assertContains(response, 'id="synopsis-narrative-page"', html=False)
        self.assertContains(response, "cePreservePageState({", html=False)
        self.assertContains(
            response,
            f'"synopsis-narrative-state-{self.project.id}"',
            html=False,
        )

    def test_narrative_workspace_renders_saved_chapter_text_preview(self):
        url = reverse("synopsis:project_synopsis_narrative", args=[self.project.id])
        SynopsisChapter.objects.create(
            project=self.project,
            title="Advisory Board",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            background_text="<i>Italic</i> chapter text",
            position=1,
        )

        response = self.client.get(url)

        self.assertContains(response, "Saved preview")
        self.assertContains(
            response,
            '<span class="ce-inline-markup-rendered"><i>Italic</i> chapter text</span>',
            html=False,
        )

    def test_narrative_workspace_updates_chapter_title(self):
        url = reverse("synopsis:project_synopsis_narrative", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="Advisory Board",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "update-chapter-title",
                "chapter_id": chapter.id,
                "title": "3. Editorial Board",
            },
        )

        self.assertEqual(response.status_code, 302)
        self.assertEqual(
            urlparse(response["Location"]).path,
            urlparse(url).path,
        )
        chapter.refresh_from_db()
        self.assertEqual(chapter.title, "Editorial Board")

    def test_narrative_workspace_drops_stray_legacy_number_from_non_intro_title(self):
        url = reverse("synopsis:project_synopsis_narrative", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="1. Advisory Board",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "update-chapter-title",
                "chapter_id": chapter.id,
                "title": "Editorial Board",
            },
        )

        self.assertEqual(response.status_code, 302)
        chapter.refresh_from_db()
        self.assertEqual(chapter.title, "Editorial Board")

    def test_narrative_workspace_renumbers_evidence_when_reserved_intro_is_removed(self):
        url = reverse("synopsis:project_synopsis_narrative", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="1. About this book",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=1,
        )
        evidence = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=2,
        )

        response = self.client.post(
            url,
            {
                "action": "update-chapter-title",
                "chapter_id": chapter.id,
                "title": "Editorial Board",
            },
        )

        self.assertEqual(response.status_code, 302)
        chapter.refresh_from_db()
        evidence.refresh_from_db()
        self.assertEqual(chapter.title, "Editorial Board")
        self.assertEqual(evidence.title, "1. Threat: Demo")

    def test_narrative_workspace_preserves_reserved_intro_number_when_renamed(self):
        url = reverse("synopsis:project_synopsis_narrative", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="1. About this book",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "update-chapter-title",
                "chapter_id": chapter.id,
                "title": "About the book",
            },
        )

        self.assertEqual(response.status_code, 302)
        chapter.refresh_from_db()
        self.assertEqual(chapter.title, "1. About the book")

    def test_narrative_workspace_rejects_evidence_type_change_posts(self):
        url = reverse("synopsis:project_synopsis_narrative", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="References",
            chapter_type=SynopsisChapter.TYPE_APPENDIX,
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "update-chapter-type",
                "chapter_id": chapter.id,
                "chapter_type": SynopsisChapter.TYPE_EVIDENCE,
            },
            follow=True,
        )

        self.assertEqual(response.status_code, 200)
        chapter.refresh_from_db()
        self.assertEqual(chapter.chapter_type, SynopsisChapter.TYPE_APPENDIX)
        self.assertContains(
            response,
            "Selected chapter type is not available in this workspace.",
        )

    def test_narrative_workspace_repositions_chapter_when_type_changes(self):
        url = reverse("synopsis:project_synopsis_narrative", args=[self.project.id])
        front = SynopsisChapter.objects.create(
            project=self.project,
            title="Advisory Board",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=1,
        )
        evidence = SynopsisChapter.objects.create(
            project=self.project,
            title="2. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=2,
        )
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="References",
            chapter_type=SynopsisChapter.TYPE_APPENDIX,
            position=3,
        )

        response = self.client.post(
            url,
            {
                "action": "update-chapter-type",
                "chapter_id": chapter.id,
                "chapter_type": SynopsisChapter.TYPE_TEXT,
            },
        )

        self.assertEqual(response.status_code, 302)
        front.refresh_from_db()
        evidence.refresh_from_db()
        chapter.refresh_from_db()
        self.assertEqual(chapter.chapter_type, SynopsisChapter.TYPE_TEXT)
        self.assertEqual(front.position, 1)
        self.assertEqual(chapter.position, 2)
        self.assertEqual(evidence.position, 3)

    def test_narrative_workspace_move_chapter_does_not_cross_into_evidence_section(self):
        url = reverse("synopsis:project_synopsis_narrative", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="Advisory Board",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=1,
        )
        evidence = SynopsisChapter.objects.create(
            project=self.project,
            title="1. Threat: Demo",
            chapter_type=SynopsisChapter.TYPE_EVIDENCE,
            position=2,
        )
        appendix = SynopsisChapter.objects.create(
            project=self.project,
            title="References",
            chapter_type=SynopsisChapter.TYPE_APPENDIX,
            position=3,
        )

        response = self.client.post(
            url,
            {
                "action": "move-chapter",
                "chapter_id": chapter.id,
                "direction": "down",
            },
            follow=True,
        )

        self.assertEqual(response.status_code, 200)
        chapter.refresh_from_db()
        evidence.refresh_from_db()
        appendix.refresh_from_db()
        self.assertEqual(chapter.position, 1)
        self.assertEqual(evidence.position, 2)
        self.assertEqual(appendix.position, 3)
        self.assertContains(response, "Already at the edge.")

    def test_narrative_workspace_requires_confirmation_to_delete_chapter(self):
        url = reverse("synopsis:project_synopsis_narrative", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="Advisory Board",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=1,
        )

        response = self.client.get(url)

        self.assertContains(
            response,
            'id="delete-narrative-chapter-modal"',
            html=False,
        )
        self.assertContains(
            response,
            f'data-chapter-id="{chapter.id}"',
            count=1,
            html=False,
        )
        self.assertContains(response, 'value="move-chapter"', count=2, html=False)
        self.assertContains(response, f'href="#chapter-{chapter.id}"', html=False)
        self.assertContains(
            response,
            'name="confirm_delete_chapter" value="1"',
            html=False,
        )
        self.assertContains(
            response,
            "This removes the chapter and its authored text from the synopsis outline and DOCX export.",
        )

    def test_narrative_workspace_requires_confirmation_flag_to_delete_chapter(self):
        url = reverse("synopsis:project_synopsis_narrative", args=[self.project.id])
        chapter = SynopsisChapter.objects.create(
            project=self.project,
            title="Advisory Board",
            chapter_type=SynopsisChapter.TYPE_TEXT,
            position=1,
        )

        response = self.client.post(
            url,
            {
                "action": "delete-chapter",
                "chapter_id": chapter.id,
            },
            follow=True,
        )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(SynopsisChapter.objects.filter(pk=chapter.id).exists())
        self.assertContains(
            response,
            "Please confirm chapter deletion before removing it.",
        )


class ProjectDeleteFormTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Wetland Recovery")

    def test_requires_matching_title(self):
        form = ProjectDeleteForm(
            data={
                "confirm_title": "Wrong title",
                "acknowledge_irreversible": True,
            },
            project=self.project,
        )
        self.assertFalse(form.is_valid())
        self.assertIn("Title does not match", form.errors["confirm_title"][0])

    def test_requires_acknowledgement(self):
        form = ProjectDeleteForm(
            data={"confirm_title": "Wetland Recovery"}, project=self.project
        )
        self.assertFalse(form.is_valid())
        self.assertIn(
            "This field is required.", form.errors["acknowledge_irreversible"][0]
        )

    def test_valid_when_all_checks_pass(self):
        form = ProjectDeleteForm(
            data={
                "confirm_title": "Wetland Recovery",
                "acknowledge_irreversible": True,
            },
            project=self.project,
        )
        self.assertTrue(form.is_valid())


class ProjectSettingsFormTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Forest Restoration")

    def test_requires_title(self):
        form = ProjectSettingsForm(data={"title": ""}, instance=self.project)
        self.assertFalse(form.is_valid())
        self.assertIn("Enter a title", form.errors["title"][0])

    def test_updates_title(self):
        form = ProjectSettingsForm(
            data={"title": "Forest Recovery"}, instance=self.project
        )
        self.assertTrue(form.is_valid())
        updated = form.save()
        self.assertEqual(updated.title, "Forest Recovery")

    def test_updates_description(self):
        form = ProjectSettingsForm(
            data={
                "title": "Forest Recovery",
                "description": "  A pilot synopsis for forest restoration.  ",
            },
            instance=self.project,
        )
        self.assertTrue(form.is_valid())
        updated = form.save()
        self.assertEqual(updated.description, "A pilot synopsis for forest restoration.")


class ViewHelperTests(TestCase):
    def setUp(self):
        self.project = Project.objects.create(title="Helper Project")

    def test_user_is_manager_for_staff(self):
        user = User.objects.create_user(
            username="staffer", password="pw", is_staff=True
        )
        self.assertTrue(_user_is_manager(user))

    def test_user_is_manager_for_group_member(self):
        group, _ = Group.objects.get_or_create(name="manager")
        user = User.objects.create_user(username="manager_user")
        user.groups.add(group)
        self.assertTrue(_user_is_manager(user))

    def test_user_is_manager_false_for_others(self):
        user = User.objects.create_user(username="regular")
        self.assertFalse(_user_is_manager(user))

    def test_log_project_change_records_entry(self):
        user = User.objects.create_user(username="logger")
        _log_project_change(self.project, user, "Edited", "Updated title")
        entry = ProjectChangeLog.objects.get(project=self.project)
        self.assertEqual(entry.action, "Edited")
        self.assertEqual(entry.details, "Updated title")
        self.assertEqual(entry.changed_by, user)

    def test_log_project_change_anonymous(self):
        anonymous = AnonymousUser()
        _log_project_change(self.project, anonymous, "Edited", "Updated title")
        entry = (
            ProjectChangeLog.objects.filter(project=self.project)
            .order_by("-id")
            .first()
        )
        self.assertIsNone(entry.changed_by)

    def test_format_value_handles_none_and_date(self):
        self.assertEqual(_format_value(None), "—")
        today = date(2025, 1, 2)
        self.assertEqual(_format_value(today), "2025-01-02")
        self.assertEqual(_format_value(42), "42")

    def test_funder_contact_label(self):
        self.assertEqual(_funder_contact_label("Ann", "Thornton"), "Ann Thornton")
        self.assertEqual(_funder_contact_label("", ""), "—")

    def test_format_deadline_formats_timezone(self):
        aware = timezone.make_aware(datetime(2025, 7, 1, 15, 0))
        formatted = _format_deadline(aware)
        self.assertEqual(
            formatted, timezone.localtime(aware).strftime("%d %b %Y %H:%M")
        )
        self.assertEqual(_format_deadline(None), "—")

    def test_split_inline_italic_markup_supports_simple_i_tags(self):
        segments = split_inline_italic_markup(
            "Gamma G. (2009) <i>Glipa</i> restoration note."
        )

        self.assertEqual(
            segments,
            [
                ("Gamma G. (2009) ", False),
                ("Glipa", True),
                (" restoration note.", False),
            ],
        )

    def test_split_inline_markup_supports_subscript_and_superscript(self):
        segments = split_inline_markup(
            "CO<sub>2</sub> rose in the 10<sup>th</sup> plot with <i>Festuca</i>."
        )

        self.assertEqual(
            segments,
            [
                InlineMarkupSegment("CO"),
                InlineMarkupSegment("2", subscript=True),
                InlineMarkupSegment(" rose in the 10"),
                InlineMarkupSegment("th", superscript=True),
                InlineMarkupSegment(" plot with "),
                InlineMarkupSegment("Festuca", italic=True),
                InlineMarkupSegment("."),
            ],
        )

    def test_format_inline_markup_html_escapes_unknown_tags_but_keeps_supported_formatting(self):
        rendered = format_inline_markup_html(
            "CO<sub>2</sub> and <i>Festuca</i><script>alert(1)</script>"
        )

        self.assertEqual(
            rendered,
            "CO<sub>2</sub> and <i>Festuca</i>&lt;script&gt;alert(1)&lt;/script&gt;",
        )

    def test_format_inline_markup_html_keeps_entity_escaped_supported_tags_literal(self):
        rendered = format_inline_markup_html("&lt;sup&gt;literal&lt;/sup&gt;")

        self.assertEqual(rendered, "&lt;sup&gt;literal&lt;/sup&gt;")

    def test_background_form_rejects_unclosed_supported_inline_tag(self):
        form = SynopsisBackgroundForm(
            data={"background_text": "CO<sup>2", "background_references": ""}
        )

        self.assertFalse(form.is_valid())
        self.assertIn(
            "matching closing tag",
            form.errors["background_text"][0],
        )

    def test_key_message_form_rejects_nested_subscript_and_superscript(self):
        form = SynopsisKeyMessageForm(
            data={
                "response_group": SynopsisInterventionKeyMessage.GROUP_RESPONSE,
                "outcome_label": "",
                "study_count": "",
                "statement": "<sub><sup>x</sup></sub>",
                "supporting_summaries": [],
            }
        )

        self.assertFalse(form.is_valid())
        self.assertIn(
            "cannot be nested inside one another",
            form.errors["statement"][0],
        )

    def test_user_can_confirm_phase(self):
        staff = User.objects.create_user(username="staff", is_staff=True)
        self.assertTrue(_user_can_confirm_phase(staff, self.project))
        author = User.objects.create_user(username="author")
        UserRole.objects.create(user=author, project=self.project, role="author")
        self.assertTrue(_user_can_confirm_phase(author, self.project))
        manager = User.objects.create_user(username="manager")
        UserRole.objects.create(user=manager, project=self.project, role="manager")
        self.assertTrue(_user_can_confirm_phase(manager, self.project))
        outsider = User.objects.create_user(username="outsider")
        self.assertFalse(_user_can_confirm_phase(outsider, self.project))
        self.assertFalse(_user_can_confirm_phase(AnonymousUser(), self.project))


class ProjectDescriptionUiTests(TestCase):
    def setUp(self):
        self.manager = User.objects.create_user(
            username="manager",
            password="pass123",
            is_staff=True,
        )
        self.project = Project.objects.create(
            title="Forest Restoration",
            description="A pilot synopsis for forest restoration.",
        )

    def test_project_hub_shows_description_when_present(self):
        self.client.login(username="manager", password="pass123")

        response = self.client.get(reverse("synopsis:project_hub", args=[self.project.id]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Synopsis overview")
        self.assertContains(response, "Synopsis settings")
        self.assertContains(response, "Description")
        self.assertContains(response, "A pilot synopsis for forest restoration.")

    def test_project_settings_shows_description_field(self):
        self.client.login(username="manager", password="pass123")

        response = self.client.get(
            reverse("synopsis:project_settings", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "optional description")
        self.assertContains(response, "Phase tracker")
        self.assertContains(response, "A pilot synopsis for forest restoration.")
        self.assertContains(response, 'value="Forest Restoration"', html=False)

    def test_project_settings_shows_protocol_and_advisory_relevance_fields(self):
        self.client.login(username="manager", password="pass123")

        response = self.client.get(
            reverse("synopsis:project_settings", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Protocol is relevant for this project")
        self.assertContains(response, "Advisory board is relevant for this project")

    def test_project_pages_show_back_to_project_button_in_nav(self):
        self.client.login(username="manager", password="pass123")

        response = self.client.get(
            reverse("synopsis:project_authors_manage", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, ">Back to project</a>", html=False)


class ProjectHomepageStatusUiTests(TestCase):
    def setUp(self):
        self.author = User.objects.create_user(
            username="status-author",
            password="pass123",
        )
        self.project = Project.objects.create(title="Status Managed Synopsis")
        UserRole.objects.create(user=self.author, project=self.project, role="author")
        self.client.login(username="status-author", password="pass123")

    def test_project_settings_shows_homepage_listing_controls(self):
        response = self.client.get(
            reverse("synopsis:project_settings", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Homepage listing")
        self.assertContains(response, "Shown under active synopses")
        self.assertContains(response, "Mark as completed / archived")
        self.assertNotContains(response, "Archive synopsis")

    def test_author_can_move_synopsis_to_completed_section_without_locking_it(self):
        response = self.client.post(
            reverse("synopsis:project_settings", args=[self.project.id]),
            {"status_action": "mark_completed"},
            follow=True,
        )

        self.assertRedirects(
            response, reverse("synopsis:project_settings", args=[self.project.id])
        )
        self.project.refresh_from_db()
        self.assertEqual(self.project.status, "completed")
        self.assertContains(response, "Shown under completed / archived synopses")
        self.assertContains(response, "Move back to active")
        self.assertTrue(
            ProjectChangeLog.objects.filter(
                project=self.project,
                action="Updated project status",
                details="Status: Planning → Completed",
            ).exists()
        )

        dashboard_response = self.client.get(reverse("synopsis:dashboard"))
        self.assertNotIn(self.project, dashboard_response.context["active_projects"])
        self.assertIn(self.project, dashboard_response.context["completed_projects"])
        self.assertContains(
            dashboard_response,
            reverse("synopsis:project_hub", args=[self.project.id]),
            html=False,
        )
        self.assertContains(dashboard_response, "Move to active")

        hub_response = self.client.get(
            reverse("synopsis:project_hub", args=[self.project.id])
        )
        self.assertEqual(hub_response.status_code, 200)
        self.assertContains(hub_response, self.project.title)

    def test_author_can_move_completed_synopsis_back_to_active_section(self):
        self.project.status = "completed"
        self.project.save(update_fields=["status"])

        response = self.client.post(
            reverse("synopsis:project_settings", args=[self.project.id]),
            {"status_action": "reactivate"},
            follow=True,
        )

        self.assertRedirects(
            response, reverse("synopsis:project_settings", args=[self.project.id])
        )
        self.project.refresh_from_db()
        self.assertEqual(self.project.status, "active")
        self.assertContains(response, "Shown under active synopses")
        self.assertTrue(
            ProjectChangeLog.objects.filter(
                project=self.project,
                action="Updated project status",
                details="Status: Completed → Active",
            ).exists()
        )

        dashboard_response = self.client.get(reverse("synopsis:dashboard"))
        self.assertIn(self.project, dashboard_response.context["active_projects"])
        self.assertNotIn(self.project, dashboard_response.context["completed_projects"])

    def test_author_can_move_completed_synopsis_back_to_active_from_dashboard_row(self):
        self.project.status = "completed"
        self.project.save(update_fields=["status"])

        response = self.client.post(
            reverse("synopsis:project_settings", args=[self.project.id]),
            {"status_action": "reactivate", "return_to": "dashboard"},
            follow=True,
        )

        self.assertRedirects(response, reverse("synopsis:dashboard"))
        self.project.refresh_from_db()
        self.assertEqual(self.project.status, "active")
        self.assertIn(self.project, response.context["active_projects"])
        self.assertNotIn(self.project, response.context["completed_projects"])


class ProjectPhaseUiTests(TestCase):
    def setUp(self):
        self.author = User.objects.create_user(
            username="phase-author",
            password="pass123",
        )
        self.project = Project.objects.create(title="Phase Tracker")
        UserRole.objects.create(user=self.author, project=self.project, role="author")
        self.client.login(username="phase-author", password="pass123")

    def test_project_hub_shows_phase_summary_and_shortcut_controls(self):
        response = self.client.get(
            reverse("synopsis:project_hub", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Move to invite advisory board")
        self.assertContains(response, "Manage phase tracker")
        self.assertContains(response, "Default starting phase")
        self.assertContains(response, "Step 1 of 8")
        self.assertNotContains(response, "Set current phase")

    def test_phase_tracker_skips_protocol_and_advisory_when_not_relevant(self):
        self.project.protocol_relevant = False
        self.project.advisory_board_relevant = False
        self.project.save(update_fields=["protocol_relevant", "advisory_board_relevant"])

        response = self.client.get(
            reverse("synopsis:project_hub", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "References screening")
        self.assertContains(response, "Move to summary writing")
        self.assertContains(response, "Step 1 of 6")
        self.assertNotContains(response, "Move to draft protocol")
        self.assertNotContains(response, "Move to invite advisory board")

    def test_protocol_and_advisory_cards_show_not_relevant_state(self):
        self.project.protocol_relevant = False
        self.project.advisory_board_relevant = False
        self.project.save(update_fields=["protocol_relevant", "advisory_board_relevant"])

        response = self.client.get(
            reverse("synopsis:project_hub", args=[self.project.id])
        )

        self.assertContains(response, "Protocol")
        self.assertContains(response, "This synopsis is not using the protocol workflow in the portal.")
        self.assertContains(response, "Advisory Board")
        self.assertContains(response, "This synopsis is not using an advisory board in the portal.")

    def test_project_settings_shows_full_phase_tracker_controls(self):
        response = self.client.get(
            reverse("synopsis:project_settings", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Phase tracker")
        self.assertContains(response, "Set current phase")
        self.assertContains(response, "Default starting phase")

    def test_author_can_set_phase_backwards_or_forwards(self):
        self.project.phase_manual = "summary_writing"
        self.project.phase_manual_updated = timezone.now()
        self.project.save(update_fields=["phase_manual", "phase_manual_updated"])

        response = self.client.post(
            reverse(
                "synopsis:project_phase_confirm",
                args=[self.project.id, "draft_protocol"],
            )
        )

        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.project.id])
        )
        self.project.refresh_from_db()
        self.assertEqual(self.project.phase_manual, "draft_protocol")
        self.assertEqual(self.project.phase, "draft_protocol")
        event = self.project.phase_events.first()
        self.assertIsNotNone(event)
        self.assertEqual(event.phase, "draft_protocol")

    def test_project_settings_shows_phase_history_entries(self):
        self.client.post(
            reverse(
                "synopsis:project_phase_confirm",
                args=[self.project.id, "invite_advisory_board"],
            )
        )

        response = self.client.get(
            reverse("synopsis:project_settings", args=[self.project.id])
        )

        self.assertContains(response, "Phase history")
        self.assertContains(response, "Invite advisory board")
        self.assertContains(
            response,
            "Phase changed from Draft protocol to Invite advisory board.",
        )

    def test_project_hub_shortcut_moves_to_next_phase(self):
        response = self.client.post(
            reverse(
                "synopsis:project_phase_confirm",
                args=[self.project.id, "invite_advisory_board"],
            )
        )

        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.project.id])
        )
        self.project.refresh_from_db()
        self.assertEqual(self.project.phase, "invite_advisory_board")

    def test_phase_updates_are_logged_in_recent_changes(self):
        response = self.client.post(
            reverse(
                "synopsis:project_phase_confirm",
                args=[self.project.id, "invite_advisory_board"],
            )
        )

        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.project.id])
        )
        change = ProjectChangeLog.objects.filter(
            project=self.project,
            action="Updated project phase",
        ).first()
        self.assertIsNotNone(change)
        self.assertIn("Draft protocol", change.details)
        self.assertIn("Invite advisory board", change.details)

    def test_project_hub_shows_revision_history_timeline(self):
        self.client.post(
            reverse(
                "synopsis:project_phase_confirm",
                args=[self.project.id, "invite_advisory_board"],
            )
        )

        response = self.client.get(
            reverse("synopsis:project_hub", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Revision history")
        self.assertContains(response, "Phase confirmed: Invite advisory board")
        self.assertContains(response, "Updated project phase")
        self.assertContains(response, "Draft protocol")
        self.assertContains(response, "phase-author")

    def test_project_hub_normalises_collaborative_history_entries(self):
        ProjectChangeLog.objects.create(
            project=self.project,
            changed_by=self.author,
            action="Protocol collaborative session closed",
            details="Session 123e4567-e89b-12d3-a456-426614174000 closed (status 3).",
        )
        ProjectChangeLog.objects.create(
            project=self.project,
            changed_by=self.author,
            action="Protocol updated via collaborative edit",
            details=(
                "Session: 123e4567-e89b-12d3-a456-426614174000 | "
                "Status: 6 | File: protocol-v2.docx | Users: phase-author | "
                "Size: 24.0 KB | Reason: Updated references section"
            ),
        )

        response = self.client.get(
            reverse("synopsis:project_hub", args=[self.project.id])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response, "Protocol collaborative session ended"
        )
        self.assertContains(
            response, "Closed without additional document changes."
        )
        self.assertContains(
            response, "Protocol revision saved from collaborative editing"
        )
        self.assertContains(response, "Saved file: protocol-v2.docx")
        self.assertContains(response, "Revision note: Updated references section")
        self.assertNotContains(
            response, "123e4567-e89b-12d3-a456-426614174000"
        )
        self.assertNotContains(response, "status 3")

    def test_manager_role_can_update_phase(self):
        manager = User.objects.create_user(username="phase-manager", password="pass123")
        UserRole.objects.create(user=manager, project=self.project, role="manager")
        self.client.login(username="phase-manager", password="pass123")

        response = self.client.post(
            reverse(
                "synopsis:project_phase_confirm",
                args=[self.project.id, "draft_synopsis"],
            )
        )

        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.project.id])
        )
        self.project.refresh_from_db()
        self.assertEqual(self.project.phase, "draft_synopsis")

    def test_cannot_set_phase_to_disabled_step(self):
        self.project.protocol_relevant = False
        self.project.save(update_fields=["protocol_relevant"])

        response = self.client.post(
            reverse(
                "synopsis:project_phase_confirm",
                args=[self.project.id, "draft_protocol"],
            )
        )

        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.project.id])
        )
        messages = [message.message for message in get_messages(response.wsgi_request)]
        self.assertIn("That phase is not available for this project.", messages)


class ProjectWorkflowApplicabilityTests(TestCase):
    def setUp(self):
        self.author = User.objects.create_user(
            username="workflow-author",
            password="pass123",
        )
        self.project = Project.objects.create(title="Workflow flexibility")
        UserRole.objects.create(user=self.author, project=self.project, role="author")
        self.client.login(username="workflow-author", password="pass123")

    def test_project_settings_can_mark_protocol_and_advisory_not_relevant(self):
        response = self.client.post(
            reverse("synopsis:project_settings", args=[self.project.id]),
            {
                "title": self.project.title,
                "description": "",
                "protocol_relevant": "",
                "advisory_board_relevant": "",
            },
        )

        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.project.id])
        )
        self.project.refresh_from_db()
        self.assertFalse(self.project.protocol_relevant)
        self.assertFalse(self.project.advisory_board_relevant)
        change = ProjectChangeLog.objects.filter(
            project=self.project, action="Updated project settings"
        ).first()
        self.assertIsNotNone(change)
        self.assertIn("Protocol: relevant → not relevant", change.details)
        self.assertIn("Advisory board: relevant → not relevant", change.details)

    def test_protocol_page_redirects_when_protocol_not_relevant(self):
        self.project.protocol_relevant = False
        self.project.save(update_fields=["protocol_relevant"])

        response = self.client.get(
            reverse("synopsis:protocol_detail", args=[self.project.id])
        )

        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.project.id])
        )
        messages = [message.message for message in get_messages(response.wsgi_request)]
        self.assertIn(
            "Protocol is marked as not relevant for this project. Update Project settings if you want to use the protocol workflow.",
            messages,
        )

    def test_advisory_board_page_redirects_when_not_relevant(self):
        self.project.advisory_board_relevant = False
        self.project.save(update_fields=["advisory_board_relevant"])

        response = self.client.get(
            reverse("synopsis:advisory_board_list", args=[self.project.id])
        )

        self.assertRedirects(
            response, reverse("synopsis:project_hub", args=[self.project.id])
        )
        messages = [message.message for message in get_messages(response.wsgi_request)]
        self.assertIn(
            "Advisory board is marked as not relevant for this project. Update Project settings if you want to use the advisory board workflow.",
            messages,
        )
